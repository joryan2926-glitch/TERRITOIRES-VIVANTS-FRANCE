from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image

ROOT = Path(r'C:\Users\jowst\Documents\TERRITOIRES VIVANTS FRANCE')
OUT = ROOT / 'documents' / 'conventions-tvf'
OUT.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / 'assets' / 'logo-territoires-vivants-france-web.png'
GREEN = RGBColor(24,57,47); DARK = RGBColor(31,41,55); GRAY = RGBColor(102,112,133)
PALE='F3F7F2'; NOTE='FBF8EF'; BORDER='D9E0DC'
TVF = {'name':'Territoires Vivants France','subtitle':'Agence Territoriale de Revitalisation Immobilière','address':'25 rue Élise Gervais, 42000 Saint-Étienne','phone':'04 65 81 54 69','email':'contact@territoiresvivantsfrance.fr','site':'www.territoiresvivantsfrance.fr','rna':'RNA W922015538','siret':'SIRET 897 226 138 00018'}
slug='convention-confidentialite-donnees-tvf'
title='CONVENTION CONFIDENTIALITÉ ET DONNÉES'
subtitle='Protection des informations sensibles, données personnelles, accès et engagements de confidentialité'
articles=[
('1. Parties et objet',["La présente convention est conclue entre Territoires Vivants France et toute personne, structure, collectivité, entreprise, association, prestataire, partenaire ou intervenant amené à accéder à des informations confidentielles ou à des données personnelles dans le cadre d’une action TVF.","Elle a pour objet d’encadrer la confidentialité, la protection des données, les conditions d’accès, les usages autorisés, les obligations de sécurité et les modalités de restitution ou suppression des informations communiquées."]),
('2. Informations couvertes par la confidentialité',["Sont notamment confidentiels les dossiers propriétaires, coordonnées, adresses précises de biens privés, photographies non publiques, documents transmis, situations personnelles, données patrimoniales, informations cadastrales, échanges, devis, stratégies de projet, données économiques, éléments de diagnostic et informations internes TVF.","La confidentialité s’applique quel que soit le support : papier, e-mail, fichier numérique, photographie, base de données, extrait TVF OS, compte rendu, message, document partagé ou information communiquée oralement."]),
('3. Données personnelles concernées',["Les données personnelles peuvent concerner les propriétaires, demandeurs, habitants, représentants, partenaires, agents, bénévoles, entreprises, contacts de collectivités ou toute personne liée à une demande, un signalement, un bien ou un dossier.","Les données traitées doivent être limitées aux informations nécessaires : identité, coordonnées, qualité, lien avec le bien, informations de suivi, documents transmis, échanges et éléments strictement utiles à la mission concernée."]),
('4. Finalité et usage autorisé',["Les informations ne peuvent être utilisées que pour la finalité définie : repérage, qualification, instruction, accompagnement, partenariat, intervention, suivi, restitution ou obligation légale applicable.","Toute réutilisation pour une autre finalité, transmission à un tiers, extraction, copie, communication publique, prospection, exploitation commerciale ou usage personnel est interdite sans autorisation écrite préalable de TVF et, lorsque nécessaire, des personnes concernées." ]),
('5. Principe de minimisation',["Chaque partie s’engage à ne collecter, consulter, transmettre ou conserver que les informations strictement nécessaires à la mission confiée. Les documents ou données non utiles ne doivent pas être demandés, copiés ou diffusés.","Lorsque des informations peuvent être anonymisées, agrégées ou limitées, cette solution doit être privilégiée, notamment pour les cartes, rapports, bilans ou présentations publiques." ]),
('6. Rôles des parties et instructions',["Selon les situations, les parties peuvent agir comme responsables de traitement distincts, responsables conjoints ou sous-traitants. Ce rôle doit être précisé lorsque le partenariat implique un traitement structuré de données personnelles.","Lorsqu’un intervenant traite des données pour le compte de TVF, il agit uniquement sur instruction documentée de TVF, ne réutilise pas les données pour son propre compte et informe TVF de toute demande, incident ou difficulté concernant les données." ]),
('7. Accès, habilitations et besoin d’en connaître',["L’accès aux informations est réservé aux personnes habilitées et uniquement dans la limite de ce qui est nécessaire à leur mission. Les accès doivent être retirés lorsque la mission prend fin ou lorsque la personne n’a plus besoin de consulter les informations.","Les identifiants, accès TVF OS, liens de partage, fichiers et documents ne doivent pas être transmis à une personne non autorisée. Chaque utilisateur reste responsable de l’usage de ses accès." ]),
('8. Sécurité des informations',["Les parties s’engagent à mettre en œuvre des mesures adaptées : stockage sécurisé, mot de passe robuste, limitation des copies, protection des appareils, vigilance sur les e-mails, absence de dépôt sur des outils non autorisés, destruction des brouillons inutiles et protection des documents papier.","Tout envoi de document sensible doit être limité aux destinataires nécessaires. Lorsqu’un canal sécurisé est disponible, il doit être privilégié." ]),
('9. Documents papier et archives',["Les documents papier contenant des informations confidentielles doivent être conservés dans des espaces adaptés, non laissés accessibles au public ou à des tiers, puis restitués, archivés ou détruits selon les consignes applicables.","Les impressions inutiles doivent être évitées. Les brouillons, copies de travail et documents obsolètes doivent être détruits de manière empêchant leur lecture ultérieure." ]),
('10. Photographies et images',["Les photographies de biens, personnes, intérieurs, documents, plaques, boîtes aux lettres, éléments identifiants ou situations sensibles ne doivent pas être diffusées sans autorisation adaptée.","Les photographies prises pour le suivi d’un dossier sont destinées à l’usage interne ou aux partenaires directement concernés. Toute communication publique nécessite une validation préalable." ]),
('11. Droits des personnes concernées',["Les personnes concernées peuvent demander l’accès, la rectification, la limitation ou, selon les cas, l’opposition au traitement de leurs données. Toute demande reçue par un partenaire doit être transmise sans délai à TVF lorsqu’elle concerne un traitement TVF.","Les parties s’engagent à coopérer pour répondre aux demandes dans des délais raisonnables et à corriger les informations manifestement inexactes lorsqu’elles sont signalées." ]),
('12. Durée de conservation',["Les données ne doivent être conservées que pendant la durée nécessaire à la finalité poursuivie, aux obligations légales, au suivi du dossier, à la preuve des actions réalisées ou aux nécessités d’archivage.","À la fin de la mission, les informations doivent être restituées, supprimées, anonymisées ou archivées selon les instructions applicables et le rôle de chaque partie." ]),
('13. Sous-traitance et tiers',["Aucun sous-traitant, prestataire secondaire, outil externe ou tiers non prévu ne peut accéder aux informations confidentielles ou données personnelles sans accord préalable et sans garanties suffisantes.","Lorsqu’un tiers intervient, ses obligations de confidentialité, sécurité, limitation d’usage et restitution doivent être équivalentes à celles prévues par la présente convention." ]),
('14. Violation, perte ou accès non autorisé',["Toute perte, accès non autorisé, divulgation accidentelle, erreur d’envoi, vol, compromission de compte, document égaré ou suspicion d’incident doit être signalé sans délai à TVF.","Le signalement doit préciser les informations concernées, les personnes ou dossiers potentiellement touchés, les circonstances, les mesures déjà prises et les actions correctives proposées." ]),
('15. Communication externe',["Aucune information confidentielle, donnée personnelle, adresse précise, photographie sensible, document interne, chiffre non validé ou détail de dossier ne peut être communiqué publiquement sans validation préalable.","Les partenaires s’interdisent de présenter un dossier TVF, un propriétaire, un bien, une collectivité ou un projet comme acquis, validé ou public lorsque cette information n’a pas été officiellement autorisée." ]),
('16. Durée de l’obligation de confidentialité',["L’obligation de confidentialité s’applique pendant toute la durée de la mission, de la convention ou de la relation entre les parties. Elle se poursuit après sa fin aussi longtemps que les informations ne sont pas devenues publiques de manière licite.","La fin de la relation n’autorise pas la conservation ou la réutilisation libre des informations obtenues dans le cadre de TVF." ]),
('17. Restitution et suppression',["À la demande de TVF ou à la fin de la mission, le partenaire restitue ou supprime les documents, fichiers, exports, photographies, copies et supports contenant des informations confidentielles, sauf obligation légale contraire.","Une attestation de suppression ou restitution peut être demandée lorsque les informations présentent une sensibilité particulière." ]),
('18. Responsabilité et sanctions contractuelles',["Tout manquement à la confidentialité, à la protection des données ou aux restrictions d’usage peut entraîner la suspension de l’accès aux informations, la fin de la relation, la demande de restitution des documents et, le cas échéant, la mise en œuvre des responsabilités applicables.","Les parties conviennent de rechercher d’abord une solution amiable, sans préjudice des recours nécessaires en cas de dommage, violation grave ou obligation légale." ])]
annexes=['Annexe 1 — Liste des informations partagées','Annexe 2 — Personnes habilitées','Annexe 3 — Durée de conservation','Annexe 4 — Procédure incident / violation','Annexe 5 — Attestation de restitution ou suppression']

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def borders(cell):
    tcPr=cell._tc.get_or_add_tcPr(); tb=tcPr.first_child_found_in('w:tcBorders')
    if tb is None: tb=OxmlElement('w:tcBorders'); tcPr.append(tb)
    for e in ('top','left','bottom','right'):
        el=tb.find(qn('w:'+e))
        if el is None: el=OxmlElement('w:'+e); tb.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'8'); el.set(qn('w:color'),BORDER)
def run(p,text,size=10,bold=False,color=DARK,font='Inter'):
    r=p.add_run(text); r.font.name=font; r.font.size=Pt(size); r.bold=bold; r.font.color.rgb=color; return r
def cell(cell,text,bold=False,fill=None):
    cell.text=''
    if fill: shade(cell,fill)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); run(p,text,8.5,bold,GREEN if bold else DARK); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; borders(cell)
def note(doc,text):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; cell(t.rows[0].cells[0],text,False,NOTE)
    for rr in t.rows[0].cells[0].paragraphs[0].runs: rr.font.size=Pt(8); rr.font.color.rgb=GRAY

def make_docx():
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.65); sec.right_margin=Inches(.65)
    h=doc.add_table(rows=1,cols=2); h.alignment=WD_TABLE_ALIGNMENT.CENTER; h.columns[0].width=Inches(1.25); h.columns[1].width=Inches(5.5)
    if LOGO.exists(): h.rows[0].cells[0].paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.0))
    p=h.rows[0].cells[1].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run(p,TVF['name']+'\n',12,True,GREEN,'Manrope'); run(p,TVF['subtitle']+'\n',8.5,False,DARK); run(p,f"{TVF['address']} | {TVF['phone']} | {TVF['email']}\n{TVF['rna']} | {TVF['siret']}",7.2,False,GRAY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(10); run(p,title,18,True,GREEN,'Manrope')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(p,subtitle,10,False,GRAY)
    note(doc,"MODÈLE INTERNE À ADAPTER — Document institutionnel de travail. Toute signature réelle doit être précédée d’une vérification juridique/RGPD adaptée aux rôles des parties, aux finalités, aux données et aux outils utilisés.")
    t=doc.add_table(rows=4,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.columns[0].width=Inches(2.25); t.columns[1].width=Inches(4.5)
    for i,(a,b) in enumerate([('Référence interne TVF','[à compléter]'),('Parties signataires','[à compléter]'),('Dossier / mission concerné','[à compléter]'),('Date d’effet','[à compléter]')]): cell(t.rows[i].cells[0],a,True,PALE); cell(t.rows[i].cells[1],b)
    for head,paras in articles:
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(9); run(p,head,12,True,GREEN,'Manrope')
        for para in paras:
            q=doc.add_paragraph(); q.paragraph_format.line_spacing=1.12; q.paragraph_format.space_after=Pt(5); run(q,para,9.6,False,DARK)
    p=doc.add_paragraph(); run(p,'Annexes recommandées',12,True,GREEN,'Manrope')
    for a in annexes:
        q=doc.add_paragraph(); q.paragraph_format.left_indent=Inches(.18); run(q,'□ '+a,9.5,False,DARK)
    p=doc.add_paragraph(); run(p,'Signatures',12,True,GREEN,'Manrope')
    sig=doc.add_table(rows=3,cols=2); sig.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,row in enumerate([('Pour Territoires Vivants France','Pour la partie cocontractante'),('Nom, qualité, date','Nom, qualité, date'),('Signature','Signature')]):
        for c,tv in enumerate(row): cell(sig.rows[i].cells[c],tv,i==0,PALE if i==0 else None)
    f=doc.sections[0].footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(f,f"{TVF['name']} — {title} — modèle interne confidentiel",7,False,GRAY)
    path=OUT/(slug+'.docx'); doc.save(path); return path

def make_pdf():
    st=getSampleStyleSheet(); st.add(ParagraphStyle(name='T',fontName='Helvetica-Bold',fontSize=17,leading=21,textColor=colors.HexColor('#18392F'),alignment=1)); st.add(ParagraphStyle(name='S',fontSize=9.5,leading=12,textColor=colors.HexColor('#667085'),alignment=1,spaceAfter=8)); st.add(ParagraphStyle(name='H',fontName='Helvetica-Bold',fontSize=11.5,leading=14,textColor=colors.HexColor('#18392F'),spaceBefore=9,spaceAfter=3)); st.add(ParagraphStyle(name='B',fontSize=9.1,leading=12,textColor=colors.HexColor('#1F2937'),spaceAfter=5)); st.add(ParagraphStyle(name='N',fontSize=8,leading=10.5,textColor=colors.HexColor('#475467'),backColor=colors.HexColor('#FBF8EF'),borderColor=colors.HexColor('#E4D7B8'),borderWidth=.5,borderPadding=6,spaceAfter=8)); st.add(ParagraphStyle(name='Small',fontSize=7.2,leading=9,textColor=colors.HexColor('#667085')))
    story=[]; logo=Image(str(LOGO), width=24*mm, height=18*mm) if LOGO.exists() else ''
    ht=Table([[logo,Paragraph(f"<b>{TVF['name']}</b><br/>{TVF['subtitle']}<br/>{TVF['address']} | {TVF['phone']} | {TVF['email']}<br/>{TVF['rna']} | {TVF['siret']}",st['Small'])]],colWidths=[30*mm,140*mm]); ht.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'MIDDLE'),('ALIGN',(1,0),(1,0),'RIGHT')])) ; story.append(ht); story.append(Spacer(1,7*mm)); story.append(Paragraph(title,st['T'])); story.append(Paragraph(subtitle,st['S'])); story.append(Paragraph('MODÈLE INTERNE À ADAPTER — Document institutionnel de travail. Toute signature réelle doit être précédée d’une vérification juridique/RGPD adaptée.',st['N']))
    tbl=Table([('Référence interne TVF','[à compléter]'),('Parties signataires','[à compléter]'),('Dossier / mission concerné','[à compléter]'),('Date d’effet','[à compléter]')],colWidths=[50*mm,120*mm]); tbl.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.45,colors.HexColor('#D9E0DC')),('BACKGROUND',(0,0),(0,-1),colors.HexColor('#F3F7F2')),('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,-1),7.8),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)])); story.append(tbl)
    for h,paras in articles:
        story.append(Paragraph(h,st['H']))
        for para in paras: story.append(Paragraph(para,st['B']))
    story.append(Paragraph('Annexes recommandées',st['H']))
    for a in annexes: story.append(Paragraph('□ '+a,st['B']))
    story.append(Paragraph('Signatures',st['H'])); sig=Table([['Pour Territoires Vivants France','Pour la partie cocontractante'],['Nom, qualité, date','Nom, qualité, date'],['Signature','Signature']],colWidths=[85*mm,85*mm],rowHeights=[9*mm,9*mm,22*mm]); sig.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.45,colors.HexColor('#D9E0DC')),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#F3F7F2')),('ALIGN',(0,0),(-1,-1),'CENTER'),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('FONTSIZE',(0,0),(-1,-1),8)])); story.append(sig)
    path=OUT/(slug+'.pdf'); SimpleDocTemplate(str(path),pagesize=A4,rightMargin=14*mm,leftMargin=14*mm,topMargin=11*mm,bottomMargin=12*mm).build(story); return path
print(make_docx()); print(make_pdf())
