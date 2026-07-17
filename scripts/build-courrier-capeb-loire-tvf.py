from pathlib import Path
from runpy import run_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

OUT_DIR = ROOT / "documents" / "courriers-lancement-saint-etienne"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "28-capeb-loire-lancement-cooperation-artisans-tvf.docx"
MD_PATH = SRC_DIR / "courrier-capeb-loire-lancement-tvf.md"

txt = BASE["txt"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_number = BASE["add_number"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
BORDER = BASE["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        (
            "Destinataire",
            "CAPEB Loire - Conf&eacute;d&eacute;ration de l'Artisanat et des Petites Entreprises du B&acirc;timent - via le r&eacute;seau CAPEB",
        ),
        (
            "Adresse / contact",
            "Formulaire officiel CAPEB - d&eacute;partement CAPEB LOIRE ; CAPEB nationale : 2 rue B&eacute;ranger, 75003 Paris - 01 53 60 50 00 - capeb@capeb.fr",
        ),
        (
            "Objet",
            "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les compl&eacute;mentarit&eacute;s possibles avec les artisans du b&acirc;timent autour de la r&eacute;novation, du r&eacute;emploi de mat&eacute;riaux et de la revitalisation du b&acirc;ti vacant",
        ),
        (
            "Pi&egrave;ces jointes",
            "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche Habitat Vivant, fiche entreprises / artisans, fiche mat&eacute;riaux et r&eacute;emploi, fiche besoins de lancement &agrave; Saint-&Eacute;tienne",
        ),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def add_complementarity_table(doc):
    doc.add_heading(txt("Compl&eacute;mentarit&eacute;s possibles avec les artisans du b&acirc;timent"), level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.5, 6.4, 6.2])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Sujet", "Apport possible de TVF", "Point &agrave; cadrer avec la CAPEB"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(head))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.1)

    rows = [
        (
            "R&eacute;novation du b&acirc;ti vacant",
            "Qualifier des logements, locaux et b&acirc;timents &agrave; remettre en usage, puis orienter les dossiers vers des acteurs comp&eacute;tents.",
            "Distinguer clairement ce qui rel&egrave;ve des artisans, des bureaux d'&eacute;tudes, des collectivit&eacute;s, des propri&eacute;taires et de TVF.",
        ),
        (
            "Mat&eacute;riaux de r&eacute;emploi",
            "Identifier des surplus, chutes, mobiliers, menuiseries, sanitaires, &eacute;quipements ou mat&eacute;riaux encore utilisables.",
            "Pr&eacute;ciser les exigences de qualit&eacute;, s&eacute;curit&eacute;, tra&ccedil;abilit&eacute;, assurance, responsabilit&eacute; et conformit&eacute; technique.",
        ),
        (
            "Artisans partenaires de dossiers",
            "Recenser les entreprises volontaires pour &ecirc;tre consult&eacute;es sur des diagnostics, travaux, conseils ou contributions encadr&eacute;es.",
            "Garantir que TVF ne se substitue pas au march&eacute; du travail artisanal et respecte les r&egrave;gles professionnelles.",
        ),
        (
            "Chantiers encadr&eacute;s",
            "Mobiliser des b&eacute;n&eacute;voles ou structures d'insertion uniquement sur des t&acirc;ches compatibles, simples et s&eacute;curis&eacute;es.",
            "D&eacute;finir les limites entre travaux professionnels, actions participatives, insertion, formation et b&eacute;n&eacute;volat.",
        ),
        (
            "Transmission et formation",
            "Relier certains projets TVF &agrave; des actions de d&eacute;couverte des m&eacute;tiers, de sensibilisation et de valorisation de l'artisanat.",
            "Identifier les passerelles possibles avec les dispositifs CAPEB et les entreprises artisanales locales.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.4)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - CAPEB Loire"))
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Demande d'un premier &eacute;change sur les passerelles entre artisans du b&acirc;timent, r&eacute;novation, r&eacute;emploi et revitalisation territoriale"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(
        doc,
        "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un ancrage op&eacute;rationnel progressif sur son territoire d'origine.",
    )
    add_paragraph(
        doc,
        "TVF intervient sur des sujets qui concernent directement les m&eacute;tiers du b&acirc;timent : logements vacants ou d&eacute;grad&eacute;s, locaux &agrave; remettre en usage, b&acirc;timents d&eacute;laiss&eacute;s, mat&eacute;riaux r&eacute;employables, besoins de diagnostic, travaux de r&eacute;habilitation, chantiers encadr&eacute;s et valorisation des comp&eacute;tences artisanales. L'association ne se substitue pas aux artisans ni aux entreprises du b&acirc;timent : elle cherche au contraire &agrave; cr&eacute;er un cadre de mise en relation clair, respectueux des comp&eacute;tences professionnelles et utile aux propri&eacute;taires comme aux territoires.",
    )
    add_paragraph(
        doc,
        "La CAPEB, Conf&eacute;d&eacute;ration de l'Artisanat et des Petites Entreprises du B&acirc;timent, repr&eacute;sente l'artisanat du b&acirc;timent et accompagne les entreprises artisanales dans leur cr&eacute;ation, leur gestion, leur d&eacute;veloppement, la formation, la repr&eacute;sentation et la d&eacute;fense de leurs int&eacute;r&ecirc;ts. Cette proximit&eacute; avec les artisans en fait un interlocuteur essentiel pour cadrer une d&eacute;marche qui touche au b&acirc;ti, aux travaux, aux mat&eacute;riaux et &agrave; la qualit&eacute; d'ex&eacute;cution.",
    )

    add_callout(
        doc,
        "Demande principale",
        "TVF sollicite un premier rendez-vous avec la CAPEB Loire ou le r&eacute;seau CAPEB afin de pr&eacute;senter sa m&eacute;thode, d'identifier les conditions d'une mise en relation avec les artisans du b&acirc;timent et de d&eacute;finir les limites &agrave; respecter pour les projets de r&eacute;novation, de r&eacute;emploi et de revitalisation du b&acirc;ti vacant.",
    )

    doc.add_heading(txt("Pourquoi solliciter la CAPEB"), level=1)
    add_paragraph(
        doc,
        "Le lancement de TVF &agrave; Saint-&Eacute;tienne ne peut pas &ecirc;tre cr&eacute;dible sans un dialogue avec les professionnels du b&acirc;timent. Les logements vacants, commerces ferm&eacute;s, friches et b&acirc;timents d&eacute;grad&eacute;s n&eacute;cessitent des comp&eacute;tences techniques, des assurances, des qualifications, une lecture des risques, des devis, des diagnostics et des entreprises capables d'intervenir dans un cadre clair.",
    )
    add_paragraph(
        doc,
        "La place des artisans est donc centrale : ils peuvent contribuer &agrave; identifier les contraintes, &eacute;valuer la faisabilit&eacute; de certains travaux, orienter les propri&eacute;taires, proposer des solutions de r&eacute;emploi lorsque cela est compatible, signaler des surplus ou ressources inutilis&eacute;es et participer &agrave; des projets territoriaux sans que leur savoir-faire soit d&eacute;valoris&eacute;.",
    )
    add_paragraph(
        doc,
        "TVF souhaite construire ce cadre avec prudence. Les chantiers participatifs ou solidaires ne doivent jamais brouiller les responsabilit&eacute;s, contourner les obligations professionnelles ou remplacer des interventions relevant des entreprises qualifi&eacute;es. L'objectif est de distinguer ce qui peut &ecirc;tre fait par des professionnels, ce qui peut &ecirc;tre accompagn&eacute; par des structures d'insertion, et ce qui peut relever d'une mobilisation citoyenne strictement encadr&eacute;e.",
    )

    add_complementarity_table(doc)

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "Organiser un rendez-vous de cadrage avec la CAPEB Loire ou le bon interlocuteur du r&eacute;seau CAPEB pour pr&eacute;senter TVF et comprendre les attentes des artisans du b&acirc;timent.")
    add_bullet(doc, "Identifier les conditions dans lesquelles des artisans volontaires pourraient &ecirc;tre inform&eacute;s des dossiers TVF : diagnostic, devis, travaux, conseil, r&eacute;emploi, sensibilisation ou contribution encadr&eacute;e.")
    add_bullet(doc, "Construire une fiche claire distinguant les travaux relevant obligatoirement d'une entreprise qualifi&eacute;e, les actions pouvant relever de l'insertion, et les interventions b&eacute;n&eacute;voles possibles sous conditions.")
    add_bullet(doc, "Pr&eacute;ciser les r&egrave;gles de responsabilit&eacute;, d'assurance, de s&eacute;curit&eacute;, de tra&ccedil;abilit&eacute;, de qualit&eacute; et de conformit&eacute; &agrave; respecter dans tout dossier TVF li&eacute; au b&acirc;timent.")
    add_bullet(doc, "Examiner si un premier circuit de mise en relation peut &ecirc;tre test&eacute; sur Saint-&Eacute;tienne, sans communication de partenariat ni engagement public avant validation &eacute;crite.")

    doc.add_heading(txt("Ce que TVF peut apporter aux artisans"), level=1)
    add_number(doc, "Une meilleure visibilit&eacute; sur des projets locaux &agrave; instruire : logements &agrave; remettre en usage, locaux vacants, commerces, petits b&acirc;timents, espaces associatifs ou projets de proximit&eacute;.")
    add_number(doc, "Des dossiers mieux pr&eacute;qualifi&eacute;s avant sollicitation : propri&eacute;taire identifi&eacute;, besoin exprim&eacute;, photos, localisation, usage envisag&eacute;, pi&egrave;ces disponibles et statut de la demande.")
    add_number(doc, "Une valorisation des savoir-faire artisanaux dans des projets utiles au territoire, sans confusion entre b&eacute;n&eacute;volat, insertion et travail professionnel.")
    add_number(doc, "Un cadre pour orienter des surplus, chutes, mobiliers ou mat&eacute;riaux encore utilisables vers des projets valid&eacute;s, avec tra&ccedil;abilit&eacute; et limites techniques.")
    add_number(doc, "Une passerelle entre propri&eacute;taires, collectivit&eacute;s, associations, structures d'insertion, entreprises artisanales, financeurs et habitants.")

    add_callout(
        doc,
        "Principe de prudence",
        "Aucune mention de partenariat, de soutien, de validation ou de financement par la CAPEB ou par la CAPEB Loire ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable. Ce courrier vise uniquement &agrave; solliciter un premier &eacute;change de cadrage.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(
        doc,
        "Nous proposons un rendez-vous d'environ une heure afin de pr&eacute;senter TVF, de comprendre les attentes de la CAPEB et des artisans, de pr&eacute;ciser les conditions d'une mise en relation responsable, et de d&eacute;finir un premier cadre de travail pour les sujets r&eacute;novation, r&eacute;emploi, mat&eacute;riaux, diagnostics et chantiers encadr&eacute;s.",
    )
    add_paragraph(
        doc,
        "&Agrave; l'issue de cet &eacute;change, TVF pourra produire une note de cadrage indiquant les types de dossiers compatibles, les pr&eacute;cautions &agrave; respecter, les informations &agrave; collecter, les responsabilit&eacute;s &agrave; distinguer et les limites de communication &agrave; appliquer.",
    )

    doc.add_heading("Conclusion", level=1)
    add_paragraph(
        doc,
        "Territoires Vivants France souhaite avancer avec les acteurs professionnels, non &agrave; leur place. La r&eacute;ussite du lancement st&eacute;phanois repose sur une coop&eacute;ration claire entre propri&eacute;taires, collectivit&eacute;s, artisans, associations, structures d'insertion, entreprises et habitants. La CAPEB peut aider TVF &agrave; construire une d&eacute;marche s&eacute;rieuse, respectueuse des m&eacute;tiers et utile aux territoires.",
    )
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous &agrave; la date qui vous conviendra et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(head))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration officielle de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche Habitat Vivant", "Pr&eacute;senter les sujets logements vacants, b&acirc;ti d&eacute;grad&eacute;, propri&eacute;taires et remise en usage.", "&Agrave; joindre"),
        ("Fiche entreprises / artisans", "D&eacute;crire les formes possibles de contribution ou de mise en relation avec les professionnels.", "&Agrave; joindre"),
        ("Fiche mat&eacute;riaux et r&eacute;emploi", "Pr&eacute;senter les principes de valorisation des ressources inutilis&eacute;es et les limites &agrave; respecter.", "&Agrave; joindre"),
        ("Fiche besoins de lancement", "Pr&eacute;ciser les besoins initiaux : artisans, stockage, transport, diagnostic, s&eacute;curit&eacute;, chantier encadr&eacute;.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(
        doc,
        "Objet d'e-mail recommand&eacute;",
        "Demande de rendez-vous - TERRITOIRES VIVANTS FRANCE / CAPEB Loire - artisans, r&eacute;novation et r&eacute;emploi",
        fill=LIGHT_BLUE,
    )
    doc.add_heading("Sources de cadrage", level=1)
    add_paragraph(
        doc,
        "Les informations utilis&eacute;es pour orienter ce courrier proviennent du site officiel de la CAPEB, notamment les pages de pr&eacute;sentation, de services et de contact, ainsi que l'existence de l'entr&eacute;e CAPEB LOIRE dans le formulaire officiel de contact.",
    )

    doc.save(DOCX_PATH)


def build_md():
    md = txt(
        """# Courrier - CAPEB Loire - lancement TVF

Destinataire : CAPEB Loire - Conf&eacute;d&eacute;ration de l'Artisanat et des Petites Entreprises du B&acirc;timent - via le r&eacute;seau CAPEB.

Objet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les compl&eacute;mentarit&eacute;s possibles avec les artisans du b&acirc;timent autour de la r&eacute;novation, du r&eacute;emploi de mat&eacute;riaux et de la revitalisation du b&acirc;ti vacant.

Date : Saint-&Eacute;tienne, le 14 juillet 2026.

## Intention du courrier

Ce courrier sollicite un premier rendez-vous avec la CAPEB Loire ou le r&eacute;seau CAPEB afin de pr&eacute;senter TVF, d'identifier les conditions d'une mise en relation avec les artisans du b&acirc;timent et de d&eacute;finir les limites &agrave; respecter pour les projets de r&eacute;novation, de r&eacute;emploi et de revitalisation du b&acirc;ti vacant.

## Demandes op&eacute;rationnelles

- Organiser un rendez-vous de cadrage avec la CAPEB Loire ou le bon interlocuteur du r&eacute;seau CAPEB.
- Identifier les conditions de mise en relation avec des artisans volontaires.
- Construire une fiche distinguant travaux professionnels, insertion et interventions b&eacute;n&eacute;voles possibles sous conditions.
- Pr&eacute;ciser les r&egrave;gles de responsabilit&eacute;, assurance, s&eacute;curit&eacute;, tra&ccedil;abilit&eacute;, qualit&eacute; et conformit&eacute;.
- Examiner un premier circuit de mise en relation &agrave; Saint-&Eacute;tienne.

## Principe de prudence

Aucune mention de partenariat, de soutien, de validation ou de financement par la CAPEB ou par la CAPEB Loire ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable.

## Signature

Edryan Rangoly  
Pr&eacute;sident fondateur  
TERRITOIRES VIVANTS FRANCE
"""
    )
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
