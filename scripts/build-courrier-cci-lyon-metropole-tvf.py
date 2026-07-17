from pathlib import Path
from runpy import run_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

OUT_DIR = ROOT / "documents" / "courriers-lancement-saint-etienne"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "27-cci-lyon-metropole-saint-etienne-roanne-lancement-tvf.docx"
MD_PATH = SRC_DIR / "courrier-cci-lyon-metropole-lancement-tvf.md"

txt = BASE["txt"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_number = BASE["add_number"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
BORDER = BASE["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        (
            "Destinataire",
            "CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne - D&eacute;l&eacute;gation de Saint-&Eacute;tienne - &agrave; l'attention de la direction et des services commerce, attractivit&eacute;, d&eacute;veloppement durable et entreprises",
        ),
        (
            "Adresse",
            "57 cours Fauriel - CS 70 374 - 42100 Saint-&Eacute;tienne - 04 77 43 04 00",
        ),
        (
            "Objet",
            "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les compl&eacute;mentarit&eacute;s possibles autour des commerces inoccup&eacute;s, des entreprises, du r&eacute;emploi de mat&eacute;riaux et de la revitalisation &eacute;conomique locale",
        ),
        (
            "Pi&egrave;ces jointes",
            "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche entreprises partenaires, fiche commerces inoccup&eacute;s, fiche mat&eacute;riaux et r&eacute;emploi, fiche besoins de lancement &agrave; Saint-&Eacute;tienne",
        ),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def add_complementarity_table(doc):
    doc.add_heading(txt("Compl&eacute;mentarit&eacute;s possibles &agrave; examiner"), level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.5, 6.4, 6.2])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Sujet", "Apport possible de TVF", "Lien possible avec la CCI"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(head))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.1)

    rows = [
        (
            "Commerces inoccup&eacute;s",
            "Rep&eacute;rer des cellules vacantes, qualifier les usages possibles et orienter vers des porteurs de projets utiles au territoire.",
            "Articuler la d&eacute;marche avec les services commerce, attractivit&eacute;, cr&eacute;ation et reprise d'entreprise.",
        ),
        (
            "Entreprises contributrices",
            "Identifier les entreprises pouvant proposer mat&eacute;riaux, mobilier, &eacute;quipements, transport, stockage, comp&eacute;tences ou m&eacute;c&eacute;nat.",
            "Mobiliser le r&eacute;seau &eacute;conomique dans un cadre clair, tra&ccedil;able et compatible avec les pratiques des entreprises.",
        ),
        (
            "R&eacute;emploi et transition",
            "Transformer des surplus ou ressources inutilis&eacute;es en apports pour des projets locaux valid&eacute;s.",
            "Relier la d&eacute;marche aux sujets de d&eacute;veloppement durable, transition &eacute;cologique et responsabilit&eacute; territoriale.",
        ),
        (
            "Porteurs de projets",
            "Orienter les besoins vers des locaux, commerces, ateliers, tiers-lieux ou activit&eacute;s de proximit&eacute; &agrave; instruire.",
            "S'appuyer sur les dispositifs CCI de cr&eacute;ation, reprise, d&eacute;veloppement commercial, strat&eacute;gie et financement.",
        ),
        (
            "Donn&eacute;es et suivi",
            "Documenter les demandes : localisation, besoin, acteur, ressource, statut, pi&egrave;ces, suites donn&eacute;es et impact attendu.",
            "Partager une lecture territoriale utile, sans exploiter ni diffuser de donn&eacute;es sensibles hors cadre autoris&eacute;.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.45)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne"))
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Demande d'un premier &eacute;change sur les passerelles entre entreprises, commerces, r&eacute;emploi et revitalisation territoriale"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(
        doc,
        "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un ancrage op&eacute;rationnel progressif sur son territoire d'origine.",
    )
    add_paragraph(
        doc,
        "TVF intervient sur des sujets directement li&eacute;s &agrave; la vie &eacute;conomique locale : commerces inoccup&eacute;s, locaux vacants, b&acirc;timents d&eacute;laiss&eacute;s, ressources inutilis&eacute;es, mat&eacute;riaux r&eacute;employables, besoins logistiques, entreprises engag&eacute;es et porteurs de projets. L'objectif n'est pas de remplacer les dispositifs existants, mais de faciliter l'identification des besoins, la mobilisation des acteurs et la transformation de ressources dormantes en projets utiles aux habitants et aux entreprises.",
    )
    add_paragraph(
        doc,
        "La CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne constitue un interlocuteur naturel pour cette d&eacute;marche. Son action couvre l'accompagnement des cr&eacute;ateurs, dirigeants, entreprises, collectivit&eacute;s et porteurs de projets, avec des services li&eacute;s notamment au commerce, &agrave; l'attractivit&eacute;, au d&eacute;veloppement durable, au financement, &agrave; la strat&eacute;gie, aux formalit&eacute;s et &agrave; la mise en r&eacute;seau.",
    )

    add_callout(
        doc,
        "Demande principale",
        "TVF sollicite un premier rendez-vous avec la CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne afin de pr&eacute;senter sa m&eacute;thode, d'identifier les services concern&eacute;s et d'&eacute;tudier les conditions d'une coop&eacute;ration de cadrage autour des entreprises, des commerces inoccup&eacute;s, du r&eacute;emploi de mat&eacute;riaux et du lancement pilote &agrave; Saint-&Eacute;tienne.",
    )

    doc.add_heading(txt("Pourquoi solliciter la CCI"), level=1)
    add_paragraph(
        doc,
        "La CCI accompagne les entreprises &agrave; toutes les &eacute;tapes de leur d&eacute;veloppement et contribue &agrave; l'attractivit&eacute; du territoire. Cette mission rejoint plusieurs besoins de TVF : comprendre les attentes des entreprises, identifier les acteurs &eacute;conomiques susceptibles de contribuer, orienter les porteurs de projets vers les bons dispositifs et favoriser la r&eacute;activation de locaux ou commerces inoccup&eacute;s.",
    )
    add_paragraph(
        doc,
        "TVF peut apporter une lecture compl&eacute;mentaire de terrain. Lorsqu'un commerce est ferm&eacute;, qu'un local reste inutilis&eacute;, qu'une entreprise dispose de surplus, de mobilier, d'&eacute;quipements ou de comp&eacute;tences, l'association peut aider &agrave; qualifier la ressource, documenter la demande, construire un cadre de mise en relation et rechercher une affectation utile au territoire.",
    )
    add_paragraph(
        doc,
        "Cette approche peut aussi r&eacute;pondre &agrave; des attentes d'entreprises engag&eacute;es : r&eacute;duction du gaspillage, valorisation RSE, contribution &agrave; l'&eacute;conomie circulaire, soutien &agrave; l'emploi local, m&eacute;c&eacute;nat de comp&eacute;tences, mise &agrave; disposition ponctuelle de moyens ou participation &agrave; des projets collectifs visibles et tra&ccedil;ables.",
    )

    add_complementarity_table(doc)

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "Organiser un rendez-vous avec la d&eacute;l&eacute;gation de Saint-&Eacute;tienne ou les services concern&eacute;s afin de pr&eacute;senter TVF et d'identifier les passerelles utiles.")
    add_bullet(doc, "Orienter TVF vers les services commerce, attractivit&eacute;, cr&eacute;ation/reprise, d&eacute;veloppement durable, financement, entreprises et collectivit&eacute;s lorsque les demandes rel&egrave;vent de leur champ.")
    add_bullet(doc, "Identifier les modalit&eacute;s par lesquelles des entreprises pourraient proposer des mat&eacute;riaux, mobiliers, &eacute;quipements, locaux, moyens logistiques, transport, stockage, comp&eacute;tences ou m&eacute;c&eacute;nat.")
    add_bullet(doc, "Examiner la possibilit&eacute; de relayer, dans un cadre valid&eacute;, un appel &agrave; contribution aupr&egrave;s d'entreprises volontaires du territoire.")
    add_bullet(doc, "Construire un circuit de mise en relation clair afin d'&eacute;viter les doublons, les attentes impr&eacute;cises et toute communication de partenariat sans accord &eacute;crit.")

    doc.add_heading(txt("Ce que TVF peut apporter aux entreprises"), level=1)
    add_number(doc, "Un cadre lisible pour contribuer : don de mat&eacute;riaux, mise &agrave; disposition ponctuelle d'un local, appui logistique, m&eacute;c&eacute;nat de comp&eacute;tences, orientation de surplus ou participation &agrave; un projet territorial.")
    add_number(doc, "Une valorisation responsable des ressources inutilis&eacute;es, sans fonctionnement de distribution libre ni de revente classique.")
    add_number(doc, "Une tra&ccedil;abilit&eacute; des contributions : dossier, pi&egrave;ces, b&eacute;n&eacute;ficiaires du projet, usage pr&eacute;vu, conditions, limites et suivi.")
    add_number(doc, "Une contribution concr&egrave;te &agrave; l'&eacute;conomie circulaire, &agrave; la revitalisation commerciale, au cadre de vie et &agrave; l'utilit&eacute; locale.")
    add_number(doc, "Une passerelle entre entreprises, associations, propri&eacute;taires, collectivit&eacute;s, artisans, porteurs de projets et habitants.")

    add_callout(
        doc,
        "Principe de prudence",
        "Aucune mention de partenariat, de soutien, de validation ou de financement par la CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable. Ce courrier vise uniquement &agrave; solliciter un premier &eacute;change de cadrage.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(
        doc,
        "Nous proposons un rendez-vous d'environ une heure afin de pr&eacute;senter TVF, d'identifier les services CCI concern&eacute;s, de comprendre les modalit&eacute;s d'orientation vers les entreprises et de d&eacute;finir un premier cadre simple pour les sujets commerce, r&eacute;emploi, contributions d'entreprises et revitalisation territoriale.",
    )
    add_paragraph(
        doc,
        "&Agrave; l'issue de cet &eacute;change, TVF pourra produire une note de cadrage pr&eacute;cisant les besoins identifi&eacute;s, les interlocuteurs &agrave; mobiliser, les pi&egrave;ces &agrave; demander, les limites de communication, les types de contributions possibles et les prochaines &eacute;tapes.",
    )

    doc.add_heading("Conclusion", level=1)
    add_paragraph(
        doc,
        "Territoires Vivants France souhaite construire son lancement &agrave; Saint-&Eacute;tienne de mani&egrave;re s&eacute;rieuse, progressive et utile aux acteurs &eacute;conomiques locaux. Une premi&egrave;re prise de contact avec la CCI permettrait de mieux articuler la d&eacute;marche TVF avec les besoins des entreprises, des commerces, des porteurs de projets et des collectivit&eacute;s.",
    )
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous &agrave; la date qui vous conviendra et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(head))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration officielle de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche entreprises partenaires", "D&eacute;crire les formes de contribution possibles : mat&eacute;riaux, comp&eacute;tences, logistique, locaux, m&eacute;c&eacute;nat.", "&Agrave; joindre"),
        ("Fiche commerces inoccup&eacute;s", "Pr&eacute;senter le p&ocirc;le Commerce Vivant et la m&eacute;thode de qualification d'un local vacant.", "&Agrave; joindre"),
        ("Fiche mat&eacute;riaux et r&eacute;emploi", "Pr&eacute;senter le fonctionnement de la valorisation des ressources inutilis&eacute;es.", "&Agrave; joindre"),
        ("Fiche besoins de lancement", "Pr&eacute;ciser les besoins initiaux : local de stockage, transport, entreprises contributrices, orientation des demandes.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(
        doc,
        "Objet d'e-mail recommand&eacute;",
        "Demande de rendez-vous - TERRITOIRES VIVANTS FRANCE / CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne",
        fill=LIGHT_BLUE,
    )
    doc.add_heading("Sources de cadrage", level=1)
    add_paragraph(
        doc,
        "Les informations utilis&eacute;es pour orienter ce courrier proviennent du site officiel de la CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne, notamment les pages relatives &agrave; la d&eacute;l&eacute;gation de Saint-&Eacute;tienne, aux missions de la CCI, aux services commerce et attractivit&eacute;, au d&eacute;veloppement durable et au contact.",
    )

    doc.save(DOCX_PATH)


def build_md():
    md = txt(
        """# Courrier - CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne - lancement TVF

Destinataire : CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne, d&eacute;l&eacute;gation de Saint-&Eacute;tienne, &agrave; l'attention de la direction et des services commerce, attractivit&eacute;, d&eacute;veloppement durable et entreprises.

Objet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les compl&eacute;mentarit&eacute;s possibles autour des commerces inoccup&eacute;s, des entreprises, du r&eacute;emploi de mat&eacute;riaux et de la revitalisation &eacute;conomique locale.

Date : Saint-&Eacute;tienne, le 14 juillet 2026.

## Intention du courrier

Ce courrier sollicite un premier rendez-vous avec la CCI afin de pr&eacute;senter TVF, d'identifier les services concern&eacute;s et d'&eacute;tudier les conditions d'une coop&eacute;ration de cadrage autour des entreprises, des commerces inoccup&eacute;s, du r&eacute;emploi de mat&eacute;riaux et du lancement pilote &agrave; Saint-&Eacute;tienne.

## Demandes op&eacute;rationnelles

- Organiser un rendez-vous avec la d&eacute;l&eacute;gation de Saint-&Eacute;tienne ou les services concern&eacute;s.
- Orienter TVF vers les services commerce, attractivit&eacute;, cr&eacute;ation/reprise, d&eacute;veloppement durable, financement, entreprises et collectivit&eacute;s.
- Identifier les modalit&eacute;s par lesquelles des entreprises pourraient proposer mat&eacute;riaux, mobiliers, &eacute;quipements, locaux, moyens logistiques, transport, stockage, comp&eacute;tences ou m&eacute;c&eacute;nat.
- Examiner la possibilit&eacute; de relayer un appel &agrave; contribution dans un cadre valid&eacute;.

## Principe de prudence

Aucune mention de partenariat, de soutien, de validation ou de financement par la CCI Lyon M&eacute;tropole Saint-&Eacute;tienne Roanne ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable.

## Signature

Edryan Rangoly  
Pr&eacute;sident fondateur  
TERRITOIRES VIVANTS FRANCE
"""
    )
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
