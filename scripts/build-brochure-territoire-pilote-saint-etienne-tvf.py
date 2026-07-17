from pathlib import Path
from runpy import run_path
from html import unescape

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

Document = BASE["Document"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
BLUE = BASE["BLUE"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
GOLD = BASE["GOLD"]
BORDER = BASE["BORDER"]

OUT_DIR = ROOT / "documents" / "brochures"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "brochure-territoire-pilote-saint-etienne-tvf.docx"
MD_PATH = SRC_DIR / "brochure-territoire-pilote-saint-etienne-tvf.md"

LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"
PHOTO_COVER = ROOT / "assets" / "photos" / "saint-etienne-panorama.jpg"
PHOTO_CITY = ROOT / "assets" / "photos" / "france-saint-etienne-chateaucreux.jpg"
PHOTO_COMMERCE = ROOT / "assets" / "photos" / "saint-etienne-centre-commerce.jpg"
PHOTO_STREET = ROOT / "assets" / "photos" / "saint-etienne-rue-resistance.jpg"


def u(text: str) -> str:
    return unescape(text)


def set_run(run, size=None, color=None, bold=None, italic=None):
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(u(text))
    set_run(r, 8.3, MUTED, italic=True)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(u(text))
    set_run(r, 8.2, MUTED, italic=True)


def table(doc, headers, rows, widths, fill=LIGHT_GREEN):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, widths)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=BORDER, size="8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(u(header))
        set_run(r, 8.7, GREEN, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i], color="DDE6DE", size="6")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(u(value))
            set_run(r, 8.15, BLUE if i == 0 else RGBColor(31, 45, 38), bold=(i == 0))
    doc.add_paragraph()
    return t


def small_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(u(text.upper()))
    set_run(r, 8.2, GOLD, bold=True)
    return p


def cover(doc):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Cm(5.6))
        p_logo.paragraph_format.space_after = Pt(12)

    small_label(doc, "Brochure territoriale")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(u("Saint-&Eacute;tienne, territoire pilote"))
    set_run(r, 25, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(u("Transformer les biens, lieux et ressources inutilis&eacute;s en projets utiles au territoire"))
    set_run(r, 13.5, GREEN, bold=True)

    add_paragraph(
        doc,
        "Ce document pr&eacute;sente la mani&egrave;re dont TERRITOIRES VIVANTS FRANCE peut structurer, "
        "depuis Saint-&Eacute;tienne, une premi&egrave;re d&eacute;marche op&eacute;rationnelle de revitalisation : "
        "rep&eacute;rer les besoins, qualifier les biens et ressources, mobiliser les acteurs, cadrer "
        "les conventions et pr&eacute;parer des projets mesurables, utiles aux habitants et compatibles "
        "avec les politiques publiques existantes.",
    )
    add_callout(
        doc,
        "Positionnement",
        "TVF est une association loi 1901 d&eacute;clar&eacute;e, implant&eacute;e &agrave; Saint-&Eacute;tienne, "
        "ayant vocation &agrave; devenir un outil national de coop&eacute;ration territoriale. "
        "Le territoire st&eacute;phanois constitue le premier terrain d'ancrage, d'apprentissage "
        "et de structuration de la m&eacute;thode.",
        fill=LIGHT_GREEN,
    )
    if PHOTO_COVER.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(PHOTO_COVER), width=Cm(15.8))
        caption(doc, "Saint-&Eacute;tienne comme territoire d'ancrage : ville, reliefs, quartiers et ressources &agrave; remettre en usage.")

    rows = [
        ("RNA", "W423016361"),
        ("SIREN / SIRET", "107 226 128 / 107 226 128 00018"),
        ("Si&egrave;ge", "25 rue &Eacute;lise Gervais, 42000 Saint-&Eacute;tienne"),
        ("Contact", "contact@territoiresvivantsfrance.fr - 04 65 81 54 69"),
        ("Site", "www.territoiresvivantsfrance.fr"),
    ]
    table(doc, ["Rep&egrave;re", "Information"], rows, [3.6, 12.8], fill=LIGHT_BLUE)


def toc(doc):
    doc.add_heading(u("Sommaire"), level=1)
    rows = [
        ("1", "Pourquoi Saint-&Eacute;tienne ?"),
        ("2", "Diagnostic territorial sourc&eacute;"),
        ("3", "Ce que les donn&eacute;es r&eacute;v&egrave;lent"),
        ("4", "La r&eacute;ponse TVF par les cinq p&ocirc;les"),
        ("5", "M&eacute;thode d'installation du pilote"),
        ("6", "Exemples de projets faisables"),
        ("7", "B&eacute;n&eacute;fices par public"),
        ("8", "Besoins concrets de lancement"),
        ("9", "Cadre de coop&eacute;ration et points de vigilance"),
        ("10", "Indicateurs de suivi, reproductibilit&eacute; et sources"),
    ]
    table(doc, ["Partie", "Contenu"], rows, [2.2, 14.2], fill=LIGHT_GREEN)


def why_saint_etienne(doc):
    doc.add_page_break()
    doc.add_heading(u("1. Pourquoi Saint-&Eacute;tienne ?"), level=1)
    add_paragraph(
        doc,
        "Saint-&Eacute;tienne r&eacute;unit plusieurs enjeux que TVF cherche &agrave; traiter ensemble : "
        "patrimoine b&acirc;ti ancien, logements vacants, besoins de remise en usage, commerce de proximit&eacute;, "
        "insertion professionnelle, valorisation de ressources et n&eacute;cessit&eacute; de coop&eacute;ration entre acteurs. "
        "Ce n'est donc pas seulement une ville de lancement : c'est un terrain coh&eacute;rent pour tester une "
        "m&eacute;thode qui pourra ensuite &ecirc;tre adapt&eacute;e &agrave; d'autres communes, intercommunalit&eacute;s ou territoires.",
    )
    add_paragraph(
        doc,
        "Le choix de Saint-&Eacute;tienne permet de partir d'un territoire r&eacute;el, avec des besoins document&eacute;s, "
        "des politiques publiques d&eacute;j&agrave; engag&eacute;es et une histoire urbaine qui rend visible la question "
        "du r&eacute;emploi : r&eacute;emploi du b&acirc;ti, r&eacute;emploi des locaux, r&eacute;emploi des mat&eacute;riaux, "
        "r&eacute;emploi des comp&eacute;tences et des espaces disponibles.",
    )
    add_callout(
        doc,
        "Id&eacute;e directrice",
        "TVF ne pr&eacute;tend pas remplacer les dispositifs existants. L'objectif est de cr&eacute;er un outil de terrain "
        "qui facilite la mise en relation, la qualification des dossiers et le passage de l'id&eacute;e au projet cadr&eacute;.",
        fill=LIGHT_BLUE,
    )
    if PHOTO_CITY.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(PHOTO_CITY), width=Cm(14.6))
        caption(doc, "Saint-&Eacute;tienne : un territoire urbain, productif et populaire o&ugrave; la revitalisation doit &ecirc;tre concr&egrave;te.")


def diagnosis(doc):
    doc.add_heading(u("2. Diagnostic territorial sourc&eacute;"), level=1)
    add_paragraph(
        doc,
        "Les premiers rep&egrave;res ci-dessous sont issus de sources publiques. Ils ne constituent pas un diagnostic "
        "exhaustif, mais une base de travail pour montrer la coh&eacute;rence d'un pilote TVF &agrave; Saint-&Eacute;tienne. "
        "Chaque projet op&eacute;rationnel devra ensuite &ecirc;tre qualifi&eacute; &agrave; l'adresse, au bien, au propri&eacute;taire, "
        "aux contraintes techniques et au besoin social ou territorial identifi&eacute;.",
    )
    rows = [
        ("Population", "173 136 habitants en 2023", "INSEE, dossier complet Commune de Saint-&Eacute;tienne, RP2023"),
        ("Parc de logements", "101 341 logements en 2023", "INSEE, LOG T1, RP2023"),
        ("Logements vacants", "12 175 logements vacants, soit 12,0 % du parc en 2023", "INSEE, LOG T2, RP2023"),
        ("B&acirc;ti ancien", "Plus de 45 000 r&eacute;sidences principales recens&eacute;es dans des p&eacute;riodes de construction ant&eacute;rieures &agrave; 1971", "INSEE, LOG T8, RP2023"),
        ("Ch&ocirc;mage", "19,0 % des 15-64 ans au sens du recensement en 2023", "INSEE, EMP T4, RP2023"),
        ("Pauvret&eacute;", "30,4 % au seuil de 60 % du niveau de vie m&eacute;dian en 2023", "INSEE-DGFiP-Cnaf-Cnav-Ccmsa, Filosofi 2023"),
        ("D&eacute;chets du b&acirc;timent", "Environ 42 millions de tonnes de d&eacute;chets du b&acirc;timent par an en France", "Minist&egrave;re de la Transition &eacute;cologique, fili&egrave;re PMCB"),
        ("Vacance nationale", "3,1 millions de logements vacants en France en 2023, plus de 8 % du parc", "Banque des Territoires / Localtis, 2024, d'apr&egrave;s donn&eacute;es publiques"),
    ]
    table(doc, ["Th&egrave;me", "Donn&eacute;e rep&egrave;re", "Source"], rows, [3.3, 6.2, 6.9], fill=LIGHT_GREEN)
    note(
        doc,
        "Les donn&eacute;es de vacance commerciale et de friches doivent &ecirc;tre consolid&eacute;es localement par rep&eacute;rage, "
        "donn&eacute;es fonci&egrave;res, observation terrain et outils publics comme Cartofriches lorsque les fiches existent.",
    )


def data_meaning(doc):
    doc.add_heading(u("3. Ce que les donn&eacute;es r&eacute;v&egrave;lent"), level=1)
    rows = [
        (
            "Vacance et b&acirc;ti ancien",
            "La pr&eacute;sence de logements vacants et d'un parc ancien cr&eacute;e un besoin de rep&eacute;rage, de dialogue avec les propri&eacute;taires, de solutions gradu&eacute;es et de travaux proportionn&eacute;s.",
            "Habitat Vivant peut aider &agrave; qualifier les situations, orienter vers les bons acteurs et pr&eacute;parer des sc&eacute;narios de remise en usage.",
        ),
        (
            "Co&ucirc;t des ressources",
            "La hausse des co&ucirc;ts de chantier et le volume national de d&eacute;chets du b&acirc;timent renforcent l'int&eacute;r&ecirc;t de mieux trier, stocker et r&eacute;employer.",
            "La Mat&eacute;riauth&egrave;que Solidaire peut devenir un outil local de valorisation, sans distribution automatique ni logique de d&eacute;chetterie.",
        ),
        (
            "Difficult&eacute;s sociales",
            "Le niveau de pauvret&eacute; et de ch&ocirc;mage observ&eacute; appelle des projets utiles, encadr&eacute;s, capables de relier activit&eacute;, formation, b&eacute;n&eacute;volat et insertion.",
            "Solidarit&eacute; & Insertion peut pr&eacute;parer des chantiers encadr&eacute;s, des ateliers et des parcours de remobilisation avec les structures comp&eacute;tentes.",
        ),
        (
            "Commerce et centralit&eacute;",
            "Les rues commer&ccedil;antes ont besoin d'usages lisibles, de locaux r&eacute;activ&eacute;s et d'activit&eacute;s compatibles avec les besoins des habitants.",
            "Commerce Vivant peut travailler sur des boutiques tests, locaux associatifs, ateliers partag&eacute;s ou activit&eacute;s de proximit&eacute; &agrave; &eacute;tudier.",
        ),
    ]
    table(doc, ["Constat", "Lecture territoriale", "R&ocirc;le TVF"], rows, [3.4, 6.2, 6.8], fill=LIGHT_BLUE)


def poles_response(doc):
    doc.add_page_break()
    doc.add_heading(u("4. La r&eacute;ponse TVF par les cinq p&ocirc;les"), level=1)
    add_paragraph(
        doc,
        "Le pilote st&eacute;phanois doit montrer que les probl&egrave;mes ne se traitent pas isol&eacute;ment. "
        "Un logement vacant peut mobiliser des mat&eacute;riaux de r&eacute;emploi. Un local commercial peut accueillir "
        "une activit&eacute; associative ou artisanale. Une friche peut devenir un espace utile. Un chantier peut "
        "servir de support &agrave; l'insertion, &agrave; la formation ou au b&eacute;n&eacute;volat.",
    )
    rows = [
        ("Habitat Vivant", "Logements vacants, biens dormants, immeubles &agrave; remettre en usage", "Rep&eacute;rage, qualification, dialogue propri&eacute;taire, sc&eacute;nario de remise en usage"),
        ("Mat&eacute;riauth&egrave;que Solidaire", "Mat&eacute;riaux, mobilier, &eacute;quipements et ressources encore exploitables", "Collecte, inventaire, tri, stockage, orientation vers projets valid&eacute;s"),
        ("Commerce Vivant", "Locaux commerciaux inoccup&eacute;s, rez-de-chauss&eacute;e, boutiques", "Usages transitoires, boutiques tests, ateliers, activit&eacute;s de proximit&eacute;"),
        ("Friches & Terrains Vivants", "Espaces abandonn&eacute;s, terrains inutilis&eacute;s, dents creuses, friches", "Diagnostic d'usage, projet temporaire, jardin, espace collectif, pr&eacute;figuration"),
        ("Solidarit&eacute; & Insertion", "B&eacute;n&eacute;voles, publics &agrave; remobiliser, acteurs de l'insertion", "Ateliers, chantiers encadr&eacute;s, orientation vers structures comp&eacute;tentes, engagement citoyen"),
    ]
    table(doc, ["P&ocirc;le", "Sujet trait&eacute;", "Contribution au pilote"], rows, [3.6, 5.7, 7.1], fill=LIGHT_GREEN)


def method(doc):
    doc.add_heading(u("5. M&eacute;thode d'installation du pilote"), level=1)
    add_paragraph(
        doc,
        "La mise en place du pilote doit &ecirc;tre progressive. TVF doit d'abord construire un cadre fiable avant "
        "d'annoncer des r&eacute;sultats : r&eacute;ception des demandes, qualification, priorisation, convention, suivi, "
        "puis retour d'exp&eacute;rience. Cette logique &eacute;vite les promesses trop rapides et permet de pr&eacute;senter "
        "aux partenaires un fonctionnement clair.",
    )
    rows = [
        ("1. Rep&eacute;rer", "Identifier un logement, un commerce, une friche, un local de stockage, des mat&eacute;riaux ou un besoin d'acteur local."),
        ("2. Qualifier", "V&eacute;rifier l'adresse, le propri&eacute;taire, l'&eacute;tat, les risques, les documents disponibles et l'utilit&eacute; territoriale."),
        ("3. Prioriser", "Classer les dossiers selon faisabilit&eacute;, urgence, besoin public, co&ucirc;t, s&eacute;curit&eacute;, partenaires et impact attendu."),
        ("4. Cadrer", "Choisir le type de convention : mise &agrave; disposition, valorisation, partenariat, diagnostic, stockage, chantier ou coop&eacute;ration."),
        ("5. Mobiliser", "Associer collectivit&eacute;, propri&eacute;taire, entreprise, artisan, association, structure d'insertion, financeur ou b&eacute;n&eacute;voles selon le sujet."),
        ("6. Suivre", "Documenter les d&eacute;cisions, photos, pi&egrave;ces, volumes, heures, destination des ressources, blocages et suites possibles."),
    ]
    table(doc, ["&Eacute;tape", "Objectif"], rows, [3.4, 13.0], fill=LIGHT_BLUE)


def project_examples(doc):
    doc.add_page_break()
    doc.add_heading(u("6. Exemples de projets faisables &agrave; &eacute;tudier"), level=1)
    add_paragraph(
        doc,
        "Les exemples ci-dessous sont des sc&eacute;narios de travail. Ils ne constituent ni des projets r&eacute;alis&eacute;s, "
        "ni des engagements de partenaires, ni des autorisations. Ils servent &agrave; montrer comment TVF peut "
        "traduire un besoin local en dossier d'instruction concret.",
    )
    rows = [
        (
            "Logement dormant",
            "Un propri&eacute;taire dispose d'un logement inutilis&eacute; mais ne sait pas comment le remettre en usage.",
            "Diagnostic, estimation du niveau de travaux, recherche de mat&eacute;riaux, convention de valorisation, orientation vers usage &eacute;tudiant, associatif, interg&eacute;n&eacute;rationnel ou solidaire.",
        ),
        (
            "Commerce ferm&eacute;",
            "Un rez-de-chauss&eacute;e vacant p&egrave;se sur une rue ou un quartier.",
            "Boutique test, atelier d'artisan, local associatif, espace de services, activit&eacute; ESS ou commerce de proximit&eacute; si le besoin local est confirm&eacute;.",
        ),
        (
            "Mat&eacute;riaux disponibles",
            "Une entreprise ou un particulier d&eacute;tient du bois, des portes, du carrelage, des sanitaires ou du mobilier inutilis&eacute;s.",
            "Fiche ressource, photos, quantification, tri, stockage, d&eacute;cision d'affectation &agrave; un projet valid&eacute; par TVF.",
        ),
        (
            "Friche ou terrain d&eacute;laiss&eacute;",
            "Un espace ne peut pas encore faire l'objet d'un projet lourd mais peut accueillir un usage transitoire.",
            "Jardin partag&eacute;, espace p&eacute;dagogique, micro-verger, espace associatif temporaire ou pr&eacute;figuration d'usage.",
        ),
        (
            "Chantier encadr&eacute;",
            "Un chantier l&eacute;ger peut devenir support d'apprentissage, de b&eacute;n&eacute;volat ou de remobilisation.",
            "Travaux de second oeuvre, remise en peinture, tri, nettoyage, am&eacute;nagement, toujours sous cadre de s&eacute;curit&eacute; et avec professionnels si n&eacute;cessaire.",
        ),
    ]
    table(doc, ["Cas d'usage", "Point de d&eacute;part", "Action TVF possible"], rows, [3.3, 5.3, 7.8], fill=LIGHT_GREEN)
    if PHOTO_COMMERCE.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(PHOTO_COMMERCE), width=Cm(14.2))
        caption(doc, "Les locaux et rues commer&ccedil;antes peuvent faire l'objet d'une qualification fine avant toute proposition d'usage.")


def stakeholders(doc):
    doc.add_heading(u("7. B&eacute;n&eacute;fices par public"), level=1)
    rows = [
        ("Collectivit&eacute;s", "Disposer d'un outil de rep&eacute;rage et de mise en relation, compl&eacute;mentaire des politiques habitat, commerce, foncier, insertion, ESS et transition &eacute;cologique."),
        ("Propri&eacute;taires", "Comprendre les options possibles pour un bien dormant : diagnostic, convention d'usage, remise en &eacute;tat, usage temporaire, valorisation patrimoniale."),
        ("Entreprises", "Transformer des stocks dormants, invendus ou &eacute;quipements inutilis&eacute;s en ressources territoriales, avec tra&ccedil;abilit&eacute; et valorisation RSE."),
        ("Associations", "Acc&eacute;der &agrave; des locaux, mat&eacute;riaux ou comp&eacute;tences sous cadre valid&eacute;, pour d&eacute;velopper des activit&eacute;s utiles aux habitants."),
        ("Financeurs / m&eacute;c&egrave;nes", "Soutenir des dossiers mieux qualifi&eacute;s, avec indicateurs, pi&egrave;ces, conventions, suivi et logique d'impact v&eacute;rifiable."),
        ("Habitants / b&eacute;n&eacute;voles", "Participer &agrave; des actions concr&egrave;tes : signaler, aider, trier, participer &agrave; des ateliers, contribuer &agrave; la remise en vie de lieux utiles."),
    ]
    table(doc, ["Public", "Int&eacute;r&ecirc;t d'une coop&eacute;ration avec TVF"], rows, [3.8, 12.6], fill=LIGHT_BLUE)


def needs(doc):
    doc.add_page_break()
    doc.add_heading(u("8. Besoins concrets de lancement &agrave; Saint-&Eacute;tienne"), level=1)
    add_paragraph(
        doc,
        "Pour passer de l'intention &agrave; l'action, TVF doit r&eacute;unir des moyens simples mais structurants. "
        "Ces besoins ne sont pas des acquis : ils constituent les objets de contact, de demande et de convention "
        "&agrave; travailler avec les acteurs publics, priv&eacute;s, associatifs et &eacute;conomiques.",
    )
    rows = [
        ("Local de stockage", "Espace couvert, accessible, assur&eacute;, permettant de trier, inventorier, photographier et conserver temporairement les ressources."),
        ("Solution de transport", "V&eacute;hicule utilitaire, partenariat logistique, pr&ecirc;t ponctuel ou convention avec un acteur disposant de moyens adapt&eacute;s."),
        ("Donn&eacute;es et rep&eacute;rage", "Mise en relation avec services habitat, foncier, commerce, insertion, ESS, transition &eacute;cologique et acteurs de quartier."),
        ("Mat&eacute;riaux et &eacute;quipements", "Bois, menuiseries, carrelage, sanitaires, mobilier, outillage, &eacute;quipements r&eacute;utilisables, sous r&eacute;serve d'&eacute;tat et de tra&ccedil;abilit&eacute;."),
        ("Financement d'amor&ccedil;age", "Moyens pour l'assurance, la logistique, les documents, les diagnostics, les premiers d&eacute;placements, l'outillage et la coordination."),
        ("Comp&eacute;tences", "Artisans, architectes, techniciens, juristes, urbanistes, structures d'insertion, experts s&eacute;curit&eacute;, b&eacute;n&eacute;voles qualifi&eacute;s."),
    ]
    table(doc, ["Besoin", "Utilit&eacute; pour le pilote"], rows, [4.0, 12.4], fill=LIGHT_GREEN)


def cooperation_frame(doc):
    doc.add_heading(u("9. Cadre de coop&eacute;ration et points de vigilance"), level=1)
    add_paragraph(
        doc,
        "La cr&eacute;dibilit&eacute; du pilote repose sur un principe simple : aucune action ne doit &ecirc;tre engag&eacute;e sans "
        "cadre clair. Une demande ne vaut pas acceptation. Une id&eacute;e ne vaut pas projet. Une mise en relation ne vaut "
        "pas convention. TVF doit donc travailler par &eacute;tapes, avec des pi&egrave;ces, des validations et des limites.",
    )
    rows = [
        ("Acceptation", "Chaque demande est examin&eacute;e. TVF peut refuser, ajourner ou r&eacute;orienter un dossier si les conditions ne sont pas r&eacute;unies."),
        ("Travaux", "TVF vise d'abord la remise en usage, l'am&eacute;nagement, le second oeuvre, le r&eacute;emploi et les chantiers encadr&eacute;s. Les travaux lourds ou dangereux rel&egrave;vent de professionnels qualifi&eacute;s."),
        ("Diagnostic", "TVF peut effectuer un premier rep&eacute;rage mais peut aussi orienter vers un professionnel, un bureau d'&eacute;tudes ou une structure comp&eacute;tente."),
        ("Convention", "Toute mise &agrave; disposition de bien, local, mat&eacute;riaux, stockage ou ressources doit &ecirc;tre formalis&eacute;e par un &eacute;crit adapt&eacute;."),
        ("Tra&ccedil;abilit&eacute;", "Photos, origine, propri&eacute;t&eacute;, &eacute;tat, quantit&eacute;, destination, responsabilit&eacute;s et d&eacute;cisions doivent &ecirc;tre conserv&eacute;s dans le dossier."),
        ("Communication", "Aucun logo, soutien, partenariat ou r&eacute;sultat ne doit &ecirc;tre annonc&eacute; sans accord &eacute;crit des parties concern&eacute;es."),
    ]
    table(doc, ["Sujet", "R&egrave;gle de prudence"], rows, [3.5, 12.9], fill=LIGHT_BLUE)


def indicators(doc):
    doc.add_page_break()
    doc.add_heading(u("10. Indicateurs de suivi et reproductibilit&eacute;"), level=1)
    add_paragraph(
        doc,
        "Le pilote ne doit pas seulement produire des actions ponctuelles. Il doit produire une m&eacute;thode mesurable. "
        "Les indicateurs ci-dessous ne sont pas des r&eacute;sultats acquis : ce sont les indicateurs &agrave; suivre lorsque "
        "les premiers dossiers auront &eacute;t&eacute; ouverts, conventionn&eacute;s et mis en oeuvre.",
    )
    rows = [
        ("Patrimoine", "Biens rep&eacute;r&eacute;s, biens qualifi&eacute;s, dossiers ouverts, dossiers conventionn&eacute;s, biens remis en usage."),
        ("Mat&eacute;riaux", "Ressources propos&eacute;es, ressources accept&eacute;es, volumes estim&eacute;s, destinations, refus motiv&eacute;s, r&eacute;emploi effectif."),
        ("Commerce", "Locaux rep&eacute;r&eacute;s, besoins d'activit&eacute;s, contacts propri&eacute;taires, usages temporaires, projets accompagn&eacute;s."),
        ("Insertion / b&eacute;n&eacute;volat", "B&eacute;n&eacute;voles mobilis&eacute;s, ateliers organis&eacute;s, heures encadr&eacute;es, structures associ&eacute;es, suites de parcours."),
        ("Impact territorial", "Services rendus, publics concern&eacute;s, quartiers touch&eacute;s, coop&eacute;rations ouvertes, documents produits, blocages lev&eacute;s."),
    ]
    table(doc, ["Famille d'indicateurs", "Suivi recommand&eacute;"], rows, [4.2, 12.2], fill=LIGHT_GREEN)
    add_callout(
        doc,
        "Duplicable demain",
        "La m&eacute;thode construite &agrave; Saint-&Eacute;tienne pourra &ecirc;tre reprise dans une autre commune ou intercommunalit&eacute; : "
        "diagnostic territorial, acteurs locaux, besoins prioritaires, ressources disponibles, conventions, indicateurs et retour d'exp&eacute;rience.",
        fill=LIGHT_BLUE,
    )


def next_step(doc):
    doc.add_heading(u("Proposition de prochaine &eacute;tape"), level=1)
    add_paragraph(
        doc,
        "TVF peut proposer un premier rendez-vous de cadrage avec les acteurs int&eacute;ress&eacute;s afin d'identifier "
        "trois &agrave; cinq dossiers tests : un sujet habitat, un sujet mat&eacute;riaux, un sujet commerce ou local vacant, "
        "un sujet friche ou terrain, et un sujet insertion ou b&eacute;n&eacute;volat. L'objectif n'est pas de tout traiter "
        "imm&eacute;diatement, mais de mettre en place une cha&icirc;ne de travail simple, tra&ccedil;able et reproductible.",
    )
    rows = [
        ("Rendez-vous 1", "Clarifier les attentes, acteurs concern&eacute;s et besoins prioritaires."),
        ("Dossier test", "Choisir un cas concret, rassembler les pi&egrave;ces et v&eacute;rifier la faisabilit&eacute;."),
        ("Note de cadrage", "R&eacute;sumer le besoin, les acteurs, les risques, les options et la convention possible."),
        ("D&eacute;cision", "Poursuivre, r&eacute;orienter, ajourner ou refuser selon les conditions r&eacute;elles."),
        ("Retour d'exp&eacute;rience", "Documenter ce qui fonctionne, ce qui bloque et ce qui doit &ecirc;tre am&eacute;lior&eacute;."),
    ]
    table(doc, ["Action", "But"], rows, [4.2, 12.2], fill=LIGHT_BLUE)


def sources(doc):
    doc.add_heading(u("Sources utilis&eacute;es"), level=1)
    for item in [
        "INSEE, Dossier complet Commune de Saint-&Eacute;tienne (42218), donn&eacute;es RP2023, paru le 9 juillet 2026 : https://www.insee.fr/fr/statistiques/2011101?geo=COM-42218",
        "INSEE-DGFiP-Cnaf-Cnav-Ccmsa, Filosofi 2023, indicateurs de niveau de vie et pauvret&eacute;.",
        "Minist&egrave;re de la Transition &eacute;cologique, fili&egrave;re REP Produits et Mat&eacute;riaux de Construction du secteur du B&acirc;timent (PMCB), mise &agrave; jour 2025 : https://www.ecologie.gouv.fr/politiques-publiques/produits-materiaux-construction-du-secteur-du-batiment-pmcb",
        "Banque des Territoires / Localtis, article du 16 janvier 2024 sur la vacance des logements en France : https://www.banquedesterritoires.fr/la-france-compte-plus-de-3-millions-de-logements-vacants",
        "Cerema, Cartofriches, outil national de connaissance des friches, &agrave; mobiliser pour les fiches locales disponibles : https://cartofriches.cerema.fr/cartofriches/",
        "Documents internes TVF : statuts, r&eacute;c&eacute;piss&eacute; RNA, avis de situation SIRENE, documents de m&eacute;thode et brochures th&eacute;matiques.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Contact",
        "TERRITOIRES VIVANTS FRANCE - 25 rue &Eacute;lise Gervais, 42000 Saint-&Eacute;tienne - "
        "contact@territoiresvivantsfrance.fr - 04 65 81 54 69 - www.territoiresvivantsfrance.fr",
        fill=LIGHT_GREEN,
    )


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    cover(doc)
    doc.add_page_break()
    toc(doc)
    why_saint_etienne(doc)
    diagnosis(doc)
    data_meaning(doc)
    poles_response(doc)
    method(doc)
    project_examples(doc)
    stakeholders(doc)
    needs(doc)
    cooperation_frame(doc)
    indicators(doc)
    next_step(doc)
    sources(doc)
    doc.save(DOCX_PATH)


def build_md():
    md = u(
        """# Brochure territoriale TVF - Saint-&Eacute;tienne, territoire pilote

Document de pr&eacute;sentation destin&eacute; aux collectivit&eacute;s, financeurs, entreprises, associations, propri&eacute;taires et partenaires souhaitant comprendre l'int&eacute;r&ecirc;t du pilote st&eacute;phanois.

## Message central

Saint-&Eacute;tienne constitue le premier territoire d'ancrage de TERRITOIRES VIVANTS FRANCE. Le pilote doit permettre de structurer une m&eacute;thode reproductible : rep&eacute;rer les biens et ressources inutilis&eacute;s, qualifier les dossiers, mobiliser les acteurs, conventionner les coop&eacute;rations, suivre les actions et documenter les impacts.

## Donn&eacute;es rep&egrave;res

- 173 136 habitants en 2023 &agrave; Saint-&Eacute;tienne, selon l'INSEE.
- 101 341 logements en 2023, selon l'INSEE.
- 12 175 logements vacants, soit 12,0 % du parc en 2023, selon l'INSEE.
- 19,0 % de ch&ocirc;mage des 15-64 ans au sens du recensement en 2023, selon l'INSEE.
- 30,4 % de pauvret&eacute; au seuil de 60 % du niveau de vie m&eacute;dian en 2023, selon Filosofi.
- Environ 42 millions de tonnes de d&eacute;chets du b&acirc;timent par an en France, selon le Minist&egrave;re de la Transition &eacute;cologique.

## Sources principales

- INSEE, Dossier complet Commune de Saint-&Eacute;tienne (42218) : https://www.insee.fr/fr/statistiques/2011101?geo=COM-42218
- Minist&egrave;re de la Transition &eacute;cologique, fili&egrave;re PMCB : https://www.ecologie.gouv.fr/politiques-publiques/produits-materiaux-construction-du-secteur-du-batiment-pmcb
- Banque des Territoires / Localtis, vacance des logements : https://www.banquedesterritoires.fr/la-france-compte-plus-de-3-millions-de-logements-vacants
- Cerema, Cartofriches : https://cartofriches.cerema.fr/cartofriches/

## Pr&eacute;cautions

Les exemples de projets sont des pistes &agrave; &eacute;tudier, pas des projets r&eacute;alis&eacute;s. Toute action doit &ecirc;tre qualifi&eacute;e, valid&eacute;e et conventionn&eacute;e. Une demande ne vaut pas acceptation ni engagement de r&eacute;alisation.

## Contact

TERRITOIRES VIVANTS FRANCE - 25 rue &Eacute;lise Gervais, 42000 Saint-&Eacute;tienne - contact@territoiresvivantsfrance.fr - 04 65 81 54 69 - www.territoiresvivantsfrance.fr
"""
    )
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
