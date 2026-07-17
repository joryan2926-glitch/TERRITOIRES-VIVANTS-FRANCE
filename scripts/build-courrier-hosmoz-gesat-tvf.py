from pathlib import Path
from runpy import run_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

OUT_DIR = ROOT / "documents" / "courriers-lancement-saint-etienne"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "32-hosmoz-reseau-gesat-lancement-achats-inclusifs-tvf.docx"
MD_PATH = SRC_DIR / "courrier-hosmoz-reseau-gesat-lancement-tvf.md"

txt = BASE["txt"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_number = BASE["add_number"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
BORDER = BASE["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        ("Destinataire", "Hosmoz / R&eacute;seau Gesat - &agrave; l'attention de la direction, du p&ocirc;le ESAT-EA et du p&ocirc;le entreprises priv&eacute;es / organismes publics"),
        ("Adresse", "Adresse op&eacute;rationnelle &agrave; confirmer avant envoi - demande &agrave; orienter via les canaux officiels Hosmoz / R&eacute;seau Gesat"),
        ("Objet", "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les passerelles possibles avec les ESAT-EA, les achats inclusifs, le r&eacute;emploi et les projets de revitalisation territoriale"),
        ("Pi&egrave;ces jointes", "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche Solidarit&eacute; & Insertion, fiche mat&eacute;riauth&egrave;que, fiche besoins prestations / achats inclusifs"),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def add_inclusive_table(doc):
    doc.add_heading(txt("Passerelles possibles avec Hosmoz / R&eacute;seau Gesat"), level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.6, 6.3, 6.2])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Sujet", "Besoin potentiel de TVF", "Apport possible du r&eacute;seau ESAT-EA"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(head))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.1)
    rows = [
        ("Tri et pr&eacute;paration de mat&eacute;riaux", "Organiser le tri, l'&eacute;tiquetage, la pr&eacute;paration et le conditionnement de mat&eacute;riaux r&eacute;employables.", "Identifier des ESAT-EA capables d'assurer des prestations adapt&eacute;es, tra&ccedil;ables et compatibles avec les contraintes de s&eacute;curit&eacute;."),
        ("Logistique et manutention", "Structurer la collecte, la pr&eacute;paration, l'inventaire et le rangement dans un futur local de stockage TVF.", "Explorer les prestations logistiques, de conditionnement, de petite manutention ou de gestion de flux propos&eacute;es par le secteur prot&eacute;g&eacute; et adapt&eacute;."),
        ("Mobilier et reconditionnement", "R&eacute;employer du mobilier, des &eacute;quipements et des ressources issues d'entreprises, collectivit&eacute;s ou particuliers.", "Orienter TVF vers des prestataires inclusifs pouvant intervenir sur nettoyage, petite remise en &eacute;tat, montage, emballage ou valorisation."),
        ("Achats responsables", "Permettre aux futurs partenaires TVF d'int&eacute;grer une logique d'achats inclusifs dans les projets de revitalisation.", "Aider &agrave; identifier les fili&egrave;res ESAT-EA pertinentes et les modalit&eacute;s de consultation ou de sourcing responsables."),
        ("Impact social", "Associer les projets TVF &agrave; une logique d'inclusion, d'emploi prot&eacute;g&eacute; ou adapt&eacute; et de d&eacute;veloppement territorial.", "Cadrer les indicateurs, les limites de communication, les preuves &agrave; conserver et les conditions de valorisation."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.25)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - Hosmoz / R&eacute;seau Gesat"))
    r.bold = True
    r.font.size = Pt(16.5)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Demande d'un premier &eacute;change sur les ESAT-EA, les achats inclusifs, le r&eacute;emploi et les besoins op&eacute;rationnels TVF"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(doc, "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF porte une d&eacute;marche appel&eacute;e &agrave; se d&eacute;velopper &agrave; l'&eacute;chelle nationale, avec un premier ancrage op&eacute;rationnel progressif sur son territoire d'origine.")
    add_paragraph(doc, "L'association agit sur les ressources inutilis&eacute;es : logements vacants, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. Pour lancer son action &agrave; Saint-&Eacute;tienne, TVF doit structurer une cha&icirc;ne op&eacute;rationnelle simple : identifier, collecter, trier, stocker, orienter, r&eacute;employer et suivre l'impact des ressources confi&eacute;es.")
    add_paragraph(doc, "Le site officiel Hosmoz indique que Hosmoz est la marque issue de la fusion du R&eacute;seau Gesat et de Handeco, t&ecirc;te de r&eacute;seau &eacute;conomique nationale des ESAT et Entreprises Adapt&eacute;es. Il pr&eacute;sente un r&ocirc;le de promotion des savoir-faire, d'accompagnement des ESAT-EA et de d&eacute;veloppement des achats responsables des clients priv&eacute;s et publics. Cette mission rejoint directement les besoins futurs de TVF autour de l'inclusion, du r&eacute;emploi et de la logistique responsable.")

    add_callout(doc, "Demande principale", "TVF sollicite un premier rendez-vous afin d'identifier les passerelles possibles avec Hosmoz / R&eacute;seau Gesat : sourcing d'ESAT-EA, achats inclusifs, prestations de tri, logistique, reconditionnement, appui aux projets territoriaux et cadrage des conditions de coop&eacute;ration.")

    doc.add_heading(txt("Pourquoi solliciter Hosmoz / R&eacute;seau Gesat"), level=1)
    add_paragraph(doc, "TVF souhaite construire une mat&eacute;riauth&egrave;que solidaire et une capacit&eacute; de valorisation des ressources inutilis&eacute;es. Ces activit&eacute;s peuvent n&eacute;cessiter des prestations r&eacute;currentes : tri, inventaire, pr&eacute;paration, manutention, nettoyage, conditionnement, petite remise en &eacute;tat, livraison ou gestion de flux.")
    add_paragraph(doc, "Les ESAT et Entreprises Adapt&eacute;es peuvent constituer des partenaires op&eacute;rationnels pertinents lorsque les besoins sont correctement cadr&eacute;s, que les volumes sont r&eacute;alistes, que les responsabilit&eacute;s sont pr&eacute;cis&eacute;es et que les prestations correspondent aux savoir-faire disponibles. TVF souhaite donc comprendre comment structurer ses futurs besoins de mani&egrave;re lisible, responsable et compatible avec les pratiques du secteur prot&eacute;g&eacute; et adapt&eacute;.")
    add_paragraph(doc, "L'objectif n'est pas de solliciter un soutien automatique ni de faire appara&icirc;tre un partenariat avant accord. Il s'agit d'obtenir un premier cadrage sur les bonnes pratiques, les interlocuteurs, les prestations envisageables et les conditions &agrave; respecter pour travailler avec des ESAT-EA dans une logique d'achats inclusifs.")

    add_inclusive_table(doc)

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "Identifier les typologies de prestations ESAT-EA compatibles avec les futurs besoins TVF : tri, logistique, reconditionnement, conditionnement, mobilier, espaces verts, nettoyage, num&eacute;risation ou services support.")
    add_bullet(doc, "Comprendre la m&eacute;thode de sourcing ou de consultation &agrave; suivre pour travailler correctement avec des ESAT-EA.")
    add_bullet(doc, "Pr&eacute;ciser les informations &agrave; fournir pour qualifier un besoin : volume, fr&eacute;quence, lieu, contraintes, d&eacute;lai, assurance, s&eacute;curit&eacute;, transport et responsabilit&eacute;s.")
    add_bullet(doc, "Identifier si un premier besoin pilote &agrave; Saint-&Eacute;tienne peut &ecirc;tre formul&eacute; de mani&egrave;re simple et r&eacute;aliste, notamment autour de la future mat&eacute;riauth&egrave;que TVF.")
    add_bullet(doc, "D&eacute;finir les r&egrave;gles de communication : aucune mention publique de partenariat, logo, soutien ou engagement sans validation &eacute;crite pr&eacute;alable.")

    doc.add_heading(txt("Ce que TVF peut apporter"), level=1)
    add_number(doc, "Des besoins territoriaux concrets li&eacute;s &agrave; la revitalisation : mat&eacute;riaux, mobilier, locaux, logistique, projets de quartier.")
    add_number(doc, "Une capacit&eacute; de pr&eacute;qualification : nature du besoin, localisation, volumes estim&eacute;s, contraintes connues, acteur demandeur, calendrier.")
    add_number(doc, "Un cadre de coop&eacute;ration &agrave; construire avec les collectivit&eacute;s, entreprises, artisans, associations, propri&eacute;taires et financeurs.")
    add_number(doc, "Une volont&eacute; d'int&eacute;grer l'inclusion et les achats responsables dans les projets de r&eacute;emploi, sans confusion entre prestation, don, b&eacute;n&eacute;volat ou partenariat.")
    add_number(doc, "Une tra&ccedil;abilit&eacute; des dossiers et des impacts : ressources orient&eacute;es, prestations mobilis&eacute;es, acteurs associ&eacute;s, suites donn&eacute;es.")

    add_callout(doc, "Principe de prudence", "TVF ne se substitue pas &agrave; Hosmoz / R&eacute;seau Gesat, aux ESAT-EA, aux donneurs d'ordre, aux acheteurs publics ou priv&eacute;s, ni aux structures sp&eacute;cialis&eacute;es du handicap. Aucune mention de partenariat, de soutien, d'adh&eacute;sion, de labellisation, de r&eacute;f&eacute;rencement ou de financement ne sera utilis&eacute;e publiquement sans accord &eacute;crit pr&eacute;alable.", fill=LIGHT_BLUE)

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(doc, "Nous proposons un rendez-vous d'environ une heure afin de pr&eacute;senter TVF, de comprendre les modalit&eacute;s de recours aux ESAT-EA, d'identifier les prestations adapt&eacute;es aux besoins de lancement et de fixer les points de vigilance avant toute sollicitation op&eacute;rationnelle.")
    add_paragraph(doc, "&Agrave; l'issue de cet &eacute;change, TVF pourra formaliser une premi&egrave;re fiche besoin : objet, prestations attendues, volumes, lieux, contraintes, calendrier, documents utiles et conditions de consultation ou d'orientation vers les acteurs comp&eacute;tents.")

    doc.add_heading("Conclusion", level=1)
    add_paragraph(doc, "Territoires Vivants France souhaite construire son lancement avec une dimension sociale et inclusive forte. Les projets de r&eacute;emploi, de logistique territoriale et de revitalisation peuvent devenir des supports utiles pour d&eacute;velopper des achats responsables, valoriser les savoir-faire du secteur prot&eacute;g&eacute; et adapt&eacute; et renforcer l'impact social des projets locaux.")
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(head))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration officielle de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche mat&eacute;riauth&egrave;que", "D&eacute;crire les besoins de tri, stockage, valorisation et orientation des ressources.", "&Agrave; joindre"),
        ("Fiche prestations / achats inclusifs", "Qualifier les prestations pouvant &ecirc;tre confi&eacute;es &agrave; des ESAT-EA selon les besoins TVF.", "&Agrave; joindre"),
        ("Fiche Solidarit&eacute; & Insertion", "Pr&eacute;senter la dimension inclusion, emploi, parcours et impact social.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(doc, "Objet d'e-mail recommand&eacute;", "Demande de rendez-vous - TERRITOIRES VIVANTS FRANCE / Hosmoz - R&eacute;seau Gesat - ESAT-EA, achats inclusifs et r&eacute;emploi", fill=LIGHT_BLUE)
    doc.add_heading("Sources de cadrage", level=1)
    add_paragraph(doc, "Le courrier s'appuie sur les informations publiques du site Hosmoz / R&eacute;seau Gesat relatives au r&eacute;seau &eacute;conomique des ESAT et Entreprises Adapt&eacute;es, aux achats inclusifs, &agrave; la promotion des savoir-faire et &agrave; l'accompagnement des clients priv&eacute;s et publics.")
    doc.add_heading("Point &agrave; confirmer", level=1)
    add_paragraph(doc, "L'adresse ou le canal d'envoi op&eacute;rationnel doit &ecirc;tre confirm&eacute; avant transmission afin de viser le bon p&ocirc;le ou le bon formulaire officiel.")

    doc.save(DOCX_PATH)


def build_md():
    md = txt("""# Courrier - Hosmoz / R&eacute;seau Gesat - lancement TVF

Objet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les passerelles possibles avec les ESAT-EA, les achats inclusifs, le r&eacute;emploi et les projets de revitalisation territoriale.

Intention : solliciter un premier rendez-vous pour comprendre les modalit&eacute;s de recours aux ESAT-EA et la structuration de besoins de prestations inclusives.

Principe : aucune mention de partenariat, de soutien, d'adh&eacute;sion, de r&eacute;f&eacute;rencement ou de financement ne sera utilis&eacute;e sans accord &eacute;crit.
""")
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
