from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Image

ROOT = Path(r'C:\Users\jowst\Documents\TERRITOIRES VIVANTS FRANCE')
OUT = ROOT / 'documents' / 'formulaires-internes-tvf'
OUT.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / 'assets' / 'logo-territoires-vivants-france-web.png'
GREEN = RGBColor(24,57,47)
GRAY = RGBColor(102,112,133)
DARK = RGBColor(31,41,55)
PALE = 'F3F7F2'
BORDER = 'D9E0DC'
TVF = {
    'name':'Territoires Vivants France',
    'subtitle':'Agence Territoriale de Revitalisation Immobilière',
    'address':'25 rue Élise Gervais, 42000 Saint-Étienne',
    'phone':'04 65 81 54 69',
    'email':'contact@territoiresvivantsfrance.fr',
    'site':'www.territoiresvivantsfrance.fr',
    'rna':'RNA W922015538',
    'siret':'SIRET 897 226 138 00018',
}

FORMS = [
    {
        'slug': 'fiche-interne-premier-contact-tvf',
        'title': 'FICHE INTERNE TVF — PREMIER CONTACT',
        'subtitle': 'Traçabilité du premier échange — usage interne',
        'note': "Cette fiche sert à tracer le premier échange avec un propriétaire, représentant ou interlocuteur. Elle ne vaut pas acceptation automatique d’un accompagnement, ouverture définitive de dossier ou engagement de TVF.",
        'footer': 'Territoires Vivants France — document interne confidentiel — premier contact',
        'sections': [
            ('1. Référence interne', 'fields', ['Numéro demande / dossier', 'Date du contact', 'Agent TVF', 'Fiche bien liée', 'Fiche propriétaire liée']),
            ('2. Personne contactée', 'checks', ['Propriétaire', 'Héritier', 'Mandataire', 'Notaire', 'Agence / gestionnaire', 'Syndic', 'Collectivité', 'Autre : ____________________'], 2),
            ('Identité de l’interlocuteur', 'fields', ['Nom / prénom', 'Qualité', 'Téléphone', 'E-mail']),
            ('3. Mode de contact', 'checks', ['Téléphone', 'E-mail', 'Courrier', 'Rendez-vous', 'Réponse formulaire site', 'TVF Mobile', 'Accueil physique', 'Autre'], 2),
            ('4. Objet du contact', 'checks', ['Présenter TVF', 'Vérifier la propriété', 'Comprendre la situation du bien', 'Proposer un rendez-vous', 'Demander des documents', 'Présenter un accompagnement possible', 'Répondre à une demande', 'Autre'], 2),
            ('5. Résumé de l’échange', 'textarea', 5),
            ('6. Position de la personne', 'checks', ['Intéressée', 'À recontacter', 'Hésitante', 'Demande des informations', 'Souhaite un rendez-vous', 'Refuse l’échange', 'Ne se considère pas concernée', 'Autre'], 2),
            ('7. Informations recueillies', 'checks', ['Durée de vacance évoquée', 'Travaux à prévoir', 'Documents disponibles', 'Succession / indivision', 'Projet envisagé', 'Blocage administratif', 'Occupation à vérifier', 'Urgence ou sensibilité signalée'], 2),
            ('8. Documents demandés ou transmis', 'checks', ['Titre de propriété', 'Taxe foncière', 'Diagnostics', 'Photos', 'Plans', 'Devis', 'Coordonnées notaire / mandataire', 'Autre'], 2),
            ('9. Cadre et consentement', 'checks', ['Accord pour être recontacté', 'Accord e-mail', 'Accord rendez-vous', 'Donnée confidentielle', 'Information à vérifier', 'Refus de contact', 'Ne pas diffuser', 'Besoin validation responsable'], 2),
            ('10. Prochaine action', 'fields', ['Action à réaliser', 'Responsable', 'Date limite', 'Priorité']),
        ],
        'legal': "Document interne TVF. Les informations recueillies doivent être utilisées uniquement pour orienter la demande ou le dossier concerné, dans le respect des droits de la personne contactée et de la protection des données personnelles."
    },
    {
        'slug': 'fiche-interne-visite-diagnostic-tvf',
        'title': 'FICHE INTERNE TVF — VISITE ET DIAGNOSTIC INITIAL',
        'subtitle': 'Constat interne non réglementaire — usage interne',
        'note': "Cette fiche est un outil interne de visite et de diagnostic initial. Elle ne remplace pas un diagnostic technique, immobilier, énergétique, structurel, juridique ou sanitaire réalisé par un professionnel habilité.",
        'footer': 'Territoires Vivants France — document interne confidentiel — visite et diagnostic initial',
        'sections': [
            ('1. Référence interne', 'fields', ['Numéro dossier', 'Date de visite', 'Agent TVF', 'Bien concerné', 'Propriétaire / interlocuteur']),
            ('2. Conditions de visite', 'checks', ['Accord propriétaire', 'Accord mandataire', 'Depuis voie publique uniquement', 'Visite extérieure', 'Visite intérieure', 'Accès partiel', 'Accès impossible', 'Photos autorisées'], 2),
            ('3. Description générale du bien', 'fields', ['Type de bien', 'Surface estimée si connue', 'Nombre de niveaux', 'Usage initial supposé', 'Environnement proche']),
            ('4. État apparent', 'checks', ['Structure à vérifier', 'Toiture visible', 'Façade dégradée', 'Menuiseries à vérifier', 'Réseaux à vérifier', 'Humidité visible', 'Accès difficile', 'Parties communes concernées', 'Extérieurs délaissés', 'Compteurs visibles'], 2),
            ('5. Niveau de dégradation', 'checks', ['Bon état apparent', 'Travaux légers', 'Travaux moyens', 'Travaux lourds', 'Très dégradé', 'Risque visible', 'Diagnostic professionnel nécessaire', 'Information insuffisante'], 2),
            ('6. Photos et pièces relevées', 'checks', ['Photos façade', 'Photos intérieur', 'Photos accès', 'Photos dégradations', 'Photos extérieurs', 'Plans remis', 'Documents remis', 'Aucune pièce transmise'], 2),
            ('7. Points de vigilance', 'checks', ['Sécurité', 'Insalubrité supposée', 'Accès dangereux', 'Copropriété', 'Succession / indivision', 'Voisinage', 'Urbanisme', 'Occupation inconnue', 'Humidité', 'Risque juridique à vérifier'], 2),
            ('8. Potentiel de remise en usage', 'checks', ['Habitat', 'Commerce', 'Activité', 'Stockage', 'Projet associatif', 'Usage temporaire', 'Rénovation locative', 'Réemploi matériaux', 'Projet mixte', 'À déterminer'], 2),
            ('9. Besoins à confirmer', 'checks', ['Diagnostic technique', 'Devis', 'Avis architecte', 'Avis mairie', 'Relevé cadastral', 'Étude juridique', 'Estimation travaux', 'Aides possibles', 'Partenaire à mobiliser', 'Visite complémentaire'], 2),
            ('10. Conclusion de visite', 'checks', ['Poursuivre l’instruction', 'Informations insuffisantes', 'Visite complémentaire', 'Orientation partenaire', 'Non prioritaire', 'Classer sans suite', 'Préparer proposition TVF', 'Avis responsable nécessaire'], 2),
            ('11. Compte rendu synthétique', 'textarea', 5),
            ('12. Suite à donner', 'fields', ['Action à réaliser', 'Responsable', 'Date limite', 'Priorité']),
        ],
        'legal': "Rappel interne : aucune entrée dans une propriété privée ne doit être réalisée sans autorisation. Les constats visuels doivent être confirmés par des professionnels habilités lorsque la situation l’exige."
    },
    {
        'slug': 'fiche-interne-proposition-tvf',
        'title': 'FICHE INTERNE TVF — PROPOSITION D’ACCOMPAGNEMENT',
        'subtitle': 'Pistes d’orientation après analyse — usage interne',
        'note': "Cette fiche présente les pistes d’accompagnement envisageables par TVF. Elle ne constitue pas une promesse de financement, de travaux, de résultat ou de prise en charge automatique.",
        'footer': 'Territoires Vivants France — document interne confidentiel — proposition TVF',
        'sections': [
            ('1. Référence interne', 'fields', ['Numéro dossier', 'Date de proposition', 'Agent TVF', 'Bien concerné', 'Propriétaire / interlocuteur']),
            ('2. Rappel de la situation', 'textarea', 4),
            ('3. Objectif recherché', 'checks', ['Remise en usage', 'Valorisation patrimoniale', 'Rénovation', 'Conventionnement', 'Projet temporaire', 'Mobilisation de partenaires', 'Orientation vers aides', 'Autre'], 2),
            ('4. Proposition d’accompagnement TVF', 'checks', ['Appui administratif', 'Recherche partenaires', 'Recherche matériaux', 'Orientation entreprises', 'Suivi du dossier', 'Préparation documents', 'Mise en relation', 'Accompagnement convention'], 2),
            ('5. Scénarios étudiés', 'fields', ['Scénario 1', 'Scénario 2', 'Scénario 3', 'Scénario 4']),
            ('6. Conditions nécessaires', 'checks', ['Accord propriétaire', 'Documents à fournir', 'Visite complète', 'Devis nécessaires', 'Diagnostics nécessaires', 'Validation réglementaire', 'Budget à confirmer', 'Partenaires à mobiliser'], 2),
            ('7. Rôle de TVF', 'checks', ['Coordination', 'Accompagnement', 'Suivi', 'Mise en relation', 'Appui documentaire', 'Orientation', 'Recherche de solutions', 'Observation territoriale'], 2),
            ('8. Limites de l’intervention', 'checks', ['Pas de garantie d’aide', 'Pas de travaux sans accord', 'Pas de décision à la place du propriétaire', 'Pas d’engagement sans convention', 'Pas de diagnostic réglementaire si non habilité', 'Validation responsable requise'], 2),
            ('9. Documents à fournir', 'checks', ['Titre de propriété', 'Pièce d’identité', 'Taxe foncière', 'Diagnostics', 'Photos', 'Plans', 'Devis', 'Coordonnées notaire / mandataire', 'RIB si nécessaire', 'Autre'], 2),
            ('10. Décision proposée', 'checks', ['Poursuivre l’instruction', 'Programmer réunion', 'Préparer convention', 'Demander pièces', 'Orienter partenaire', 'Suspendre', 'Classer sans suite', 'Avis direction nécessaire'], 2),
            ('11. Avis et validation interne', 'fields', ['Avis chargé de mission', 'Avis responsable', 'Décision retenue', 'Date de validation']),
            ('12. Commentaire final', 'textarea', 4),
        ],
        'legal': "Cette proposition reste soumise à vérification, accord des parties, cadre légal applicable, moyens disponibles et formalisation éventuelle par convention ou document adapté."
    },
]


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right'):
        tag = 'w:' + edge
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), BORDER)


def set_cell(cell, text, bold=False, color=DARK, size=8.4, fill=None):
    cell.text = ''
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Inter'
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(cell)


def add_docx_header(doc, form):
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.columns[0].width = Inches(1.25)
    header.columns[1].width = Inches(5.5)
    if LOGO.exists():
        run = header.rows[0].cells[0].paragraphs[0].add_run()
        run.add_picture(str(LOGO), width=Inches(1.0))
    txt = header.rows[0].cells[1]
    txt.text = ''
    p = txt.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(TVF['name'] + '\n')
    r.font.name = 'Manrope'; r.font.size = Pt(12); r.bold = True; r.font.color.rgb = GREEN
    r = p.add_run(TVF['subtitle'] + '\n')
    r.font.name = 'Inter'; r.font.size = Pt(8.5); r.font.color.rgb = DARK
    r = p.add_run(f"{TVF['address']} | {TVF['phone']} | {TVF['email']}\n{TVF['rna']} | {TVF['siret']}")
    r.font.name = 'Inter'; r.font.size = Pt(7.2); r.font.color.rgb = GRAY
    for row in header.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run(form['title'])
    r.font.name = 'Manrope'; r.font.size = Pt(14.5); r.bold = True; r.font.color.rgb = GREEN
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run(form['subtitle'])
    r.font.name = 'Inter'; r.font.size = Pt(8.3); r.font.color.rgb = GRAY
    add_docx_note(doc, form['note'])


def section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = 'Manrope'; r.font.size = Pt(11.2); r.bold = True; r.font.color.rgb = GREEN


def docx_fields(doc, labels):
    for label in labels:
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        t.columns[0].width = Inches(2.2)
        t.columns[1].width = Inches(4.55)
        set_cell(t.rows[0].cells[0], label, True, GREEN, 8.4, PALE)
        set_cell(t.rows[0].cells[1], '', False, DARK, 8.4)


def docx_checks(doc, items, cols=2):
    rows = (len(items) + cols - 1) // cols
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for col in t.columns:
        col.width = Inches(6.75 / cols)
    idx = 0
    for r in range(rows):
        for c in range(cols):
            txt = '☐ ' + items[idx] if idx < len(items) else ''
            set_cell(t.rows[r].cells[c], txt, False, DARK, 8.0)
            idx += 1


def docx_textarea(doc, lines=4):
    t = doc.add_table(rows=lines, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        row.height = Inches(0.34)
        borders(row.cells[0])


def add_docx_note(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell(t.rows[0].cells[0], text, False, GRAY, 7.7, 'FBF8EF')


def build_docx(form):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.42)
    sec.bottom_margin = Inches(0.42)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)
    add_docx_header(doc, form)
    for s in form['sections']:
        title, kind = s[0], s[1]
        section(doc, title)
        if kind == 'fields':
            docx_fields(doc, s[2])
        elif kind == 'checks':
            docx_checks(doc, s[2], s[3])
        elif kind == 'textarea':
            docx_textarea(doc, s[2])
    add_docx_note(doc, form['legal'])
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run(form['footer'])
    rf.font.name = 'Inter'; rf.font.size = Pt(7); rf.font.color.rgb = GRAY
    path = OUT / (form['slug'] + '.docx')
    doc.save(path)
    return path


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='TitleTVF', fontName='Helvetica-Bold', fontSize=13.6, leading=16, textColor=colors.HexColor('#18392F'), alignment=1, spaceAfter=3))
    styles.add(ParagraphStyle(name='Sub', fontName='Helvetica', fontSize=8, leading=10, textColor=colors.HexColor('#667085'), alignment=1, spaceAfter=4))
    styles.add(ParagraphStyle(name='H', fontName='Helvetica-Bold', fontSize=10, leading=12, textColor=colors.HexColor('#18392F'), spaceBefore=5, spaceAfter=3))
    styles.add(ParagraphStyle(name='Small', fontName='Helvetica', fontSize=7.2, leading=9, textColor=colors.HexColor('#667085'), spaceAfter=4))
    styles.add(ParagraphStyle(name='Note', fontName='Helvetica', fontSize=7.4, leading=9.5, textColor=colors.HexColor('#475467'), backColor=colors.HexColor('#FBF8EF'), borderColor=colors.HexColor('#E4D7B8'), borderWidth=0.5, borderPadding=5, spaceAfter=5))
    return styles


def add_pdf_header(story, form, styles):
    if LOGO.exists():
        logo = Image(str(LOGO), width=24*mm, height=18*mm)
    else:
        logo = ''
    ht = Table([[logo, Paragraph(f"<b>{TVF['name']}</b><br/>{TVF['subtitle']}<br/>{TVF['address']} | {TVF['phone']} | {TVF['email']}<br/>{TVF['rna']} | {TVF['siret']}", styles['Small'])]], colWidths=[30*mm, 140*mm])
    ht.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'MIDDLE'),('ALIGN',(1,0),(1,0),'RIGHT')]))
    story.append(ht)
    story.append(Paragraph(form['title'], styles['TitleTVF']))
    story.append(Paragraph(form['subtitle'], styles['Sub']))
    story.append(Paragraph(form['note'], styles['Note']))


def pdf_fields(story, labels):
    t = Table([[l, ''] for l in labels], colWidths=[56*mm, 114*mm])
    t.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),0.45,colors.HexColor('#D9E0DC')),
        ('BACKGROUND',(0,0),(0,-1),colors.HexColor('#F3F7F2')),
        ('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),
        ('TEXTCOLOR',(0,0),(0,-1),colors.HexColor('#18392F')),
        ('FONTNAME',(1,0),(1,-1),'Helvetica'),
        ('FONTSIZE',(0,0),(-1,-1),7.15),
        ('LEADING',(0,0),(-1,-1),8.6),
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ('TOPPADDING',(0,0),(-1,-1),3.5),
        ('BOTTOMPADDING',(0,0),(-1,-1),3.5),
        ('LEFTPADDING',(0,0),(-1,-1),5),
        ('RIGHTPADDING',(0,0),(-1,-1),5),
    ]))
    story.append(t)


def pdf_checks(story, items, cols=2):
    rows=[]
    for i in range(0, len(items), cols):
        rows.append(['□ ' + x for x in items[i:i+cols]] + ['']*(cols-len(items[i:i+cols])))
    t = Table(rows, colWidths=[170*mm/cols]*cols)
    t.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),0.45,colors.HexColor('#D9E0DC')),
        ('FONTNAME',(0,0),(-1,-1),'Helvetica'),
        ('FONTSIZE',(0,0),(-1,-1),7.05),
        ('LEADING',(0,0),(-1,-1),8.6),
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ('TOPPADDING',(0,0),(-1,-1),3.5),
        ('BOTTOMPADDING',(0,0),(-1,-1),3.5),
        ('LEFTPADDING',(0,0),(-1,-1),5),
    ]))
    story.append(t)


def pdf_textarea(story, lines=4):
    t = Table([[''] for _ in range(lines)], colWidths=[170*mm])
    t.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),0.45,colors.HexColor('#D9E0DC')),
        ('TOPPADDING',(0,0),(-1,-1),7),
        ('BOTTOMPADDING',(0,0),(-1,-1),7),
    ]))
    story.append(t)


def build_pdf(form):
    styles = pdf_styles()
    story = []
    add_pdf_header(story, form, styles)
    for s in form['sections']:
        title, kind = s[0], s[1]
        story.append(Paragraph(title, styles['H']))
        if kind == 'fields':
            pdf_fields(story, s[2])
        elif kind == 'checks':
            pdf_checks(story, s[2], s[3])
        elif kind == 'textarea':
            pdf_textarea(story, s[2])
    story.append(Paragraph(form['legal'], styles['Note']))
    path = OUT / (form['slug'] + '.pdf')
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=13*mm, leftMargin=13*mm, topMargin=10*mm, bottomMargin=10*mm)
    doc.build(story)
    return path

for form in FORMS:
    print(build_docx(form))
    print(build_pdf(form))
