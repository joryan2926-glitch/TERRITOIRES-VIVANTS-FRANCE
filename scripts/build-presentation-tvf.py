from __future__ import annotations

import json
import math
import shutil
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
PHOTOS = ASSETS / "photos"
OUT_DOCS = ROOT / "output" / "documents"
OUT_PDF = ROOT / "output" / "pdf"
OUT_WORK = ROOT / "output" / "documents" / "presentation-tvf-assets"

DOCX_PATH = OUT_DOCS / "dossier-presentation-territoires-vivants-france.docx"
PDF_PATH = OUT_PDF / "dossier-presentation-territoires-vivants-france.pdf"
MD_PATH = OUT_DOCS / "dossier-presentation-territoires-vivants-france-source.md"

NAVY = "09243B"
GREEN = "2E7D32"
GREEN_DARK = "1F5A2D"
GREEN_LIGHT = "EAF5EB"
GOLD = "D8A21D"
BLUE = "0B3A5B"
INK = "13202E"
MUTED = "61717F"
OFFWHITE = "F7F5EF"
CARD = "F4F8F2"
BORDER = "D7E3D3"

def u(text: str) -> str:
    return text.encode("utf-8").decode("unicode_escape")


CONTACT = {
    "nom": "TERRITOIRES VIVANTS FRANCE",
    "adresse": "25 rue Élise Gervais, 42000 Saint-Étienne",
    "email": "contact@territoiresvivantsfrance.fr",
    "telephone": "04 65 81 54 69",
    "site": "https://www.territoiresvivantsfrance.fr",
    "rna": "W423016361",
    "siret": "107 226 128 00018",
}

SOURCES = [
    {
        "label": "INSEE, dossier complet Commune de Saint-Étienne (42218), paru le 09/07/2026",
        "url": "https://www.insee.fr/fr/statistiques/2011101?geo=COM-42218",
        "usage": "population, logement, emploi, revenus, établissements et équipements 2023-2026",
    },
    {
        "label": "INSEE, base permanente des équipements 2025",
        "url": "https://www.insee.fr/fr/statistiques/2011101?geo=COM-42218",
        "usage": "nombre de commerces et services à Saint-Étienne",
    },
    {
        "label": "Cour des comptes, rapport 2025 sur la lutte contre les logements vacants, cité par Le Monde",
        "url": "https://www.lemonde.fr/societe/article/2025/05/22/la-cour-des-comptes-deplore-les-faibles-resultats-de-la-lutte-contre-les-logements-vacants_6607818_3224.html",
        "usage": "ordre de grandeur national des logements vacants et besoin de coordination",
    },
]

DATA = {
    "population_2023": "173 136",
    "population_variation": "+0,1 % par an entre 2017 et 2023",
    "logements_vacants_2022": "12,2 %",
    "france_logements_vacants_2022": "8,0 %",
    "pauvreté_2023": "30,4 %",
    "niveau_vie_median_2023": "20 880 €",
    "taux_chomage_2023": "19,0 %",
    "chomage_15_24_2023": "27,1 %",
    "actifs_15_64_2023": "68,3 %",
    "etablissements_2023": "14 870",
    "unites_legales_2023": "13 131",
    "commerces_grandes_surfaces_2025": "47",
    "epiceries_superettes_2025": "123",
    "boulangeries_2025": "142",
    "coiffures_2025": "233",
}

PHOTOS_USED = {
    "cover": PHOTOS / "saint-etienne-panorama.jpg",
    "saint": PHOTOS / "france-saint-etienne-chateaucreux.jpg",
    "commerce": PHOTOS / "saint-etienne-centre-commerce.jpg",
    "habitat": PHOTOS / "immeuble-renovation-meudon.jpg",
    "materiaux": PHOTOS / "materiaux-reemploi-echantillons.jpg",
    "friche": PHOTOS / "france-friche-pcuk.jpg",
    "solidarite": PHOTOS / "chantier-renovation-lyon.jpg",
    "jardin": PHOTOS / "community-garden-paris.webp",
    "artisan": PHOTOS / "artisan-menuiserie-bois.jpg",
}


def ensure_dirs() -> None:
    OUT_DOCS.mkdir(parents=True, exist_ok=True)
    OUT_PDF.mkdir(parents=True, exist_ok=True)
    OUT_WORK.mkdir(parents=True, exist_ok=True)


def font_path() -> Path:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/segoeui.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return path
    raise FileNotFoundError("Aucune police système compatible trouvée.")


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf")
    if not path.exists():
        path = font_path()
    return ImageFont.truetype(str(path), size=size)


def hex_rgb(hex_color: str) -> tuple[int, int, int]:
    hex_color = hex_color.strip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def crop_image(src: Path, dst: Path, size: tuple[int, int]) -> Path:
    img = Image.open(src).convert("RGB")
    img = ImageOps.exif_transpose(img)
    img.thumbnail((size[0] * 2, size[1] * 2))
    img = ImageOps.fit(img, size, method=Image.Resampling.LANCZOS)
    img.save(dst, quality=88, optimize=True)
    return dst


def draw_round_rect(draw: ImageDraw.ImageDraw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def make_key_figures() -> Path:
    path = OUT_WORK / "graphique-chiffres-cles-saint-etienne.png"
    w, h = 1500, 720
    img = Image.new("RGB", (w, h), hex_rgb(OFFWHITE))
    d = ImageDraw.Draw(img)
    title = load_font(54, True)
    sub = load_font(30)
    big = load_font(52, True)
    small = load_font(25)
    d.text((70, 50), "Saint-Étienne : repères territoriaux sourcés", font=title, fill=hex_rgb(NAVY))
    d.text((72, 120), "Données publiques utilisées comme point de départ du diagnostic TVF.", font=sub, fill=hex_rgb(MUTED))
    cards = [
        ("173 136", "habitants en 2023", GREEN),
        ("12,2 %", "logements vacants en 2022", GOLD),
        ("30,4 %", "taux de pauvreté en 2023", BLUE),
        ("19,0 %", "taux de chômage en 2023", GREEN_DARK),
        ("14 870", "établissements actifs en 2023", NAVY),
    ]
    x = 70
    y = 220
    card_w = 260
    for value, label, color in cards:
        draw_round_rect(d, (x, y, x + card_w, y + 310), 34, "white", "#D7E3D3", 2)
        d.rectangle((x, y, x + card_w, y + 12), fill=hex_rgb(color))
        d.text((x + 28, y + 78), value, font=big, fill=hex_rgb(color))
        d.text((x + 28, y + 165), label, font=small, fill=hex_rgb(INK))
        x += card_w + 28
    d.text((72, 620), "Sources : INSEE, dossier complet Commune de Saint-Étienne (42218), publication 09/07/2026 ; logement 2022 repris dans les données INSEE.", font=load_font(22), fill=hex_rgb(MUTED))
    img.save(path, quality=92)
    return path


def make_process_diagram() -> Path:
    path = OUT_WORK / "schema-methode-tvf.png"
    w, h = 1600, 520
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title = load_font(48, True)
    step_font = load_font(26, True)
    body_font = load_font(22)
    d.text((60, 42), "Méthode TVF : du signalement à la remise en usage", font=title, fill=hex_rgb(NAVY))
    steps = [
        ("1", "Repérer", "Bien, matériau,\nbesoin local"),
        ("2", "Qualifier", "Données, visite,\npriorités"),
        ("3", "Structurer", "Dossier,\npartenaires"),
        ("4", "Conventionner", "Cadre écrit,\nresponsabilités"),
        ("5", "Agir", "Travaux,\nréemploi"),
        ("6", "Mesurer", "Impact,\nsuivi"),
    ]
    start_x = 70
    y = 170
    box_w = 220
    gap = 38
    for i, (num, title_s, body) in enumerate(steps):
        x = start_x + i * (box_w + gap)
        draw_round_rect(d, (x, y, x + box_w, y + 230), 26, hex_rgb(GREEN_LIGHT), "#BFD4BD", 2)
        d.ellipse((x + 22, y + 22, x + 74, y + 74), fill=hex_rgb(GREEN), outline=hex_rgb(GREEN))
        d.text((x + 39, y + 31), num, font=load_font(26, True), fill="white")
        d.text((x + 24, y + 92), title_s, font=step_font, fill=hex_rgb(NAVY))
        for line_idx, line in enumerate(body.split("\n")):
            d.text((x + 24, y + 135 + line_idx * 30), line, font=body_font, fill=hex_rgb(INK))
        if i < len(steps) - 1:
            ax = x + box_w + 8
            ay = y + 115
            d.line((ax, ay, ax + gap - 16, ay), fill=hex_rgb(GOLD), width=5)
            d.polygon([(ax + gap - 16, ay - 9), (ax + gap - 16, ay + 9), (ax + gap, ay)], fill=hex_rgb(GOLD))
    d.text((70, 445), "Principe : TVF ne remplace pas les acteurs existants ; l'association prépare, coordonne et rend lisible un parcours d'action local.", font=load_font(23), fill=hex_rgb(MUTED))
    img.save(path, quality=92)
    return path


def make_ecosystem_diagram() -> Path:
    path = OUT_WORK / "schema-ecosysteme-tvf.png"
    w, h = 1500, 950
    img = Image.new("RGB", (w, h), hex_rgb(OFFWHITE))
    d = ImageDraw.Draw(img)
    title = load_font(50, True)
    d.text((70, 50), "TVF : une plateforme de coopération territoriale", font=title, fill=hex_rgb(NAVY))
    center = (750, 500)
    d.ellipse((center[0] - 160, center[1] - 160, center[0] + 160, center[1] + 160), fill=hex_rgb(GREEN), outline=hex_rgb(GREEN_DARK), width=6)
    d.text((center[0] - 86, center[1] - 35), "TVF", font=load_font(72, True), fill="white")
    d.text((center[0] - 120, center[1] + 45), "coordination", font=load_font(30, True), fill="white")
    actors = [
        ("Collectivités", 750, 180),
        ("Propriétaires", 1110, 310),
        ("Entreprises", 1110, 690),
        ("Associations", 750, 825),
        ("Bénévoles", 390, 690),
        ("Artisans", 390, 310),
        ("Mécènes", 990, 160),
        ("Habitants", 510, 160),
    ]
    actor_font = load_font(27, True)
    for label, x, y in actors:
        d.line((center[0], center[1], x, y), fill="#C9D6C4", width=4)
        draw_round_rect(d, (x - 145, y - 45, x + 145, y + 45), 22, "white", "#C9D6C4", 2)
        tw = d.textlength(label, font=actor_font)
        d.text((x - tw / 2, y - 16), label, font=actor_font, fill=hex_rgb(NAVY))
    d.text((76, 890), "Objectif : relier les besoins, les ressources et les acteurs pour transformer les biens et matériaux inutilisés en usages utiles au territoire.", font=load_font(24), fill=hex_rgb(MUTED))
    img.save(path, quality=92)
    return path


def prepare_visuals() -> dict[str, Path]:
    visuals = {
        "key_figures": make_key_figures(),
        "process": make_process_diagram(),
        "ecosystem": make_ecosystem_diagram(),
    }
    for key, src in PHOTOS_USED.items():
        if src.exists():
            visuals[key] = crop_image(src, OUT_WORK / f"{key}.jpg", (1400, 780))
    logo = ASSETS / "logo-territoires-vivants-france-web.png"
    if not logo.exists():
        logo = ASSETS / "logo-territoires-vivants-france.png"
    visuals["logo"] = logo
    return visuals


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D7E3D3", sz="8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_pr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_text(paragraph, text: str, size=10.5, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    styles["Normal"].paragraph_format.space_after = Pt(10)
    styles["Normal"].paragraph_format.line_spacing = 1.28
    for name, size, color in [
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, GREEN_DARK),
        ("Heading 3", 12, BLUE),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(8)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    set_paragraph_text(
        p,
        f"{CONTACT['nom']} - {CONTACT['site']} - {CONTACT['email']} - {CONTACT['telephone']}",
        size=8.5,
        color=MUTED,
    )


def add_doc_title(doc: Document, text: str, subtitle: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.space_after = Pt(10)
    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_text(sp, subtitle, size=12.5, color=GREEN_DARK, bold=True)
        sp.paragraph_format.space_after = Pt(20)


def add_callout(doc: Document, text: str, title: str | None = None, fill: str = GREEN_LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=BORDER, sz="6")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if title:
        p = cell.paragraphs[0]
        set_paragraph_text(p, title, size=11, color=GREEN_DARK, bold=True)
        p.paragraph_format.space_after = Pt(5)
        p2 = cell.add_paragraph()
        set_paragraph_text(p2, text, size=10, color=INK)
    else:
        p = cell.paragraphs[0]
        set_paragraph_text(p, text, size=10, color=INK)
    doc.add_paragraph()



def add_metric_cards_doc(doc: Document, cards: list[tuple[str, str, str]], fill: str = OFFWHITE) -> None:
    """Add compact institutional cards for executive pages."""
    cols = min(3, max(1, len(cards)))
    rows = math.ceil(len(cards) / cols)
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    idx = 0
    for row in range(rows):
        for col in range(cols):
            cell = table.cell(row, col)
            set_cell_shading(cell, fill if idx < len(cards) else "FFFFFF")
            set_cell_border(cell, color=BORDER, sz="6")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if idx < len(cards):
                value, label, detail = cards[idx]
                p = cell.paragraphs[0]
                set_paragraph_text(p, value, size=15, color=GREEN_DARK, bold=True)
                p.paragraph_format.space_after = Pt(1)
                p2 = cell.add_paragraph()
                set_paragraph_text(p2, label, size=9.5, color=NAVY, bold=True)
                p2.paragraph_format.space_after = Pt(1)
                p3 = cell.add_paragraph()
                set_paragraph_text(p3, detail, size=8.7, color=MUTED)
            idx += 1
    doc.add_paragraph()


def add_signature_block_doc(doc: Document, name: str, role: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_text(p, name, size=10.5, color=NAVY, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_text(p2, role, size=9.2, color=MUTED, italic=True)
    doc.add_paragraph()


def add_page_label_doc(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_text(p, label.upper(), size=8.5, color=GOLD, bold=True)
    p.paragraph_format.space_after = Pt(2)

def add_image_doc(doc: Document, path: Path, width_cm: float = 16.5, caption: str | None = None) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_text(cp, caption, size=8.5, color=MUTED, italic=True)


def add_table_doc(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        set_cell_shading(hdr[i], GREEN_DARK)
        set_cell_border(hdr[i], "BFD4BD")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_text(p, head, size=9.2, color="FFFFFF", bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i], "D7E3D3")
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_text(p, value, size=9.2, color=INK)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_bullets_doc(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.18)
        set_paragraph_text(p, "• ", size=10.5, color=GREEN_DARK, bold=True)
        set_paragraph_text(p, item, size=10.5, color=INK)


def write_source_markdown() -> None:
    lines = [
        "# Dossier de présentation - Territoires Vivants France",
        "",
        "Ce fichier source accompagne le DOCX et le PDF générés automatiquement.",
        "",
        "## Données utilisées",
    ]
    for key, value in DATA.items():
        lines.append(f"- {key}: {value}")
    lines.append("")
    lines.append("## Sources")
    for source in SOURCES:
        lines.append(f"- {source['label']} - {source['url']} ({source['usage']})")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx(visuals: dict[str, Path]) -> None:
    doc = Document()
    configure_doc(doc)
    add_footer(doc.sections[0])

    # Cover
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.add_run().add_picture(str(visuals["logo"]), width=Cm(10.5))
    add_doc_title(
        doc,
        "DOSSIER DE PRÉSENTATION",
        "Territoires Vivants France - association implantée à Saint-Étienne, à vocation nationale progressive",
    )
    add_image_doc(doc, visuals["cover"], width_cm=16.5, caption="Saint-Étienne, territoire de lancement de la démarche TVF.")
    add_callout(
        doc,
        "Remettre en usage les logements, commerces, bâtiments, terrains et matériaux inutilisés grâce à une coopération structurée entre propriétaires, collectivités, entreprises, associations, citoyens et financeurs.",
        "Objet du dossier",
        fill=OFFWHITE,
    )
    meta = [
        ["Siège", CONTACT["adresse"]],
        ["RNA", CONTACT["rna"]],
        ["SIRET", CONTACT["siret"]],
        ["Contact", f"{CONTACT['email']} - {CONTACT['telephone']}"],
        ["Site internet", CONTACT["site"]],
    ]
    add_table_doc(doc, ["Repère", "Information"], meta, widths=[4.0, 12.0])
    doc.add_page_break()

    add_page_label_doc(doc, "Synthèse institutionnelle")
    doc.add_heading("TVF en une page", level=1)
    doc.add_paragraph(
        "Territoires Vivants France est une association nouvelle, implantée à Saint-Étienne, construite pour devenir progressivement un outil national de coopération territoriale. Sa mission est de transformer des ressources inutilisées - biens vacants, locaux fermés, terrains, friches et matériaux - en projets utiles, documentés et suivis."
    )
    add_metric_cards_doc(
        doc,
        [
            ("Mission", "Redonner un usage", "Biens, matériaux et espaces inutilisés deviennent des ressources territoriales."),
            ("Méthode", "Instruire avant d'agir", "Chaque demande passe par une fiche, des pièces, une qualification et un cadre écrit."),
            ("Impact", "Mesurer après action", "Aucun chiffre d'impact n'est annoncé sans preuve, convention et suivi réel."),
        ],
        fill=GREEN_LIGHT,
    )
    add_table_doc(
        doc,
        ["TVF relie", "TVF produit", "TVF mesure"],
        [
            ["Collectivités, propriétaires, entreprises, associations, bénévoles et financeurs.", "Dossiers instruits, conventions, mises en relation, affectation de ressources.", "Biens repérés, matériaux qualifiés, projets lancés, preuves et bilans."],
            ["Besoins publics, ressources privées et capacités citoyennes.", "Parcours de coopération clair pour passer de l'idée au projet.", "Indicateurs uniquement renseignés après action réelle."],
        ],
        widths=[5.3, 5.3, 5.3],
    )
    add_callout(
        doc,
        "Le dossier doit servir en rendez-vous : il présente l'objet, le sérieux de la méthode, les besoins de lancement et les coopérations recherchées, sans promettre de résultats fictifs.",
        "À retenir",
        fill="FFF7E3",
    )
    doc.add_page_break()

    add_page_label_doc(doc, "Note d'intention")
    doc.add_heading("Pourquoi lancer TVF depuis Saint-Étienne ?", level=1)
    doc.add_paragraph(
        "Saint-Étienne concentre plusieurs enjeux que TVF souhaite traiter de manière coordonnée : habitat ancien, logements vacants, locaux commerciaux fragilisés, besoins sociaux, matériaux réemployables et nécessité de redonner des usages concrets à des espaces parfois sous-utilisés. Le territoire offre donc un terrain de lancement pertinent pour tester une méthode simple : repérer, qualifier, conventionner, agir et mesurer."
    )
    doc.add_paragraph(
        "L'ambition de TVF n'est pas de se substituer aux collectivités, aux services publics ou aux opérateurs existants. L'association souhaite au contraire devenir un facilitateur : un point d'entrée capable d'écouter les besoins, de structurer les demandes, de rapprocher les acteurs et de transformer des opportunités dispersées en dossiers exploitables."
    )
    doc.add_paragraph(
        "Cette première étape doit permettre de construire une preuve locale sérieuse avant toute extension progressive. La priorité est donc de créer des partenariats simples, vérifiables et utiles : un lieu de stockage, des premières ressources à qualifier, des services à rencontrer, des propriétaires à accompagner et des acteurs économiques à mobiliser."
    )
    add_signature_block_doc(doc, "Edryan Rangoly", "Président fondateur - Territoires Vivants France")
    doc.add_page_break()

    add_page_label_doc(doc, "Demande partenariale")
    doc.add_heading("Ce que TVF recherche pour démarrer", level=1)
    doc.add_paragraph(
        "Pour passer de la structuration à l'action de terrain, TVF doit sécuriser quelques moyens de départ. Ces demandes sont volontairement concrètes : elles permettent de tester la méthode sans surdimensionner le projet."
    )
    add_table_doc(
        doc,
        ["Besoin prioritaire", "Utilité opérationnelle", "Forme de coopération possible"],
        [
            ["Local de stockage", "Recevoir, trier et conserver temporairement des matériaux ou petits équipements.", "Mise à disposition, bail précaire, convention d'occupation ou partenariat."],
            ["Transport / logistique", "Collecter des matériaux, déplacer du mobilier, organiser de petites tournées.", "Prêt ponctuel, mécénat en nature, tarif solidaire ou convention."],
            ["Mise en relation", "Identifier les services, propriétaires, entreprises et associations concernées.", "Rendez-vous, orientation, réunion technique ou comité de lancement."],
            ["Premiers dossiers", "Tester le parcours d'instruction sur des cas réels mais simples.", "Signalement qualifié, visite, fiche projet, accord écrit."],
            ["Soutien financier", "Couvrir les coûts de démarrage : assurance, outils, communication, déplacements.", "Don, mécénat, subvention compatible, financement fléché."],
        ],
        widths=[4.2, 6.4, 5.4],
    )
    add_callout(
        doc,
        "La demande n'est pas de financer une promesse abstraite, mais de permettre à TVF de tester un cadre de travail traçable, utile et reproductible.",
        "Message pour les partenaires",
        fill=GREEN_LIGHT,
    )
    doc.add_page_break()

    add_page_label_doc(doc, "Cadre de confiance")
    doc.add_heading("Un fonctionnement cadré dès le départ", level=1)
    doc.add_paragraph(
        "Pour inspirer confiance aux institutions et aux partenaires privés, TVF doit fonctionner avec une logique de dossier. Toute demande importante doit pouvoir être comprise, suivie, justifiée et archivée. Cette rigueur protège l'association, le partenaire et le bénéficiaire du projet."
    )
    add_table_doc(
        doc,
        ["Garantie", "Application concrète"],
        [
            ["Traçabilité", "Numéro de dossier, fiche contact, pièces reçues, décisions et historique des échanges."],
            ["Convention", "Aucune action sensible sans accord écrit, responsabilités définies et durée précisée."],
            ["Sincérité", "Aucun résultat, partenaire ou financement n'est affiché s'il n'est pas réel ou officialisé."],
            ["Sécurité", "État du bien ou du matériau, conditions d'accès, assurance et limites d'intervention vérifiées."],
            ["Impact", "Indicateurs renseignés après action, avec preuves : photos, fiches, bilans, attestations."],
        ],
        widths=[4.2, 11.8],
    )
    doc.add_page_break()

    # Sommaire
    doc.add_heading("Sommaire", level=1)
    add_table_doc(
        doc,
        ["Partie", "Contenu"],
        [
            ["1", "Synthèse exécutive"],
            ["2", "Constat territorial et données sourcées"],
            ["3", "Mission, positionnement et valeur ajoutée de TVF"],
            ["4", "Les cinq pôles d'action"],
            ["5", "Méthode opérationnelle et parcours projet"],
            ["6", "Saint-Étienne : territoire de lancement"],
            ["7", "Bénéfices par public et coopérations recherchées"],
            ["8", "Indicateurs, besoins de lancement et sources"],
        ],
        widths=[2.0, 14.0],
    )
    doc.add_page_break()

    doc.add_heading("1. Synthèse exécutive", level=1)
    doc.add_paragraph(
        "Territoires Vivants France (TVF) est une association implantée à Saint-Étienne, construite pour devenir un outil national de coopération au service de la revitalisation des territoires. Son rôle n'est pas de remplacer les politiques publiques existantes, mais de relier ce qui reste trop souvent dispersé : biens vacants, propriétaires, collectivités, entreprises, matériaux réemployables, associations, bénévoles et financeurs."
    )
    doc.add_paragraph(
        "Le projet répond à une difficulté très concrète : des logements restent vacants, des commerces se ferment, des bâtiments perdent leur usage, des friches restent sans projet, tandis que des matériaux encore utilisables sont détruits ou mal orientés. Dans le même temps, les collectivités et les acteurs locaux recherchent des solutions opérationnelles, mesurables et juridiquement cadrées."
    )
    add_image_doc(doc, visuals["key_figures"], width_cm=16.8)
    add_callout(
        doc,
        "Le dossier ne présente aucun résultat fictif. Les chiffres d'impact futurs devront être renseignés uniquement après instruction des projets, validation des conventions et suivi réel des actions.",
        "Principe de sincérité",
        fill="FFF7E3",
    )


    doc.add_heading("1.1 Une probl\u00e9matique de coordination autant que de r\u00e9novation", level=2)
    doc.add_paragraph("La vacance immobili\u00e8re, la fermeture de commerces, l'abandon de certains terrains ou la perte de mat\u00e9riaux encore utilisables ne rel\u00e8vent pas uniquement d'un manque de volont\u00e9 locale. Dans de nombreux cas, le probl\u00e8me vient d'une succession de blocages plus discrets : absence d'interlocuteur identifi\u00e9, manque de temps pour qualifier les situations, complexit\u00e9 administrative, difficult\u00e9 \u00e0 r\u00e9unir les propri\u00e9taires, absence de lieu de stockage, incertitude sur les responsabilit\u00e9s, faiblesse du financement initial ou manque d'ing\u00e9nierie pour transformer une opportunit\u00e9 en dossier concret.")
    doc.add_paragraph("TVF part de ce constat : les ressources existent souvent d\u00e9j\u00e0 sur le territoire, mais elles ne se rencontrent pas. Un local vide peut int\u00e9resser une association, des mat\u00e9riaux peuvent r\u00e9pondre \u00e0 un besoin de chantier, une entreprise peut vouloir contribuer localement, une collectivit\u00e9 peut chercher \u00e0 r\u00e9activer un secteur, mais aucun acteur ne dispose toujours du temps ou du cadre pour relier ces \u00e9l\u00e9ments. L'enjeu n'est donc pas seulement de r\u00e9nover ; il est de rendre les coop\u00e9rations possibles, compr\u00e9hensibles et s\u00e9curis\u00e9es.")

    doc.add_heading("1.2 Pourquoi les situations restent souvent bloqu\u00e9es", level=2)
    add_table_doc(
        doc,
        ["Blocage observ\u00e9", "Cons\u00e9quence pour le territoire", "R\u00e9ponse que TVF veut structurer"],
        [
            ["Information dispers\u00e9e", "Les biens, mat\u00e9riaux et besoins sont connus localement mais rarement centralis\u00e9s.", "Cr\u00e9er des fiches, des signalements qualifi\u00e9s et un suivi de dossier."],
            ["Absence de porteur clair", "Une opportunit\u00e9 reste sans suite faute d'interlocuteur responsable du montage.", "Assurer une premi\u00e8re instruction associative et organiser les \u00e9changes."],
            ["Manque de cadre \u00e9crit", "Les partenaires h\u00e9sitent par crainte des responsabilit\u00e9s, co\u00fbts ou risques.", "Pr\u00e9parer conventions, pi\u00e8ces \u00e0 fournir, limites d'intervention et suivi."],
            ["Difficult\u00e9 logistique", "Les mat\u00e9riaux disponibles ne sont pas utilis\u00e9s faute de stockage, transport ou tri.", "Rechercher local, transport, inventaire et affectation \u00e0 des projets valid\u00e9s."],
            ["Financement initial fragile", "Les premiers projets restent au stade de l'id\u00e9e faute de moyens d'amor\u00e7age.", "Mobiliser m\u00e9c\u00e9nat, dons, subventions compatibles et partenariats en nature."],
        ],
        widths=[4.2, 6.0, 5.8],
    )

    doc.add_heading("1.3 Ce que TVF veut apporter : une ing\u00e9nierie associative de proximit\u00e9", level=2)
    doc.add_paragraph("La valeur de TVF n'est pas de se pr\u00e9senter comme un op\u00e9rateur public suppl\u00e9mentaire. Sa valeur est d'assumer un r\u00f4le d'interface : \u00e9couter les besoins, traduire une situation en dossier, r\u00e9unir les bons interlocuteurs, formaliser une coop\u00e9ration et conserver une m\u00e9moire de ce qui a \u00e9t\u00e9 d\u00e9cid\u00e9. Cette ing\u00e9nierie de proximit\u00e9 est essentielle pour des sujets qui sont souvent trop petits pour mobiliser de grands op\u00e9rateurs, mais trop complexes pour \u00eatre trait\u00e9s uniquement par bonne volont\u00e9.")
    doc.add_paragraph("L'association doit donc se construire avec une exigence de m\u00e9thode : ne pas confondre intention et r\u00e9alisation, ne pas annoncer d'impact sans preuve, ne pas afficher de partenaires non officialis\u00e9s, ne pas distribuer de ressources sans affectation v\u00e9rifi\u00e9e. Cette prudence n'affaiblit pas le projet ; elle en constitue au contraire la cr\u00e9dibilit\u00e9.")

    doc.add_page_break()
    doc.add_heading("2. Constat territorial et données sourcées", level=1)
    doc.add_paragraph(
        "Saint-Étienne constitue un territoire de lancement cohérent pour TVF : ville de patrimoine industriel, de centralités commerciales, d'habitat ancien et de besoins sociaux importants. Les données INSEE montrent un territoire où la revitalisation peut produire un effet direct sur le logement, l'activité locale, l'insertion et le cadre de vie."
    )
    add_table_doc(
        doc,
        ["Indicateur", "Valeur", "Lecture TVF"],
        [
            ["Population", DATA["population_2023"], "Un bassin urbain suffisamment dense pour expérimenter un dispositif structuré."],
            ["Évolution démographique", DATA["population_variation"], "Une dynamique stable qui renforce l'intérêt de remettre en usage le patrimoine existant."],
            ["Logements vacants", DATA["logements_vacants_2022"], "Un gisement potentiel de repérage, d'accompagnement des propriétaires et de remise en usage."],
            ["Taux de pauvreté", DATA["pauvreté_2023"], "Un enjeu social fort pour relier revitalisation, insertion et accès à des usages utiles."],
            ["Taux de chômage", DATA["taux_chomage_2023"], "Un besoin de chantiers, de formation et de parcours vers l'emploi."],
            ["Établissements actifs", DATA["etablissements_2023"], "Un tissu économique mobilisable autour du réemploi, du mécénat, de la logistique et des compétences."],
        ],
        widths=[4.2, 3.1, 8.7],
    )
    doc.add_paragraph(
        "À l'échelle nationale, la Cour des comptes a pointé en 2025 l'ampleur de la vacance de logement et les limites d'une action publique insuffisamment coordonnée. Cette lecture conforte l'approche de TVF : observer, qualifier, conventionner, mobiliser et suivre."
    )
    add_image_doc(doc, visuals["ecosystem"], width_cm=16.6)

    doc.add_heading("3. Mission, positionnement et valeur ajoutée", level=1)
    doc.add_paragraph(
        "TVF se positionne comme une plateforme de coordination territoriale. L'association identifie des ressources inutilisées, aide à qualifier les situations, prépare les dossiers, recherche les coopérations possibles et oriente les moyens vers des projets utiles. Cette logique est particulièrement adaptée aux territoires qui disposent déjà de dispositifs publics mais manquent de temps, d'ingénierie, de coordination ou de mobilisation citoyenne pour passer de l'intention à l'action."
    )
    add_table_doc(
        doc,
        ["Ce que TVF fait", "Ce que TVF ne fait pas"],
        [
            ["Coordonner les acteurs autour d'un bien, d'un matériau ou d'un besoin territorial.", "Se substituer à une collectivité, à un service de l'État ou à un opérateur public."],
            ["Préparer des conventions, dossiers d'instruction et parcours de décision.", "Annoncer des résultats sans preuve ou sans suivi documenté."],
            ["Valoriser les ressources inutilisées dans des projets validés.", "Distribuer librement des matériaux sans analyse, traçabilité ou affectation."],
            ["Mesurer l'impact réel après action.", "Inventer des chiffres, partenaires, financeurs ou projets réalisés."],
        ],
        widths=[8.0, 8.0],
    )
    add_callout(
        doc,
        "Proposition de valeur : faire gagner du temps aux territoires en reliant les besoins, les biens, les matériaux, les acteurs et les financements dans un parcours lisible et traçable.",
        "Ce qui rend TVF utile",
        fill=GREEN_LIGHT,
    )


    doc.add_heading("3.1 Doctrine d'intervention : intervenir seulement lorsque le cadre est clair", level=2)
    doc.add_paragraph("TVF doit se d\u00e9velopper avec une r\u00e8gle simple : aucune action ne doit \u00eatre engag\u00e9e si le cadre n'est pas suffisamment clair. Cette exigence concerne autant les biens immobiliers que les mat\u00e9riaux, les partenariats ou les chantiers participatifs. Un projet utile peut rapidement devenir fragile si les responsabilit\u00e9s, les pi\u00e8ces disponibles, les conditions d'acc\u00e8s, l'assurance, la dur\u00e9e d'usage ou la propri\u00e9t\u00e9 des ressources ne sont pas \u00e9tablies d\u00e8s le d\u00e9part.")
    doc.add_paragraph("Cette doctrine prot\u00e8ge l'ensemble des parties. Elle prot\u00e8ge le propri\u00e9taire, qui conserve la ma\u00eetrise de son bien. Elle prot\u00e8ge la collectivit\u00e9, qui peut soutenir une d\u00e9marche sans porter seule l'instruction. Elle prot\u00e8ge l'entreprise contributrice, qui sait comment sa ressource est valoris\u00e9e. Elle prot\u00e8ge enfin l'association, qui \u00e9vite de se disperser dans des promesses non ma\u00eetris\u00e9es.")
    add_table_doc(
        doc,
        ["Principe", "Traduction op\u00e9rationnelle", "Effet attendu"],
        [
            ["Prudence", "Ne pas annoncer de projet, partenaire ou r\u00e9sultat tant qu'il n'est pas document\u00e9.", "Renforcer la confiance institutionnelle."],
            ["Tra\u00e7abilit\u00e9", "Ouvrir un dossier, conserver les pi\u00e8ces, dater les \u00e9changes et formaliser les d\u00e9cisions.", "Rendre chaque coop\u00e9ration v\u00e9rifiable."],
            ["Utilit\u00e9 territoriale", "Affecter les ressources \u00e0 un besoin local identifi\u00e9, pas \u00e0 une distribution indiff\u00e9renci\u00e9e.", "\u00c9viter la dispersion et privil\u00e9gier les projets utiles."],
            ["Convention", "\u00c9crire les engagements, responsabilit\u00e9s, dur\u00e9es et limites d'intervention.", "S\u00e9curiser le passage \u00e0 l'action."],
        ],
        widths=[3.1, 7.4, 5.5],
    )

    doc.add_heading("3.2 Une association de lancement, mais une exigence de niveau institutionnel", level=2)
    doc.add_paragraph("TVF est en phase de lancement. Cette r\u00e9alit\u00e9 doit \u00eatre assum\u00e9e clairement. Mais \u00eatre une association nouvelle ne signifie pas travailler de mani\u00e8re approximative. Au contraire, la cr\u00e9dibilit\u00e9 du projet d\u00e9pend de sa capacit\u00e9 \u00e0 adopter d\u00e8s le d\u00e9part des standards proches de ceux attendus par les collectivit\u00e9s, les financeurs, les entreprises et les acteurs publics : dossiers num\u00e9rot\u00e9s, documents dat\u00e9s, demandes structur\u00e9es, sources identifi\u00e9es, conventions \u00e9crites, transparence sur les limites et mesure progressive de l'impact.")
    doc.add_paragraph("Le dossier ne doit donc pas chercher \u00e0 donner l'illusion d'une structure d\u00e9j\u00e0 d\u00e9ploy\u00e9e partout. Il doit montrer autre chose : une capacit\u00e9 \u00e0 penser correctement le probl\u00e8me, \u00e0 organiser les premi\u00e8res coop\u00e9rations, \u00e0 travailler avec m\u00e9thode et \u00e0 construire une preuve locale s\u00e9rieuse depuis Saint-\u00c9tienne.")

    doc.add_heading("4. Les cinq pôles d'action", level=1)
    add_table_doc(
        doc,
        ["Pôle", "Rôle", "Premiers publics concernés"],
        [
            ["Habitat Vivant", "Repérer et accompagner la remise en usage de logements vacants ou dégradés.", "Propriétaires, communes, bailleurs, associations."],
            ["Matériauthèque Solidaire", "Identifier, qualifier, stocker et affecter des matériaux réemployables à des projets validés.", "Entreprises, collectivités, artisans, associations."],
            ["Commerce Vivant", "Réactiver les locaux commerciaux vacants et préparer de nouveaux usages de proximité.", "Commerçants, propriétaires, communes, porteurs de projet."],
            ["Friches & Terrains Vivants", "Transformer des espaces délaissés en lieux utiles, végétalisés, associatifs ou économiques.", "Collectivités, propriétaires, habitants, acteurs ESS."],
            ["Solidarité & Insertion", "Relier les projets à des parcours d'engagement, de bénévolat, de formation et d'insertion.", "Structures d'insertion, bénévoles, jeunes, demandeurs d'emploi."],
        ],
        widths=[3.9, 7.3, 4.8],
    )
    for title, img_key, text in [
        ("Habitat Vivant", "habitat", "Le pôle Habitat Vivant agit sur les situations où le logement existe déjà mais ne répond plus à un usage. TVF peut repérer, documenter et orienter les propriétaires vers un parcours de remise en état, en lien avec les dispositifs publics et les acteurs techniques compétents."),
        ("Matériauthèque Solidaire", "materiaux", "La matériauthèque n'est pas une distribution gratuite. Elle vise à transformer des matériaux dormants, surplus ou équipements inutilisés en ressources affectées à des projets utiles, après diagnostic, traçabilité et validation."),
        ("Commerce Vivant", "commerce", "Le pôle Commerce Vivant travaille sur les cellules vacantes, les vitrines fermées et les usages économiques de proximité. Il peut faciliter le lien entre propriétaire, collectivité, porteur de projet et acteurs de l'accompagnement."),
        ("Friches & Terrains Vivants", "friche", "Ce pôle traite les terrains et espaces délaissés comme des réserves d'usages possibles : jardins, espaces associatifs, lieux pédagogiques, micro-activités, biodiversité ou occupation temporaire encadrée."),
        ("Solidarité & Insertion", "solidarite", "La revitalisation doit créer de l'utilité sociale. TVF souhaite relier les projets à des chantiers participatifs, des parcours bénévoles, des formations et des coopérations avec les structures d'insertion existantes."),
    ]:
        doc.add_heading(title, level=2)
        add_image_doc(doc, visuals[img_key], width_cm=13.5)
        doc.add_paragraph(text)

    doc.add_heading("5. Méthode opérationnelle", level=1)
    doc.add_paragraph(
        "La méthode TVF repose sur une chaîne simple : repérer, qualifier, structurer, conventionner, agir et mesurer. Cette méthode permet d'éviter les annonces générales et de concentrer le travail sur des dossiers concrets, instruits et exploitables."
    )
    add_image_doc(doc, visuals["process"], width_cm=16.6)
    add_table_doc(
        doc,
        ["Étape", "Objectif", "Livrable"],
        [
            ["Repérage", "Identifier un bien, un matériau, un commerce ou un terrain sans usage.", "Fiche de signalement ou fiche contact."],
            ["Qualification", "Vérifier l'intérêt territorial, le contexte, les contraintes et les acteurs.", "Note de diagnostic et pièces à demander."],
            ["Instruction", "Construire le dossier : besoin, acteurs, convention, budget indicatif, risques.", "Dossier d'instruction TVF."],
            ["Coopération", "Associer propriétaire, collectivité, entreprise ou financeur selon le besoin.", "Projet de convention ou courrier de partenariat."],
            ["Mise en oeuvre", "Lancer les actions validées : collecte, stockage, chantier, mise en relation.", "Planning et tableau de suivi."],
            ["Impact", "Suivre ce qui a été réellement produit.", "Indicateurs, preuves, bilan."],
        ],
        widths=[3.0, 7.0, 6.0],
    )

    doc.add_heading("6. Saint-Étienne : territoire de lancement", level=1)
    add_image_doc(doc, visuals["saint"], width_cm=16.2, caption="Saint-Étienne offre un terrain pertinent pour relier habitat, commerce, matériaux, insertion et requalification urbaine.")
    doc.add_paragraph(
        "TVF peut commencer par Saint-Étienne car le territoire permet de tester l'ensemble de la méthode sur un périmètre lisible : habitat ancien, locaux vacants, espaces délaissés, tissu d'entreprises, besoin de chantiers utiles et réseau associatif. L'objectif n'est pas d'annoncer des réalisations avant leur existence, mais de proposer un cadre de travail crédible pour identifier les premiers dossiers."
    )
    add_table_doc(
        doc,
        ["Type de projet faisable", "Besoin local auquel il répond", "Premiers acteurs à mobiliser"],
        [
            ["Local de stockage pour matériaux", "Centraliser les matériaux réemployables avant affectation.", "Ville, Métropole, bailleurs, entreprises propriétaires de locaux."],
            ["Fiche logement vacant", "Qualifier une situation et orienter le propriétaire.", "Propriétaire, service habitat, opérateurs techniques."],
            ["Commerce vacant expérimental", "Tester une activité de proximité, associative ou artisanale.", "Propriétaire, commerce, ville, porteur de projet."],
            ["Jardin ou espace partagé", "Transformer un terrain en usage collectif léger.", "Collectivité, habitants, associations, écoles."],
            ["Chantier participatif", "Créer une action visible mêlant réemploi, bénévolat et insertion.", "Artisans, structures d'insertion, bénévoles, entreprises."],
        ],
        widths=[4.2, 6.3, 5.5],
    )

    doc.add_heading("7. Bénéfices par public", level=1)
    add_table_doc(
        doc,
        ["Public", "Ce qu'il apporte", "Ce qu'il peut obtenir"],
        [
            ["Collectivité", "Connaissance territoriale, mise à disposition possible, orientation des besoins.", "Un outil d'animation, des dossiers structurés, des indicateurs et une valorisation des politiques publiques."],
            ["Propriétaire", "Un bien vacant, dégradé ou difficile à remettre en usage.", "Un accompagnement, une convention adaptée, une valorisation patrimoniale et un projet utile."],
            ["Entreprise", "Matériaux, compétences, logistique, mécénat ou locaux.", "Une action RSE concrète, traçable, locale et valorisable."],
            ["Association", "Besoin d'espace, bénévolat, expertise terrain, animation locale.", "Accès à des projets, lieux, matériaux ou coopérations encadrées."],
            ["Mécène / financeur", "Soutien financier, ingénierie, réseau.", "Une lecture claire de l'usage des moyens et des impacts suivis."],
            ["Citoyen / bénévole", "Signalement, temps, participation, relais local.", "Une façon concrète d'améliorer son territoire."],
        ],
        widths=[3.2, 6.2, 6.6],
    )
    add_callout(
        doc,
        "Le besoin immédiat de lancement porte sur trois leviers : un local de stockage, des solutions de transport/logistique et des premiers partenaires d'instruction pour traiter les demandes sérieusement.",
        "Besoins de démarrage",
        fill="FFF7E3",
    )



    doc.add_page_break()

    doc.add_heading("7.1 Argumentaire d\u00e9velopp\u00e9 par interlocuteur", level=2)
    doc.add_paragraph("Pour une collectivit\u00e9, l'int\u00e9r\u00eat de TVF r\u00e9side dans la capacit\u00e9 \u00e0 transformer des situations diffuses en dossiers lisibles. Un signalement, un local vide, une demande associative ou un lot de mat\u00e9riaux peuvent \u00eatre instruits dans un cadre commun. La collectivit\u00e9 n'est pas d\u00e9poss\u00e9d\u00e9e de ses comp\u00e9tences ; elle dispose d'un relais capable de pr\u00e9parer le terrain, d'organiser les informations et de restituer des indicateurs exploitables.")
    doc.add_paragraph("Pour un propri\u00e9taire, TVF peut repr\u00e9senter un interlocuteur de confiance lorsque le bien est devenu difficile \u00e0 g\u00e9rer : local inoccup\u00e9, logement ancien, commerce ferm\u00e9, b\u00e2timent \u00e0 remettre en \u00e9tat ou terrain sans usage. L'enjeu n'est pas de lui imposer un projet, mais d'\u00e9tudier avec lui les conditions d'une valorisation encadr\u00e9e : usage possible, dur\u00e9e, responsabilit\u00e9s, travaux \u00e9ventuels, limites d'intervention et sortie de convention.")
    doc.add_paragraph("Pour une entreprise, la coop\u00e9ration avec TVF peut donner une traduction concr\u00e8te \u00e0 une d\u00e9marche RSE locale. Les mat\u00e9riaux, \u00e9quipements, comp\u00e9tences ou moyens logistiques mis \u00e0 disposition ne sont pas simplement donn\u00e9s ; ils sont orient\u00e9s vers des projets suivis, document\u00e9s et utiles au territoire. Cette tra\u00e7abilit\u00e9 permet de d\u00e9passer la communication symbolique pour aller vers une contribution mesurable.")
    doc.add_paragraph("Pour un financeur ou un m\u00e9c\u00e8ne, TVF doit apporter une garantie de m\u00e9thode : l'argent ou l'appui apport\u00e9 doit \u00eatre rattach\u00e9 \u00e0 un besoin identifi\u00e9, \u00e0 un dossier, \u00e0 un cadre d'utilisation et \u00e0 un bilan. La promesse n'est pas de produire imm\u00e9diatement de grands volumes d'impact, mais de mettre en place un outil capable de documenter progressivement les r\u00e9sultats r\u00e9els.")
    doc.add_paragraph("Pour une association ou un citoyen, TVF doit rendre l'engagement plus accessible. Beaucoup de personnes souhaitent agir localement mais ne savent pas par o\u00f9 commencer. En structurant les besoins, les lieux, les mat\u00e9riaux et les chantiers, TVF peut transformer une disponibilit\u00e9 individuelle en participation utile, sans improvisation ni dispersion.")

    doc.add_heading("7.2 Ce que la coop\u00e9ration doit produire concr\u00e8tement", level=2)
    add_table_doc(
        doc,
        ["Interlocuteur", "Ce qui doit \u00eatre clarifi\u00e9", "D\u00e9cision attendue \u00e0 court terme"],
        [
            ["Collectivit\u00e9", "Services concern\u00e9s, besoins prioritaires, locaux ou terrains mobilisables, cadre de rendez-vous.", "Identifier un r\u00e9f\u00e9rent et un premier sujet d'exp\u00e9rimentation."],
            ["Propri\u00e9taire", "Situation du bien, droits, contraintes, souhaits, disponibilit\u00e9 et niveau d'intervention possible.", "Autoriser une instruction ou une visite exploratoire."],
            ["Entreprise", "Nature de la contribution : mat\u00e9riaux, mat\u00e9riel, transport, comp\u00e9tence, m\u00e9c\u00e9nat ou local.", "Qualifier une premi\u00e8re contribution utile et tra\u00e7able."],
            ["Association", "Besoin r\u00e9el, capacit\u00e9 d'animation, publics concern\u00e9s, s\u00e9curit\u00e9 et disponibilit\u00e9.", "D\u00e9finir un usage compatible avec un bien, un chantier ou une ressource."],
            ["Financeur", "Montant ou appui envisageable, affectation, indicateurs, reporting et conditions.", "Valider un cadre de soutien au lancement ou \u00e0 un dossier pilote."],
        ],
        widths=[3.4, 7.0, 5.6],
    )

    doc.add_heading("8. Parcours par public", level=1)
    doc.add_paragraph(
        "Le dossier de présentation doit permettre à chaque interlocuteur de se reconnaître rapidement. TVF doit donc parler en parcours : ce que la personne apporte, ce qui est vérifié, ce qui est instruit, puis ce qui peut être conventionné."
    )
    add_table_doc(
        doc,
        ["Public", "Point d'entrée", "Instruction TVF", "Sortie possible"],
        [
            ["Collectivité", "Besoin de revitalisation, local disponible, demande d'animation ou diagnostic.", "Qualification territoriale, services concernés, contraintes, objectifs publics.", "Convention de coopération, fiche projet, tableau de suivi."],
            ["Propriétaire", "Bien vacant, commerce fermé, immeuble dégradé, terrain sans usage.", "Situation juridique, état du bien, usage possible, niveau d'intervention.", "Convention d'utilisation ou orientation vers partenaires compétents."],
            ["Entreprise", "Matériaux, local, véhicule, compétences, mécénat ou matériel.", "Nature de la ressource, état, disponibilité, traçabilité, usage compatible.", "Convention de valorisation territoriale et attestation de contribution."],
            ["Association", "Besoin de lieu, de matériaux, de chantier ou de bénévoles.", "Objet de l'association, capacité de portage, utilité locale, sécurité.", "Affectation à un projet, mise en relation, chantier ou usage encadré."],
            ["Bénévole", "Temps disponible, compétence, envie d'aider localement.", "Disponibilité, mission adaptée, cadre de sécurité, référent.", "Mission ponctuelle, chantier participatif, appui administratif."],
            ["Financeur", "Soutien financier, mécénat ou investissement solidaire.", "Affectation, traçabilité, indicateurs, gouvernance du dossier.", "Convention, reporting, bilan d'impact documenté."],
        ],
        widths=[3.0, 4.6, 4.7, 3.7],
    )

    doc.add_heading("9. Programme Bien Solidaire à Usage Partagé", level=1)
    doc.add_paragraph(
        "Le programme Bien Solidaire à Usage Partagé constitue un dispositif phare à structurer. Il s'adresse aux propriétaires qui possèdent un bien vacant, dégradé ou sans usage et qui souhaitent contribuer à une solution locale sans perdre la propriété du bien. Le principe repose sur une convention définissant les engagements de chacun, la durée d'utilisation, les travaux ou aménagements envisagés, les responsabilités et les modalités de sortie."
    )
    add_table_doc(
        doc,
        ["Élément", "Contenu à cadrer dans le dossier"],
        [
            ["Bien concerné", "Logement, commerce, immeuble, local, bâtiment, terrain ou espace sans usage."],
            ["Diagnostic", "État général, photos, contraintes techniques, sécurité, accès, usage possible."],
            ["Projet", "Logement solidaire, atelier, commerce de proximité, espace associatif, formation, tiers-lieu ou usage temporaire."],
            ["Convention", "Durée, responsabilités, assurances, travaux, entretien, suivi, fin d'utilisation."],
            ["Bénéfices", "Patrimoine valorisé, usage utile, visibilité territoriale, contribution sociale et écologique."],
        ],
        widths=[4.0, 12.0],
    )
    add_callout(
        doc,
        "TVF ne doit jamais promettre une rénovation sans diagnostic, financement, convention et validation des parties. Le dispositif doit rester cadré, traçable et proportionné au bien.",
        "Point de vigilance",
        fill="FFF7E3",
    )

    doc.add_heading("10. Banque de matériaux TVF", level=1)
    doc.add_paragraph(
        "La Banque de matériaux TVF est un outil de valorisation territoriale. Elle ne fonctionne pas comme une plateforme de dons libres. Les matériaux sont proposés, qualifiés, triés si nécessaire, puis affectés à des projets validés. L'enjeu est double : éviter la perte de ressources utiles et réduire les coûts de projets locaux lorsque le réemploi est techniquement possible."
    )
    add_image_doc(doc, visuals["materiaux"], width_cm=13.6)
    add_table_doc(
        doc,
        ["Flux possible", "Exemples", "Contrôle nécessaire"],
        [
            ["Menuiseries", "Portes, fenêtres, volets, escaliers.", "Dimensions, état, conformité, stockage."],
            ["Bois et second œuvre", "Planches, panneaux, parquets, cloisons, carrelage.", "Quantité, humidité, casse, sécurité sanitaire."],
            ["équipements", "Sanitaires, luminaires, mobilier professionnel.", "Fonctionnement, démontage, nettoyage, traçabilité."],
            ["Matériel", "Outils, consommables, équipements de chantier.", "Sécurité, propriété, disponibilité, usage prévu."],
        ],
        widths=[3.8, 6.2, 6.0],
    )

    doc.add_heading("11. Besoins de lancement", level=1)
    doc.add_paragraph(
        "Pour transformer l'idée en fonctionnement réel, TVF doit disposer d'un minimum d'outils opérationnels. Ces besoins ne sont pas des résultats : ils constituent les conditions de départ à rechercher auprès des partenaires publics et privés."
    )
    add_table_doc(
        doc,
        ["Besoin", "Pourquoi c'est nécessaire", "Interlocuteurs possibles"],
        [
            ["Local de stockage", "Recevoir, trier et sécuriser les matériaux avant affectation.", "Ville, Métropole, bailleurs, entreprises, propriétaires de locaux."],
            ["Solution de transport", "Collecter ou déplacer des matériaux, petit mobilier et équipements.", "Entreprises de transport, loueurs, mécènes, collectivités."],
            ["Matériel de manutention", "Organiser les lots, éviter les pertes, assurer la sécurité.", "Entreprises BTP, plateformes logistiques, ressourceries."],
            ["Appui technique", "Évaluer les biens, matériaux, contraintes et usages.", "Architectes, artisans, bureaux d'études, services techniques."],
            ["Financement de démarrage", "Assurer assurance, communication, outils, déplacements, premières actions.", "Mécènes, fondations, entreprises, dispositifs publics compatibles."],
        ],
        widths=[3.3, 6.5, 6.2],
    )

    doc.add_heading("12. Documents et preuves à produire", level=1)
    doc.add_paragraph(
        "La crédibilité de TVF reposera sur la qualité des dossiers. Chaque demande doit pouvoir être suivie, justifiée et archivée. Les documents ci-dessous structurent une base de travail professionnelle pour les futurs échanges avec les collectivités, entreprises et propriétaires."
    )
    add_table_doc(
        doc,
        ["Document", "Utilité", "Moment d'utilisation"],
        [
            ["Fiche contact", "Identifier l'interlocuteur, ses coordonnées, son rôle et son besoin.", "Dès le premier échange."],
            ["Fiche signalement", "Décrire un bien, un commerce, un terrain ou un matériau.", "Avant toute instruction."],
            ["Liste des pièces", "Demander les documents utiles selon le type de dossier.", "Avant rendez-vous ou visite."],
            ["Dossier d'instruction", "Analyser la demande, les risques, les acteurs et la suite possible.", "Après qualification."],
            ["Projet de convention", "Cadrer les engagements réciproques et les responsabilités.", "Avant toute action opérationnelle."],
            ["Bilan d'impact", "Rendre compte de ce qui a été réellement fait.", "Après action."],
        ],
        widths=[3.8, 7.0, 5.2],
    )

    doc.add_heading("13. Coopérations à rechercher à Saint-Étienne", level=1)
    doc.add_paragraph(
        "Le lancement doit commencer par des rencontres ciblées, avec un message simple : TVF peut aider à transformer des ressources inutilisées en projets utiles, à condition de construire un cadre de coopération sérieux. Les premières demandes doivent rester concrètes : local, logistique, matériaux, expertise, signalements et orientation vers les bons services."
    )
    add_table_doc(
        doc,
        ["Famille d'acteurs", "Objet du contact", "Ce que TVF peut proposer"],
        [
            ["Ville / Métropole", "Identifier les besoins, les locaux possibles, les services concernés.", "Un dispositif pilote, des fiches projets et des indicateurs."],
            ["Bailleurs / propriétaires", "Repérer des locaux, logements ou bâtiments sans usage.", "Une instruction cadrée et une convention adaptée."],
            ["Entreprises BTP / artisans", "Matériaux, compétences, diagnostics, chantiers.", "Valorisation RSE locale et traçabilité des contributions."],
            ["Associations / insertion", "Missions, chantiers participatifs, publics à accompagner.", "Projets concrets, encadrement et valorisation des parcours."],
            ["Écoles / formations", "Ateliers, sensibilisation, métiers, projets pédagogiques.", "Support terrain et cas d'étude."],
        ],
        widths=[4.0, 6.0, 6.0],
    )

    doc.add_heading("14. Messages clés pour les rendez-vous", level=1)
    add_bullets_doc(
        doc,
        [
            "TVF est une association nouvelle, implantée à Saint-Étienne, en voie de structuration nationale progressive.",
            "Le projet ne revendique aucun résultat fictif : il propose une méthode, des outils, des documents et une capacité de coordination.",
            "La priorité est de lancer des dossiers simples, vérifiables et utiles : matériaux, biens vacants, commerces, espaces et chantiers solidaires.",
            "Chaque coopération doit être formalisée par écrit : fiche, dossier, convention, suivi et bilan.",
            "Le site internet permet de comprendre l'objet de l'association, mais les partenariats doivent être construits par rendez-vous et instruction.",
        ]
    )


    doc.add_heading("14.1 Pourquoi nous rencontrer maintenant ?", level=2)
    doc.add_paragraph("Le premier rendez-vous avec TVF ne doit pas \u00eatre compris comme une demande g\u00e9n\u00e9rale de soutien. Il doit \u00eatre compris comme une proposition de travail : identifier les sujets r\u00e9ellement exploitables, v\u00e9rifier les besoins du territoire, distinguer ce qui rel\u00e8ve d'une simple id\u00e9e de ce qui peut devenir un dossier, et d\u00e9terminer les conditions minimales d'un lancement s\u00e9rieux.")
    doc.add_paragraph("La rencontre doit permettre de r\u00e9pondre \u00e0 quatre questions : quels biens, mat\u00e9riaux ou besoins peuvent \u00eatre \u00e9tudi\u00e9s en priorit\u00e9 ; quels services ou acteurs doivent \u00eatre associ\u00e9s ; quels moyens de d\u00e9part peuvent \u00eatre mobilis\u00e9s ; et quel cadre \u00e9crit permettra de s\u00e9curiser une premi\u00e8re exp\u00e9rimentation. Cette approche rend la discussion imm\u00e9diatement op\u00e9rationnelle et \u00e9vite de rester dans un discours g\u00e9n\u00e9ral sur la revitalisation.")
    add_table_doc(
        doc,
        ["Question de rendez-vous", "R\u00e9ponse recherch\u00e9e", "Suite possible"],
        [
            ["Existe-t-il un local de stockage mobilisable ?", "Identifier une solution temporaire, s\u00e9curis\u00e9e et conventionnable.", "Visite, conditions d'usage, convention ou orientation."],
            ["Quels services doivent \u00eatre associ\u00e9s ?", "Habitat, commerce, \u00e9conomie circulaire, insertion, patrimoine, politique de la ville.", "R\u00e9union technique ou d\u00e9signation d'un r\u00e9f\u00e9rent."],
            ["Quels premiers dossiers tester ?", "Un bien, un lot de mat\u00e9riaux, un commerce ou un besoin associatif r\u00e9aliste.", "Fiche d'instruction et liste de pi\u00e8ces."],
            ["Quels moyens d'amor\u00e7age sont possibles ?", "Transport, mat\u00e9riel, assurance, appui financier ou m\u00e9c\u00e9nat en nature.", "Plan de lancement limit\u00e9 et tra\u00e7able."],
        ],
        widths=[5.1, 6.2, 4.7],
    )

    doc.add_heading("14.2 Formulation de demande institutionnelle", level=2)
    add_callout(
        doc,
        "TVF sollicite un rendez-vous afin d'\u00e9tudier les conditions de lancement op\u00e9rationnel de l'association \u00e0 Saint-\u00c9tienne : identification d'un lieu de stockage, mise en relation avec les services comp\u00e9tents, rep\u00e9rage de premiers dossiers test et examen des coop\u00e9rations possibles avec les acteurs publics, \u00e9conomiques et associatifs du territoire.",
        "Demande \u00e0 porter en rendez-vous",
        fill="FFF7E3",
    )

    doc.add_heading("15. Indicateurs de suivi", level=1)
    doc.add_paragraph(
        "TVF doit être évalué sur des faits vérifiables. Les indicateurs ci-dessous ne sont pas des résultats actuels ; ils constituent le tableau de bord à renseigner uniquement lorsque les dossiers seront instruits et les actions réalisées."
    )
    add_table_doc(
        doc,
        ["Famille", "Indicateurs à suivre", "Preuves attendues"],
        [
            ["Habitat", "Logements repérés, dossiers propriétaires ouverts, logements remis en usage.", "Fiches dossiers, conventions, photos, attestations."],
            ["Commerce", "Locaux identifiés, porteurs orientés, occupations ou réactivations validées.", "Fiches locaux, accords propriétaires, bilans d'usage."],
            ["Matériaux", "Lots proposés, lots qualifiés, matériaux affectés à un projet.", "Inventaire, photos, bordereaux, fiches de sortie."],
            ["Insertion", "Bénévoles mobilisés, heures de chantier, parcours accompagnés.", "Feuilles de présence, conventions, attestations."],
            ["Territoire", "Communes ou services rencontrés, dossiers structurés, coopérations signées.", "Comptes rendus, courriers, conventions."],
        ],
        widths=[3.0, 7.0, 6.0],
    )

    doc.add_heading("16. Sources et limites", level=1)
    doc.add_paragraph(
        "Les données chiffrées du dossier reposent sur les sources publiques disponibles à la date de préparation. Elles doivent être actualisées avant toute transmission officielle à un financeur ou une collectivité si de nouvelles données sont publiées."
    )
    add_table_doc(
        doc,
        ["Source", "Utilisation"],
        [[s["label"], f"{s['usage']}\n{s['url']}"] for s in SOURCES],
        widths=[7.0, 9.0],
    )
    doc.add_heading("Contact", level=2)
    add_table_doc(
        doc,
        ["Élément", "Coordonnée"],
        [
            ["Adresse", CONTACT["adresse"]],
            ["Email", CONTACT["email"]],
            ["Téléphone", CONTACT["telephone"]],
            ["Site", CONTACT["site"]],
            ["RNA / SIRET", f"{CONTACT['rna']} / {CONTACT['siret']}"],
        ],
        widths=[4.0, 12.0],
    )
    doc.save(DOCX_PATH)


def para_style(styles, name, size=10, color=INK, leading=13, bold=False, alignment=TA_LEFT):
    return ParagraphStyle(
        name,
        parent=styles["Normal"],
        fontName="Arial" if not bold else "Arial-Bold",
        fontSize=size,
        leading=leading,
        textColor=colors.HexColor("#" + color),
        alignment=alignment,
        spaceAfter=6,
    )


def pdf_header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Arial", 7.8)
    canvas.setFillColor(colors.HexColor("#" + MUTED))
    canvas.drawCentredString(A4[0] / 2, 0.65 * cm, f"{CONTACT['nom']} - {CONTACT['site']} - page {doc.page}")
    canvas.restoreState()


def pdf_table(rows, col_widths, header=True):
    t = Table(rows, colWidths=col_widths, repeatRows=1 if header else 0)
    style = [
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#" + BORDER)),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#" + BORDER)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]
    if header:
        style += [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + GREEN_DARK)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Arial-Bold"),
        ]
    t.setStyle(TableStyle(style))
    return t


def build_pdf(visuals: dict[str, Path]) -> None:
    pdfmetrics.registerFont(TTFont("Arial", str(font_path())))
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    pdfmetrics.registerFont(TTFont("Arial-Bold", str(bold_path if bold_path.exists() else font_path())))
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.3 * cm,
    )
    styles = getSampleStyleSheet()
    H1 = para_style(styles, "TVF-H1", 20, NAVY, 24, True)
    H2 = para_style(styles, "TVF-H2", 14, GREEN_DARK, 18, True)
    BODY = para_style(styles, "TVF-Body", 9.6, INK, 13, False, TA_JUSTIFY)
    BODYL = para_style(styles, "TVF-BodyL", 9.6, INK, 13, False, TA_LEFT)
    SMALL = para_style(styles, "TVF-Small", 8.2, MUTED, 10, False, TA_LEFT)
    CENTER = para_style(styles, "TVF-Center", 10.5, NAVY, 13, True, TA_CENTER)

    story = []
    story.append(Spacer(1, 0.6 * cm))
    story.append(PdfImage(str(visuals["logo"]), width=9.0 * cm, height=3.2 * cm, kind="proportional"))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("DOSSIER DE PRÉSENTATION", H1))
    story.append(Paragraph("Territoires Vivants France - association implantée à Saint-Étienne, à vocation nationale progressive", CENTER))
    story.append(Spacer(1, 0.3 * cm))
    story.append(PdfImage(str(visuals["cover"]), width=16.4 * cm, height=8.8 * cm, kind="proportional"))
    story.append(Spacer(1, 0.25 * cm))
    story.append(
        pdf_table(
            [
                [Paragraph("Objet", CENTER), Paragraph("Remettre en usage les logements, commerces, bâtiments, terrains et matériaux inutilisés grâce à une coopération structurée entre acteurs publics, privés, associatifs et citoyens.", BODYL)],
                [Paragraph("Contact", CENTER), Paragraph(f"{CONTACT['email']} - {CONTACT['telephone']}<br/>{CONTACT['site']}<br/>RNA {CONTACT['rna']} - SIRET {CONTACT['siret']}", BODYL)],
            ],
            [3.0 * cm, 13.4 * cm],
            header=False,
        )
    )
    story.append(PageBreak())

    story.append(Paragraph("Sommaire", H1))
    story.append(
        pdf_table(
            [[Paragraph("Partie", CENTER), Paragraph("Contenu", CENTER)]]
            + [[str(i), c] for i, c in enumerate([
                "Synthèse exécutive",
                "Constat territorial et données sourcées",
                "Mission et valeur ajoutée",
                "Les cinq pôles d'action",
                "Méthode opérationnelle",
                "Saint-Étienne : territoire de lancement",
                "Bénéfices par public",
                "Indicateurs, besoins et sources",
            ], start=1)],
            [2.2 * cm, 14.2 * cm],
        )
    )
    story.append(PageBreak())

    def section(title: str, body: str):
        story.append(Paragraph(title, H1))
        story.append(Paragraph(body, BODY))
        story.append(Spacer(1, 0.2 * cm))

    section(
        "1. Synthèse exécutive",
        "Territoires Vivants France (TVF) est une association implantée à Saint-Étienne, construite pour devenir un outil national de coopération au service de la revitalisation des territoires. Son rôle consiste à relier les biens vacants, les matériaux réemployables, les propriétaires, les collectivités, les entreprises, les associations, les bénévoles et les financeurs dans un parcours lisible, documenté et mesurable.",
    )
    story.append(PdfImage(str(visuals["key_figures"]), width=16.4 * cm, height=7.9 * cm, kind="proportional"))
    story.append(Paragraph("Aucun résultat fictif n'est présenté : les impacts seront renseignés après instruction réelle des dossiers.", SMALL))
    story.append(PageBreak())

    section(
        "2. Constat territorial et données sourcées",
        "Saint-Étienne constitue un territoire pertinent pour lancer TVF : ville dense, marquée par l'habitat ancien, les centralités commerciales, la requalification urbaine et des besoins sociaux élevés. Les indicateurs ci-dessous sont issus des données publiques disponibles.",
    )
    rows = [[Paragraph("Indicateur", CENTER), Paragraph("Valeur", CENTER), Paragraph("Lecture TVF", CENTER)]]
    for indicator, value, reading in [
        ("Population", DATA["population_2023"], "Un bassin urbain dense pour expérimenter un outil opérationnel."),
        ("Logements vacants", DATA["logements_vacants_2022"], "Un gisement de repérage et d'accompagnement des propriétaires."),
        ("Taux de pauvreté", DATA["pauvreté_2023"], "Un enjeu social qui justifie une approche insertion et utilité collective."),
        ("Taux de chômage", DATA["taux_chomage_2023"], "Un besoin de chantiers, de formation et de parcours vers l'emploi."),
        ("Établissements actifs", DATA["etablissements_2023"], "Un tissu local mobilisable pour le réemploi et les compétences."),
    ]:
        rows.append([Paragraph(indicator, BODYL), Paragraph(value, BODYL), Paragraph(reading, BODYL)])
    story.append(pdf_table(rows, [4.0 * cm, 3.0 * cm, 9.4 * cm]))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("Source : INSEE, dossier complet Commune de Saint-Étienne (42218), paru le 09/07/2026.", SMALL))
    story.append(PageBreak())

    section(
        "3. Mission et valeur ajoutée",
        "TVF se distingue par sa capacité à créer une chaîne de travail : observer, qualifier, instruire, conventionner, mobiliser, agir et mesurer. Cette approche complète les dispositifs existants sans les remplacer. Elle donne aux collectivités et partenaires un cadre de coopération plus simple à suivre.",
    )
    story.append(PdfImage(str(visuals["ecosystem"]), width=16.4 * cm, height=10.4 * cm, kind="proportional"))
    story.append(PageBreak())

    section("4. Les cinq pôles d'action", "Les pôles structurent le travail sans disperser l'association : chacun traite un besoin précis et peut mobiliser les autres lorsqu'un projet l'exige.")
    poles = [
        ("Habitat Vivant", "Repérer, qualifier et remettre en usage des logements vacants ou dégradés.", "Propriétaires, communes, bailleurs, associations."),
        ("Matériauthèque Solidaire", "Transformer les matériaux inutilisés en ressources affectées à des projets validés.", "Entreprises, collectivités, artisans, associations."),
        ("Commerce Vivant", "Réactiver les locaux commerciaux et soutenir les usages de proximité.", "Propriétaires, commerçants, collectivités, porteurs."),
        ("Friches & Terrains Vivants", "Réorienter des espaces délaissés vers des usages utiles, verts ou associatifs.", "Collectivités, habitants, propriétaires, ESS."),
        ("Solidarité & Insertion", "Relier les projets à des parcours d'engagement, de formation et d'emploi.", "Bénévoles, structures d'insertion, jeunes, associations."),
    ]
    story.append(pdf_table([[Paragraph("Pôle", CENTER), Paragraph("Mission", CENTER), Paragraph("Publics", CENTER)]] + [[Paragraph(a, BODYL), Paragraph(b, BODYL), Paragraph(c, BODYL)] for a, b, c in poles], [4.0 * cm, 7.3 * cm, 5.1 * cm]))
    story.append(PageBreak())

    # Pole pages with visuals.
    for title, key, body in [
        ("Habitat Vivant", "habitat", "Le logement vacant ou dégradé représente un enjeu à la fois patrimonial, social et urbain. TVF peut aider à transformer une situation bloquée en dossier lisible : coordonnées du propriétaire, état du bien, contraintes, possibilités de convention et acteurs techniques à mobiliser."),
        ("Matériauthèque Solidaire", "materiaux", "La banque de matériaux TVF n'est pas une plateforme de dons libres. Elle organise une valorisation territoriale : chaque lot doit être identifié, qualifié, stocké si nécessaire, puis affecté à un projet compatible et suivi."),
        ("Commerce Vivant", "commerce", "Les locaux commerciaux vacants affaiblissent les centralités. TVF peut faciliter les liens entre propriétaires, collectivités et porteurs de projets pour tester de nouveaux usages : commerce de proximité, atelier, espace associatif ou activité temporaire."),
        ("Friches & Terrains Vivants", "friche", "Les friches et terrains délaissés peuvent devenir des espaces d'expérimentation : jardin partagé, biodiversité, atelier, tiers-lieu léger ou espace événementiel encadré. TVF apporte une méthode de qualification et de coopération."),
        ("Solidarité & Insertion", "solidarite", "La revitalisation doit créer de l'utilité humaine : bénévolat, chantiers participatifs, premières expériences, découverte des métiers, soutien aux structures d'insertion et valorisation des compétences locales."),
    ]:
        story.append(Paragraph(title, H1))
        story.append(PdfImage(str(visuals[key]), width=16.4 * cm, height=7.8 * cm, kind="proportional"))
        story.append(Paragraph(body, BODY))
        story.append(PageBreak())

    section(
        "5. Méthode opérationnelle",
        "TVF doit fonctionner par dossiers, non par annonces. Chaque sujet est qualifié avant action : besoin, acteurs concernés, pièces disponibles, risques, convention possible et indicateurs de suivi.",
    )
    story.append(PdfImage(str(visuals["process"]), width=16.4 * cm, height=5.4 * cm, kind="proportional"))
    story.append(PageBreak())

    section(
        "6. Saint-Étienne : territoire de lancement",
        "Le lancement à Saint-Étienne doit servir de preuve méthodologique. Les premiers dossiers peuvent porter sur un local de stockage, un circuit matériaux, des signalements de biens vacants, une coopération avec les services concernés et des chantiers participatifs encadrés.",
    )
    story.append(PdfImage(str(visuals["saint"]), width=16.4 * cm, height=7.8 * cm, kind="proportional"))
    rows = [[Paragraph("Projet faisable", CENTER), Paragraph("Besoin", CENTER), Paragraph("Acteurs à mobiliser", CENTER)]]
    for r in [
        ("Local de stockage", "Centraliser les matériaux avant affectation.", "Ville, Métropole, bailleurs, propriétaires de locaux."),
        ("Signalement logement", "Qualifier une situation et orienter le propriétaire.", "Propriétaire, habitat, opérateurs techniques."),
        ("Commerce vacant", "Tester un usage de proximité.", "Propriétaire, commerce, porteur de projet, collectivité."),
        ("Chantier participatif", "Lier réemploi, bénévolat et insertion.", "Artisans, insertion, bénévoles, entreprises."),
    ]:
        rows.append([Paragraph(x, BODYL) for x in r])
    story.append(pdf_table(rows, [4.2 * cm, 6.0 * cm, 6.2 * cm]))
    story.append(PageBreak())

    section("7. Bénéfices par public", "Le dossier doit permettre à chaque interlocuteur de comprendre ce qu'il apporte, ce qu'il obtient et pourquoi le cadre TVF sécurise la coopération.")
    rows = [[Paragraph("Public", CENTER), Paragraph("Bénéfice principal", CENTER), Paragraph("Exemple d'engagement", CENTER)]]
    for r in [
        ("Collectivité", "Un outil de coordination, des dossiers lisibles et des indicateurs.", "Mise à disposition, orientation des services, diagnostic partagé."),
        ("Propriétaire", "Un accompagnement vers la remise en usage et la valorisation patrimoniale.", "Convention d'utilisation, étude technique, projet compatible."),
        ("Entreprise", "Une action RSE concrète et traçable.", "Matériaux, compétence, logistique, mécénat."),
        ("Association", "Des espaces, ressources ou chantiers utiles.", "Animation locale, bénévolat, appui projet."),
        ("Financeur", "Un suivi d'impact clair et documenté.", "Soutien financier ou en ingénierie."),
    ]:
        rows.append([Paragraph(x, BODYL) for x in r])
    story.append(pdf_table(rows, [3.1 * cm, 7.0 * cm, 6.3 * cm]))
    story.append(PageBreak())

    section("8. Indicateurs, besoins et sources", "La crédibilité de TVF dépendra de sa capacité à documenter les résultats. Les indicateurs ci-dessous sont à renseigner après actions réelles, avec preuves associées.")
    rows = [[Paragraph("Famille", CENTER), Paragraph("Indicateurs à suivre", CENTER), Paragraph("Preuves attendues", CENTER)]]
    for r in [
        ("Habitat", "Logements repérés, dossiers ouverts, remises en usage.", "Fiches, photos, conventions, attestations."),
        ("Commerce", "Locaux identifiés, porteurs orientés, réactivations.", "Accords, bilans d'usage, photos."),
        ("Matériaux", "Lots proposés, qualifiés, réemployés.", "Inventaires, bordereaux, photos."),
        ("Insertion", "Bénévoles, heures de chantier, parcours.", "Feuilles de présence, conventions."),
    ]:
        rows.append([Paragraph(x, BODYL) for x in r])
    story.append(pdf_table(rows, [3.0 * cm, 7.0 * cm, 6.4 * cm]))
    story.append(Spacer(1, 0.35 * cm))
    story.append(Paragraph("Sources", H2))
    for s in SOURCES:
        story.append(Paragraph(f"<b>{s['label']}</b><br/>{s['url']}<br/>{s['usage']}", SMALL))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph(f"Contact : {CONTACT['email']} - {CONTACT['telephone']} - {CONTACT['site']}", CENTER))

    doc.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)


def basic_quality_checks() -> None:
    for path in [DOCX_PATH, PDF_PATH, MD_PATH]:
        if not path.exists() or path.stat().st_size < 1000:
            raise RuntimeError(f"Fichier non généré ou trop léger : {path}")
    text = MD_PATH.read_text(encoding="utf-8")
    q = chr(63)
    bad = [f"Pr{q}sentation", f"Saint-{q}tienne", f"Territoires Vivants France {q}"]
    for pattern in bad:
        if pattern in text:
            raise RuntimeError(f"Caractère cassé détecté dans la source : {pattern}")


def main() -> None:
    ensure_dirs()
    visuals = prepare_visuals()
    write_source_markdown()
    build_docx(visuals)
    build_pdf(visuals)
    basic_quality_checks()
    print(json.dumps({
        "docx": str(DOCX_PATH),
        "pdf": str(PDF_PATH),
        "source": str(MD_PATH),
        "assets": str(OUT_WORK),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
