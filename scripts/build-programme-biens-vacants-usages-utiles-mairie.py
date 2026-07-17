from pathlib import Path
from datetime import date
from html import unescape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents" / "programmes-territoriaux"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "programme-biens-vacants-usages-utiles-saint-etienne-mairie.docx"
MD_PATH = SRC_DIR / "programme-biens-vacants-usages-utiles-saint-etienne-mairie.md"
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"

GREEN = RGBColor(24, 63, 34)
BLUE = RGBColor(7, 30, 48)
MUTED = RGBColor(84, 96, 90)
GOLD = RGBColor(178, 132, 24)
LIGHT_GREEN = "EAF3EA"
LIGHT_BLUE = "EEF4F7"
LIGHT_GOLD = "FBF4E3"
BORDER = "B8C8BC"


def txt(value: str) -> str:
    return unescape(value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8E0D9", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = BLUE
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in [
        ("Heading 1", 18.0, GREEN, 15, 8),
        ("Heading 2", 13.4, GREEN, 12, 5),
        ("Heading 3", 11.6, BLUE, 7, 3),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Cm(17.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.4, 12.1])
    for cell in table.row_cells(0):
        set_cell_border(cell, color="FFFFFF", size="0")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left = table.cell(0, 0)
    p = left.paragraphs[0]
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(4.8))

    right = table.cell(0, 1)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("TERRITOIRES VIVANTS FRANCE\n")
    r.bold = True
    r.font.size = Pt(9.8)
    r.font.color.rgb = GREEN
    r2 = p.add_run(
        "Association loi 1901 déclarée - RNA W423016361\n"
        "SIREN 107 226 128 - SIRET 107 226 128 00018\n"
        "25 rue Élise Gervais, 42000 Saint-Étienne\n"
        "contact@territoiresvivantsfrance.fr - 04 65 81 54 69"
    )
    r2.font.size = Pt(8.2)
    r2.font.color.rgb = MUTED

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Programme territorial TVF - Biens vacants, usages utiles - "
        "Document de travail pour échange institutionnel"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED


def add_para(doc, text, bold_first=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5.5)
    p.paragraph_format.line_spacing = 1.08
    if bold_first and text.startswith(bold_first):
        r = p.add_run(bold_first)
        r.bold = True
        r.font.color.rgb = GREEN
        p.add_run(text[len(bold_first):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4.8)
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4.8)
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.first_line_indent = Cm(-0.28)
    p.add_run(text)
    return p


def add_callout(doc, title, body, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17.4)
    cell = table.cell(0, 0)
    cell.width = Cm(17.4)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=BORDER, size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = GREEN
    r.font.size = Pt(11.2)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    p2.add_run(body)
    return table


def add_table(doc, headers, rows, widths_cm, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths_cm)
    repeat_header(table.rows[0])
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.2)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_border(cells[idx], color="D8E0D9", size="6")
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.03
            r = p.add_run(value)
            r.font.size = Pt(8.8)
            r.font.color.rgb = BLUE
    doc.add_paragraph()
    return table


def add_title_page(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Cm(7.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("PROGRAMME TERRITORIAL PILOTE\n")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = GOLD

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Biens vacants, usages utiles")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Proposition de programme opérationnel pour Saint-Étienne\n"
        "à destination de la Ville de Saint-Étienne"
    )
    r.font.size = Pt(14)
    r.font.color.rgb = BLUE

    add_callout(
        doc,
        "Objet du document",
        "Présenter une méthode concrète permettant de transformer des logements, commerces, bâtiments, terrains et matériaux inutilisés en projets utiles pour les habitants, les associations et le territoire, dans un cadre légal sécurisé et sans se substituer aux compétences de la collectivité ni aux dispositifs publics existants.",
        fill=LIGHT_GREEN,
    )

    meta = [
        ("Version", "Document de travail - première proposition structurée"),
        ("Date", date.today().strftime("%d/%m/%Y")),
        ("Porteur", "TERRITOIRES VIVANTS FRANCE"),
        ("Territoire pilote", "Saint-Étienne"),
        ("Contact", "contact@territoiresvivantsfrance.fr - 04 65 81 54 69"),
    ]
    add_table(doc, ["Élément", "Information"], meta, [4.2, 13.2], header_fill=LIGHT_GOLD)
    doc.add_page_break()


def build_doc():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_page(doc)

    doc.add_heading("1. Synthèse exécutive", level=1)
    add_para(
        doc,
        "TERRITOIRES VIVANTS FRANCE propose à la Ville de Saint-Étienne la mise en place d’un programme pilote intitulé « Biens vacants, usages utiles ». Ce programme vise à identifier des ressources aujourd’hui sous-utilisées - logements vacants, commerces fermés, bâtiments inutilisés, terrains délaissés, matériaux réemployables - puis à étudier, au cas par cas, les usages pouvant répondre aux besoins du territoire."
    )
    add_para(
        doc,
        "Le programme ne consiste pas à promettre une rénovation automatique de chaque bien signalé. Il repose sur une méthode progressive : repérage, qualification, vérification juridique, diagnostic technique, recherche d’usage, conventionnement, mobilisation des ressources et suivi. Cette approche permet de distinguer les projets immédiatement mobilisables, les projets nécessitant une ingénierie complémentaire et les situations non adaptées au cadre d’intervention de TVF."
    )
    add_callout(
        doc,
        "Positionnement proposé",
        "TVF agit comme plateforme de coordination territoriale : elle met en relation des propriétaires, collectivités, entreprises, associations, artisans, bénévoles et financeurs autour de biens ou ressources inutilisés. Elle ne remplace ni les services municipaux, ni les bailleurs, ni les opérateurs publics, ni les professionnels réglementés ; elle facilite la transformation de ressources dormantes en projets utiles.",
    )

    doc.add_heading("2. Pourquoi un programme pilote à Saint-Étienne ?", level=1)
    add_para(
        doc,
        "Saint-Étienne présente plusieurs caractéristiques qui rendent pertinente une expérimentation locale : un parc ancien important, une vacance de logements supérieure à la moyenne nationale, des enjeux sociaux marqués, des quartiers en renouvellement, des besoins d’insertion et un tissu d’entreprises, d’artisans et d’associations susceptible d’être mobilisé autour du réemploi et de la revitalisation."
    )
    add_para(
        doc,
        "Le programme pilote permettrait d’expérimenter une méthode simple et reproductible : partir des besoins concrets du territoire, identifier les biens ou ressources disponibles, puis construire des usages réalistes, sécurisés et utiles. L’intérêt pour la Ville est de disposer d’un interlocuteur associatif capable de capter des informations de terrain, d’orienter les propriétaires et de contribuer à la mise en mouvement de projets compatibles avec les politiques publiques locales."
    )

    key_rows = [
        ("Population communale", "173 136 habitants en 2023", "INSEE, dossier complet commune de Saint-Étienne"),
        ("Logements vacants", "12 175 logements vacants en 2023, soit 12,0 % du parc de logements", "INSEE, dossier complet commune de Saint-Étienne"),
        ("Pauvreté", "Taux de pauvreté de 30,4 %", "INSEE, dossier complet commune de Saint-Étienne"),
        ("Chômage", "Taux de chômage au sens du recensement : 19,0 %", "INSEE, dossier complet commune de Saint-Étienne"),
        ("Quartiers prioritaires", "36 016 habitants en QPV, soit 20,7 % de la population communale", "SIG Ville / politique de la ville"),
        ("Déchets et ressources", "Le secteur du bâtiment et des travaux publics reste un gisement majeur de déchets et de matériaux valorisables", "Ministère de la Transition écologique / filière PMCB, ADEME, politiques locales déchets"),
    ]
    add_table(doc, ["Indicateur", "Constat utilisé", "Source"], key_rows, [4.2, 7.2, 6.0])
    add_para(
        doc,
        "Ces données ne suffisent pas à désigner automatiquement les activités manquantes. Elles permettent en revanche d’orienter une enquête territoriale : où les besoins sont-ils les plus forts ? Quels locaux sont réellement disponibles ? Quels usages sont juridiquement possibles ? Quels acteurs peuvent être associés ?"
    )

    doc.add_heading("3. Objectifs du programme", level=1)
    for item in [
        "Recenser les besoins locaux pouvant être traités par la remise en usage de biens ou ressources inutilisés.",
        "Identifier des biens compatibles : logements vacants, locaux commerciaux, bâtiments, terrains, espaces de stockage ou ressources matérielles.",
        "Créer un processus de qualification unique permettant de trier les situations selon leur faisabilité réelle.",
        "Proposer des usages utiles au territoire : habitat temporaire, local associatif, atelier partagé, commerce test, matériauthèque, chantier d’insertion, jardin partagé ou espace de formation.",
        "Structurer les conventions entre TVF, les propriétaires, les collectivités, les entreprises ou les associations concernées.",
        "Mesurer l’impact : biens remis en usage, matériaux réemployés, projets accompagnés, bénéficiaires, bénévoles mobilisés et coûts évités lorsque ceux-ci sont documentables.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("4. Publics concernés et bénéfices attendus", level=1)
    rows = [
        ("Ville / services municipaux", "Disposer d’un outil de repérage, d’orientation et de mobilisation complémentaire aux politiques publiques.", "Information de terrain, mise en relation, projets légers, suivi des besoins."),
        ("Propriétaires privés", "Étudier un usage possible pour un bien dormant sans perdre la propriété.", "Diagnostic, orientation, convention, valorisation progressive du patrimoine."),
        ("Entreprises", "Valoriser des stocks, matériaux ou équipements inutilisés.", "Réduction du gaspillage, contribution territoriale, démarche RSE documentable."),
        ("Associations", "Accéder à des locaux, matériaux ou appuis pour développer des actions utiles.", "Espaces d’activité, ressources matérielles, partenariat local."),
        ("Habitants", "Bénéficier de services, lieux et activités utiles dans leur quartier.", "Amélioration du cadre de vie, participation citoyenne, proximité."),
        ("Artisans / structures d’insertion", "Participer à des projets de rénovation légère ou de réemploi.", "Activité, transmission, formation, insertion et coopération locale."),
    ]
    add_table(doc, ["Acteur", "Intérêt du programme", "Bénéfice concret"], rows, [4.2, 6.6, 6.6])

    doc.add_heading("5. Activités à qualifier dans des biens vacants ou sous-utilisés", level=1)
    add_para(
        doc,
        "Le programme ne part pas d’une liste figée d’activités. Il propose une grille d’analyse permettant de vérifier quelles activités répondent réellement aux besoins locaux. Les pistes suivantes constituent des hypothèses de travail à qualifier avec les services de la Ville, les acteurs économiques, les associations et les propriétaires concernés."
    )
    rows = [
        ("Logement vacant", "Logement temporaire, colocation solidaire, logement étudiant, habitat intergénérationnel", "Propriétaire, urbanisme, sécurité, diagnostics, assurance, convention d’usage"),
        ("Local commercial fermé", "Boutique test, atelier artisan, commerce de proximité, café associatif, point de services", "Destination commerciale, bail ou convention, ERP, accessibilité, sécurité incendie"),
        ("Bâtiment inutilisé", "Local associatif, espace de formation, atelier partagé, stockage léger, permanence sociale", "État structurel, sécurité, usage autorisé, capacité d’accueil, assurance"),
        ("Ancien atelier / dépôt", "Matériauthèque, tri de matériaux, reconditionnement, réparation, logistique solidaire", "Accès livraison, stockage, sécurité, risques, assurance, traçabilité des matériaux"),
        ("Terrain ou friche légère", "Jardin partagé, espace pédagogique, biodiversité, compostage encadré, micro-projet nourricier", "Propriété, pollution potentielle, autorisations, clôture, gestion, convention"),
        ("Rez-de-chaussée vacant", "Point numérique, atelier vélo, ressourcerie spécialisée, accueil associatif", "ERP, accessibilité, nuisances, convention, gestion des horaires"),
    ]
    add_table(doc, ["Type de bien", "Usages possibles", "Conditions à vérifier"], rows, [4.2, 7.0, 6.2])

    doc.add_heading("6. Matrice opérationnelle proposée à la Ville", level=1)
    rows = [
        ("Vacance de logements", "Logement vacant ou dégradé léger", "Occupation temporaire, logement étudiant, logement associatif", "Habitat Vivant", "Propriétaire, service habitat, juriste, assureur, diagnostiqueur"),
        ("Locaux commerciaux fermés", "Boutique, cellule commerciale, rez-de-chaussée", "Commerce test, atelier artisan, activité ESS", "Commerce Vivant", "Ville, CCI, CMA, commerçants, propriétaires"),
        ("Besoin d’espaces associatifs", "Bâtiment ou local inutilisé", "Local associatif, permanence, atelier de quartier", "Solidarité & Insertion", "Associations, services municipaux, bailleurs, propriétaires"),
        ("Matériaux inutilisés", "Dépôt, stock dormant, surplus entreprise", "Banque de matériaux, reconditionnement, affectation projet", "Matériauthèque Solidaire", "Entreprises, artisans, services déchets, assureur"),
        ("Terrain délaissé", "Terrain vacant, friche légère", "Jardin partagé, biodiversité, espace pédagogique", "Friches & Terrains Vivants", "Urbanisme, environnement, associations, riverains"),
        ("Besoin d’insertion", "Chantier léger, local atelier", "Chantier participatif, formation, découverte métiers", "Solidarité & Insertion", "Mission Locale, PLIE, structures IAE, entreprises"),
    ]
    add_table(doc, ["Besoin identifié", "Support potentiel", "Usage à étudier", "Pôle TVF", "Acteurs à associer"], rows, [3.3, 3.7, 4.3, 2.8, 3.3])

    doc.add_heading("7. Méthode de traitement d’un dossier", level=1)
    steps = [
        "Signalement ou identification du besoin : bien vacant, local fermé, terrain, matériau disponible ou besoin d’un acteur local.",
        "Création d’une fiche dossier : coordonnées, catégorie, adresse, propriétaire connu, photos, description, urgence, premier usage envisagé.",
        "Vérification de recevabilité : accord du propriétaire ou interlocuteur légitime, cohérence avec l’objet TVF, absence de danger manifeste, niveau de travaux compatible.",
        "Pré-diagnostic : état apparent, accès, usage possible, risques, besoin de professionnel compétent si le dossier dépasse la rénovation légère.",
        "Orientation juridique et administrative : urbanisme, destination, ERP, assurances, autorisations, diagnostics, convention adaptée.",
        "Recherche de ressources : matériaux de réemploi, artisans, bénévoles, entreprises, associations, financements ou mécénat si nécessaire.",
        "Conventionnement : définition de l’usage, durée, responsabilités, assurances, accès, suivi, fin de convention.",
        "Mise en œuvre progressive : travaux légers, aménagement, ouverture encadrée, communication validée avec les parties.",
        "Suivi d’impact : indicateurs simples, bilan, points de blocage, suites possibles.",
    ]
    for step in steps:
        add_number(doc, step)

    doc.add_heading("8. Cadre légal et limites d’intervention", level=1)
    add_callout(
        doc,
        "Principe de prudence",
        "Chaque projet devra faire l’objet d’une vérification au cas par cas. Le programme ne vaut ni autorisation administrative, ni diagnostic technique, ni engagement financier automatique, ni promesse de rénovation. Les interventions lourdes, structurelles ou réglementées devront être assurées par des professionnels compétents.",
        fill=LIGHT_GOLD,
    )
    for item in [
        "TVF ne peut intervenir qu’avec l’accord explicite du propriétaire ou d’un interlocuteur habilité.",
        "Les changements d’usage, ouvertures au public, travaux, installations et occupations doivent respecter le droit de l’urbanisme, les règles ERP le cas échéant, les assurances et les exigences de sécurité.",
        "Les biens présentant un risque structurel, sanitaire, électrique, incendie ou environnemental doivent être orientés vers des professionnels compétents avant toute mobilisation.",
        "Le programme vise prioritairement la rénovation légère, l’orientation, la coordination, le réemploi, la mise en relation et les occupations utiles juridiquement encadrées.",
        "Chaque convention devra préciser les responsabilités, la durée, les assurances, la fin d’usage, la gestion des clés, les conditions financières éventuelles et les modalités de suivi.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("9. Demandes proposées à la Ville de Saint-Étienne", level=1)
    add_para(
        doc,
        "Pour lancer le programme dans de bonnes conditions, TVF pourrait solliciter de la Ville un premier cadre de coopération léger, sans engagement automatique de financement ni transfert de responsabilité. L’objectif est d’ouvrir un dialogue opérationnel et d’identifier les conditions d’expérimentation."
    )
    rows = [
        ("Rendez-vous institutionnel", "Présentation du programme aux services concernés", "Valider l’intérêt, identifier les interlocuteurs et cadrer les priorités."),
        ("Mise en relation", "Habitat, commerce, urbanisme, politique de la ville, transition écologique, vie associative", "Éviter les doublons et inscrire TVF dans les priorités locales."),
        ("Local de stockage", "Étude d’un local municipal ou d’un partenaire pour matériaux", "Permettre la réception, le tri et l’affectation de ressources utiles."),
        ("Biens à qualifier", "Orientation vers locaux, friches ou bâtiments pouvant faire l’objet d’une étude", "Constituer les premières fiches projet sans promesse préalable."),
        ("Appui documentaire", "Accès aux dispositifs publics, contacts, procédures et contraintes locales", "Sécuriser le cadre légal et administratif."),
        ("Expérimentation limitée", "Démarrage sur quelques cas pilotes simples", "Prouver la méthode avant élargissement."),
    ]
    add_table(doc, ["Demande", "Forme possible", "Utilité"], rows, [4.5, 6.0, 6.9])

    doc.add_heading("10. Exemples de projets pilotes à étudier", level=1)
    rows = [
        ("Matériauthèque TVF", "Ancien dépôt, atelier, local logistique", "Réceptionner et qualifier des matériaux utiles aux projets locaux.", "Entreprises BTP, artisans, services déchets, associations"),
        ("Boutique test ESS", "Commerce vacant en rez-de-chaussée", "Tester une activité de proximité avant installation durable.", "Ville, CCI, CMA, propriétaires, porteurs de projet"),
        ("Atelier rénovation légère", "Local atelier ou bâtiment municipal disponible", "Former, réparer, reconditionner, mobiliser des bénévoles.", "Mission Locale, PLIE, structures insertion, artisans"),
        ("Logement utile temporaire", "Logement vacant nécessitant peu de travaux", "Étudier un usage temporaire encadré pour étudiant, association ou public accompagné.", "Propriétaire, service habitat, assureur, diagnostiqueur"),
        ("Jardin partagé de quartier", "Terrain vacant non pollué ou espace en attente", "Créer un espace utile, pédagogique et fédérateur.", "Urbanisme, associations, habitants, écoles"),
        ("Point de services de proximité", "Local vacant accessible", "Accueillir permanence associative, aide administrative ou médiation.", "Associations, services publics, bailleurs, habitants"),
    ]
    add_table(doc, ["Projet pilote", "Support adapté", "Objectif", "Acteurs mobilisables"], rows, [3.6, 4.2, 5.2, 4.4])

    doc.add_heading("11. Calendrier de lancement proposé", level=1)
    rows = [
        ("0 à 30 jours", "Rendez-vous Ville / TVF, validation des interlocuteurs, accord sur une méthode de travail.", "Feuille de route courte, liste des services à rencontrer."),
        ("30 à 60 jours", "Recueil des besoins, premières fiches biens/ressources, identification d’un local de stockage potentiel.", "Matrice besoins-biens-usages, premiers dossiers de qualification."),
        ("60 à 90 jours", "Pré-diagnostics, rencontres propriétaires/entreprises/associations, vérification juridique des cas simples.", "Liste de projets recevables et non recevables, priorisation."),
        ("3 à 6 mois", "Conventionnement des premiers dossiers simples, collecte matériaux, mobilisation bénévoles et partenaires.", "Premiers projets pilotes préparés ou lancés."),
        ("6 à 12 mois", "Bilan, indicateurs, amélioration méthode, décision sur poursuite ou extension.", "Rapport d’étape et recommandations pour la suite."),
    ]
    add_table(doc, ["Période", "Actions", "Livrables"], rows, [3.0, 8.0, 6.4])

    doc.add_heading("12. Indicateurs de suivi", level=1)
    add_para(
        doc,
        "Les indicateurs ci-dessous ne constituent pas des résultats annoncés. Ils définissent les éléments que TVF propose de suivre si le programme est expérimenté."
    )
    for item in [
        "Nombre de biens signalés, qualifiés et orientés.",
        "Nombre de propriétaires rencontrés.",
        "Nombre de locaux commerciaux ou bâtiments étudiés.",
        "Nombre de ressources matérielles proposées et réellement réemployées.",
        "Nombre de conventions signées.",
        "Nombre d’associations, entreprises ou porteurs de projet mobilisés.",
        "Nombre de projets rejetés ou réorientés pour raison juridique, technique ou financière.",
        "Volume de matériaux réemployés, lorsque la mesure est possible.",
        "Bénéficiaires directs et indirects documentés.",
        "Difficultés rencontrées et besoins d’ingénierie complémentaire.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("13. Documents à préparer pour chaque dossier", level=1)
    rows = [
        ("Bien immobilier", "Adresse, propriétaire, photos, description, état apparent, usage actuel, titre ou autorisation, contraintes connues."),
        ("Matériaux", "Type, quantité, état, localisation, photos, conditions de retrait, disponibilité, éventuels risques ou précautions."),
        ("Projet d’usage", "Public visé, porteur, besoin local, durée envisagée, contraintes, ressources nécessaires, modèle de convention."),
        ("Cadre légal", "Urbanisme, assurance, ERP si nécessaire, diagnostics, sécurité, accessibilité, accord écrit."),
        ("Partenariat", "Rôle de chaque partie, engagements réciproques, durée, responsabilités, modalités de suivi."),
    ]
    add_table(doc, ["Catégorie", "Pièces ou informations à réunir"], rows, [4.2, 13.2])

    doc.add_heading("14. Sources mobilisées", level=1)
    sources = [
        ("INSEE", "Dossier complet de la commune de Saint-Étienne, code commune 42218", "https://www.insee.fr/fr/statistiques/2011101?geo=COM-42218"),
        ("SIG Ville", "Données politique de la ville et quartiers prioritaires - Saint-Étienne", "https://sig.ville.gouv.fr/territoire/42218"),
        ("Saint-Étienne Métropole", "Plan Climat Air Énergie Territorial et politiques de transition", "https://www.saint-etienne-metropole.fr/preserver-recycler/energie-climat/plan-climat"),
        ("Saint-Étienne Métropole", "Réduction des déchets et prévention", "https://www.saint-etienne-metropole.fr/preserver-recycler/gestion-des-dechets/reduire-vos-dechets-oui-mais-comment"),
        ("Cerema / data.gouv.fr", "Cartofriches - sites référencés dans la base nationale", "https://www.data.gouv.fr/datasets/sites-references-dans-cartofriches/"),
        ("Ville de Saint-Étienne", "Actualités et données publiques relatives au commerce et à l’attractivité", "https://www.saint-etienne.fr/"),
        ("CCI Lyon Métropole Saint-Étienne Roanne", "Plan commerce Rhône-Loire et accompagnement des commerçants", "https://www.lyon-metropole.cci.fr/"),
    ]
    add_table(doc, ["Source", "Utilisation", "Lien"], sources, [3.2, 7.2, 7.0])

    doc.add_heading("15. Conclusion proposée", level=1)
    add_para(
        doc,
        "Le programme « Biens vacants, usages utiles » propose à Saint-Étienne une méthode opérationnelle pour passer du constat à l’action. Il ne s’agit pas d’ajouter un dispositif isolé, mais de créer une interface de terrain capable de relier des besoins locaux, des biens dormants, des ressources matérielles et des acteurs prêts à agir."
    )
    add_para(
        doc,
        "La première étape proposée est simple : organiser une rencontre entre TVF et les services municipaux concernés afin de vérifier les priorités, identifier les premiers cas d’usage et définir les conditions d’une expérimentation courte, prudente et mesurable."
    )
    add_callout(
        doc,
        "Demande de TVF",
        "Obtenir un rendez-vous de travail avec la Ville de Saint-Étienne afin de présenter le programme, identifier les services référents et étudier les conditions d’un premier pilote local : local de stockage, fiches biens à qualifier, partenaires à mobiliser et cadre de conventionnement.",
        fill=LIGHT_GREEN,
    )

    doc.save(DOCX_PATH)


def build_markdown():
    md = """# Programme territorial pilote - Biens vacants, usages utiles

**Destinataire :** Ville de Saint-Étienne  
**Porteur :** TERRITOIRES VIVANTS FRANCE  
**Objet :** Proposition de programme opérationnel pour transformer des biens vacants et ressources inutilisées en usages utiles au territoire.

## Synthèse

TERRITOIRES VIVANTS FRANCE propose à la Ville de Saint-Étienne la mise en place d’un programme pilote intitulé « Biens vacants, usages utiles ». Ce programme vise à identifier des ressources aujourd’hui sous-utilisées - logements vacants, commerces fermés, bâtiments inutilisés, terrains délaissés, matériaux réemployables - puis à étudier, au cas par cas, les usages pouvant répondre aux besoins du territoire.

Le programme repose sur une méthode progressive : repérage, qualification, vérification juridique, diagnostic technique, recherche d’usage, conventionnement, mobilisation des ressources et suivi.

## Données de cadrage

- Population communale : 173 136 habitants en 2023 (INSEE).
- Logements vacants : 12 175 logements vacants en 2023, soit 12,0 % du parc (INSEE).
- Taux de pauvreté : 30,4 % (INSEE).
- Taux de chômage au sens du recensement : 19,0 % (INSEE).
- QPV : 36 016 habitants, soit 20,7 % de la population communale (SIG Ville).

## Sources

- INSEE, dossier complet commune 42218 : https://www.insee.fr/fr/statistiques/2011101?geo=COM-42218
- SIG Ville, Saint-Étienne : https://sig.ville.gouv.fr/territoire/42218
- Saint-Étienne Métropole, Plan Climat : https://www.saint-etienne-metropole.fr/preserver-recycler/energie-climat/plan-climat
- Saint-Étienne Métropole, déchets : https://www.saint-etienne-metropole.fr/preserver-recycler/gestion-des-dechets/reduire-vos-dechets-oui-mais-comment
- Cerema / data.gouv.fr, Cartofriches : https://www.data.gouv.fr/datasets/sites-references-dans-cartofriches/
"""
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_doc()
    build_markdown()
    print(DOCX_PATH)
    print(MD_PATH)
