from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "documents" / "sources" / "courrier-ville-saint-etienne-lancement-tvf.md"
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"
NAVY = RGBColor(9, 31, 49)
GREEN = RGBColor(46, 125, 50)
DARK_GREEN = RGBColor(24, 63, 34)
MUTED = RGBColor(84, 96, 103)
BORDER = "D9E4DA"


def parse_source(path):
    text = path.read_text(encoding="utf-8")
    if not text.startswith("---\n"):
        raise SystemExit(f"Source front matter manquant: {path}")
    _, meta_raw, body = text.split("---\n", 2)
    meta = {}
    current = None
    for line in meta_raw.splitlines():
        if not line.strip():
            continue
        if line.startswith("  - ") and current:
            meta.setdefault(current, []).append(line[4:].strip())
            continue
        if ":" in line:
            key, value = line.split(":", 1)
            key = key.strip()
            value = value.strip()
            current = key
            if value:
                meta[key] = value
            else:
                meta[key] = []
    return meta, body.strip()


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E4DA", size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=10.5, color=None, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_text(paragraph, text, size=10.5, color=None, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_paragraph(doc, text="", size=10.5, color=None, bold=False, italic=False, after=6, before=0, align=None, line=1.12):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align
    if text:
        add_text(paragraph, text, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    add_text(paragraph, text, size=13, color=DARK_GREEN, bold=True)


def bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.08
    add_text(paragraph, "- ", size=10.2, color=GREEN, bold=True)
    add_text(paragraph, text, size=10.2, color=NAVY)


def add_rule(paragraph, color="2E7D32", size="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def info_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [2.05, 4.25])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "EEF7EE")
        set_cell_border(cell, BORDER, "6")
        set_cell_margins(cell, 120, 160, 120, 160)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left, right = table.rows[0].cells
    p = left.paragraphs[0]
    add_text(p, "Références TVF", size=9.5, color=DARK_GREEN, bold=True)
    for label, value in [
        ("RNA", "W423016361"),
        ("SIRET", "107 226 128 00018"),
        ("Siège", "25 rue Élise Gervais, 42000 Saint-Étienne"),
    ]:
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_text(p, f"{label} : ", size=8.8, color=NAVY, bold=True)
        add_text(p, value, size=8.8, color=NAVY)
    p = right.paragraphs[0]
    add_text(p, "Contact", size=9.5, color=DARK_GREEN, bold=True)
    for label, value in [
        ("Téléphone", "04 65 81 54 69"),
        ("E-mail", "contact@territoiresvivantsfrance.fr"),
        ("Site", "https://www.territoiresvivantsfrance.fr"),
    ]:
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_text(p, f"{label} : ", size=8.8, color=NAVY, bold=True)
        add_text(p, value, size=8.8, color=NAVY)


def add_object_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0].cells[0].width = Inches(6.3)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EEF4F8")
    set_cell_border(cell, "C8D7E3", "8")
    set_cell_margins(cell, 130, 170, 130, 170)
    paragraph = cell.paragraphs[0]
    add_text(paragraph, "Objet : ", size=10.3, color=NAVY, bold=True)
    add_text(paragraph, text, size=10.3, color=NAVY, bold=True)


def add_attachments(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0].cells[0].width = Inches(6.3)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8FAF5")
    set_cell_border(cell, BORDER, "6")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    add_text(p, "Pièces pouvant être jointes : ", size=9, color=DARK_GREEN, bold=True)
    add_text(p, text, size=9, color=NAVY)


def add_markdown_body(doc, body):
    for block in body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("## "):
            heading(doc, block[3:].strip())
        elif block.startswith("- "):
            for line in block.splitlines():
                if line.startswith("- "):
                    bullet(doc, line[2:].strip())
        else:
            add_paragraph(doc, block.replace("\n", " "), size=10.35, color=NAVY, after=5, line=1.12)


def build():
    meta, body = parse_source(SOURCE)
    output = ROOT / meta["output"]
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.12

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(hp, "TERRITOIRES VIVANTS FRANCE - Courrier institutionnel", size=8.5, color=MUTED)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(fp, "TERRITOIRES VIVANTS FRANCE - 25 rue Élise Gervais, 42000 Saint-Étienne - contact@territoiresvivantsfrance.fr - 04 65 81 54 69", size=7.8, color=MUTED)

    if LOGO.exists():
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(2.35))
        p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_text(p, "Association loi 1901 déclarée - Plateforme nationale de coopération pour la revitalisation des territoires", size=9.2, color=MUTED, bold=True)
    add_rule(p)
    info_table(doc)
    add_paragraph(doc, 'Saint-Étienne, le [date à compléter]', size=10.2, color=NAVY, after=8, before=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for idx, line in enumerate(meta.get("recipient", [])):
        add_paragraph(doc, line, size=10.3, color=NAVY, bold=(idx == 0), after=1)
    add_object_box(doc, meta["object"])
    add_markdown_body(doc, body)
    add_paragraph(doc, "Pour TERRITOIRES VIVANTS FRANCE,", size=10.4, color=NAVY, after=18, before=8)
    add_paragraph(doc, meta.get("signature_name", ""), size=10.6, color=NAVY, bold=True, after=1)
    add_paragraph(doc, meta.get("signature_role", ""), size=10.2, color=MUTED, after=10)
    add_attachments(doc, meta.get("attachments", ""))
    props = doc.core_properties
    props.title = meta.get("title", "Document TVF")
    props.subject = meta.get("object", "")
    props.author = "Territoires Vivants France"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    build()
