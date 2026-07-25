from pathlib import Path
from runpy import run_path
from html import unescape
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))
Document = BASE["Document"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
BLUE = BASE["BLUE"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
GOLD = BASE["GOLD"]
BORDER = BASE["BORDER"]

OUT_DIR = ROOT / "documents" / "brochures" / "poles"
SRC_DIR = ROOT / "documents" / "sources" / "brochures-poles"
for d in (OUT_DIR, SRC_DIR):
    d.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"

def u(text):
    return unescape(text)

def set_run(run, size=None, color=None, bold=None, italic=None):
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(u(text))
    set_run(r, 8.5, MUTED, italic=True)

def table(doc, headers, rows, widths, fill=LIGHT_GREEN):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, widths)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=BORDER, size="8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(u(header))
        set_run(r, 8.7, GREEN, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i], color="DDE6DE", size="6")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(u(str(value)))
            set_run(r, 8.25, BLUE if i == 0 else RGBColor(31, 45, 38), bold=(i == 0))
    doc.add_paragraph()
    return t

def image_block(doc, image_path, caption):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(15.4))
    p.paragraph_format.space_after = Pt(2)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(9)
    r = cp.add_run(u(caption))
    set_run(r, 8.4, MUTED, italic=True)

def section_intro(doc, kicker, title, lead):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(u(kicker.upper()))
    set_run(r, 8.2, GOLD, bold=True)
    doc.add_heading(u(title), level=1)
    add_paragraph(doc, lead)

def metadata(doc, pole):
    rows = [("Document", "Brochure pôle TVF"), ("Pôle", pole["name"]), ("Public", pole["public"]), ("Version", "Juillet 2026")]
    t = doc.add_table(rows=len(rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, [4.3, 12.8])
    for i, (label, value) in enumerate(rows):
        for cell in t.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(t.cell(i, 0), LIGHT_BLUE)
        p1 = t.cell(i, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(u(label))
        set_run(r1, 8.6, GREEN, bold=True)
        p2 = t.cell(i, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(u(value))
        set_run(r2, 8.7, BLUE)
    doc.add_paragraph()

POLES = [
    {"slug":"habitat-vivant","name":"Habitat Vivant","subtitle":"Remettre en usage les logements vacants ou dégradés lorsque le cadre technique, humain et juridique le permet.","public":"propriétaires, collectivités, bailleurs, associations, habitants","photo":"immeuble-renovation-meudon.jpg","caption":"Réhabilitation d'un bâtiment ancien : le pôle Habitat Vivant étudie les conditions de remise en usage, sans engagement automatique de travaux.","why":["La vacance résidentielle n'a pas une seule cause : succession non réglée, coût de rénovation, copropriété fragile, absence de projet, difficultés administratives, état technique incertain ou manque de porteur d'usage.","Pour TVF, un logement vacant ne doit jamais être traité seulement comme une adresse. Il faut comprendre son état, son propriétaire, son environnement, les contraintes de sécurité, les usages possibles et l'intérêt territorial réel.","Le pôle Habitat Vivant vise donc à transformer une situation bloquée en dossier lisible : pièces disponibles, besoin public, niveau de travaux, partenaires à consulter et scénario réaliste de remise en usage."],"repere":"Repères publics mobilisables : INSEE pour le parc de logements et la vacance ; ministère chargé du Logement pour les politiques de rénovation ; données locales des collectivités lorsque disponibles.","missions":[("Repérer","Identifier les biens vacants ou sous-utilisés à partir de signalements, de données publiques ou d'échanges avec les acteurs locaux."),("Qualifier","Vérifier le type de bien, l'état apparent, les contraintes connues, les pièces disponibles et le niveau de faisabilité."),("Orienter","Diriger vers un parcours propriétaire, collectivité, association, diagnostic ou classement si la demande n'est pas compatible."),("Construire","Préparer un scénario d'usage : logement associatif, étudiant, temporaire, intergénérationnel ou autre usage utile."),("Conventionner","Encadrer par écrit les rôles, durées, responsabilités, limites et conditions de suivi.")],"documents":["Fiche propriétaire","Photos du bien","Adresse précise","Éléments de propriété ou mandat","Contraintes connues","Diagnostics existants","Projet ou usage envisagé"],"conventions":[("Autorisation de visite","Permettre une première observation encadrée sans démarrer de travaux."),("Convention de mise à disposition","Définir une durée, un usage, les responsabilités et les conditions de restitution."),("Convention de valorisation partagée","Organiser un projet où le propriétaire conserve son bien et où l'usage sert un objectif territorial.")],"limits":["TVF ne se substitue pas à un bureau d'études, un maître d'œuvre, un diagnostiqueur ou une entreprise qualifiée.","Une demande ne vaut pas acceptation du bien ni engagement de rénovation.","Les travaux lourds, dangereux ou non assurables doivent être orientés vers les professionnels compétents."],"benefits":["Clarification du dossier","Orientation vers le bon parcours","Valorisation du patrimoine dormant","Utilité sociale et territoriale","Traçabilité des décisions"],"cta":"Proposer un bien vacant ou demander un premier échange : contact@territoiresvivantsfrance.fr"},
    {"slug":"materiautheque-solidaire","name":"Matériauthèque Solidaire","subtitle":"Transformer des matériaux inutilisés en ressources utiles pour des projets territoriaux validés.","public":"entreprises, artisans, collectivités, particuliers, associations","photo":"materiaux-reemploi-echantillons.jpg","caption":"Matériaux de réemploi : TVF cherche à qualifier l'état, l'usage possible, la traçabilité et l'affectation vers des projets utiles.","why":["Les matériaux inutilisés représentent souvent un coût de stockage, une perte économique ou un futur déchet, alors que certains peuvent encore servir à des projets locaux.","La Matériauthèque Solidaire TVF n'est pas une plateforme de distribution gratuite. Elle fonctionne comme un outil de valorisation territoriale : chaque ressource doit être identifiée, contrôlée, acceptée ou refusée, puis affectée à un projet cohérent.","L'objectif est de relier les besoins des territoires avec les ressources disponibles, dans un cadre documenté : origine, état, quantité, localisation, contraintes, destination et preuve de réemploi."],"repere":"Repère public : le ministère de la Transition écologique rappelle que le bâtiment génère environ 42 millions de tonnes de déchets par an en France, dans le cadre de la filière PMCB.","missions":[("Recenser","Recevoir les propositions de matériaux, équipements, mobilier, menuiseries, sanitaires ou outillage."),("Diagnostiquer","Vérifier l'état, la quantité, l'accessibilité, la sécurité, la provenance et la compatibilité d'usage."),("Stocker","Identifier si un stockage temporaire, un retrait rapide ou une affectation directe est possible."),("Affecter","Réserver les matériaux à un projet validé par TVF, plutôt qu'à une distribution libre."),("Tracer","Conserver un registre : donateur, nature, volume, état, destination et preuve de remise.")],"documents":["Bordereau matériaux","Photos","Quantité","Dimensions","État","Lieu de retrait","Délai disponible","Propriété ou autorisation de céder"],"conventions":[("Convention de valorisation de matériaux","Cadrer l'origine, la responsabilité, l'état, la remise et l'affectation."),("Convention de stockage","Définir les conditions d'accueil temporaire, sécurité, assurance et durée."),("Procès-verbal de remise","Tracer le transfert physique et l'utilisation prévue des ressources.")],"limits":["TVF n'est pas une déchetterie, un site de petites annonces ou une plateforme de dons automatiques.","Les matériaux dangereux, non identifiés ou non réemployables peuvent être refusés.","L'acceptation dépend des besoins, de l'état, des capacités logistiques et du cadre juridique."],"benefits":["Réduction du gaspillage","Valorisation RSE","Soutien aux projets locaux","Traçabilité","Réduction des coûts d'évacuation lorsque le cadre le permet"],"cta":"Proposer des matériaux ou une solution de stockage : contact@territoiresvivantsfrance.fr"},
    {"slug":"commerce-vivant","name":"Commerce Vivant","subtitle":"Réactiver les locaux commerciaux vacants par des usages utiles, progressifs et adaptés au territoire.","public":"propriétaires de locaux, commerçants, artisans, collectivités, CCI, CMA, associations","photo":"saint-etienne-centre-commerce.jpg","caption":"Rue commerçante française : le pôle Commerce Vivant étudie les locaux vacants, les usages possibles et les acteurs à mobiliser.","why":["Un commerce fermé fragilise une rue entière : baisse de fréquentation, perte d'image, sentiment d'abandon, locaux dégradés et difficulté à attirer de nouveaux porteurs de projet.","La réactivation commerciale ne consiste pas seulement à rouvrir une vitrine. Il faut vérifier le bail, l'état du local, l'accessibilité, les besoins du quartier, le modèle économique, les travaux nécessaires et les partenaires économiques compétents.","Commerce Vivant positionne TVF comme un outil de préqualification : repérer, comprendre, orienter et préparer des scénarios d'usage compatibles avec les besoins locaux."],"repere":"Sources mobilisables : INSEE Base permanente des équipements, CCI, CMA, données communales, observations de terrain et diagnostics de centre-ville.","missions":[("Repérer les cellules","Identifier les vitrines fermées, locaux à louer durablement, rez-de-chaussée délaissés ou espaces commerciaux sous-utilisés."),("Analyser le besoin","Comprendre les manques d'activité : services de proximité, artisanat, économie sociale, commerce test, association, formation."),("Préparer un scénario","Boutique éphémère, atelier partagé, local associatif, activité artisanale, tiers-lieu de proximité ou commerce adapté."),("Mobiliser","Mettre autour de la table propriétaire, collectivité, acteur économique, porteur de projet et financeur potentiel."),("Suivre","Documenter l'état d'avancement, les freins, les décisions et les indicateurs d'activité.")],"documents":["Adresse du local","Photos de façade et intérieur","Surface","État apparent","Conditions d'accès","Situation bail/propriété","Usage souhaité","Contraintes techniques"],"conventions":[("Autorisation de visite","Observer le local et réunir les premières informations."),("Convention d'occupation temporaire","Tester un usage limité dans le temps lorsque le propriétaire et le cadre juridique l'autorisent."),("Convention de coopération économique","Organiser l'appui entre TVF, propriétaire, porteur de projet et acteur économique local.")],"limits":["TVF ne garantit pas la rentabilité commerciale d'un projet.","TVF ne remplace pas les chambres consulaires, experts-comptables, juristes, bailleurs ou services économiques.","Un local peut être classé sans suite si l'état, le coût ou le contexte rend l'usage irréaliste."],"benefits":["Meilleure lecture des besoins","Réduction de la vacance visible","Test d'usages","Lien avec porteurs de projet","Contribution à l'animation de quartier"],"cta":"Signaler un local vacant ou proposer un usage : contact@territoiresvivantsfrance.fr"},
    {"slug":"friches-terrains-vivants","name":"Friches & Terrains Vivants","subtitle":"Qualifier les espaces délaissés et préparer des usages sobres, utiles et compatibles avec le site.","public":"collectivités, propriétaires fonciers, aménageurs, associations, habitants","photo":"france-friche-pcuk.jpg","caption":"Friche ou terrain délaissé : avant toute transformation, TVF aide à clarifier propriété, accès, sécurité, usages possibles et limites.","why":["Les friches et terrains inutilisés peuvent concentrer plusieurs enjeux : sécurité, image urbaine, biodiversité, foncier, pollution potentielle, usages non autorisés, attentes d'habitants et contraintes réglementaires.","Tous les espaces abandonnés ne peuvent pas devenir immédiatement un jardin, un parc ou un équipement. Il faut d'abord qualifier le site : propriété, risques, accès, sol, environnement, voisinage, contraintes et temporalité.","Le pôle Friches & Terrains Vivants prépare cette lecture opérationnelle afin d'éviter les projets flous et de favoriser des usages proportionnés : renaturation, jardin partagé, espace associatif, pédagogie, occupation transitoire ou simple veille."],"repere":"Sources mobilisables : Cartofriches du Cerema, données foncières publiques lorsque accessibles, documents d'urbanisme, objectifs de sobriété foncière et ZAN.","missions":[("Identifier","Localiser les friches, dents creuses, terrains abandonnés ou espaces sous-utilisés signalés."),("Sécuriser la lecture","Vérifier les premiers risques : accès, clôture, pollution supposée, voisinage, statut, usages interdits."),("Définir les usages","Étudier des usages légers ou transitoires : jardin, biodiversité, pédagogie, rencontre, culture, stockage encadré."),("Orienter","Renvoyer vers les acteurs compétents lorsque le sujet relève du foncier, de l'aménagement, de la dépollution ou d'une maîtrise d'ouvrage publique."),("Documenter","Produire une fiche de site, une grille de risques, une note d'opportunité et un suivi des décisions.")],"documents":["Localisation cadastrale ou adresse","Photos","Propriétaire connu","Accès","État apparent","Risques visibles","Usage envisagé","Documents d'urbanisme disponibles"],"conventions":[("Autorisation de visite de site","Permettre l'observation sans travaux ni usage."),("Convention d'usage temporaire","Encadrer un usage léger et limité dans le temps."),("Convention de coopération territoriale","Associer collectivité, propriétaire, association et TVF autour d'un objectif commun.")],"limits":["TVF ne réalise pas de dépollution, diagnostic environnemental réglementaire ou maîtrise d'ouvrage d'aménagement.","Les sites dangereux ou juridiquement incertains doivent être orientés vers les professionnels compétents.","L'usage temporaire ne doit jamais masquer les obligations de sécurité, d'assurance et d'autorisation."],"benefits":["Requalification progressive","Meilleure connaissance du foncier délaissé","Usages temporaires utiles","Sobriété foncière","Cadre clair avant mobilisation citoyenne"],"cta":"Signaler une friche, un terrain ou un espace délaissé : contact@territoiresvivantsfrance.fr"},
    {"slug":"solidarite-insertion","name":"Solidarité & Insertion","subtitle":"Transformer les projets de revitalisation en occasions d'engagement, d'apprentissage et de remobilisation.","public":"bénévoles, associations, structures d'insertion, jeunes, habitants, entreprises engagées","photo":"france-ressourcerie-vichy.jpg","caption":"Ressourcerie et engagement local : le pôle Solidarité & Insertion relie utilité territoriale, bénévolat encadré et parcours de remobilisation.","why":["La revitalisation d'un territoire ne repose pas seulement sur des bâtiments ou des matériaux. Elle dépend aussi des personnes capables de participer, d'apprendre, de transmettre, de s'engager et de retrouver une place dans un projet concret.","Les chantiers participatifs, ateliers de tri, actions de sensibilisation, missions bénévoles ou parcours d'insertion doivent être encadrés. Sans cadre, l'engagement devient fragile ; avec une méthode, il peut devenir formateur et utile.","Le pôle Solidarité & Insertion vise à créer des missions réalistes autour des autres pôles : repérage, accueil, tri, logistique, petits travaux encadrés, médiation, animation, documentation et suivi."],"repere":"Sources mobilisables : INSEE pour le contexte emploi et pauvreté ; DARES et acteurs de l'insertion pour les cadres de parcours ; France Travail, missions locales, PLIE et structures locales compétentes.","missions":[("Accueillir","Recevoir les candidatures bénévoles, associations et structures d'accompagnement."),("Qualifier","Identifier disponibilités, compétences, besoins d'accompagnement, limites et cadre d'intervention."),("Orienter","Diriger vers une mission utile : tri, collecte, visite encadrée, administratif, communication, atelier, chantier léger."),("Encadrer","Formaliser consignes, sécurité, présence, référent, assurance, horaires et compte rendu."),("Valoriser","Documenter les heures, compétences mobilisées, apprentissages et contributions aux projets.")],"documents":["Fiche bénévole","Disponibilités","Compétences","Autorisation image si besoin","Feuille d'émargement","Consignes sécurité","Compte rendu mission"],"conventions":[("Charte bénévole","Définir le cadre de participation, les règles et les limites."),("Convention association ou insertion","Organiser l'intervention d'une structure autour d'une mission encadrée."),("Feuille de mission","Décrire précisément l'action, le lieu, la durée, le référent et les consignes.")],"limits":["TVF ne remplace pas les structures d'insertion, organismes de formation ou employeurs compétents.","Aucun chantier ne doit être ouvert sans cadre sécurité, assurance et référent identifié.","Les missions doivent être proportionnées et compatibles avec les capacités des participants."],"benefits":["Engagement citoyen concret","Découverte de métiers","Remobilisation","Lien social","Soutien aux projets de terrain"],"cta":"Proposer une mission, devenir bénévole ou coopérer avec TVF : contact@territoiresvivantsfrance.fr"},
]

def cover(doc, pole):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Cm(5.7))
        p_logo.paragraph_format.space_after = Pt(9)
    p = doc.add_paragraph()
    r = p.add_run("BROCHURE PÔLE TVF")
    set_run(r, 9.4, GOLD, bold=True)
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    r = p.add_run(pole["name"])
    set_run(r, 25, BLUE, bold=True)
    p.paragraph_format.space_after = Pt(5)
    p = doc.add_paragraph()
    r = p.add_run(pole["subtitle"])
    set_run(r, 12.6, GREEN, bold=True)
    p.paragraph_format.space_after = Pt(10)
    metadata(doc, pole)
    image_block(doc, ROOT / "assets" / "photos" / pole["photo"], pole["caption"])
    add_callout(doc, "Positionnement", "Cette brochure présente le rôle du pôle, les situations concernées, les documents à préparer et les cadres de coopération possibles. Elle ne vaut pas acceptation automatique d'une demande, partenariat officiel ou engagement de réalisation.", fill=LIGHT_GREEN)

def build_doc(pole):
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    cover(doc, pole)
    doc.add_page_break()
    section_intro(doc, "Comprendre", "Pourquoi ce pôle existe", pole["why"][0])
    for paragraph in pole["why"][1:]:
        add_paragraph(doc, paragraph)
    add_callout(doc, "Repère public", pole["repere"], fill=LIGHT_BLUE)
    section_intro(doc, "Action", "Ce que TVF peut faire", "Le pôle transforme une situation vague en dossier instruit : informations collectées, besoin clarifié, pièces demandées, partenaires identifiés, décision préparée.")
    table(doc, ["Mission", "Rôle opérationnel"], pole["missions"], [4.3, 12.6], fill=LIGHT_GREEN)
    section_intro(doc, "Parcours", "Déroulé type d'un dossier", "Le traitement d'une demande suit une logique progressive. Chaque étape peut aboutir à une poursuite, une demande de pièces, une orientation vers un acteur compétent ou un classement sans suite.")
    steps = [("1. Réception", "Formulaire, e-mail, appel, signalement mobile ou contact direct."),("2. Qualification", "Vérification du profil, du lieu, du besoin, des pièces et du niveau d'urgence."),("3. Instruction", "Analyse de faisabilité, risques, partenaires, documents et modalités possibles."),("4. Décision", "Poursuite, rendez-vous, visite, convention, orientation ou classement."),("5. Suivi", "Historique, pièces, relances, preuves, impacts et clôture du dossier.")]
    table(doc, ["Étape", "Objectif"], steps, [4.2, 12.7], fill=LIGHT_BLUE)
    section_intro(doc, "Pièces", "Documents à préparer", "Ces pièces permettent à TVF de comprendre la demande sans créer d'engagement immédiat. Elles servent à instruire le dossier, vérifier sa maturité et éviter les échanges incomplets.")
    for item in pole["documents"]:
        add_bullet(doc, item)
    note(doc, "TVF peut demander des pièces complémentaires selon le profil, la nature du bien, l'état apparent, le statut juridique, les contraintes de sécurité et le niveau d'avancement du dossier.")
    section_intro(doc, "Cadre", "Conventions et modalités possibles", "Les conventions servent à préciser les rôles, durées, responsabilités, limites, assurances, conditions de communication, suivi et modalités de fin d'action.")
    table(doc, ["Document", "Utilité"], pole["conventions"], [5.2, 11.7], fill=LIGHT_GREEN)
    section_intro(doc, "Bénéfices", "Ce que le pôle apporte aux acteurs", "L'utilité du pôle dépend de sa capacité à rendre les demandes lisibles, traçables et orientées vers une suite réaliste.")
    table(doc, ["Bénéfice", "Lecture concrète"], [(b, "Apport direct pour le porteur de demande, le territoire ou les partenaires mobilisés.") for b in pole["benefits"]], [5.2, 11.7], fill=LIGHT_BLUE)
    section_intro(doc, "Limites", "Points à verrouiller avant toute action", "TVF doit rester clair sur ce qu'elle peut faire et sur ce qui relève d'acteurs techniques, juridiques, financiers ou publics compétents.")
    for item in pole["limits"]:
        add_bullet(doc, item)
    add_callout(doc, "Contact", pole["cta"], fill=LIGHT_GREEN)
    doc.add_heading("Sources et références à mobiliser", level=1)
    for source in ["INSEE : données publiques locales, population, logements, emploi, équipements.", "Ministère de la Transition écologique : filière PMCB, économie circulaire, sobriété foncière.", "Cerema : Cartofriches et ressources méthodologiques territoriales.", "Collectivités locales : documents d'urbanisme, diagnostics, programmes et données de terrain lorsque disponibles."]:
        add_bullet(doc, source)
    note(doc, "Les sources servent de cadrage. Les chiffres, partenaires et résultats TVF ne doivent être publiés que lorsqu'ils sont vérifiés, documentés et validés.")
    return doc

def md_for(pole):
    lines = [f"# Brochure pôle - {pole['name']}", "", pole["subtitle"], "", "## Pourquoi ce pôle existe", ""]
    lines += pole["why"]
    lines += ["", "## Repère public", pole["repere"], "", "## Missions"]
    lines += [f"- **{a}** : {b}" for a, b in pole["missions"]]
    lines += ["", "## Documents à préparer"] + [f"- {x}" for x in pole["documents"]]
    lines += ["", "## Conventions possibles"] + [f"- **{a}** : {b}" for a, b in pole["conventions"]]
    lines += ["", "## Limites"] + [f"- {x}" for x in pole["limits"]]
    lines += ["", "## Contact", pole["cta"], ""]
    return "\n".join(lines)

def main():
    index_rows = []
    for pole in POLES:
        doc = build_doc(pole)
        docx_path = OUT_DIR / f"brochure-pole-{pole['slug']}-tvf.docx"
        doc.save(docx_path)
        md_path = SRC_DIR / f"brochure-pole-{pole['slug']}-tvf.md"
        md_path.write_text(md_for(pole), encoding="utf-8")
        index_rows.append((pole["name"], str(docx_path.relative_to(ROOT)), str(md_path.relative_to(ROOT))))
    readme = ["# Brochures des pôles TVF", "", "Documents prêts à utiliser pour présenter chaque pôle de Territoires Vivants France.", "", "| Pôle | DOCX | Source |", "|---|---|---|"]
    for name, docx, md in index_rows:
        readme.append(f"| {name} | `{docx}` | `{md}` |")
    readme.append("")
    (OUT_DIR / "README.md").write_text("\n".join(readme), encoding="utf-8")
    print("BROCHURES_POLES_CREATED")
    for name, docx, md in index_rows:
        print(f"- {name}: {docx}")

if __name__ == "__main__":
    main()
