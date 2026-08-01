from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

ROOT = Path(r'C:\Users\jowst\Documents\TERRITOIRES VIVANTS FRANCE')
OUT = ROOT / 'documents' / 'formulaires-internes-tvf'
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / 'fiche-interne-identification-proprietaire-tvf.docx'
PDF = OUT / 'fiche-interne-identification-proprietaire-tvf.pdf'
LOGO = ROOT / 'assets' / 'logo-territoires-vivants-france-web.png'

GREEN = RGBColor(24,57,47)
GRAY = RGBColor(102,112,133)
DARK = RGBColor(31,41,55)
PALE = 'F3F7F2'
BORDER = 'D9E0DC'

TVF = {
    'name':'Territoires Vivants France',
    'subtitle':'Agence Territoriale de Revitalisation Immobilière',
    'address':'25 rue Élise Gervais, 42000 Saint-Étienne',
    'phone':'04 65 81 54 69',
    'email':'contact@territoiresvivantsfrance.fr',
    'site':'www.territoiresvivantsfrance.fr',
    'rna':'RNA W922015538',
    'siret':'SIRET 897 226 138 00018',
}


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right'):
        tag = 'w:' + edge
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), BORDER)


def set_cell(cell, text, bold=False, color=DARK, size=8.8, fill=None):
    cell.text = ''
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Inter'
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(cell)


def section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = 'Manrope'
    r.font.size = Pt(11.5)
    r.bold = True
    r.font.color.rgb = GREEN


def label_line(doc, label, width_label=2.15, width_value=4.6):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(width_label)
    t.columns[1].width = Inches(width_value)
    set_cell(t.rows[0].cells[0], label, True, GREEN, 8.6, PALE)
    set_cell(t.rows[0].cells[1], '', False, DARK, 8.6)


def checkbox_grid(doc, items, cols=2):
    rows = (len(items) + cols - 1) // cols
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for col in t.columns:
        col.width = Inches(6.75 / cols)
    idx = 0
    for r in range(rows):
        for c in range(cols):
            text = '☐ ' + items[idx] if idx < len(items) else ''
            set_cell(t.rows[r].cells[c], text, False, DARK, 8.3)
            idx += 1


def add_note(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell(cell, text, False, GRAY, 7.8, 'FBF8EF')


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.42)
    sec.bottom_margin = Inches(0.42)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.columns[0].width = Inches(1.25)
    header.columns[1].width = Inches(5.5)
    if LOGO.exists():
        run = header.rows[0].cells[0].paragraphs[0].add_run()
        run.add_picture(str(LOGO), width=Inches(1.0))
    txt = header.rows[0].cells[1]
    txt.text = ''
    p = txt.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(TVF['name'] + '\n')
    r.font.name = 'Manrope'; r.font.size = Pt(12); r.bold = True; r.font.color.rgb = GREEN
    r = p.add_run(TVF['subtitle'] + '\n')
    r.font.name = 'Inter'; r.font.size = Pt(8.5); r.font.color.rgb = DARK
    r = p.add_run(f"{TVF['address']} | {TVF['phone']} | {TVF['email']}\n{TVF['rna']} | {TVF['siret']}")
    r.font.name = 'Inter'; r.font.size = Pt(7.2); r.font.color.rgb = GRAY
    for row in header.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run('FICHE INTERNE TVF — IDENTIFICATION DU PROPRIÉTAIRE')
    r.font.name = 'Manrope'; r.font.size = Pt(15); r.bold = True; r.font.color.rgb = GREEN
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run('Document de repérage et de vérification préalable — usage interne')
    r.font.name = 'Inter'; r.font.size = Pt(8.3); r.font.color.rgb = GRAY

    add_note(doc, "Cette fiche sert uniquement à organiser l’identification d’un propriétaire ou titulaire de droits. Elle ne confirme pas à elle seule la propriété, ne vaut pas décision d’accompagnement et doit être utilisée dans le respect du droit de propriété et de la protection des données personnelles.")

    section(doc, '1. Référence interne')
    label_line(doc, 'Numéro fiche propriétaire')
    label_line(doc, 'Date de création')
    label_line(doc, 'Agent / bénévole TVF')
    label_line(doc, 'Numéro fiche bien liée')

    section(doc, '2. Bien concerné')
    label_line(doc, 'Adresse du bien')
    label_line(doc, 'Commune')
    label_line(doc, 'Code postal')
    label_line(doc, 'Référence cadastrale si connue')

    section(doc, '3. Propriétaire identifié ou supposé')
    checkbox_grid(doc, ['Personne physique', 'Indivision', 'Succession possible', 'SCI', 'Société', 'Collectivité', 'Bailleur / organisme', 'Propriétaire inconnu'], 2)
    label_line(doc, 'Nom / prénom ou raison sociale')
    label_line(doc, 'Qualité connue')

    section(doc, '4. Coordonnées disponibles')
    label_line(doc, 'Adresse postale')
    label_line(doc, 'Téléphone')
    label_line(doc, 'E-mail')
    label_line(doc, 'Autre coordonnée utile')

    section(doc, '5. Source de l’information')
    checkbox_grid(doc, ['Propriétaire déclarant', 'Document transmis', 'Cadastre / donnée accessible', 'Collectivité', 'Notaire', 'Syndic / gestionnaire', 'Voisinage / tiers', 'Courrier reçu', 'E-mail', 'Autre : ____________________'], 2)

    section(doc, '6. Niveau de vérification')
    checkbox_grid(doc, ['Confirmé par document', 'Confirmé par propriétaire', 'Supposé', 'À vérifier', 'Information incomplète', 'Donnée contradictoire'], 2)

    section(doc, '7. Représentant ou interlocuteur')
    checkbox_grid(doc, ['Notaire', 'Mandataire', 'Agence / gestionnaire', 'Membre de famille', 'Syndic', 'Représentant légal', 'Interlocuteur collectivité', 'Aucun identifié'], 2)
    label_line(doc, 'Nom de l’interlocuteur')
    label_line(doc, 'Coordonnées')

    section(doc, '8. Démarches effectuées')
    checkbox_grid(doc, ['Recherche cadastrale', 'Courrier préparé', 'Courrier envoyé', 'Appel effectué', 'E-mail envoyé', 'Rendez-vous proposé', 'Retour reçu', 'Absence de réponse'], 2)
    label_line(doc, 'Date de dernière démarche')

    section(doc, '9. Cadre et précautions')
    checkbox_grid(doc, ['Donnée confidentielle', 'Ne pas diffuser', 'Accord de contact obtenu', 'Accord à vérifier', 'Identité à confirmer', 'Besoin validation responsable'], 2)
    add_note(doc, "Les informations personnelles ne doivent être consultées que par les personnes habilitées et uniquement pour le traitement de la demande ou du dossier TVF concerné.")

    section(doc, '10. Suite à donner')
    checkbox_grid(doc, ['Contacter le propriétaire', 'Envoyer un courrier', 'Demander un document', 'Rattacher au dossier', 'Créer / compléter le contact', 'Vérifier l’identité', 'Transmettre au responsable', 'Classer sans suite'], 2)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run('Territoires Vivants France — document interne confidentiel — identification propriétaire')
    rf.font.name = 'Inter'; rf.font.size = Pt(7); rf.font.color.rgb = GRAY

    doc.save(DOCX)


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='TitleTVF', fontName='Helvetica-Bold', fontSize=14, leading=17, textColor=colors.HexColor('#18392F'), alignment=1, spaceAfter=3))
    styles.add(ParagraphStyle(name='Sub', fontName='Helvetica', fontSize=8, leading=10, textColor=colors.HexColor('#667085'), alignment=1, spaceAfter=4))
    styles.add(ParagraphStyle(name='H', fontName='Helvetica-Bold', fontSize=10.5, leading=13, textColor=colors.HexColor('#18392F'), spaceBefore=6, spaceAfter=3))
    styles.add(ParagraphStyle(name='Small', fontName='Helvetica', fontSize=7.2, leading=9, textColor=colors.HexColor('#667085'), spaceAfter=4))
    styles.add(ParagraphStyle(name='Note', fontName='Helvetica', fontSize=7.6, leading=10, textColor=colors.HexColor('#475467'), backColor=colors.HexColor('#FBF8EF'), borderColor=colors.HexColor('#E4D7B8'), borderWidth=0.5, borderPadding=5, spaceAfter=5))

    story = []
    header_data = []
    if LOGO.exists():
        from reportlab.platypus import Image
        logo = Image(str(LOGO), width=24*mm, height=18*mm)
        header_data.append([logo, Paragraph(f"<b>{TVF['name']}</b><br/>{TVF['subtitle']}<br/>{TVF['address']} | {TVF['phone']} | {TVF['email']}<br/>{TVF['rna']} | {TVF['siret']}", styles['Small'])])
    else:
        header_data.append(['', Paragraph(f"<b>{TVF['name']}</b><br/>{TVF['subtitle']}<br/>{TVF['address']} | {TVF['phone']} | {TVF['email']}<br/>{TVF['rna']} | {TVF['siret']}", styles['Small'])])
    ht = Table(header_data, colWidths=[30*mm, 140*mm])
    ht.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'MIDDLE'),('ALIGN',(1,0),(1,0),'RIGHT')]))
    story.append(ht)
    story.append(Paragraph('FICHE INTERNE TVF — IDENTIFICATION DU PROPRIÉTAIRE', styles['TitleTVF']))
    story.append(Paragraph('Document de repérage et de vérification préalable — usage interne', styles['Sub']))
    story.append(Paragraph('Cette fiche sert uniquement à organiser l’identification d’un propriétaire ou titulaire de droits. Elle ne confirme pas à elle seule la propriété, ne vaut pas décision d’accompagnement et doit être utilisée dans le respect du droit de propriété et de la protection des données personnelles.', styles['Note']))

    def pdf_table(rows, widths=[56*mm, 114*mm]):
        t = Table(rows, colWidths=widths)
        t.setStyle(TableStyle([
            ('GRID',(0,0),(-1,-1),0.45,colors.HexColor('#D9E0DC')),
            ('BACKGROUND',(0,0),(0,-1),colors.HexColor('#F3F7F2')),
            ('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),
            ('TEXTCOLOR',(0,0),(0,-1),colors.HexColor('#18392F')),
            ('FONTNAME',(1,0),(1,-1),'Helvetica'),
            ('FONTSIZE',(0,0),(-1,-1),7.4),
            ('LEADING',(0,0),(-1,-1),9),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('TOPPADDING',(0,0),(-1,-1),4),
            ('BOTTOMPADDING',(0,0),(-1,-1),4),
            ('LEFTPADDING',(0,0),(-1,-1),5),
            ('RIGHTPADDING',(0,0),(-1,-1),5),
        ]))
        story.append(t)

    def pdf_checks(title, items, cols=2):
        story.append(Paragraph(title, styles['H']))
        rows=[]
        for i in range(0, len(items), cols):
            rows.append(['□ ' + x for x in items[i:i+cols]] + ['']*(cols-len(items[i:i+cols])))
        t = Table(rows, colWidths=[170*mm/cols]*cols)
        t.setStyle(TableStyle([
            ('GRID',(0,0),(-1,-1),0.45,colors.HexColor('#D9E0DC')),
            ('FONTNAME',(0,0),(-1,-1),'Helvetica'),
            ('FONTSIZE',(0,0),(-1,-1),7.3),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('TOPPADDING',(0,0),(-1,-1),4),
            ('BOTTOMPADDING',(0,0),(-1,-1),4),
            ('LEFTPADDING',(0,0),(-1,-1),5),
        ]))
        story.append(t)

    for title, labels in [
        ('1. Référence interne', ['Numéro fiche propriétaire','Date de création','Agent / bénévole TVF','Numéro fiche bien liée']),
        ('2. Bien concerné', ['Adresse du bien','Commune','Code postal','Référence cadastrale si connue'])
    ]:
        story.append(Paragraph(title, styles['H']))
        pdf_table([[l,''] for l in labels])

    pdf_checks('3. Propriétaire identifié ou supposé', ['Personne physique','Indivision','Succession possible','SCI','Société','Collectivité','Bailleur / organisme','Propriétaire inconnu'], 2)
    story.append(Paragraph('Coordonnées d’identification', styles['H']))
    pdf_table([['Nom / prénom ou raison sociale',''], ['Qualité connue','']])
    story.append(Paragraph('4. Coordonnées disponibles', styles['H']))
    pdf_table([['Adresse postale',''], ['Téléphone',''], ['E-mail',''], ['Autre coordonnée utile','']])
    pdf_checks('5. Source de l’information', ['Propriétaire déclarant','Document transmis','Cadastre / donnée accessible','Collectivité','Notaire','Syndic / gestionnaire','Voisinage / tiers','Courrier reçu','E-mail','Autre : ____________________'], 2)
    pdf_checks('6. Niveau de vérification', ['Confirmé par document','Confirmé par propriétaire','Supposé','À vérifier','Information incomplète','Donnée contradictoire'], 2)
    pdf_checks('7. Représentant ou interlocuteur', ['Notaire','Mandataire','Agence / gestionnaire','Membre de famille','Syndic','Représentant légal','Interlocuteur collectivité','Aucun identifié'], 2)
    pdf_table([['Nom de l’interlocuteur',''], ['Coordonnées','']])
    pdf_checks('8. Démarches effectuées', ['Recherche cadastrale','Courrier préparé','Courrier envoyé','Appel effectué','E-mail envoyé','Rendez-vous proposé','Retour reçu','Absence de réponse'], 2)
    pdf_table([['Date de dernière démarche','']])
    pdf_checks('9. Cadre et précautions', ['Donnée confidentielle','Ne pas diffuser','Accord de contact obtenu','Accord à vérifier','Identité à confirmer','Besoin validation responsable'], 2)
    story.append(Paragraph('Les informations personnelles ne doivent être consultées que par les personnes habilitées et uniquement pour le traitement de la demande ou du dossier TVF concerné.', styles['Note']))
    pdf_checks('10. Suite à donner', ['Contacter le propriétaire','Envoyer un courrier','Demander un document','Rattacher au dossier','Créer / compléter le contact','Vérifier l’identité','Transmettre au responsable','Classer sans suite'], 2)

    doc = SimpleDocTemplate(str(PDF), pagesize=A4, rightMargin=13*mm, leftMargin=13*mm, topMargin=10*mm, bottomMargin=10*mm)
    doc.build(story)

build_docx()
build_pdf()
print(DOCX)
print(PDF)
