from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image

ROOT = Path(r'C:\Users\jowst\Documents\TERRITOIRES VIVANTS FRANCE')
OUT = ROOT / 'documents' / 'brochures-tvf'
DOCX = OUT / 'catalogue-conventions-cadres-tvf.docx'
PDF = OUT / 'catalogue-conventions-cadres-tvf.pdf'
LOGO = ROOT / 'assets' / 'logo-territoires-vivants-france-web.png'
GREEN = RGBColor(24,57,47); DARK = RGBColor(31,41,55); GRAY = RGBColor(102,112,133)

DETAILS = [
('Convention cadre propriétaire',
 'À utiliser dès qu’un propriétaire accepte que TVF étudie un bien, échange sur sa situation ou prépare une orientation. Elle est le socle de confiance entre TVF et le détenteur du bien.',
 ['Le propriétaire reste propriétaire et conserve seul la décision finale sur son bien.', 'TVF peut étudier, orienter et accompagner, mais ne devient pas gestionnaire du bien.', 'L’autorisation de visite, les photos et la transmission des documents doivent être intégrées comme clauses ou annexes.', 'La convention doit contenir une clause d’absence de mandat immobilier tant que TVF n’intervient pas dans un cadre professionnel réglementé adapté.'],
 ['Ne jamais promettre une aide, un financement, un locataire, une vente, une mise en location ou un rendement automatique.', 'Ne jamais laisser penser que TVF peut entrer dans un bien sans autorisation claire.']),
('Convention de mise à disposition d’un bien',
 'À utiliser lorsqu’un bien est réellement confié pour un usage déterminé : stockage, action associative, occupation temporaire, atelier, expérimentation ou projet territorial.',
 ['Le propriétaire conserve la propriété du bien.', 'L’usage autorisé doit être décrit précisément : durée, bénéficiaire, accès, horaires, zones, interdictions.', 'La convention doit préciser les charges, fluides, entretien courant, assurance, responsabilité, sécurité et restitution.', 'Si une participation financière, indemnité, redevance ou rendement est envisagé, sa nature juridique doit être vérifiée avant signature.'],
 ['Ne pas transformer une mise à disposition en bail ou gestion locative sans analyse juridique.', 'Ne pas promettre de rendement si la convention ne définit pas clairement la contrepartie et le cadre légal.']),
('Convention cadre collectivité',
 'À utiliser avec une commune, intercommunalité, département ou établissement public pour organiser une coopération territoriale.',
 ['TVF agit comme appui de repérage, qualification, coordination, observatoire et suivi.', 'La collectivité conserve ses compétences, ses décisions administratives et ses pouvoirs propres.', 'Le partage de données doit être encadré : finalité, accès, confidentialité, durée, sécurité.', 'Les résultats publics doivent être vérifiés, anonymisés ou agrégés lorsque les données concernent des biens privés.'],
 ['Ne pas présenter TVF comme autorité publique.', 'Ne pas affirmer juridiquement qu’un bien est vacant, abandonné ou sans propriétaire sur simple signalement.']),
('Convention cadre entreprise / partenaire',
 'À utiliser pour les entreprises, fondations, artisans, experts et partenaires économiques qui veulent contribuer par expertise, moyens, financement, matériaux ou action territoriale.',
 ['Distinguer don, mécénat, prestation, parrainage, mise à disposition de compétences et partenariat institutionnel.', 'Chaque action doit préciser livrables, responsabilités, assurances, coûts éventuels et communication.', 'L’usage du nom ou du logo du partenaire nécessite un accord.', 'Le mécénat doit rester sans contrepartie équivalente.'],
 ['Ne pas mélanger mécénat et sponsoring.', 'Ne pas annoncer une réduction fiscale sans vérification des conditions légales.']),
('Convention ressources et matériaux',
 'À utiliser pour la collecte, le stockage, le contrôle, le réemploi et l’affectation des matériaux, mobiliers, équipements ou stocks inutilisés.',
 ['TVF conserve un droit de refus : ressource dangereuse, trop dégradée, non conforme, inutile, impossible à transporter ou stocker.', 'Chaque ressource acceptée doit être inventoriée : quantité, état, origine, photos, lieu, affectation.', 'Le transport, la manutention, le stockage et la responsabilité doivent être définis.', 'Un bon de remise ou une fiche de sortie permet de tracer l’usage.'],
 ['Ne pas distribuer librement les ressources sans cadre.', 'Ne pas délivrer de reçu fiscal si les conditions du mécénat ne sont pas vérifiées.']),
('Convention projet solidaire / association',
 'À utiliser avec une association, structure d’insertion, organisme de formation ou porteur de projet pour une action utile au territoire.',
 ['Décrire le projet, les publics, l’objectif, le lieu, les participants, les encadrants et le calendrier.', 'Vérifier assurances, sécurité, encadrement, outillage, accueil du public et autorisations.', 'Si le projet utilise un bien, prévoir une mise à disposition ou autorisation séparée.', 'Prévoir une restitution : bilan, photos, indicateurs qualitatifs, retour d’expérience.'],
 ['Ne pas confondre soutien au projet et responsabilité directe de TVF sur tous les participants.', 'Ne pas organiser de chantier sans encadrement et assurance adaptés.']),
('Convention confidentialité et données',
 'À utiliser seule ou en annexe chaque fois qu’un partenaire accède à des informations sensibles, données personnelles, dossiers, photos, adresses ou documents.',
 ['Définir les informations couvertes : propriétaires, adresses, photos, documents, données patrimoniales, devis, échanges.', 'Limiter l’accès aux personnes habilitées et au besoin d’en connaître.', 'Préciser finalité, durée de conservation, sécurité, restitution et suppression.', 'Prévoir une procédure en cas de perte, erreur d’envoi ou accès non autorisé.'],
 ['Ne pas transmettre d’adresse précise, photo ou document sensible sans nécessité.', 'Ne pas réutiliser les données pour une finalité différente de la mission TVF.'])
]

UNIVERSAL_ARTICLES = [
('Clause d’objet', 'Décrit précisément la finalité de la convention et évite les interprétations trop larges.'),
('Clause de capacité', 'Vérifie que le signataire peut réellement engager le propriétaire, l’entreprise, la collectivité ou l’association.'),
('Clause d’absence de mandat immobilier', 'Protège TVF contre toute confusion avec transaction, location, gestion locative ou administration de biens.'),
('Clause de visite', 'Encadre accès, présence, sécurité, zones visitées, photos, report ou refus d’intervention.'),
('Clause de mise à disposition', 'À intégrer seulement si un usage réel du bien est accordé ; elle doit être très précise.'),
('Clause de rendement ou participation', 'À manier prudemment : participation aux charges, indemnité, redevance ou rendement doivent être qualifiés juridiquement.'),
('Clause matériaux', 'Précise inventaire, état, transport, stockage, refus, transfert de propriété et affectation.'),
('Clause fiscale', 'Rappelle qu’un avantage fiscal n’est jamais automatique et dépend de conditions légales.'),
('Clause RGPD', 'Indique finalité, minimisation, durée, sécurité, droits des personnes et destinataires.'),
('Clause de restitution', 'Organise la fin : restitution du bien, documents, clés, données, matériaux ou accès.'),
]

def r(p,text,size=10,bold=False,color=DARK,font='Inter'):
    run=p.add_run(text); run.font.name=font; run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=color

def add_docx():
    doc = Document(DOCX)
    doc.add_page_break()
    p=doc.add_paragraph(); r(p,'14. Fiches détaillées par convention',17,True,GREEN,'Manrope')
    p=doc.add_paragraph(); r(p,'Cette partie sert de guide opérationnel pour choisir, adapter et assembler les conventions cadres selon chaque situation concrète.',10,False,DARK)
    for title, intro, clauses, avoids in DETAILS:
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); r(p,title,13.5,True,GREEN,'Manrope')
        p=doc.add_paragraph(); r(p,intro,9.8,False,DARK)
        p=doc.add_paragraph(); r(p,'Clauses et règles à prévoir',10,True,GREEN)
        for c in clauses:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); r(p,'• '+c,9.4,False,DARK)
        p=doc.add_paragraph(); r(p,'Points à éviter',10,True,GREEN)
        for a in avoids:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); r(p,'• '+a,9.4,False,DARK)
    doc.add_page_break()
    p=doc.add_paragraph(); r(p,'15. Articles universels à intégrer dans les conventions',17,True,GREEN,'Manrope')
    table=doc.add_table(rows=1, cols=2); table.style='Table Grid'
    table.rows[0].cells[0].text='Article universel'; table.rows[0].cells[1].text='Utilité'
    for a,b in UNIVERSAL_ARTICLES:
        row=table.add_row().cells; row[0].text=a; row[1].text=b
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); r(p,'Lecture pratique',12,True,GREEN,'Manrope')
    p=doc.add_paragraph(); r(p,'Les conventions cadres doivent rester lisibles. Les clauses universelles sont reprises dans les modèles principaux, puis renforcées par des annexes lorsque la situation devient plus sensible : mise à disposition réelle d’un bien, échange de données personnelles, collecte de matériaux, intervention d’une entreprise ou présence de publics accompagnés.',9.8,False,DARK)
    doc.save(DOCX)

def styles():
    st=getSampleStyleSheet(); st.add(ParagraphStyle(name='T',fontName='Helvetica-Bold',fontSize=20,leading=24,textColor=colors.HexColor('#18392F'),alignment=1)); st.add(ParagraphStyle(name='Sub',fontSize=10.5,leading=14,textColor=colors.HexColor('#667085'),alignment=1,spaceAfter=8)); st.add(ParagraphStyle(name='H1',fontName='Helvetica-Bold',fontSize=14.5,leading=18,textColor=colors.HexColor('#18392F'),spaceBefore=9,spaceAfter=5)); st.add(ParagraphStyle(name='H2',fontName='Helvetica-Bold',fontSize=11.5,leading=14,textColor=colors.HexColor('#18392F'),spaceBefore=7,spaceAfter=3)); st.add(ParagraphStyle(name='B',fontSize=9,leading=12,textColor=colors.HexColor('#1F2937'),spaceAfter=4)); st.add(ParagraphStyle(name='Small',fontSize=7.4,leading=9.2,textColor=colors.HexColor('#667085'))); st.add(ParagraphStyle(name='Note',fontSize=8.1,leading=10.6,textColor=colors.HexColor('#475467'),backColor=colors.HexColor('#FBF8EF'),borderColor=colors.HexColor('#E4D7B8'),borderWidth=.5,borderPadding=6,spaceAfter=7)); return st

def add_pdf():
    st=styles(); story=[]
    if LOGO.exists(): story.append(Image(str(LOGO),width=32*mm,height=24*mm))
    story.append(Spacer(1,6*mm)); story.append(Paragraph('CATALOGUE DES CONVENTIONS CADRES TVF',st['T'])); story.append(Paragraph('Guide interne d’utilisation, clauses universelles et règles de prudence',st['Sub']))
    story.append(Paragraph('Cette brochure présente les grandes conventions cadres TVF, leur usage, leurs clauses structurantes, les articles universels à intégrer et les limites à respecter pour sécuriser propriétaires, collectivités, entreprises, associations et partenaires.',st['Note']))
    story.append(PageBreak())
    story.append(Paragraph('Les 7 conventions cadres',st['H1']))
    for title, intro, clauses, avoids in DETAILS:
        story.append(Paragraph(title,st['H2'])); story.append(Paragraph(intro,st['B']))
        story.append(Paragraph('<b>Clauses et règles à prévoir</b>',st['B']))
        for c in clauses: story.append(Paragraph('• '+c,st['B']))
        story.append(Paragraph('<b>Points à éviter</b>',st['B']))
        for a in avoids: story.append(Paragraph('• '+a,st['B']))
    story.append(PageBreak())
    story.append(Paragraph('Articles universels à intégrer',st['H1']))
    rows=[[Paragraph('Article universel',st['Small']),Paragraph('Utilité',st['Small'])]]
    for a,b in UNIVERSAL_ARTICLES: rows.append([Paragraph(a,st['Small']),Paragraph(b,st['Small'])])
    table=Table(rows,colWidths=[52*mm,117*mm],repeatRows=1); table.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.35,colors.HexColor('#D9E0DC')),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#F3F7F2')),('VALIGN',(0,0),(-1,-1),'TOP'),('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4)])); story.append(table)
    story.append(Spacer(1,5*mm)); story.append(Paragraph('Références de prudence',st['H1']))
    for ref in ['Service-Public Entreprendre : mécénat d’entreprise, dons, formes de don et conditions de réduction d’impôt.', 'CNIL : principes RGPD de finalité, minimisation, transparence, droits, durée de conservation et sécurité.', 'Service-Public : prudence sur les activités immobilières réglementées et la nécessité d’un cadre adapté.']:
        story.append(Paragraph('• '+ref,st['Small']))
    SimpleDocTemplate(str(PDF),pagesize=A4,rightMargin=14*mm,leftMargin=14*mm,topMargin=13*mm,bottomMargin=13*mm).build(story)

add_docx(); add_pdf(); print(DOCX); print(PDF)
