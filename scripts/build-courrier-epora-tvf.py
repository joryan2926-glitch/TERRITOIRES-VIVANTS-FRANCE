from pathlib import Path
from runpy import run_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

OUT_DIR = ROOT / "documents" / "courriers-lancement-saint-etienne"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "25-epora-lancement-cooperation-fonciere-tvf.docx"
MD_PATH = SRC_DIR / "courrier-epora-lancement-tvf.md"

txt = BASE["txt"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_number = BASE["add_number"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
BORDER = BASE["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        (
            "Destinataire",
            "EPORA - &Eacute;tablissement public foncier de l'Ouest Rh&ocirc;ne-Alpes - &agrave; l'attention de la Direction territoriale Loire et de la Direction g&eacute;n&eacute;rale",
        ),
        (
            "Adresse",
            "Si&egrave;ge social et direction territoriale Loire - Saint-&Eacute;tienne - 04 77 47 47 50",
        ),
        (
            "Objet",
            "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change technique sur les compl&eacute;mentarit&eacute;s fonci&egrave;res, les friches, les biens vacants et la valorisation des ressources inutilis&eacute;es &agrave; Saint-&Eacute;tienne",
        ),
        (
            "Pi&egrave;ces jointes",
            "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche m&eacute;thode d'instruction, fiche friches / biens vacants, fiche mat&eacute;riaux et r&eacute;emploi",
        ),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def add_axes_table(doc):
    doc.add_heading(txt("Compl&eacute;mentarit&eacute;s possibles &agrave; examiner"), level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.7, 6.2, 6.2])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Sujet", "Apport possible de TVF", "Point &agrave; cadrer avec EPORA"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(head))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.1)

    rows = [
        (
            "Observation de terrain",
            "Rep&eacute;rage citoyen et associatif de biens vacants, friches, locaux d&eacute;laiss&eacute;s ou ressources inutilis&eacute;es.",
            "Identifier les limites de la remont&eacute;e d'information et les modalit&eacute;s de transmission utiles.",
        ),
        (
            "Pr&eacute;qualification d'usages",
            "Lecture des besoins locaux : associations, artisans, services aux habitants, activit&eacute;s de proximit&eacute;, usages transitoires.",
            "V&eacute;rifier les cas o&ugrave; cette contribution peut compl&eacute;ter les &eacute;tudes et strat&eacute;gies fonci&egrave;res existantes.",
        ),
        (
            "R&eacute;emploi des mat&eacute;riaux",
            "Identification de mat&eacute;riaux, mobiliers ou &eacute;quipements potentiellement valorisables dans des projets valid&eacute;s.",
            "Respecter les cadres techniques, assurantiels, sanitaires, s&eacute;curitaires et juridiques des op&eacute;rations.",
        ),
        (
            "Mobilisation locale",
            "Mise en relation progressive avec propri&eacute;taires, entreprises, associations, b&eacute;n&eacute;voles et structures d'insertion.",
            "D&eacute;finir le bon niveau d'intervention de TVF sans empi&eacute;ter sur les missions fonci&egrave;res d'EPORA.",
        ),
        (
            "Suivi d'impact",
            "Suivi de l'utilit&eacute; territoriale : usage retrouv&eacute;, ressources valoris&eacute;es, acteurs mobilis&eacute;s, besoins couverts.",
            "Convenir des indicateurs partageables, des limites de communication et du calendrier de suivi.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.55)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - EPORA"))
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Demande d'un premier &eacute;change technique sur les compl&eacute;mentarit&eacute;s fonci&egrave;res et territoriales"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(
        doc,
        "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un ancrage op&eacute;rationnel progressif sur son territoire d'origine.",
    )
    add_paragraph(
        doc,
        "Notre objet est de contribuer &agrave; remettre en usage des ressources aujourd'hui insuffisamment mobilis&eacute;es : logements vacants ou d&eacute;grad&eacute;s, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains inutilis&eacute;s, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. La m&eacute;thode TVF repose sur une logique de terrain : observer, qualifier, documenter, mettre en relation, conventionner lorsque le cadre le permet, puis suivre l'utilit&eacute; produite.",
    )
    add_paragraph(
        doc,
        "EPORA occupe une place particuli&egrave;re dans cet &eacute;cosyst&egrave;me. En tant qu'&eacute;tablissement public foncier d'&Eacute;tat intervenant au coeur de la r&eacute;gion Auvergne-Rh&ocirc;ne-Alpes, EPORA dispose d'une expertise fonci&egrave;re, patrimoniale, juridique, op&eacute;rationnelle et de coordination particuli&egrave;rement importante pour les territoires confront&eacute;s aux friches, aux mutations urbaines, au foncier &eacute;conomique et &agrave; la sobri&eacute;t&eacute; fonci&egrave;re.",
    )

    add_callout(
        doc,
        "Demande principale",
        "TVF sollicite un premier rendez-vous technique avec EPORA afin de pr&eacute;senter sa m&eacute;thode, de comprendre les cadres d'intervention d'un &eacute;tablissement public foncier et d'identifier les sujets sur lesquels TVF pourrait, le cas &eacute;ch&eacute;ant, compl&eacute;ter les dispositifs existants sans se substituer aux missions d'EPORA.",
    )

    doc.add_heading(txt("Pourquoi solliciter EPORA"), level=1)
    add_paragraph(
        doc,
        "Les champs d'intervention de TVF rejoignent plusieurs enjeux trait&eacute;s par les acteurs fonciers : recyclage de biens d&eacute;laiss&eacute;s, transformation de sites sous-utilis&eacute;s, r&eacute;activation de foncier existant, adaptation des territoires aux mutations &eacute;conomiques et climatiques, ma&icirc;trise de l'artificialisation et recherche d'usages utiles apr&egrave;s intervention.",
    )
    add_paragraph(
        doc,
        "TVF ne revendique aucune comp&eacute;tence de portage foncier, d'acquisition, de pr&eacute;emption, de d&eacute;pollution, de d&eacute;molition ou de ma&icirc;trise d'ouvrage fonci&egrave;re. Son positionnement est diff&eacute;rent : faire remonter des besoins, qualifier des situations, mobiliser des ressources locales, structurer des dossiers pr&eacute;alables, orienter les propri&eacute;taires ou porteurs de projets et faciliter l'&eacute;mergence d'usages utiles aux habitants.",
    )
    add_paragraph(
        doc,
        "Cette compl&eacute;mentarit&eacute; pourrait &ecirc;tre utile &agrave; Saint-&Eacute;tienne et dans la Loire, notamment lorsque des biens, sites ou ressources ne rel&egrave;vent pas imm&eacute;diatement d'une op&eacute;ration fonci&egrave;re lourde, mais pourraient &ecirc;tre mieux document&eacute;s, orient&eacute;s, valoris&eacute;s ou mis en relation avec des acteurs locaux.",
    )

    add_axes_table(doc)

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "Organiser un rendez-vous d'environ une heure avec la direction territoriale Loire ou le service comp&eacute;tent afin de pr&eacute;senter TVF et de comprendre les cadres d'intervention d'EPORA.")
    add_bullet(doc, "Identifier les situations dans lesquelles une association comme TVF peut utilement contribuer &agrave; l'observation, &agrave; la remont&eacute;e de besoins, &agrave; la pr&eacute;qualification d'usages ou &agrave; la mobilisation locale.")
    add_bullet(doc, "Pr&eacute;ciser les limites &agrave; respecter : confidentialit&eacute;, propri&eacute;t&eacute; des donn&eacute;es, statut des biens, proc&eacute;dures fonci&egrave;res, responsabilit&eacute;s, assurances, communication et validation pr&eacute;alable.")
    add_bullet(doc, "Examiner la possibilit&eacute; d'identifier un premier cas de travail ou une premi&egrave;re m&eacute;thode de remont&eacute;e d'information sur Saint-&Eacute;tienne, sans engagement public avant accord &eacute;crit.")
    add_bullet(doc, "Orienter TVF, si cela est pertinent, vers les services ou partenaires avec lesquels un cadre d'&eacute;change plus pr&eacute;cis pourrait &ecirc;tre construit.")

    doc.add_heading(txt("Ce que TVF peut apporter"), level=1)
    add_number(doc, "Une capacit&eacute; de rep&eacute;rage de terrain : habitants, associations, propri&eacute;taires, artisans et entreprises peuvent signaler des situations qui m&eacute;ritent une qualification.")
    add_number(doc, "Une lecture des usages : TVF peut aider &agrave; identifier les besoins concrets du territoire avant ou apr&egrave;s une intervention : local associatif, atelier, tiers-lieu, commerce utile, stockage, formation, logement solidaire ou espace de proximit&eacute;.")
    add_number(doc, "Une logique de r&eacute;emploi : lorsque le cadre technique le permet, TVF peut contribuer &agrave; orienter des mat&eacute;riaux, mobiliers ou &eacute;quipements vers des projets valid&eacute;s, afin d'&eacute;viter leur perte et de soutenir des actions locales.")
    add_number(doc, "Une mobilisation des acteurs : TVF peut relier propri&eacute;taires, entreprises, associations, structures d'insertion, b&eacute;n&eacute;voles, collectivit&eacute;s et financeurs autour d'un dossier qualifi&eacute;.")
    add_number(doc, "Un suivi d'utilit&eacute; : chaque dossier peut &ecirc;tre document&eacute; par des pi&egrave;ces, des &eacute;tapes, des besoins, des d&eacute;cisions, des indicateurs et des limites de communication.")

    add_callout(
        doc,
        "Principe de prudence",
        "Aucune mention de partenariat, de soutien, de validation ou de financement par EPORA ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable. Le courrier vise uniquement &agrave; solliciter un premier &eacute;change de cadrage et &agrave; v&eacute;rifier les compl&eacute;mentarit&eacute;s possibles.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(
        doc,
        "Nous proposons un rendez-vous de cadrage afin de vous pr&eacute;senter TVF, d'&eacute;couter les conditions d'intervention d'EPORA, d'identifier les pr&eacute;cautions n&eacute;cessaires et de d&eacute;terminer si une contribution associative de type observation, pr&eacute;qualification, r&eacute;emploi ou mobilisation locale peut &ecirc;tre utile dans le respect des cadres existants.",
    )
    add_paragraph(
        doc,
        "&Agrave; l'issue de ce rendez-vous, TVF pourra transmettre une note synth&eacute;tique pr&eacute;cisant les sujets compatibles, les limites &agrave; respecter, les pi&egrave;ces utiles, les interlocuteurs concern&eacute;s et les suites &eacute;ventuelles &agrave; envisager.",
    )

    doc.add_heading("Conclusion", level=1)
    add_paragraph(
        doc,
        "Territoires Vivants France souhaite construire son lancement &agrave; Saint-&Eacute;tienne de mani&egrave;re s&eacute;rieuse, progressive et tra&ccedil;able. L'objectif n'est pas de se substituer aux acteurs publics, fonciers ou techniques, mais de cr&eacute;er un outil de coordination capable de rendre plus visibles les ressources inutilis&eacute;es, de mieux orienter les besoins et de faciliter l'&eacute;mergence de projets utiles aux territoires.",
    )
    add_paragraph(
        doc,
        "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous &agrave; la date qui vous conviendra et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.",
    )
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(head))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration officielle de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche friches / biens vacants", "D&eacute;crire le type de situations que TVF souhaite observer et qualifier.", "&Agrave; joindre"),
        ("Fiche m&eacute;thode d'instruction", "Montrer comment TVF enregistre, qualifie, suit et documente un dossier.", "&Agrave; joindre"),
        ("Fiche mat&eacute;riaux et r&eacute;emploi", "Pr&eacute;senter le fonctionnement de la valorisation des ressources inutilis&eacute;es.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(
        doc,
        "Objet d'e-mail recommand&eacute;",
        "Demande de rendez-vous - TERRITOIRES VIVANTS FRANCE / EPORA - compl&eacute;mentarit&eacute;s fonci&egrave;res et territoriales &agrave; Saint-&Eacute;tienne",
        fill=LIGHT_BLUE,
    )
    doc.add_heading("Sources de cadrage", level=1)
    add_paragraph(
        doc,
        "Les informations utilis&eacute;es pour orienter ce courrier proviennent du site officiel d'EPORA, notamment les mentions relatives &agrave; son statut d'&eacute;tablissement public foncier d'&Eacute;tat, &agrave; ses domaines d'intervention, &agrave; ses solutions, &agrave; sa rubrique travailler ensemble et &agrave; la direction territoriale Loire situ&eacute;e &agrave; Saint-&Eacute;tienne.",
    )

    doc.save(DOCX_PATH)


def build_md():
    md = txt(
        """# Courrier - EPORA - lancement TVF

Destinataire : EPORA - &Eacute;tablissement public foncier de l'Ouest Rh&ocirc;ne-Alpes, &agrave; l'attention de la Direction territoriale Loire et de la Direction g&eacute;n&eacute;rale.

Objet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change technique sur les compl&eacute;mentarit&eacute;s fonci&egrave;res, les friches, les biens vacants et la valorisation des ressources inutilis&eacute;es &agrave; Saint-&Eacute;tienne.

Date : Saint-&Eacute;tienne, le 14 juillet 2026.

## Intention du courrier

Ce courrier sollicite un premier rendez-vous technique avec EPORA afin de pr&eacute;senter TVF, comprendre les cadres d'intervention d'un &eacute;tablissement public foncier et identifier les sujets sur lesquels TVF pourrait compl&eacute;ter les dispositifs existants sans se substituer aux missions d'EPORA.

## Demandes op&eacute;rationnelles

- Organiser un rendez-vous avec la direction territoriale Loire ou le service comp&eacute;tent.
- Identifier les limites de contribution d'une association sur l'observation, la pr&eacute;qualification d'usages, le r&eacute;emploi et la mobilisation locale.
- Cadrer les questions de confidentialit&eacute;, responsabilit&eacute;, communication et validation pr&eacute;alable.
- Examiner si un premier cas de travail ou une premi&egrave;re m&eacute;thode peut &ecirc;tre envisag&eacute;e &agrave; Saint-&Eacute;tienne.

## Principe de prudence

Aucune mention de partenariat, de soutien, de validation ou de financement par EPORA ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable.

## Signature

Edryan Rangoly  
Pr&eacute;sident fondateur  
TERRITOIRES VIVANTS FRANCE
"""
    )
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
