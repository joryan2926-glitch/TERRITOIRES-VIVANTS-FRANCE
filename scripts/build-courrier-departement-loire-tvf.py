from pathlib import Path
import runpy
from html import unescape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
TPL = runpy.run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

DOCX_PATH = ROOT / "documents" / "courriers-lancement-saint-etienne" / "23-departement-loire-lancement-cooperation-tvf.docx"
MD_PATH = ROOT / "documents" / "sources" / "courrier-departement-loire-lancement-tvf.md"
DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
MD_PATH.parent.mkdir(parents=True, exist_ok=True)

txt = TPL["txt"]
style_document = TPL["style_document"]
add_header_footer = TPL["add_header_footer"]
add_paragraph = TPL["add_paragraph"]
add_bullet = TPL["add_bullet"]
add_number = TPL["add_number"]
add_callout = TPL["add_callout"]
set_table_width = TPL["set_table_width"]
set_cell_border = TPL["set_cell_border"]
set_cell_shading = TPL["set_cell_shading"]
set_repeat_table_header = TPL["set_repeat_table_header"]
GREEN = TPL["GREEN"]
MUTED = TPL["MUTED"]
LIGHT_GREEN = TPL["LIGHT_GREEN"]
LIGHT_BLUE = TPL["LIGHT_BLUE"]
BORDER = TPL["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        ("Destinataire", "D&eacute;partement de la Loire - &agrave; l'attention de Monsieur le Pr&eacute;sident du Conseil d&eacute;partemental et des services comp&eacute;tents"),
        ("Adresse", "2 rue Charles de Gaulle, 42000 Saint-&Eacute;tienne"),
        ("Objet", "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change pour le lancement d'une action pilote dans la Loire"),
        ("Pi&egrave;ces jointes", "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche besoins de lancement, fiche insertion / chantiers solidaires, fiche local de stockage / mat&eacute;riaux"),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - D&eacute;partement de la Loire"))
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'un premier rendez-vous de cadrage"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(doc, "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e, dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un ancrage op&eacute;rationnel progressif dans la Loire.")
    add_paragraph(doc, "Notre objet est de contribuer &agrave; remettre en usage des ressources aujourd'hui insuffisamment mobilis&eacute;es : logements vacants ou d&eacute;grad&eacute;s, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains inutilis&eacute;s, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. La d&eacute;marche repose sur une logique simple : identifier, qualifier, mettre en relation, conventionner, suivre et mesurer l'utilit&eacute; territoriale.")

    add_callout(doc, "Demande principale", "TVF sollicite un premier rendez-vous de cadrage avec le D&eacute;partement de la Loire afin de pr&eacute;senter sa m&eacute;thode, d'identifier les services concern&eacute;s et d'&eacute;tudier les conditions d'un lancement pilote articul&eacute; avec les besoins sociaux, territoriaux et environnementaux du d&eacute;partement.")

    doc.add_heading(txt("Pourquoi s'adresser au D&eacute;partement de la Loire"), level=1)
    add_paragraph(doc, "Le D&eacute;partement occupe une place essentielle dans les politiques de solidarit&eacute;, d'insertion, d'accompagnement social, d'habitat, d'appui aux territoires, d'environnement et de services aux communes. Ces sujets croisent directement les besoins auxquels TVF souhaite contribuer : remettre en usage des lieux, mobiliser des ressources existantes, soutenir des projets utiles et favoriser des parcours d'engagement ou d'insertion.")
    add_paragraph(doc, "TVF ne se substitue pas aux dispositifs d&eacute;partementaux existants. L'association peut au contraire agir comme un outil de terrain compl&eacute;mentaire : faire remonter des situations, structurer les premiers dossiers, orienter les ressources disponibles, relier les acteurs locaux et pr&eacute;parer des projets compatibles avec les cadres publics et associatifs existants.")

    doc.add_heading(txt("Axes de coop&eacute;ration propos&eacute;s"), level=1)
    add_number(doc, "Insertion et chantiers solidaires : identifier des actions simples pouvant associer des structures d'insertion, des b&eacute;n&eacute;voles, des associations locales ou des publics accompagn&eacute;s, dans un cadre encadr&eacute; et progressif.")
    add_number(doc, "Habitat et lieux vacants : contribuer &agrave; la qualification de situations, &agrave; l'orientation des propri&eacute;taires et &agrave; la pr&eacute;paration de dossiers utiles aux acteurs comp&eacute;tents.")
    add_number(doc, "Mat&eacute;riaux et r&eacute;emploi : organiser la collecte, la qualification et l'affectation de ressources r&eacute;employables vers des projets valid&eacute;s, sans distribution libre ni usage non trac&eacute;.")
    add_number(doc, "Territoires et communes : proposer une m&eacute;thode duplicable pour aider une commune, un quartier ou une association &agrave; transformer une ressource inutilis&eacute;e en projet utile.")
    add_number(doc, "Solidarit&eacute; et cadre de vie : soutenir des initiatives qui am&eacute;liorent concr&egrave;tement le cadre de vie, l'acc&egrave;s &agrave; des locaux, l'engagement citoyen et la coop&eacute;ration locale.")

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "La d&eacute;signation d'un interlocuteur ou d'un service r&eacute;f&eacute;rent pour orienter TVF vers les services d&eacute;partementaux concern&eacute;s.")
    add_bullet(doc, "Une mise en relation avec les services ou partenaires d&eacute;partementaux li&eacute;s &agrave; l'insertion, l'habitat-logement, l'action sociale, l'environnement, les communes, les associations et les projets territoriaux.")
    add_bullet(doc, "L'identification de structures d'insertion, associations ou acteurs locaux susceptibles de contribuer &agrave; des chantiers solidaires ou &agrave; des actions de r&eacute;emploi.")
    add_bullet(doc, "L'orientation vers des lieux, locaux ou espaces disponibles pouvant permettre le stockage temporaire, le tri et la qualification de mat&eacute;riaux r&eacute;employables.")
    add_bullet(doc, "L'&eacute;tude d'un cadre de coop&eacute;ration progressif, sans communication de partenariat ni engagement public avant validation formelle par le D&eacute;partement et par TVF.")

    add_callout(doc, "Principe de prudence", "Aucune mention de partenariat, de soutien, de financement ou de validation par le D&eacute;partement de la Loire ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable. Le premier &eacute;change demand&eacute; vise uniquement &agrave; cadrer, orienter et identifier les suites possibles.", fill=LIGHT_BLUE)

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(doc, "Nous proposons un rendez-vous d'environ une heure afin de pr&eacute;senter TVF, de comprendre les priorit&eacute;s d&eacute;partementales concern&eacute;es, d'identifier les services &agrave; associer et de d&eacute;terminer si un premier dossier pilote peut &ecirc;tre instruit &agrave; Saint-&Eacute;tienne ou dans une commune volontaire de la Loire.")
    add_paragraph(doc, "&Agrave; l'issue de cet &eacute;change, TVF pourra transmettre une note de cadrage plus pr&eacute;cise comprenant les besoins identifi&eacute;s, les acteurs &agrave; mobiliser, les pi&egrave;ces utiles, les limites d'intervention de l'association et les suites envisageables.")

    doc.add_heading("Conclusion", level=1)
    add_paragraph(doc, "Territoires Vivants France souhaite avancer de mani&egrave;re s&eacute;rieuse, progressive et tra&ccedil;able. Le lancement dans la Loire n'a pas pour objectif de multiplier les annonces, mais de construire une premi&egrave;re m&eacute;thode utile, sobre, partenariale et reproductible, au service des habitants et des territoires.")
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous &agrave; la date qui vous conviendra et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(text))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche insertion / chantiers solidaires", "Pr&eacute;senter les premi&egrave;res pistes de coop&eacute;ration avec des structures locales.", "&Agrave; joindre"),
        ("Fiche local de stockage / mat&eacute;riaux", "D&eacute;crire le besoin minimal : accessibilit&eacute;, s&eacute;curit&eacute;, stockage, tri, dur&eacute;e.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(doc, "Objet d'e-mail recommand&eacute;", "Demande de rendez-vous - lancement pilote TERRITOIRES VIVANTS FRANCE avec le D&eacute;partement de la Loire", fill=LIGHT_BLUE)
    doc.save(DOCX_PATH)


def build_md():
    md = txt("""# Courrier - D&eacute;partement de la Loire - lancement TVF\n\nDestinataire : D&eacute;partement de la Loire, &agrave; l'attention de Monsieur le Pr&eacute;sident du Conseil d&eacute;partemental et des services comp&eacute;tents.\n\nAdresse : 2 rue Charles de Gaulle, 42000 Saint-&Eacute;tienne.\n\nObjet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change pour le lancement d'une action pilote dans la Loire.\n\nDate : Saint-&Eacute;tienne, le 14 juillet 2026.\n\n## Corps du courrier\n\nMadame, Monsieur,\n\nJ'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e, dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un ancrage op&eacute;rationnel progressif dans la Loire.\n\nTVF sollicite un premier rendez-vous de cadrage avec le D&eacute;partement de la Loire afin de pr&eacute;senter sa m&eacute;thode, d'identifier les services concern&eacute;s et d'&eacute;tudier les conditions d'un lancement pilote articul&eacute; avec les besoins sociaux, territoriaux et environnementaux du d&eacute;partement.\n\n## Demandes op&eacute;rationnelles\n\n- D&eacute;signation d'un interlocuteur ou d'un service r&eacute;f&eacute;rent.\n- Mise en relation avec les services insertion, habitat-logement, action sociale, environnement, communes, associations et projets territoriaux.\n- Identification de structures d'insertion, associations ou acteurs locaux susceptibles de contribuer &agrave; des chantiers solidaires ou &agrave; des actions de r&eacute;emploi.\n- Orientation vers des lieux, locaux ou espaces disponibles pouvant permettre le stockage temporaire, le tri et la qualification de mat&eacute;riaux r&eacute;employables.\n- &Eacute;tude d'un cadre de coop&eacute;ration progressif, sans communication de partenariat ni engagement public avant validation formelle.\n\n## Signature\n\nEdryan Rangoly  \nPr&eacute;sident fondateur  \nTERRITOIRES VIVANTS FRANCE\n""")
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)