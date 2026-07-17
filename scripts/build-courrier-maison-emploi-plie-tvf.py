from pathlib import Path
from runpy import run_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

OUT_DIR = ROOT / "documents" / "courriers-lancement-saint-etienne"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "30-maison-emploi-plie-saint-etienne-lancement-insertion-tvf.docx"
MD_PATH = SRC_DIR / "courrier-maison-emploi-plie-saint-etienne-lancement-tvf.md"

txt = BASE["txt"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_number = BASE["add_number"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
BORDER = BASE["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        ("Destinataire", "Maison de l'Emploi / PLIE de Saint-&Eacute;tienne - &agrave; l'attention de la direction, de la coordination PLIE et des services insertion / emploi"),
        ("Adresse", "Coordonn&eacute;es op&eacute;rationnelles &agrave; confirmer avant envoi - Saint-&Eacute;tienne / Saint-&Eacute;tienne M&eacute;tropole"),
        ("Objet", "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les passerelles possibles entre revitalisation territoriale, chantiers encadr&eacute;s, r&eacute;emploi et parcours d'insertion"),
        ("Pi&egrave;ces jointes", "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche Solidarit&eacute; & Insertion, fiche chantiers encadr&eacute;s, fiche mat&eacute;riaux / r&eacute;emploi, fiche besoins Saint-&Eacute;tienne"),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def add_insertion_table(doc):
    doc.add_heading(txt("Passerelles &agrave; cadrer avec la Maison de l'Emploi / PLIE"), level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.6, 6.3, 6.2])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Sujet", "Apport possible de TVF", "Point &agrave; cadrer avec les acteurs emploi / insertion"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(head))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.1)

    rows = [
        ("Chantiers solidaires et encadr&eacute;s", "Rep&eacute;rer des supports de travaux utiles : nettoyage, tri, r&eacute;emploi, petite remise en &eacute;tat, pr&eacute;paration logistique ou animation de chantier participatif.", "V&eacute;rifier le cadre juridique, les responsabilit&eacute;s, l'encadrement professionnel, les assurances et l'articulation avec les structures habilit&eacute;es."),
        ("Orientation des publics", "Faire &eacute;merger des situations de terrain pouvant servir de supports &agrave; des parcours progressifs vers l'activit&eacute;.", "D&eacute;finir qui oriente, qui prescrit, qui accompagne, qui suit et quelles informations peuvent &ecirc;tre partag&eacute;es."),
        ("M&eacute;tiers du b&acirc;timent et du r&eacute;emploi", "Valoriser les m&eacute;tiers li&eacute;s &agrave; la r&eacute;novation, au tri, &agrave; la logistique, &agrave; la mat&eacute;riauth&egrave;que et &agrave; l'entretien d'espaces.", "Identifier les organismes de formation, les SIAE, les entreprises et les artisans pouvant intervenir dans un cadre s&eacute;curis&eacute;."),
        ("Besoins des collectivit&eacute;s et entreprises", "Structurer des demandes de terrain : local de stockage, collecte de mat&eacute;riaux, chantiers d'int&eacute;r&ecirc;t local, appui logistique.", "Croiser ces besoins avec les clauses sociales, les dispositifs existants et les priorit&eacute;s locales d'insertion."),
        ("Suivi d'impact", "Suivre les dossiers instruits, les ressources mobilis&eacute;es, les acteurs impliqu&eacute;s et les suites donn&eacute;es.", "D&eacute;finir des indicateurs prudents, tra&ccedil;ables et compatibles avec les obligations de chaque acteur."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.25)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - Maison de l'Emploi / PLIE Saint-&Eacute;tienne"))
    r.bold = True
    r.font.size = Pt(16.2)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Demande d'un premier &eacute;change sur les parcours d'insertion, les chantiers encadr&eacute;s et les besoins op&eacute;rationnels TVF"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(doc, "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF porte une d&eacute;marche appel&eacute;e &agrave; se d&eacute;velopper &agrave; l'&eacute;chelle nationale, avec un premier ancrage op&eacute;rationnel sur son territoire d'origine.")
    add_paragraph(doc, "L'association agit sur les ressources inutilis&eacute;es des territoires : logements vacants ou d&eacute;grad&eacute;s, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. Sa m&eacute;thode consiste &agrave; identifier les besoins, qualifier les situations, orienter les dossiers, mobiliser les acteurs comp&eacute;tents et suivre les suites donn&eacute;es, sans se substituer aux institutions, aux employeurs, aux structures d'insertion ou aux professionnels habilit&eacute;s.")
    add_paragraph(doc, "Dans cette perspective, le lien avec la Maison de l'Emploi / PLIE est essentiel. Les projets de revitalisation port&eacute;s ou accompagn&eacute;s par TVF peuvent devenir des supports utiles pour l'insertion : d&eacute;couverte de m&eacute;tiers, chantiers encadr&eacute;s, tri et valorisation de mat&eacute;riaux, appui logistique, remise en usage de lieux, mobilisation b&eacute;n&eacute;vole et mise en relation avec des acteurs &eacute;conomiques locaux.")

    add_callout(doc, "Demande principale", "TVF sollicite un premier rendez-vous afin de pr&eacute;senter sa m&eacute;thode, d'identifier les interlocuteurs comp&eacute;tents et de d&eacute;terminer comment les futurs dossiers TVF pourraient &ecirc;tre articul&eacute;s avec les parcours d'insertion, les structures d'accompagnement, les entreprises et les dispositifs locaux existants.")

    doc.add_heading(txt("Pourquoi cette rencontre est importante"), level=1)
    add_paragraph(doc, "La revitalisation territoriale ne peut pas &ecirc;tre seulement technique ou immobili&egrave;re. Un logement remis en &eacute;tat, un local r&eacute;activ&eacute;, une friche transform&eacute;e ou une mat&eacute;riauth&egrave;que locale peuvent aussi devenir des supports de parcours : apprendre un geste professionnel, reprendre confiance, participer &agrave; un projet utile, rencontrer des entreprises et se projeter vers une activit&eacute;.")
    add_paragraph(doc, "TVF souhaite donc construire d&egrave;s le lancement un cadre s&eacute;rieux avec les acteurs de l'emploi et de l'insertion. Il ne s'agit pas de cr&eacute;er un dispositif parall&egrave;le, mais de relier les futurs projets TVF aux cadres existants, aux structures habilit&eacute;es et aux comp&eacute;tences d&eacute;j&agrave; pr&eacute;sentes sur le territoire st&eacute;phanois.")
    add_paragraph(doc, "Cette rencontre permettrait aussi de s&eacute;curiser les limites : publics concern&eacute;s, modalit&eacute;s d'orientation, obligations d'encadrement, tra&ccedil;abilit&eacute;, assurances, responsabilit&eacute;s, suivi, donn&eacute;es personnelles et communication. Ces points doivent &ecirc;tre clarifi&eacute;s avant toute action de terrain.")

    add_insertion_table(doc)

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "Identifier le bon interlocuteur emploi / insertion pour le lancement de TVF &agrave; Saint-&Eacute;tienne.")
    add_bullet(doc, "Comprendre les conditions de mobilisation de publics en insertion dans le cadre de chantiers encadr&eacute;s ou d'actions de terrain.")
    add_bullet(doc, "Cartographier les structures pouvant intervenir : SIAE, chantiers d'insertion, organismes de formation, entreprises, artisans, associations et prescripteurs.")
    add_bullet(doc, "D&eacute;finir un circuit prudent entre demande TVF, qualification du besoin, orientation vers les acteurs comp&eacute;tents et suivi du dossier.")
    add_bullet(doc, "Identifier les premi&egrave;res actions compatibles : tri de mat&eacute;riaux, aide logistique, remise en &eacute;tat l&eacute;g&egrave;re, animation de chantier participatif, sensibilisation au r&eacute;emploi.")

    doc.add_heading(txt("Ce que TVF peut apporter"), level=1)
    add_number(doc, "Des supports concrets de terrain, li&eacute;s aux besoins r&eacute;els du territoire : biens vacants, locaux, mat&eacute;riaux, espaces &agrave; remettre en usage.")
    add_number(doc, "Une capacit&eacute; de pr&eacute;qualification : objet de la demande, localisation, acteurs concern&eacute;s, niveau d'urgence, contraintes connues, pi&egrave;ces disponibles.")
    add_number(doc, "Une passerelle vers les entreprises, propri&eacute;taires, collectivit&eacute;s, associations, b&eacute;n&eacute;voles et artisans pouvant contribuer &agrave; des projets utiles.")
    add_number(doc, "Une logique de r&eacute;emploi et de sobri&eacute;t&eacute; qui peut servir de support &agrave; des parcours professionnels dans le b&acirc;timent, la logistique, la maintenance, le tri, la collecte et l'animation locale.")
    add_number(doc, "Un suivi des dossiers et des contributions, afin de garder une trace claire des demandes, des orientations et des suites donn&eacute;es.")

    add_callout(doc, "Principe de prudence", "TVF ne se substitue ni &agrave; la Maison de l'Emploi, ni au PLIE, ni &agrave; France Travail, ni aux prescripteurs, ni aux structures d'insertion par l'activit&eacute; &eacute;conomique, ni aux employeurs. Aucune orientation, promesse d'emploi, validation de parcours, convention d'insertion ou communication publique ne sera engag&eacute;e sans accord et cadre &eacute;crit des acteurs comp&eacute;tents.", fill=LIGHT_BLUE)

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(doc, "Nous proposons un rendez-vous d'environ une heure afin de vous pr&eacute;senter TVF, d'identifier les points de convergence, de comprendre les exigences applicables et de pr&eacute;parer une m&eacute;thode d'orientation simple pour les futurs dossiers de terrain.")
    add_paragraph(doc, "&Agrave; l'issue de cet &eacute;change, TVF pourra formaliser une note de cadrage pr&eacute;cisant les types de dossiers compatibles, les pi&egrave;ces &agrave; r&eacute;unir, les interlocuteurs &agrave; solliciter, les limites &agrave; respecter et les conditions minimales avant toute action.")

    doc.add_heading("Conclusion", level=1)
    add_paragraph(doc, "Territoires Vivants France souhaite d&eacute;marrer &agrave; Saint-&Eacute;tienne avec une logique responsable, partenariale et utile. Le lien avec la Maison de l'Emploi / PLIE permettrait de relier la revitalisation des lieux aux parcours humains, aux comp&eacute;tences locales et aux besoins d'insertion du territoire.")
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(head))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration officielle de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche Solidarit&eacute; & Insertion", "Expliquer les passerelles entre revitalisation, chantiers et parcours d'insertion.", "&Agrave; joindre"),
        ("Fiche chantiers encadr&eacute;s", "Pr&eacute;senter les actions de terrain possibles et les limites &agrave; respecter.", "&Agrave; joindre"),
        ("Fiche mat&eacute;riaux / r&eacute;emploi", "D&eacute;crire les besoins autour du tri, de la collecte, du stockage et de la valorisation.", "&Agrave; joindre"),
        ("Fiche besoins Saint-&Eacute;tienne", "Lister les premiers besoins : local de stockage, transport, acteurs insertion, entreprises, projets pilotes.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(doc, "Objet d'e-mail recommand&eacute;", "Demande de rendez-vous - TERRITOIRES VIVANTS FRANCE / Maison de l'Emploi - PLIE Saint-&Eacute;tienne - insertion et chantiers encadr&eacute;s", fill=LIGHT_BLUE)
    doc.add_heading("Point &agrave; confirmer", level=1)
    add_paragraph(doc, "Les coordonn&eacute;es op&eacute;rationnelles exactes du destinataire doivent &ecirc;tre confirm&eacute;es avant envoi afin d'&eacute;viter toute erreur d'adressage. Aucune adresse non v&eacute;rifi&eacute;e n'a &eacute;t&eacute; invent&eacute;e dans ce document.")

    doc.save(DOCX_PATH)


def build_md():
    md = txt("""# Courrier - Maison de l'Emploi / PLIE Saint-&Eacute;tienne - lancement TVF

Destinataire : Maison de l'Emploi / PLIE de Saint-&Eacute;tienne - coordonn&eacute;es &agrave; confirmer avant envoi.

Objet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les passerelles possibles entre revitalisation territoriale, chantiers encadr&eacute;s, r&eacute;emploi et parcours d'insertion.

Date : Saint-&Eacute;tienne, le 14 juillet 2026.

## Intention

TVF sollicite un premier rendez-vous afin d'identifier les interlocuteurs comp&eacute;tents et de d&eacute;terminer comment les futurs dossiers TVF pourraient &ecirc;tre articul&eacute;s avec les parcours d'insertion, les structures d'accompagnement, les entreprises et les dispositifs locaux existants.

## Principe de prudence

TVF ne se substitue ni &agrave; la Maison de l'Emploi, ni au PLIE, ni aux prescripteurs, ni aux structures d'insertion par l'activit&eacute; &eacute;conomique, ni aux employeurs. Aucune promesse ou communication publique ne sera engag&eacute;e sans accord &eacute;crit.
""")
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
