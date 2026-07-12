from __future__ import annotations

from datetime import date
from pathlib import Path
import shutil
import unicodedata

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "documents" / "kit-formulaires-conventions-tvf"
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"

TVF = {
    "name": "TERRITOIRES VIVANTS FRANCE",
    "sigle": "TVF",
    "address": "25 rue Elise Gervais, 42000 Saint-Etienne",
    "email": "contact@territoiresvivantsfrance.fr",
    "phone": "04 65 81 54 69",
    "site": "www.territoiresvivantsfrance.fr",
    "rna": "W423016361",
    "siren": "107 226 128",
    "siret": "107 226 128 00018",
    "president": "Edryan Rangoly, president fondateur",
    "secretary": "M. Lambeau Jordan, secretaire et tresorier",
}

NAVY = RGBColor(9, 34, 58)
GREEN = RGBColor(46, 125, 50)
GOLD = RGBColor(184, 139, 26)
DARK = RGBColor(30, 39, 46)
MUTED = RGBColor(90, 100, 108)
LIGHT_GREEN = "EEF7EE"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF8E6"
LIGHT_GRAY = "F4F6F8"
BORDER = "C8D2DC"
WHITE = "FFFFFF"


def safe_name(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    normalized = normalized.lower().replace("'", "").replace(" ", "-")
    keep = "abcdefghijklmnopqrstuvwxyz0123456789-_"
    return "".join(ch for ch in normalized if ch in keep).strip("-")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def paragraph(doc, text="", style=None, before=0, after=6, line=1.15, align=None, color=None, bold=False, italic=False, size=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text, bold=bold, italic=italic, color=color, size=size)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    return p.add_run(text)


def field_line(label: str, width: int = 80) -> str:
    return f"{label} : " + "_" * width


def checkbox(label: str) -> str:
    return f"[ ] {label}"


def add_note_box(doc, title: str, body: str, fill: str = LIGHT_GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.25])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    add_run(p, title + " - ", bold=True, color=NAVY)
    add_run(p, body, color=DARK)
    paragraph(doc, "", after=2)


def add_bullets(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if isinstance(item, tuple):
            add_run(p, item[0], bold=True, color=DARK)
            add_run(p, item[1], color=DARK)
        else:
            add_run(p, item, color=DARK)


def add_numbered(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        add_run(p, item, color=DARK)


def add_kv_table(doc, rows: list[tuple[str, str]], widths=(1.75, 4.5), header=None, fill=WHITE):
    if header:
        paragraph(doc, header, after=4, bold=True, color=NAVY, size=10.5)
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, list(widths))
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        set_cell_shading(c0, LIGHT_BLUE)
        set_cell_shading(c1, fill)
        p0 = c0.paragraphs[0]
        p1 = c1.paragraphs[0]
        add_run(p0, label, bold=True, color=NAVY, size=9.5)
        add_run(p1, value, color=DARK, size=9.5)
    paragraph(doc, "", after=3)
    return table


def add_check_table(doc, title: str, items: list[tuple[str, str, str]]):
    paragraph(doc, title, after=4, bold=True, color=NAVY, size=10.5)
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [0.45, 2.3, 1.0, 2.5])
    headers = ["OK", "Piece / controle", "Priorite", "Commentaire"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT_BLUE)
        add_run(cell.paragraphs[0], h, bold=True, color=NAVY, size=9)
    for piece, priority, comment in items:
        row = table.add_row().cells
        values = ["[ ]", piece, priority, comment]
        for idx, value in enumerate(values):
            set_cell_border(row[idx])
            set_cell_margins(row[idx])
            p = row[idx].paragraphs[0]
            add_run(p, value, color=DARK, size=8.8)
    paragraph(doc, "", after=3)


def add_signature_block(doc, left_label="Pour TVF", right_label="Pour le demandeur / structure"):
    paragraph(doc, "Signatures", style="Heading 2")
    rows = [
        ("Lieu et date", "Fait a : ____________________________    Le : ____ / ____ / ______"),
        (left_label, "Nom : ____________________________\nFonction : ________________________\nSignature :\n\n"),
        (right_label, "Nom : ____________________________\nFonction : ________________________\nSignature :\n\n"),
    ]
    table = doc.add_table(rows=3, cols=2)
    set_table_width(table, [1.45, 4.8])
    for i, (label, value) in enumerate(rows):
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        add_run(table.cell(i, 0).paragraphs[0], label, bold=True, color=NAVY, size=9.5)
        add_run(table.cell(i, 1).paragraphs[0], value, color=DARK, size=9.5)
    paragraph(doc, "Le present modele doit etre adapte au dossier, aux pieces transmises et au cadre juridique applicable avant signature definitive.", italic=True, color=MUTED, size=9, after=0)


def setup_doc(title: str, reference: str, subtitle: str, public: str, status="Modele operationnel pret a completer") -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for s in doc.sections:
        s.top_margin = Inches(0.72)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        s.header_distance = Inches(0.35)
        s.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for level, size, color, before, after in [
        (1, 16, NAVY, 16, 8),
        (2, 13, GREEN, 12, 6),
        (3, 11.5, NAVY, 8, 4),
    ]:
        st = styles[f"Heading {level}"]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = section.header
    htable = header.add_table(rows=1, cols=2, width=Inches(6.9))
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER
    htable.autofit = False
    htable.columns[0].width = Inches(2.1)
    htable.columns[1].width = Inches(4.8)
    if LOGO.exists():
        hp = htable.cell(0, 0).paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.add_run().add_picture(str(LOGO), width=Inches(1.65))
    rp = htable.cell(0, 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(rp, TVF["name"], bold=True, color=NAVY, size=9.5)
    rp.add_run("\n")
    add_run(rp, f"RNA {TVF['rna']} - SIREN {TVF['siren']} - {TVF['email']}", color=MUTED, size=8.2)
    for cell in htable.rows[0].cells:
        for paragraph_ in cell.paragraphs:
            paragraph_.paragraph_format.space_after = Pt(0)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, f"{TVF['name']} - Document modifiable - {TVF['site']} - {TVF['email']}", color=MUTED, size=8)

    p = paragraph(doc, reference, after=2, color=GOLD, bold=True, size=9)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p = paragraph(doc, title, after=4, color=NAVY, bold=True, size=20)
    title_p.paragraph_format.keep_with_next = True
    paragraph(doc, subtitle, after=12, color=MUTED, size=11)
    add_kv_table(
        doc,
        [
            ("Reference", reference),
            ("Version", f"V1.0 - {date.today().strftime('%d/%m/%Y')}"),
            ("Public cible", public),
            ("Statut", status),
            ("Contact TVF", f"{TVF['email']} - {TVF['phone']}"),
        ],
        widths=(1.4, 4.85),
    )
    add_note_box(
        doc,
        "Cadre d'utilisation",
        "Document pret a remplir pour preparer l'instruction TVF. Les clauses engageantes doivent etre adaptees au dossier, aux pieces disponibles et au cadre juridique applicable avant signature.",
        fill=LIGHT_GOLD,
    )
    return doc


FORMS = [
    {
        "ref": "TVF-F-01",
        "folder": "01-contact-general",
        "title": "Formulaire de contact general",
        "public": "Tout public",
        "subtitle": "Qualifier rapidement une demande et orienter vers le bon parcours TVF.",
        "need": "Demande generale, demande de rendez-vous, question institutionnelle ou premiere orientation.",
        "fields": [
            "Nom et prenom / structure",
            "Profil du demandeur",
            "E-mail",
            "Telephone",
            "Commune / territoire concerne",
            "Objet principal",
            "Message synthetique",
            "Pieces deja disponibles",
        ],
        "pieces": [
            ("Message ou note de contexte", "Obligatoire", "Decrire la demande en quelques lignes."),
            ("Coordonnees completes", "Obligatoire", "Indiquer au moins un e-mail de reponse."),
            ("Photos ou documents utiles", "Si disponibles", "Ne transmettre que des documents autorises."),
            ("Urgence ou echeance", "Recommande", "Preciser la date limite si elle existe."),
        ],
        "clauses": [
            "TVF qualifie la demande avant toute orientation ou prise d'engagement.",
            "Aucune action terrain, communication publique ou depense n'est engagee sans validation ecrite.",
            "Les donnees transmises sont utilisees uniquement pour traiter la demande et organiser la suite.",
        ],
        "next": "Orientation vers le formulaire specialise, proposition de rendez-vous ou reponse ecrite.",
    },
    {
        "ref": "TVF-F-02",
        "folder": "02-collectivite-territoire",
        "title": "Demande collectivite - territoire de cooperation",
        "public": "Commune, EPCI, departement, region, service public",
        "subtitle": "Preparer un rendez-vous de cadrage, un diagnostic territorial ou une cooperation TVF.",
        "need": "Diagnostic de vacance, reemploi de ressources, animation territoriale, projet pilote ou convention de cooperation.",
        "fields": [
            "Nom de la collectivite / EPCI",
            "Service ou elu referent",
            "Perimetre geographique",
            "Besoin public identifie",
            "Programmes publics concernes",
            "Biens, friches ou ressources disponibles",
            "Attentes envers TVF",
            "Calendrier souhaite",
        ],
        "pieces": [
            ("Contact referent", "Obligatoire", "Elu, direction, service technique ou chef de projet."),
            ("Perimetre territorial", "Obligatoire", "Commune, quartier, rue ou secteur prioritaire."),
            ("Donnees disponibles", "Si disponibles", "Vacance, friches, commerces, habitat, foncier."),
            ("Programmes existants", "Recommande", "PLH, OPAH, Action Coeur de Ville, NPNRU, PCAET, ZAN."),
            ("Biens ou ressources cibles", "Recommande", "Liste, carte, photos, contexte administratif."),
        ],
        "clauses": [
            "La cooperation est formalisee par une convention territoriale definissant le perimetre, les objectifs, les referents, les modalites de suivi et les usages autorises.",
            "La collectivite reste responsable de ses competences, decisions et obligations administratives.",
            "TVF intervient comme outil de coordination, d'instruction, de mobilisation et de suivi, sans se substituer aux dispositifs publics existants.",
            "Toute communication publique est validee conjointement avant diffusion.",
        ],
        "next": "Rendez-vous de cadrage, fiche territoire, calendrier de travail, puis convention si le dossier est retenu.",
    },
    {
        "ref": "TVF-F-03",
        "folder": "03-proprietaire-bien-vacant",
        "title": "Proposition de bien vacant ou inutilise",
        "public": "Proprietaire prive, indivision, bailleur, personne morale",
        "subtitle": "Etudier un logement, commerce, batiment, local, terrain ou friche pouvant etre remis en usage.",
        "need": "Bien vacant, degrade, inutilise ou sous-utilise pouvant faire l'objet d'un diagnostic et d'un scenario de remise en usage.",
        "fields": [
            "Identite du proprietaire ou mandataire",
            "Adresse complete du bien",
            "Nature du bien",
            "Etat apparent",
            "Surface estimee",
            "Acces, reseaux, contraintes connues",
            "Usage envisageable",
            "Duree de cooperation possible",
            "Photos disponibles",
        ],
        "pieces": [
            ("Identite du proprietaire / representant", "Obligatoire", "Nom, prenom, structure, coordonnees."),
            ("Adresse complete du bien", "Obligatoire", "Numero, rue, commune, code postal."),
            ("Justificatif de propriete ou mandat", "A verifier", "Titre, taxe fonciere, mandat, attestation."),
            ("Photos recentes", "Obligatoire", "Facade, acces, interieur si possible et autorise."),
            ("Documents techniques", "Si disponibles", "Diagnostics, plans, devis, etudes, copropriete."),
        ],
        "clauses": [
            "Le proprietaire conserve la pleine propriete du bien.",
            "Toute intervention fait l'objet d'une convention ecrite precisant le droit d'usage, la duree, les responsabilites, l'assurance, les conditions de sortie et les modalites de valorisation.",
            "TVF ne s'engage pas a financer ou realiser des travaux sans instruction complete et decision interne.",
            "Aucune visite, publication ou intervention n'a lieu sans accord explicite du proprietaire ou de son representant habilite.",
        ],
        "next": "Prequalification, demande de pieces, visite eventuelle, scenario d'usage puis convention si le dossier est retenu.",
    },
    {
        "ref": "TVF-F-04",
        "folder": "04-materiaux-reemploi",
        "title": "Proposition de materiaux reemployables",
        "public": "Entreprise, particulier, collectivite, association, artisan",
        "subtitle": "Qualifier des materiaux, mobiliers ou equipements pouvant etre valorises dans un projet TVF.",
        "need": "Materiaux, mobilier, outils ou equipements pouvant etre integres dans une strategie de valorisation territoriale.",
        "fields": [
            "Identite du contributeur",
            "Categorie des materiaux",
            "Quantite estimee",
            "Etat apparent",
            "Localisation",
            "Date de disponibilite",
            "Conditions d'enlevement",
            "Restrictions d'usage connues",
        ],
        "pieces": [
            ("Description detaillee", "Obligatoire", "Nature, quantite, etat, dimensions si possible."),
            ("Photos", "Recommande", "Vue generale et details."),
            ("Localisation et acces", "Obligatoire", "Adresse, horaires, contraintes de chargement."),
            ("Disponibilite", "Obligatoire", "Date limite, urgence, duree possible de stockage."),
            ("Informations de securite", "A verifier", "Amiante, produits dangereux, electricite, charges lourdes."),
        ],
        "clauses": [
            "La Banque de Materiaux TVF n'est pas une plateforme de distribution gratuite.",
            "TVF decide de l'acceptation, de la priorisation et de l'affectation des ressources selon les projets valides.",
            "Un bordereau de remise ou une convention de valorisation precise les materiaux, les responsabilites, les conditions logistiques et la destination prevue.",
            "Les materiaux non conformes, dangereux, pollues ou impossibles a valoriser peuvent etre refuses.",
        ],
        "next": "Verification d'acceptabilite, faisabilite logistique, affectation possible a un projet, bordereau de remise si retenu.",
    },
    {
        "ref": "TVF-F-05",
        "folder": "05-entreprise-partenariat",
        "title": "Demande entreprise - cooperation territoriale",
        "public": "Entreprise, artisan, commerce, promoteur, bailleur, logisticien",
        "subtitle": "Construire une cooperation RSE : materiaux, local, transport, competences, financement ou mecanat.",
        "need": "Contribution materielle, technique, logistique, fonciere, financiere ou de competences au service d'un projet territorial.",
        "fields": [
            "Raison sociale",
            "SIRET",
            "Contact referent",
            "Type de contribution",
            "Objectifs RSE ou territoriaux",
            "Ressources proposees",
            "Conditions et contraintes",
            "Valorisation attendue",
        ],
        "pieces": [
            ("Identification de l'entreprise", "Obligatoire", "Nom, SIRET, contact habilite."),
            ("Description de la contribution", "Obligatoire", "Materiaux, local, vehicule, competences, financement."),
            ("Contraintes logistiques", "Si applicable", "Enlevement, assurance, acces, delais."),
            ("Attentes RSE", "Recommande", "Reporting, visibilite, impact, communication."),
            ("Elements de conformite", "A verifier", "Regles internes, assurances, autorisations."),
        ],
        "clauses": [
            "La cooperation est formalisee par une convention de partenariat ou de contribution definissant l'objet, les ressources, les responsabilites, la duree et les modalites de communication.",
            "TVF peut valoriser l'engagement de l'entreprise uniquement dans les conditions validees par les deux parties.",
            "La contribution ne cree pas de droit automatique sur l'affectation finale des ressources sauf mention expresse dans la convention.",
            "Les engagements financiers, materiels ou de competence doivent etre traces par ecrit.",
        ],
        "next": "Qualification, rendez-vous, definition du cadre puis convention de partenariat ou de contribution.",
    },
    {
        "ref": "TVF-F-06",
        "folder": "06-local-stockage",
        "title": "Mise a disposition potentielle d'un local de stockage",
        "public": "Collectivite, entreprise, bailleur, proprietaire, association",
        "subtitle": "Identifier un local temporaire pouvant accueillir des materiaux ou equipements TVF.",
        "need": "Local sec, accessible et securisable pour stocker, trier ou preparer des ressources de reemploi.",
        "fields": [
            "Proprietaire ou gestionnaire du local",
            "Adresse du local",
            "Surface approximative",
            "Accessibilite vehicule",
            "Etat general",
            "Disponibilite",
            "Conditions d'usage",
            "Assurance / securite",
        ],
        "pieces": [
            ("Adresse et plan d'acces", "Obligatoire", "Preciser les contraintes de circulation."),
            ("Photos du local", "Obligatoire", "Interieur, exterieur, acces, zones sensibles."),
            ("Surface et caracteristiques", "Recommande", "Hauteur, sol, electricite, eau, ventilation."),
            ("Conditions de mise a disposition", "A verifier", "Duree, cout, charges, assurance, horaires."),
            ("Autorisations / reglement", "A verifier", "Copropriete, bail, ERP, securite incendie."),
        ],
        "clauses": [
            "L'usage du local est limite aux finalites prevues dans la convention de mise a disposition.",
            "TVF tient un registre des entrees, sorties et affectations des ressources stockees.",
            "Les conditions d'assurance, de securite, d'acces, d'entretien et de restitution doivent etre precisees avant toute utilisation.",
            "Aucun produit dangereux ou non conforme n'est stocke sans validation specifique.",
        ],
        "next": "Visite technique, evaluation securite, protocole de stockage, convention de mise a disposition.",
    },
    {
        "ref": "TVF-F-07",
        "folder": "07-transport-logistique",
        "title": "Mise a disposition transport et logistique",
        "public": "Entreprise, collectivite, loueur, transporteur, association",
        "subtitle": "Recenser les moyens de transport pouvant aider a collecter ou deplacer materiaux et equipements.",
        "need": "Vehicule, chauffeur, manutention, stockage temporaire, livraison ou appui logistique ponctuel.",
        "fields": [
            "Structure proposant le moyen",
            "Type de vehicule / capacite",
            "Conducteur disponible",
            "Perimetre d'intervention",
            "Disponibilites",
            "Assurance",
            "Contraintes de chargement",
            "Conditions de prise en charge",
        ],
        "pieces": [
            ("Description du moyen logistique", "Obligatoire", "Vehicule, volume, charge utile, equipements."),
            ("Disponibilites", "Obligatoire", "Dates, horaires, preavis necessaire."),
            ("Assurance et habilitations", "A verifier", "Permis, assurance, responsabilite."),
            ("Conditions d'utilisation", "A verifier", "Carburant, chauffeur, cout, zone couverte."),
            ("Contact operationnel", "Obligatoire", "Personne joignable le jour de l'operation."),
        ],
        "clauses": [
            "Chaque operation logistique doit etre planifiee et validee avant execution.",
            "Les responsabilites de chargement, transport, dechargement et assurance sont precisees dans un accord ecrit.",
            "TVF ne mobilise aucun vehicule sans verification minimale de l'assurance, du conducteur et du cadre d'usage.",
        ],
        "next": "Qualification logistique, verification assurance, accord ponctuel ou convention avant usage.",
    },
    {
        "ref": "TVF-F-08",
        "folder": "08-benevole-engagement",
        "title": "Candidature benevole ou citoyenne",
        "public": "Habitant, benevole, etudiant, retraite, volontaire",
        "subtitle": "Identifier les competences, disponibilites et missions possibles pour participer aux actions TVF.",
        "need": "Participation citoyenne, appui terrain, communication, logistique, tri, evenements ou chantiers encadres.",
        "fields": [
            "Nom et prenom",
            "Coordonnees",
            "Commune",
            "Disponibilites",
            "Competences utiles",
            "Types de missions souhaitees",
            "Contraintes de sante ou securite",
            "Personne a prevenir si besoin",
        ],
        "pieces": [
            ("Coordonnees completes", "Obligatoire", "E-mail et telephone si possible."),
            ("Disponibilites", "Recommande", "Jours, horaires, frequence."),
            ("Competences ou experiences", "Recommande", "Technique, social, communication, logistique."),
            ("Contraintes particulieres", "A signaler", "Sante, mobilite, equipements, assurance."),
            ("Autorisation parentale", "Si mineur", "Obligatoire selon l'age et l'action."),
        ],
        "clauses": [
            "Le benevole intervient dans un cadre encadre par TVF ou un partenaire identifie.",
            "Aucune mission dangereuse, technique ou sensible n'est confiee sans verification des competences, consignes et conditions de securite.",
            "Le benevole s'engage a respecter les consignes, les lieux, les personnes et la confidentialite des informations recueillies.",
        ],
        "next": "Entretien court, orientation mission, charte benevole, feuille d'emargement lors des actions.",
    },
    {
        "ref": "TVF-F-09",
        "folder": "09-association-insertion-chantier",
        "title": "Cooperation association, insertion ou chantier participatif",
        "public": "Association, SIAE, structure insertion, organisme formation",
        "subtitle": "Construire une action encadree autour du reemploi, de la renovation, de l'animation locale ou de l'insertion.",
        "need": "Projet social, chantier participatif, formation, insertion, animation de quartier ou accompagnement de publics.",
        "fields": [
            "Nom de la structure",
            "Objet social",
            "Public accompagne",
            "Referent operationnel",
            "Projet ou besoin",
            "Competences disponibles",
            "Cadre d'intervention",
            "Contraintes securite",
        ],
        "pieces": [
            ("Presentation de la structure", "Obligatoire", "Objet, publics, territoire, contacts."),
            ("Cadre d'intervention", "Obligatoire", "Insertion, formation, benevolat, animation."),
            ("Assurance et encadrement", "A verifier", "Responsables, qualification, securite."),
            ("Description de l'action", "Obligatoire", "Lieu, objectifs, calendrier, besoins."),
            ("Pieces chantier", "Si applicable", "Plan de prevention, autorisations, equipements, consignes."),
        ],
        "clauses": [
            "Toute action de chantier ou d'insertion doit etre encadree, tracee et adaptee aux publics concernes.",
            "Les roles de TVF, de la structure partenaire et des encadrants sont precises avant action.",
            "Les regles de securite, d'assurance, de prevention et d'emargement sont obligatoires pour les operations terrain.",
        ],
        "next": "Rendez-vous technique, verification securite, convention partenariale, plan de prevention si chantier.",
    },
    {
        "ref": "TVF-F-10",
        "folder": "10-financeur-mecene",
        "title": "Demande financeur, mecene ou investisseur solidaire",
        "public": "Fondation, entreprise mecene, banque, financeur public ou prive",
        "subtitle": "Structurer un echange sur un soutien financier, materiel ou en competences.",
        "need": "Soutien au lancement, financement de projet, mecanat, contribution en nature ou appui aux investissements solidaires.",
        "fields": [
            "Nom du financeur / mecene",
            "Contact referent",
            "Type de soutien envisage",
            "Axes d'interet",
            "Territoire concerne",
            "Montant ou nature du soutien",
            "Modalites de reporting attendues",
            "Contraintes administratives",
        ],
        "pieces": [
            ("Identification du financeur", "Obligatoire", "Structure, contact, statut."),
            ("Objet du soutien", "Obligatoire", "Projet, territoire, theme, public."),
            ("Cadre financier", "A verifier", "Don, subvention, convention, mecanat, appel a projets."),
            ("Attentes de reporting", "Recommande", "Indicateurs, calendrier, livrables."),
            ("Regles de communication", "A verifier", "Logo, mention publique, validation prealable."),
        ],
        "clauses": [
            "Tout soutien financier ou materiel fait l'objet d'une convention precisant l'objet, les montants, les affectations, les obligations de reporting et les conditions de communication.",
            "TVF ne garantit aucun avantage fiscal sans verification du cadre juridique applicable.",
            "Les fonds sont affectes selon la convention signee et font l'objet d'une tracabilite interne.",
        ],
        "next": "Note d'opportunite, rendez-vous, dossier de financement, budget previsionnel, convention de soutien.",
    },
    {
        "ref": "TVF-F-11",
        "folder": "11-signalement-citoyen",
        "title": "Signalement citoyen d'un lieu ou d'une ressource",
        "public": "Citoyen, habitant, association locale",
        "subtitle": "Signaler un logement vacant, commerce ferme, batiment abandonne, friche, terrain ou materiaux disponibles.",
        "need": "Information terrain utile a l'observatoire TVF, a verifier avant toute qualification publique.",
        "fields": [
            "Type de signalement",
            "Adresse ou localisation",
            "Description",
            "Photos disponibles",
            "Danger apparent",
            "Source de l'information",
            "Contact facultatif",
            "Autorisation d'etre recontacte",
        ],
        "pieces": [
            ("Localisation precise", "Obligatoire", "Adresse, rue, repere ou commune."),
            ("Description", "Obligatoire", "Ce qui est observe, depuis quand si connu."),
            ("Photos", "Si autorisees", "Uniquement depuis l'espace public ou avec accord."),
            ("Informations de securite", "A signaler", "Danger, acces interdit, risque apparent."),
            ("Contact", "Facultatif", "Permettre a TVF de demander des precisions."),
        ],
        "clauses": [
            "Un signalement ne vaut pas preuve juridique de vacance, d'abandon ou de disponibilite.",
            "TVF ne publie pas de donnee nominative ou sensible sans verification, base legale et autorisation lorsque necessaire.",
            "Le signalant ne doit pas entrer dans une propriete privee ni prendre de risque pour collecter l'information.",
        ],
        "next": "Verification prudente, qualification territoriale, aucune publication sans controle et sans base legale.",
    },
    {
        "ref": "TVF-F-12",
        "folder": "12-presse-institutionnel",
        "title": "Demande presse, institution ou communication",
        "public": "Journaliste, institution, media, service communication",
        "subtitle": "Traiter les demandes d'information, interview, kit media ou elements institutionnels.",
        "need": "Demande de presentation, interview, visuel officiel, citation, communique ou prise de parole.",
        "fields": [
            "Nom du media / institution",
            "Journaliste ou contact",
            "Objet de la demande",
            "Format souhaite",
            "Date limite",
            "Angle editorial",
            "Elements demandes",
            "Autorisation de citation",
        ],
        "pieces": [
            ("Presentation de la demande", "Obligatoire", "Objet, angle, format, date limite."),
            ("Contact presse", "Obligatoire", "Nom, e-mail, telephone."),
            ("Questions ou brief", "Recommande", "Permettre une reponse officielle preparee."),
            ("Conditions de publication", "A verifier", "Citation, relecture, images, logo."),
            ("Support final", "Si disponible", "Lien ou exemplaire apres publication."),
        ],
        "clauses": [
            "Les elements officiels TVF doivent etre utilises sans alteration et dans leur contexte.",
            "Les citations, logos, photos et documents transmis doivent respecter les conditions de diffusion convenues.",
            "Aucune annonce de partenariat, financement ou resultat ne doit etre publiee sans validation formelle.",
        ],
        "next": "Qualification, transmission kit media, preparation reponse officielle ou rendez-vous.",
    },
]


CONVENTIONS = [
    {
        "ref": "TVF-CONV-01",
        "folder": "13-conventions-types",
        "title": "Convention type de cooperation territoriale",
        "public": "Collectivite, EPCI, service public",
        "subtitle": "Formaliser une cooperation entre TVF et un territoire autour d'un diagnostic, d'un projet ou d'une ressource.",
        "object": "Definir les conditions de cooperation autour d'un diagnostic territorial, d'une action pilote, d'une ressource publique ou d'un projet de revitalisation.",
        "clauses": [
            "Perimetre territorial : la convention precise la commune, le quartier, le secteur ou le theme concerne.",
            "Referents : chaque partie designe un referent operationnel et un referent institutionnel.",
            "Methode : observation, qualification, mobilisation des acteurs, instruction, conventionnement, suivi.",
            "Donnees : les donnees territoriales transmises sont utilisees uniquement pour les finalites prevues.",
            "Communication : toute mention publique de la cooperation est validee par les deux parties.",
            "Sortie : la convention peut prevoir une fin anticipee en cas de changement de contexte ou de non-respect des obligations.",
        ],
        "annexes": ["Fiche territoire", "Liste des ressources ciblees", "Calendrier", "Contacts referents", "Indicateurs de suivi"],
    },
    {
        "ref": "TVF-CONV-02",
        "folder": "13-conventions-types",
        "title": "Convention type de mise a disposition d'un bien",
        "public": "Proprietaire, bailleur, personne morale, collectivite",
        "subtitle": "Encadrer l'etude ou l'usage temporaire d'un logement, commerce, local, batiment, terrain ou friche.",
        "object": "Definir le cadre d'etude, de visite, de securisation, de remise en usage ou d'occupation temporaire d'un bien.",
        "clauses": [
            "Propriete : le proprietaire conserve la pleine propriete du bien.",
            "Usage : le droit d'usage eventuel est limite a l'objet, a la duree et aux conditions fixees dans la convention.",
            "Etat initial : un etat descriptif, des photos et les contraintes connues sont annexes.",
            "Travaux : aucune intervention technique n'est realisee sans accord ecrit, devis, assurance et validation des autorisations necessaires.",
            "Assurance : les assurances applicables sont precisees avant tout acces ou usage.",
            "Restitution : les conditions de sortie, de remise des cles et de restitution du bien sont prevues.",
        ],
        "annexes": ["Justificatif de propriete ou mandat", "Photos", "Etat des lieux", "Diagnostics disponibles", "Scenario d'usage"],
    },
    {
        "ref": "TVF-CONV-03",
        "folder": "13-conventions-types",
        "title": "Convention type de valorisation de materiaux",
        "public": "Entreprise, collectivite, artisan, particulier, association",
        "subtitle": "Encadrer la collecte, le tri, la valorisation et l'affectation de ressources reutilisables.",
        "object": "Qualifier et affecter des materiaux, mobiliers ou equipements a des projets TVF valides.",
        "clauses": [
            "Selection : TVF peut accepter ou refuser les ressources selon leur etat, leur securite, leur utilite et la faisabilite logistique.",
            "Affectation : les ressources acceptees sont integrees a une strategie de valorisation territoriale et non distribuees automatiquement.",
            "TraÃ§abilite : un bordereau liste les quantites, categories, etats, lieux et dates de remise.",
            "Securite : les materiaux dangereux, pollues ou non conformes ne sont pas acceptes sans controle specifique.",
            "Transport : les responsabilites de chargement, transport et dechargement sont precisees.",
            "Communication : la valorisation RSE est possible uniquement selon les termes valides par les parties.",
        ],
        "annexes": ["Bordereau materiaux", "Photos", "Conditions d'enlevement", "Fiche securite si applicable", "Affectation projet"],
    },
    {
        "ref": "TVF-CONV-04",
        "folder": "13-conventions-types",
        "title": "Convention type de cooperation entreprise",
        "public": "Entreprise, artisan, bailleur, logisticien, commerce",
        "subtitle": "Encadrer un partenariat RSE, materiel, logistique, technique ou territorial.",
        "object": "Definir les contributions d'une entreprise et les modalites de cooperation avec TVF.",
        "clauses": [
            "Contribution : nature, volume, calendrier et conditions de mobilisation sont decrits.",
            "Responsabilites : chaque partie reste responsable de ses obligations legales, sociales, fiscales et assurantielles.",
            "Visibilite : la mention de l'entreprise partenaire est encadree et ne vaut pas label officiel.",
            "Confidentialite : les informations non publiques transmises pendant l'instruction restent confidentielles.",
            "Reporting : TVF peut fournir une synthese d'utilisation ou d'impact lorsque le projet le permet.",
            "Fin : les conditions de suspension, renouvellement ou fin de partenariat sont prevues.",
        ],
        "annexes": ["Fiche contribution", "Contacts", "Calendrier", "Elements de communication valides", "Reporting attendu"],
    },
    {
        "ref": "TVF-CONV-05",
        "folder": "13-conventions-types",
        "title": "Convention type de soutien financier ou mecanat",
        "public": "Mecene, fondation, entreprise, financeur public ou prive",
        "subtitle": "Encadrer un soutien financier, materiel ou en competences affecte a une action TVF.",
        "object": "Definir l'objet du soutien, son affectation, les obligations de suivi et les conditions de communication.",
        "clauses": [
            "Affectation : le soutien est affecte a l'objet prevu ou, si la convention le permet, a une action d'interet general proche.",
            "Versement : les montants, echeances, justificatifs et modalites de paiement sont precises.",
            "Suivi : TVF tient une trace des depenses et produit les livrables prevus.",
            "Fiscalite : aucun avantage fiscal n'est garanti sans verification du cadre applicable.",
            "Communication : logo, citation et annonce publique sont soumis a validation prealable.",
            "Restitution : les conditions de remboursement ou reaffectation en cas d'impossibilite sont prevues si necessaire.",
        ],
        "annexes": ["Budget previsionnel", "Note projet", "Calendrier", "Indicateurs", "Modalites de reporting"],
    },
    {
        "ref": "TVF-CONV-06",
        "folder": "13-conventions-types",
        "title": "Charte type benevole et chantier participatif",
        "public": "Benevole, citoyen, structure encadrante, association",
        "subtitle": "Encadrer la participation citoyenne, les missions benevoles et les actions collectives.",
        "object": "Fixer les principes d'engagement, de securite, de respect des personnes et de traÃ§abilite des participations.",
        "clauses": [
            "Engagement : le benevole agit dans le cadre defini par TVF ou par la structure encadrante.",
            "Securite : les consignes, equipements et limites d'intervention doivent etre respectes.",
            "Confidentialite : les informations relatives aux biens, proprietaires, personnes accompagnees ou partenaires ne sont pas diffusees sans autorisation.",
            "Image : toute photo ou video identifiable necessite une autorisation adaptee.",
            "Emargement : les presences peuvent etre tracees pour l'organisation, l'assurance et le suivi d'impact.",
            "Arret : TVF peut interrompre une mission si les conditions de securite, de respect ou de cadre ne sont pas reunies.",
        ],
        "annexes": ["Fiche mission", "Feuille d'emargement", "Consignes securite", "Autorisation image si necessaire", "Referent chantier"],
    },
]


LETTERS = [
    ("Demande de rendez-vous institutionnel", "Collectivite / EPCI", "Presenter TVF et solliciter un rendez-vous de cadrage territorial."),
    ("Demande de local de stockage", "Collectivite, bailleur, entreprise", "Obtenir une mise a disposition ou une visite d'un local compatible avec le stockage de materiaux."),
    ("Demande de contribution materiaux", "Entreprise, artisan, promoteur, commerce", "Proposer une filiere de valorisation pour surplus, invendus, deposes ou mobiliers."),
    ("Demande de vehicule ou appui logistique", "Transporteur, loueur, entreprise, collectivite", "Obtenir un appui ponctuel pour collecte, manutention ou livraison."),
    ("Demande de mecanat ou soutien financier", "Fondation, entreprise, financeur", "Presenter une demande de soutien au lancement ou a un projet qualifie."),
    ("Demande de partenariat insertion", "SIAE, association, organisme formation", "Construire une action encadree autour d'un chantier, d'une formation ou d'un parcours vers l'emploi."),
    ("Relance pieces manquantes", "Tout demandeur", "Demander les documents indispensables pour poursuivre l'instruction."),
    ("Confirmation de rendez-vous", "Tout demandeur", "Confirmer date, objet, pieces attendues et participants."),
]


def add_form_text_area(doc: Document, title: str, prompts: list[str], lines_per_prompt: int = 2) -> None:
    paragraph(doc, title, style="Heading 2")
    for prompt in prompts:
        paragraph(doc, prompt, bold=True, color=NAVY, size=9.7, after=2)
        for _ in range(lines_per_prompt):
            paragraph(doc, "________________________________________________________________________________", after=1)
    paragraph(doc, "", after=2)


def add_common_form_sections(doc: Document, form: dict):
    paragraph(doc, "Objet du formulaire", style="Heading 1")
    paragraph(doc, form["need"])
    add_note_box(
        doc,
        "Important",
        "Ce document ouvre ou complete un dossier TVF. Il ne vaut pas accord, financement, partenariat, convention ou engagement d'intervention. La decision appartient a TVF apres analyse des informations et pieces transmises.",
        fill=LIGHT_GOLD,
    )

    paragraph(doc, "1. Identification administrative du dossier", style="Heading 1")
    add_kv_table(
        doc,
        [
            ("Numero dossier TVF", "TVF-____-____"),
            ("Date de reception", "____ / ____ / ______"),
            ("Canal d'entree", "[ ] Site web   [ ] E-mail   [ ] Telephone   [ ] Rendez-vous   [ ] Courrier   [ ] Autre"),
            ("Referent TVF", "____________________________________________________________"),
            ("Priorite", "[ ] Standard   [ ] A qualifier vite   [ ] Urgent terrain   [ ] Institutionnel"),
            ("Statut", "[ ] Recu   [ ] Pieces demandees   [ ] En instruction   [ ] Rendez-vous   [ ] En attente   [ ] Classe"),
        ],
        header="Zone reservee a TVF",
    )

    paragraph(doc, "2. Identification du demandeur", style="Heading 1")
    rows = [(field, "________________________________________________________________________") for field in form["fields"]]
    add_kv_table(doc, rows, widths=(2.35, 3.9))
    add_kv_table(
        doc,
        [
            ("Qualite du signataire", "[ ] Proprietaire   [ ] Representant habilite   [ ] Service competent   [ ] Dirigeant   [ ] Benevole   [ ] Autre"),
            ("Pouvoir d'agir", "[ ] A verifier   [ ] Justificatif fourni   [ ] Mandat fourni   [ ] Non applicable"),
            ("Mode de reponse souhaite", "[ ] E-mail   [ ] Telephone   [ ] Rendez-vous   [ ] Courrier"),
        ],
        widths=(2.35, 3.9),
    )

    detail_prompts = [
        "Contexte general de la demande : pourquoi la demande est-elle formulee maintenant ?",
        "Besoin principal : que faut-il resoudre, debloquer, verifier ou organiser ?",
        "Lieu, bien, ressource ou territoire concerne : adresse, perimetre, surface, usage actuel, etat connu.",
        "Objectif recherche : diagnostic, rendez-vous, mise en relation, convention, collecte, stockage, financement, action terrain.",
        "Contraintes connues : delai, securite, assurance, acces, autorisations, copropriete, budget, occupation, communication.",
        "Attentes vis-a-vis de TVF : coordination, instruction, orientation, recherche de ressources, mise en relation, preparation de convention.",
    ]
    paragraph(doc, "3. Description detaillee de la demande", style="Heading 1")
    paragraph(doc, "Cette partie doit etre remplie avec precision. Elle permet a TVF de comprendre le sujet, d'eviter les malentendus et de classer le dossier dans le bon parcours.")
    add_form_text_area(doc, "Details a renseigner", detail_prompts, lines_per_prompt=2)

    specific_prompts = {
        "TVF-F-02": ["Programmes publics ou politiques locales concernes", "Services ou acteurs publics deja mobilises", "Decision ou echeance administrative attendue"],
        "TVF-F-03": ["Origine de la vacance ou de l'inutilisation du bien", "Etat technique apparent et travaux connus", "Usage envisageable apres remise en etat"],
        "TVF-F-04": ["Categorie et origine des ressources proposees", "Quantites, dimensions, etat, contraintes de retrait", "Destination ou usage souhaite si connu"],
        "TVF-F-05": ["Type de contribution proposee par l'entreprise", "Objectifs RSE ou territoriaux", "Conditions, limites ou attentes de communication"],
        "TVF-F-06": ["Description du local et conditions d'acces", "Capacite de stockage et contraintes techniques", "Conditions de mise a disposition envisagees"],
        "TVF-F-07": ["Moyens de transport ou manutention disponibles", "Perimetre geographique et disponibilites", "Conditions d'assurance, conducteur et responsabilites"],
        "TVF-F-08": ["Competences, disponibilites et missions souhaitees", "Contraintes de sante, mobilite ou securite", "Experience associative ou professionnelle utile"],
        "TVF-F-09": ["Publics accompagnes et cadre d'intervention", "Encadrement, assurance et competences mobilisables", "Action ou chantier envisage"],
        "TVF-F-10": ["Type de soutien envisage", "Objet ou territoire a soutenir", "Attentes de reporting, justificatifs ou communication"],
        "TVF-F-11": ["Ce qui a ete observe", "Depuis quand la situation semble exister", "Precautions prises pour ne pas entrer dans un lieu prive"],
        "TVF-F-12": ["Objet editorial ou institutionnel de la demande", "Date limite et support de diffusion", "Elements a valider avant publication"],
    }.get(form["ref"], ["Informations complementaires utiles a l'analyse", "Personnes ou structures deja contactees", "Documents ou preuves disponibles"])
    add_form_text_area(doc, "4. Informations specifiques au sujet", specific_prompts, lines_per_prompt=2)

    paragraph(doc, "5. Pieces a fournir ou a verifier", style="Heading 1")
    add_check_table(doc, "Checklist de recevabilite du dossier", form["pieces"])
    add_form_text_area(doc, "Pieces deja transmises / pieces manquantes", ["Liste des pieces recues", "Liste des pieces manquantes", "Observations sur la conformite ou la lisibilite des pieces"], lines_per_prompt=1)

    paragraph(doc, "6. Analyse et qualification TVF", style="Heading 1")
    add_kv_table(
        doc,
        [
            ("Categorie retenue", "[ ] Habitat   [ ] Commerce   [ ] Materiaux   [ ] Friche / terrain   [ ] Insertion   [ ] Collectivite   [ ] Financement   [ ] Autre"),
            ("Niveau de maturite", "[ ] Idee   [ ] Informations partielles   [ ] Dossier complet   [ ] Visite necessaire   [ ] Convention envisageable"),
            ("Interet territorial", "[ ] Fort   [ ] Moyen   [ ] A verifier   [ ] Hors perimetre TVF"),
            ("Risque principal", "[ ] Juridique   [ ] Securite   [ ] Assurance   [ ] Budget   [ ] Donnees personnelles   [ ] Image   [ ] Aucun identifie"),
            ("Suites possibles", form["next"]),
        ],
        widths=(2.0, 4.25),
    )
    add_form_text_area(doc, "Notes internes TVF", ["Analyse du besoin", "Points a clarifier", "Acteurs a contacter", "Decision proposee"], lines_per_prompt=2)

    paragraph(doc, "7. Conditions de traitement et engagements", style="Heading 1")
    add_bullets(doc, form["clauses"])
    add_bullets(
        doc,
        [
            "Le demandeur declare transmettre des informations exactes a sa connaissance et s'engage a signaler toute evolution importante du dossier.",
            "TVF peut demander des pieces complementaires, refuser une demande insuffisamment documentee ou orienter vers un autre interlocuteur si le sujet ne releve pas de son perimetre.",
            "Toute visite, collecte, mise a disposition, action terrain, communication publique ou convention fait l'objet d'une validation ecrite specifique.",
            "Les informations et documents sont conserves uniquement pour l'instruction, le suivi et l'archivage du dossier TVF, dans le respect des regles applicables de confidentialite et de protection des donnees.",
        ],
    )

    paragraph(doc, "8. Decision et validation", style="Heading 1")
    add_kv_table(
        doc,
        [
            ("Decision TVF", "[ ] Poursuivre   [ ] Demander pieces   [ ] Rendez-vous   [ ] Visite   [ ] Convention   [ ] Orienter   [ ] Classer"),
            ("Motif de decision", "____________________________________________________________"),
            ("Pieces a demander", "____________________________________________________________"),
            ("Responsable suite", "____________________________________________________________"),
            ("Date de prochaine action", "____ / ____ / ______"),
        ],
        widths=(1.85, 4.4),
    )

    add_signature_block(doc, left_label="Pour TVF", right_label="Demandeur / structure")

def save_doc(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_form(form: dict) -> Path:
    doc = setup_doc(form["title"], form["ref"], form["subtitle"], form["public"])
    add_common_form_sections(doc, form)
    out = OUT / form["folder"] / f"{form['ref'].lower()}-{safe_name(form['title'])}.docx"
    save_doc(doc, out)
    return out


def build_convention(conv: dict) -> Path:
    doc = setup_doc(conv["title"], conv["ref"], conv["subtitle"], conv["public"], status="Modele de convention a adapter avant signature")
    add_note_box(
        doc,
        "Portee du modele",
        "Ce document est une base de travail. Il doit etre relu, adapte au dossier, complete par les annexes et valide par les personnes habilitees avant signature. Il ne remplace pas un conseil juridique lorsque le dossier comporte un risque particulier.",
        fill=LIGHT_GOLD,
    )

    paragraph(doc, "1. Parties a la convention", style="Heading 1")
    add_kv_table(
        doc,
        [
            ("TERRITOIRES VIVANTS FRANCE", f"Association loi 1901 declaree - RNA {TVF['rna']} - SIREN {TVF['siren']} - siege : {TVF['address']}"),
            ("Representant TVF", "Nom : ____________________________   Fonction : ____________________________"),
            ("Autre partie signataire", "Nom / structure : _________________________________________________"),
            ("Representant habilite", "Nom : ____________________________   Fonction : ____________________________"),
            ("Adresse", "____________________________________________________________"),
            ("Contact operationnel", "E-mail : ____________________________   Telephone : ____________________________"),
        ],
        widths=(1.9, 4.35),
    )

    paragraph(doc, "2. Objet et contexte", style="Heading 1")
    paragraph(doc, conv["object"])
    add_form_text_area(
        doc,
        "Contexte specifique du dossier",
        [
            "Situation initiale : bien, local, ressource, territoire, action ou besoin concerne.",
            "Objectif concret recherche par les parties.",
            "Benefices attendus pour le territoire, les habitants, la structure ou le projet.",
        ],
        lines_per_prompt=2,
    )

    paragraph(doc, "3. Perimetre operationnel", style="Heading 1")
    add_kv_table(
        doc,
        [
            ("Lieu / territoire concerne", "____________________________________________________________"),
            ("Objet precis", "____________________________________________________________"),
            ("Duree", "Du ____ / ____ / ______ au ____ / ____ / ______"),
            ("Duree renouvelable", "[ ] Non   [ ] Oui, selon accord ecrit   [ ] A definir"),
            ("Referent TVF", "____________________________________________________________"),
            ("Referent autre partie", "____________________________________________________________"),
            ("Frequence de suivi", "[ ] Hebdomadaire   [ ] Mensuelle   [ ] A chaque etape   [ ] Sur demande   [ ] Autre"),
        ],
        widths=(1.9, 4.35),
    )

    paragraph(doc, "4. Obligations et engagements de TVF", style="Heading 1")
    add_bullets(
        doc,
        [
            "Instruire le dossier avec prudence, sur la base des informations et pieces disponibles.",
            "Designer un referent operationnel et conserver une trace des echanges, pieces et decisions dans TVF OS.",
            "Informer l'autre partie des points bloquants connus : pieces manquantes, securite, assurance, autorisations, budget ou calendrier.",
            "Ne pas engager de depense, intervention, communication publique, collecte, stockage ou chantier hors du perimetre valide par ecrit.",
            "Produire, lorsque cela est prevu, un compte rendu, un bordereau, un bilan ou une synthese de suivi.",
        ],
    )

    paragraph(doc, "5. Obligations et engagements de l'autre partie", style="Heading 1")
    add_bullets(
        doc,
        [
            "Transmettre des informations exactes, actualisees et utiles a l'instruction du dossier.",
            "Fournir les autorisations, justificatifs, assurances, plans, photos, devis, diagnostics ou documents necessaires lorsque le dossier le requiert.",
            "Designer un interlocuteur habilite a repondre, valider les informations et suivre l'execution de la convention.",
            "Ne pas presenter TVF comme financeur, executant ou garant d'un resultat tant qu'aucun engagement ecrit ne le prevoit.",
            "Informer TVF de tout changement pouvant modifier le perimetre, le calendrier, la securite, le budget ou les responsabilites.",
        ],
    )

    paragraph(doc, "6. Clauses particulieres au sujet", style="Heading 1")
    add_bullets(doc, conv["clauses"])
    add_form_text_area(
        doc,
        "Conditions particulieres a completer",
        [
            "Modalites precises d'acces, d'usage, de collecte, de stockage, de transport, de financement ou d'intervention.",
            "Limites, reserves, exclusions ou conditions suspensives.",
            "Pieces indispensables avant execution.",
        ],
        lines_per_prompt=2,
    )

    paragraph(doc, "7. Responsabilites, assurance et securite", style="Heading 1")
    add_check_table(
        doc,
        "Controle prealable obligatoire",
        [
            ("Pouvoir de signature / autorisation", "Obligatoire", "Verifier la personne habilitee a signer."),
            ("Assurance de TVF", "A verifier", "Identifier les garanties mobilisables."),
            ("Assurance de l'autre partie", "Obligatoire", "Bien, local, vehicule, chantier, activite ou responsabilite civile."),
            ("Securite / acces", "Obligatoire", "Aucune action sans conditions de securite suffisantes."),
            ("Autorisations administratives", "A verifier", "Urbanisme, ERP, copropriete, voirie, occupation, travaux."),
            ("Budget / couts", "A verifier", "Identifier depenses, charges, fournitures, transport et prise en charge."),
            ("Donnees personnelles", "A verifier", "Limiter les donnees au strict necessaire et proteger les pieces sensibles."),
        ],
    )
    paragraph(doc, "Chaque partie demeure responsable de ses obligations legales, fiscales, sociales, assurantielles, administratives et techniques. La convention ne transfere aucune responsabilite non expressement prevue.")

    paragraph(doc, "8. Suivi, tracabilite et documents", style="Heading 1")
    add_kv_table(
        doc,
        [
            ("Documents a produire", "[ ] Compte rendu   [ ] Bordereau   [ ] Photos   [ ] Inventaire   [ ] Budget   [ ] Bilan   [ ] Autre"),
            ("Lieu d'archivage", "TVF OS - Dossier n degre : TVF-____-____"),
            ("Indicateurs suivis", "____________________________________________________________"),
            ("Livrables attendus", "____________________________________________________________"),
            ("Bilan intermediaire", "Date : ____ / ____ / ______"),
            ("Bilan final", "Date : ____ / ____ / ______"),
        ],
        widths=(1.9, 4.35),
    )

    paragraph(doc, "9. Confidentialite, donnees et communication", style="Heading 1")
    add_bullets(
        doc,
        [
            "Les informations non publiques transmises dans le cadre du dossier restent confidentielles, sauf accord ecrit ou obligation legale.",
            "Les donnees personnelles sont traitees uniquement pour la gestion, l'instruction, le suivi et l'archivage du dossier.",
            "Toute utilisation de nom, logo, image, citation, photo, video ou reference publique doit etre validee avant diffusion.",
            "Aucune mention de partenariat, soutien, financement ou projet commun ne doit etre publiee sans accord ecrit des parties concernees.",
        ],
    )

    paragraph(doc, "10. Suspension, resiliation et cloture", style="Heading 1")
    add_bullets(
        doc,
        [
            "La convention peut etre suspendue si les conditions de securite, d'assurance, d'autorisation, de financement ou de confiance ne sont plus reunies.",
            "Chaque partie peut demander une revision ou une cloture anticipee selon les modalites precisees aux conditions particulieres.",
            "La cloture donne lieu, lorsque necessaire, a un proces-verbal, un bilan, un inventaire, une restitution de cles, materiaux, documents ou equipements.",
            "Les obligations de confidentialite, de protection des donnees, de restitution et de communication continuent a s'appliquer apres la fin de la convention lorsque leur nature l'exige.",
        ],
    )

    paragraph(doc, "11. Reglement des difficultes", style="Heading 1")
    paragraph(doc, "Les parties recherchent prioritairement une solution amiable en cas de difficulte d'interpretation, d'execution ou de cloture. Toute difficulte persistante doit etre documentee par ecrit avec les pieces utiles avant decision de suite.")

    paragraph(doc, "12. Annexes", style="Heading 1")
    add_bullets(doc, conv["annexes"])
    add_form_text_area(doc, "Annexes jointes au present dossier", ["Liste des annexes jointes", "Pieces attendues mais non encore recues", "Reservations ou conditions de validite"], lines_per_prompt=1)

    add_signature_block(doc, left_label="Pour TERRITOIRES VIVANTS FRANCE", right_label="Pour l'autre partie signataire")
    out = OUT / conv["folder"] / f"{conv['ref'].lower()}-{safe_name(conv['title'])}.docx"
    save_doc(doc, out)
    return out

def build_piece_list() -> Path:
    doc = setup_doc(
        "Liste des pieces a fournir selon le type de demande",
        "TVF-LP-01",
        "Checklist transversale pour preparer un dossier complet avant instruction.",
        "Demandeurs et equipe TVF",
    )
    paragraph(doc, "Utilisation", style="Heading 1")
    paragraph(doc, "Cette liste permet de verifier rapidement si un dossier est complet avant rendez-vous, visite, convention ou classement.")
    for form in FORMS:
        paragraph(doc, f"{form['ref']} - {form['title']}", style="Heading 2")
        add_check_table(doc, "Pieces attendues", form["pieces"])
    add_note_box(doc, "Regle TVF", "Un dossier incomplet peut etre conserve en brouillon, mais il ne doit pas etre presente comme retenu tant que les pieces essentielles ne sont pas fournies.", fill=LIGHT_GOLD)
    out = OUT / "14-listes-pieces" / "tvf-lp-01-liste-pieces-par-demande.docx"
    save_doc(doc, out)
    return out


def build_letters() -> Path:
    doc = setup_doc(
        "Courriers types prets a remplir",
        "TVF-COUR-01",
        "Modeles de courriers pour contacter collectivites, entreprises, financeurs, structures d'insertion et partenaires.",
        "Equipe TVF",
    )
    paragraph(doc, "Mode d'emploi", style="Heading 1")
    add_bullets(
        doc,
        [
            "Remplacer les champs entre crochets avant envoi.",
            "Joindre la fiche TVF adaptee au besoin si necessaire.",
            "Eviter toute promesse de partenariat, financement ou resultat tant qu'aucun accord n'est signe.",
            "Archiver le courrier envoye, la reponse et les pieces jointes dans TVF OS.",
        ],
    )
    for idx, (title, target, purpose) in enumerate(LETTERS, 1):
        paragraph(doc, f"Modele {idx} - {title}", style="Heading 1")
        add_kv_table(
            doc,
            [
                ("Destinataire", target),
                ("Objectif", purpose),
                ("Piece jointe conseillee", "Fiche presentation TVF, formulaire adapte, liste des pieces ou note projet."),
            ],
            widths=(1.7, 4.55),
        )
        paragraph(doc, "[Ville], le ____ / ____ / ______", after=4)
        paragraph(doc, "Objet : " + title, bold=True, color=NAVY)
        paragraph(doc, "Madame, Monsieur,")
        paragraph(
            doc,
            "Territoires Vivants France est une association nationale declaree, basee a Saint-Etienne, dont l'objet est de contribuer a la remise en usage de biens, ressources et espaces aujourd'hui inutilises : logements vacants, commerces fermes, friches, locaux, materiaux, equipements ou competences disponibles.",
        )
        paragraph(
            doc,
            "Dans le cadre du lancement de TVF sur le territoire, nous souhaitons echanger avec vous afin d'etudier une cooperation possible autour du besoin suivant : [decrire le besoin precis : local, materiaux, financement, chantier, expertise, rendez-vous, etc.].",
        )
        paragraph(
            doc,
            "L'objectif est de verifier, de maniere concrete et responsable, si une contribution peut etre mobilisee au service d'un projet utile aux habitants, aux associations, aux collectivites ou aux acteurs economiques locaux.",
        )
        paragraph(
            doc,
            "Nous vous proposons un premier echange de cadrage afin de presenter la demarche, identifier les conditions de faisabilite, lister les pieces necessaires et definir les suites possibles.",
        )
        paragraph(doc, "Nous restons disponibles aux coordonnees suivantes : contact@territoiresvivantsfrance.fr - 04 65 81 54 69.")
        paragraph(doc, "Veuillez agreer, Madame, Monsieur, l'expression de nos salutations distinguees.")
        paragraph(doc, "\nPour TERRITOIRES VIVANTS FRANCE\nNom : ____________________________\nFonction : ________________________\nSignature :\n\n")
    out = OUT / "15-courriers-prets-a-envoyer" / "tvf-cour-01-courriers-types-demandes.docx"
    save_doc(doc, out)
    return out


def build_index(paths: list[Path]) -> Path:
    doc = setup_doc(
        "Index de la bibliotheque operationnelle TVF",
        "TVF-INDEX-01",
        "Repertoire des formulaires, conventions, listes de pieces et courriers prets a utiliser.",
        "Equipe TVF",
        status="Repertoire de classement",
    )
    paragraph(doc, "Objectif de la bibliotheque", style="Heading 1")
    paragraph(
        doc,
        "Cette bibliotheque regroupe les modeles documentaires utiles pour receptionner une demande, qualifier un besoin, demander les pieces, instruire le dossier, preparer une convention et transmettre un courrier adapte.",
    )
    paragraph(doc, "Classement recommande", style="Heading 1")
    table = doc.add_table(rows=1, cols=5)
    set_table_width(table, [0.75, 1.8, 1.75, 1.15, 1.05])
    headers = ["Ref.", "Document", "Public", "Usage", "Dossier"]
    for i, h in enumerate(headers):
        set_cell_shading(table.cell(0, i), LIGHT_BLUE)
        add_run(table.cell(0, i).paragraphs[0], h, bold=True, color=NAVY, size=8.5)
    for form in FORMS:
        row = table.add_row().cells
        values = [form["ref"], form["title"], form["public"], form["next"], form["folder"]]
        for i, value in enumerate(values):
            add_run(row[i].paragraphs[0], value, color=DARK, size=8)
            set_cell_margins(row[i], top=80, bottom=80)
            set_cell_border(row[i])
    for conv in CONVENTIONS:
        row = table.add_row().cells
        values = [conv["ref"], conv["title"], conv["public"], "Convention type", conv["folder"]]
        for i, value in enumerate(values):
            add_run(row[i].paragraphs[0], value, color=DARK, size=8)
            set_cell_margins(row[i], top=80, bottom=80)
            set_cell_border(row[i])
    paragraph(doc, "Regle de gestion", style="Heading 1")
    add_bullets(
        doc,
        [
            "Un formulaire ouvre un dossier TVF.",
            "La liste des pieces verifie la recevabilite.",
            "La fiche d'instruction trace la decision.",
            "La convention formalise uniquement les dossiers retenus.",
            "Les courriers servent aux prises de contact et aux relances.",
        ],
    )
    paragraph(doc, "Documents disponibles dans la bibliotheque", style="Heading 2")
    for path in sorted(paths):
        rel = path.relative_to(OUT)
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, str(rel), color=DARK, size=9)
    out = OUT / "00-index" / "tvf-index-bibliotheque-formulaires-conventions.docx"
    save_doc(doc, out)
    return out


def build_readme(paths: list[Path]) -> None:
    lines = [
        "# Bibliotheque formulaires et conventions TVF",
        "",
        "Ce dossier regroupe les documents prets a remplir pour la gestion operationnelle de Territoires Vivants France.",
        "",
        "## Organisation",
        "",
        "- `00-index` : repertoire general de la bibliotheque.",
        "- `01` a `12` : formulaires par type de demande.",
        "- `13-conventions-types` : conventions types par acteur ou besoin.",
        "- `14-listes-pieces` : checklist globale des pieces a fournir.",
        "- `15-courriers-prets-a-envoyer` : courriers types a adapter et envoyer.",
        "",
        "## Regle d'utilisation",
        "",
        "Les documents sont modifiables. Avant signature ou transmission engageante, adapter le contenu au dossier, verifier les pieces, l'assurance, les autorisations et, si necessaire, demander une validation juridique.",
        "",
        "## Fichiers",
        "",
    ]
    for path in sorted(paths):
        lines.append(f"- `{path.relative_to(OUT)}`")
    (OUT / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    generated: list[Path] = []
    for form in FORMS:
        generated.append(build_form(form))
    for conv in CONVENTIONS:
        generated.append(build_convention(conv))
    generated.append(build_piece_list())
    generated.append(build_letters())
    index = build_index(generated)
    generated.append(index)
    build_readme(generated)
    print("Generated documents:")
    for path in generated:
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()

