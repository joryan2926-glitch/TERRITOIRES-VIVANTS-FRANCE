from pathlib import Path
from runpy import run_path
from html import unescape
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
BLUE = BASE["BLUE"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
GOLD = BASE["GOLD"]
BORDER = BASE["BORDER"]

OUT_DIR = ROOT / "documents" / "brochures"
SRC_DIR = ROOT / "documents" / "sources"
PUBLIC_DIR = ROOT / "output" / "documents"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)
PUBLIC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "liste-pieces-a-fournir-particuliers-tvf.docx"
PDF_PATH = PUBLIC_DIR / "liste-pieces-a-fournir-particuliers-tvf.pdf"
MD_PATH = SRC_DIR / "liste-pieces-a-fournir-particuliers-tvf.md"
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"


def u(text: str) -> str:
    return unescape(text)


def run_style(run, size=None, color=None, bold=None, italic=None):
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def small_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(u(text))
    run_style(r, 8.4, MUTED, italic=True)


def section_tag(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(u(label.upper()))
    run_style(r, 8.2, GOLD, bold=True)
    return p


def field_line(doc, label, width_hint=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4.2)
    r = p.add_run(u(f"{label} : "))
    run_style(r, 9.2, GREEN, bold=True)
    r2 = p.add_run("_" * (68 if not width_hint else 50))
    run_style(r2, 9.2, MUTED)
    return p


def checklist_table(doc, title, rows, intro=None):
    if title:
        doc.add_heading(u(title), level=2)
    if intro:
        add_paragraph(doc, intro)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [0.85, 5.35, 7.4, 3.2])
    headers = ["OK", "Pièce", "À quoi sert-elle ?", "Statut TVF"]
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(u(header))
        run_style(r, 8.4, GREEN, bold=True)

    for item, purpose, status in rows:
        cells = table.add_row().cells
        values = ["[ ]", item, purpose, status]
        for idx, value in enumerate(values):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, color="DDE6DE", size="6")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.03
            r = p.add_run(u(value))
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_style(r, 9.3, GREEN, bold=True)
            elif idx == 1:
                run_style(r, 8.25, BLUE, bold=True)
            else:
                run_style(r, 7.85, RGBColor(31, 45, 38))
    doc.add_paragraph()
    return table


def two_col_table(doc, rows, widths=(5.2, 11.9), fill=LIGHT_BLUE):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, list(widths))
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, color="DDE6DE", size="6")
        set_cell_shading(table.cell(i, 0), fill)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(u(label))
        run_style(r, 8.6, GREEN, bold=True)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(u(value))
        run_style(r, 8.45, BLUE)
    doc.add_paragraph()
    return table


def cover(doc):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Cm(5.4))
        p_logo.paragraph_format.space_after = Pt(10)

    section_tag(doc, "Brochure opérationnelle")
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run(u("Particuliers et propriétaires"))
    run_style(r, 22, BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run(u("Liste des pièces à fournir pour l'étude d'une demande TVF"))
    run_style(r, 15, GREEN, bold=True)

    add_paragraph(
        doc,
        "Ce document aide les particuliers à préparer un dossier complet lorsqu'ils souhaitent proposer "
        "un logement vacant, un local, un bâtiment, un terrain, une friche, des matériaux ou un équipement "
        "pouvant être étudié par Territoires Vivants France.",
    )
    add_callout(
        doc,
        "Important",
        "Le dépôt d'un dossier ne vaut pas acceptation automatique, ne vaut pas décision de rénovation "
        "et ne crée aucun engagement financier ou opérationnel de TVF. Chaque demande fait l'objet "
        "d'une phase de qualification préalable.",
        fill=LIGHT_GREEN,
    )
    two_col_table(
        doc,
        [
            ("Document", "Liste des pièces à fournir - particuliers"),
            ("Usage", "Préparation d'un dossier de pré-étude, de visite ou d'orientation"),
            ("Version", "Juillet 2026"),
            ("Contact", "contact@territoiresvivantsfrance.fr - 04 65 81 54 69"),
            ("Siège", "25 rue Élise Gervais, 42000 Saint-Étienne"),
            ("Identification", "RNA W423016361 - SIREN 107 226 128 - SIRET 107 226 128 00018"),
        ],
    )
    small_note(
        doc,
        "Ne transmettez jamais d'originaux. TVF peut demander des informations complémentaires si le dossier "
        "est incomplet ou si une orientation technique, juridique ou administrative est nécessaire.",
    )


def intro(doc):
    doc.add_page_break()
    doc.add_heading("1. À quoi sert cette liste ?", level=1)
    add_paragraph(
        doc,
        "La qualité d'un dossier conditionne la capacité de TVF à comprendre la situation, à identifier les "
        "bons interlocuteurs, à évaluer la faisabilité et à décider si une suite peut être donnée. Un dossier "
        "bien préparé permet d'éviter les échanges incomplets, de gagner du temps et de sécuriser les premières "
        "étapes de dialogue.",
    )
    add_paragraph(
        doc,
        "Cette liste n'a pas pour objet de compliquer la démarche. Elle sert à vérifier les éléments essentiels : "
        "qui est habilité à présenter la demande, quel est le bien ou la ressource concernée, dans quel état il se "
        "trouve, quelles contraintes existent, quels documents sont déjà disponibles et quelles suites peuvent être "
        "raisonnablement étudiées.",
    )
    checklist_table(
        doc,
        "Lecture rapide du parcours",
        [
            ("Identification du demandeur", "Vérifier que TVF échange avec la bonne personne.", "Obligatoire"),
            ("Description du bien ou de la ressource", "Comprendre le sujet avant toute visite.", "Obligatoire"),
            ("Documents de propriété ou d'autorisation", "S'assurer que le demandeur peut engager la discussion.", "Obligatoire"),
            ("Photos et informations techniques", "Évaluer le niveau de faisabilité et les risques.", "Fortement recommandé"),
            ("Pièces complémentaires", "Préparer une instruction plus complète si le dossier avance.", "Selon le cas"),
        ],
    )
    add_callout(
        doc,
        "Principe TVF",
        "TVF étudie les dossiers de manière progressive : une première demande peut être simple, mais une suite "
        "opérationnelle exige des pièces plus précises, des autorisations claires et un cadre écrit.",
        fill=LIGHT_BLUE,
    )


def requester_form(doc):
    doc.add_heading("2. Fiche d'identification du demandeur", level=1)
    add_paragraph(
        doc,
        "Cette fiche permet d'ouvrir un premier échange. Elle doit être remplie par le propriétaire, un représentant "
        "habilité, un mandataire, un membre d'indivision autorisé ou une personne disposant d'un accord écrit.",
    )
    for label in [
        "Nom et prénom",
        "Adresse postale",
        "Téléphone",
        "Adresse e-mail",
        "Lien avec le bien ou la ressource",
        "Commune concernée",
        "Adresse du bien ou lieu de stockage",
    ]:
        field_line(doc, label)
    two_col_table(
        doc,
        [
            ("Situation du demandeur", "[ ] Propriétaire  [ ] Indivisaire  [ ] Héritier  [ ] Mandataire  [ ] Autre : __________________"),
            ("Type de demande", "[ ] Logement  [ ] Immeuble  [ ] Commerce/local  [ ] Terrain/friche  [ ] Matériaux  [ ] Équipement"),
            ("Disponibilité pour échange", "[ ] Téléphone  [ ] Rendez-vous  [ ] Visite sur site  [ ] Visioconférence"),
            ("Urgence signalée", "[ ] Aucune  [ ] Sécurité  [ ] Dégradation  [ ] Succession  [ ] Vente envisagée  [ ] Autre"),
        ],
    )


def common_documents(doc):
    doc.add_heading("3. Pièces communes à toutes les demandes", level=1)
    add_paragraph(
        doc,
        "Les pièces communes permettent d'établir un minimum de traçabilité. Elles sont demandées quelle que soit "
        "la nature du dossier : bien immobilier, local, terrain, matériaux ou équipement.",
    )
    checklist_table(
        doc,
        None,
        [
            ("Pièce d'identité du demandeur", "Identifier l'interlocuteur principal du dossier.", "Obligatoire"),
            ("Justificatif de propriété ou d'autorisation", "Vérifier que la personne est habilitée à déposer la demande.", "Obligatoire"),
            ("Adresse exacte du bien ou du lieu", "Localiser le sujet et préparer l'analyse territoriale.", "Obligatoire"),
            ("Description écrite de la situation", "Comprendre l'historique, l'état actuel et la demande formulée.", "Obligatoire"),
            ("Photos récentes", "Visualiser le bien, l'accès, l'environnement et les points de vigilance.", "Obligatoire si possible"),
            ("Coordonnées d'une personne disponible", "Organiser les échanges et, si nécessaire, une visite.", "Obligatoire"),
            ("Contraintes connues", "Sécurité, humidité, copropriété, voisinage, servitude, occupation, accès.", "Recommandé"),
            ("Documents déjà disponibles", "Plans, diagnostics, devis, factures, courriers administratifs, PV.", "Selon le cas"),
        ],
    )
    small_note(
        doc,
        "TVF peut refuser d'étudier un dossier si l'identité du demandeur, l'autorisation de discussion ou la localisation "
        "du bien ne sont pas suffisamment établies.",
    )


def real_estate_sections(doc):
    doc.add_page_break()
    doc.add_heading("4. Pièces selon le type de bien", level=1)
    checklist_table(
        doc,
        "A. Logement vacant, maison ou immeuble",
        [
            ("Titre de propriété, attestation notariale ou taxe foncière", "Confirmer le lien juridique avec le bien.", "Obligatoire"),
            ("Surface approximative et nombre de pièces", "Évaluer les usages possibles et le niveau de travaux.", "Obligatoire"),
            ("État d'occupation", "Vacant, occupé, squatté, partiellement utilisé, en succession.", "Obligatoire"),
            ("Diagnostics disponibles", "DPE, amiante, plomb, électricité, gaz, termites si déjà réalisés.", "Si disponible"),
            ("Photos intérieures et extérieures", "Identifier l'état général, les accès et les risques visibles.", "Obligatoire si accès possible"),
            ("Plans, croquis ou surfaces cadastrales", "Préparer la lecture technique et fonctionnelle.", "Recommandé"),
            ("Devis, factures ou travaux déjà réalisés", "Comprendre l'historique et les coûts déjà engagés.", "Si disponible"),
            ("Documents de copropriété", "Syndic, charges, règlement, PV d'assemblée, impayés éventuels.", "Si copropriété"),
            ("Arrêtés ou procédures administratives", "Péril, insalubrité, mise en sécurité, contentieux.", "Si concerné"),
        ],
    )
    checklist_table(
        doc,
        "B. Commerce, local d'activité ou rez-de-chaussée vacant",
        [
            ("Justificatif de propriété ou autorisation du bailleur", "Vérifier la capacité à étudier un usage.", "Obligatoire"),
            ("Adresse et emplacement", "Comprendre le contexte de rue, quartier, centre-ville ou zone d'activité.", "Obligatoire"),
            ("Surface, vitrine, accès et état du local", "Qualifier la compatibilité avec des usages possibles.", "Obligatoire"),
            ("Ancienne activité exercée", "Comprendre les contraintes, installations et usages antérieurs.", "Recommandé"),
            ("Destination du local", "Commerce, bureau, atelier, stockage, ERP éventuel.", "Si connue"),
            ("Informations techniques", "Électricité, eau, extraction, chauffage, sanitaires, accessibilité.", "Si disponible"),
            ("Montant de loyer envisagé ou charges", "Évaluer les scénarios économiques réalistes.", "Si souhaité"),
            ("Photos de la façade et de l'intérieur", "Apprécier visibilité, état et contraintes de remise en usage.", "Obligatoire si possible"),
        ],
    )
    checklist_table(
        doc,
        "C. Terrain, cour, friche ou espace extérieur",
        [
            ("Référence cadastrale ou plan de situation", "Localiser précisément le terrain.", "Obligatoire"),
            ("Justificatif de propriété ou autorisation", "Vérifier la capacité à déposer une demande.", "Obligatoire"),
            ("Surface approximative", "Évaluer les usages possibles : jardin, stockage, animation, biodiversité.", "Obligatoire"),
            ("Accès et clôture", "Identifier les conditions de sécurité et d'intervention.", "Recommandé"),
            ("Réseaux disponibles", "Eau, électricité, assainissement, voirie, accès véhicule.", "Si connu"),
            ("Contraintes connues", "Pollution, servitudes, voisinage, pente, risques, urbanisme.", "Si connu"),
            ("Photos générales et détails", "Voir l'état, les limites, les accès et les points sensibles.", "Obligatoire si possible"),
        ],
    )


def materials_sections(doc):
    doc.add_heading("5. Pièces pour matériaux, mobilier ou équipements", level=1)
    add_paragraph(
        doc,
        "Les matériaux proposés ne sont pas distribués automatiquement. Ils sont qualifiés par TVF afin de vérifier "
        "leur intérêt, leur état, leurs contraintes de stockage, leur origine et leur affectation possible vers un projet utile.",
    )
    checklist_table(
        doc,
        "Matériaux ou équipements proposés par un particulier",
        [
            ("Liste détaillée", "Indiquer le type : bois, portes, fenêtres, sanitaires, luminaires, mobilier, outils.", "Obligatoire"),
            ("Quantité et dimensions", "Prévoir transport, stockage et affectation possible.", "Obligatoire"),
            ("État général", "Neuf, très bon état, réutilisable, à réparer, incomplet.", "Obligatoire"),
            ("Photos nettes", "Vérifier la nature, l'état et l'intérêt des éléments.", "Obligatoire"),
            ("Lieu de stockage", "Préparer un enlèvement ou une visite éventuelle.", "Obligatoire"),
            ("Date limite de disponibilité", "Éviter les pertes et organiser la logistique.", "Recommandé"),
            ("Contraintes de manutention", "Étage, poids, accès, véhicule nécessaire, démontage.", "Recommandé"),
            ("Origine ou facture si disponible", "Rassurer sur la provenance et la traçabilité.", "Si disponible"),
        ],
    )
    add_callout(
        doc,
        "Matériaux non acceptés sans vérification spécifique",
        "Produits dangereux, matériaux amiantés ou suspectés, déchets souillés, éléments non démontables, produits chimiques, "
        "équipements électriques non identifiables ou tout élément présentant un risque de sécurité peuvent être refusés.",
        fill=LIGHT_GREEN,
    )


def project_details(doc):
    doc.add_page_break()
    doc.add_heading("6. Décrire le sujet de la demande", level=1)
    add_paragraph(
        doc,
        "Au-delà des pièces administratives, TVF a besoin de comprendre la demande. Le propriétaire ou le particulier "
        "doit expliquer ce qu'il souhaite, ce qu'il accepte d'étudier et ce qu'il ne souhaite pas. Cette partie permet "
        "d'éviter les malentendus avant toute visite ou instruction.",
    )
    for label in [
        "Résumé de la situation en quelques lignes",
        "Depuis quand le bien ou la ressource est inutilisé",
        "Pourquoi la situation est bloquée aujourd'hui",
        "Ce que vous aimeriez voir devenir possible",
        "Ce que vous ne souhaitez pas envisager",
        "Contraintes importantes à connaître",
        "Personnes à associer ou à informer",
    ]:
        field_line(doc, label)
    two_col_table(
        doc,
        [
            ("Usages possibles à étudier", "[ ] Logement  [ ] Association  [ ] Étudiant  [ ] Atelier  [ ] Commerce  [ ] Formation  [ ] Jardin  [ ] Autre"),
            ("Cadre envisageable", "[ ] Mise à disposition  [ ] Convention d'usage  [ ] Loyer solidaire  [ ] Partage de recettes  [ ] À étudier"),
            ("Durée envisageable", "[ ] Moins d'un an  [ ] 1 à 3 ans  [ ] 3 à 5 ans  [ ] 5 ans et plus  [ ] À définir"),
            ("Niveau de travaux accepté", "[ ] Aucun  [ ] Entretien  [ ] Rénovation légère  [ ] Rénovation moyenne  [ ] À diagnostiquer"),
        ],
    )
    small_note(
        doc,
        "TVF n'intervient pas comme entreprise générale du bâtiment. Les travaux lourds, structurels ou réglementés "
        "nécessitent des professionnels compétents, des autorisations et un cadre spécifique.",
    )


def completeness(doc):
    doc.add_heading("7. Grille de complétude du dossier", level=1)
    add_paragraph(
        doc,
        "Cette grille peut être utilisée par le demandeur avant l'envoi, puis par TVF lors de la réception du dossier. "
        "Elle permet de savoir rapidement si une demande est exploitable ou si des informations doivent être complétées.",
    )
    checklist_table(
        doc,
        None,
        [
            ("Identité du demandeur complète", "Nom, coordonnées et lien avec le bien ou la ressource.", "À vérifier"),
            ("Autorisation ou justificatif transmis", "Preuve de propriété, mandat ou accord écrit.", "À vérifier"),
            ("Adresse et localisation précises", "Adresse, commune, photos ou plan de situation.", "À vérifier"),
            ("Description du sujet suffisante", "État, historique, blocage, demande formulée.", "À vérifier"),
            ("Photos exploitables", "Photos récentes et lisibles, intérieur/extérieur si possible.", "À vérifier"),
            ("Documents techniques disponibles", "Diagnostics, plans, devis, factures, PV, arrêtés, selon le cas.", "À compléter"),
            ("Contraintes identifiées", "Sécurité, accès, voisinage, copropriété, servitudes, délais.", "À compléter"),
            ("Suite demandée claire", "Information, visite, orientation, étude, convention, matériaux.", "À vérifier"),
        ],
    )
    two_col_table(
        doc,
        [
            ("Décision de réception", "[ ] Dossier incomplet  [ ] Dossier recevable  [ ] Visite à étudier  [ ] Orientation vers professionnel"),
            ("Numéro interne TVF", "TVF-__________ / Année : ______ / Commune : __________________________"),
            ("Référent TVF", "Nom : __________________________  Date de réception : ____ / ____ / ______"),
        ],
    )


def attestation(doc):
    doc.add_page_break()
    doc.add_heading("8. Attestation du demandeur", level=1)
    add_paragraph(
        doc,
        "Cette attestation permet de confirmer que le dossier est déposé de bonne foi et que le demandeur comprend "
        "le cadre de pré-étude. Elle peut être signée au moment du dépôt ou lors d'un premier rendez-vous.",
    )
    checklist_table(
        doc,
        "Déclarations à cocher",
        [
            ("Je certifie l'exactitude des informations transmises.", "Les informations remises à TVF doivent être sincères.", "À cocher"),
            ("Je suis propriétaire ou autorisé à déposer cette demande.", "TVF ne peut pas instruire une demande sans base d'autorisation.", "À cocher"),
            ("Je comprends que le dépôt ne vaut pas acceptation.", "Aucune rénovation ou convention n'est automatique.", "À cocher"),
            ("J'accepte d'être recontacté par TVF pour compléter le dossier.", "Les échanges servent uniquement à l'étude de la demande.", "À cocher"),
            ("J'accepte qu'une visite puisse être proposée si le dossier le justifie.", "La visite reste soumise à accord préalable et conditions de sécurité.", "À cocher"),
        ],
    )
    field_line(doc, "Fait à")
    field_line(doc, "Le")
    field_line(doc, "Nom et prénom")
    field_line(doc, "Signature")
    add_callout(
        doc,
        "Protection des informations",
        "Les documents transmis sont utilisés pour l'étude de la demande et la gestion du dossier TVF. Les pièces sensibles "
        "ne doivent pas être transmises si elles ne sont pas nécessaires. Le demandeur peut solliciter TVF pour connaître "
        "les informations conservées et demander leur mise à jour ou leur suppression dans les conditions applicables.",
        fill=LIGHT_BLUE,
    )


def transmission(doc):
    doc.add_heading("9. Transmission du dossier", level=1)
    add_paragraph(
        doc,
        "Le dossier peut être transmis par voie numérique ou remis lors d'un rendez-vous. Pour faciliter le traitement, "
        "il est recommandé de regrouper les pièces dans un seul dossier et de nommer les fichiers clairement.",
    )
    two_col_table(
        doc,
        [
            ("Objet conseillé du mail", "Dossier particulier - commune - type de demande"),
            ("Adresse e-mail", "contact@territoiresvivantsfrance.fr"),
            ("Téléphone", "04 65 81 54 69"),
            ("Adresse postale", "Territoires Vivants France, 25 rue Élise Gervais, 42000 Saint-Étienne"),
            ("Site", "www.territoiresvivantsfrance.fr"),
        ],
    )
    doc.add_heading("Nommage conseillé des fichiers", level=2)
    for item in [
        "01-identite-nom-prenom.pdf",
        "02-justificatif-propriete.pdf",
        "03-description-bien-commune.pdf",
        "04-photos-exterieur.zip",
        "05-photos-interieur.zip",
        "06-diagnostics-si-disponibles.pdf",
        "07-plans-devis-factures-si-disponibles.pdf",
    ]:
        add_bullet(doc, item)
    small_note(
        doc,
        "En cas de doute sur une pièce, il est préférable de transmettre une question à TVF plutôt que d'envoyer des "
        "documents inutiles ou sensibles.",
    )


def write_source_markdown():
    MD_PATH.write_text(
        "# Particuliers et propriétaires - Liste des pièces à fournir\n\n"
        "Document opérationnel TVF destiné à préparer l'étude d'une demande : logement vacant, local, bâtiment, terrain, friche, matériaux ou équipement.\n\n"
        "Sections : identification du demandeur, pièces communes, pièces par type de bien, matériaux, description du sujet, grille de complétude, attestation, transmission.\n",
        encoding="utf-8",
    )


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    cover(doc)
    intro(doc)
    requester_form(doc)
    common_documents(doc)
    real_estate_sections(doc)
    materials_sections(doc)
    project_details(doc)
    completeness(doc)
    attestation(doc)
    transmission(doc)
    doc.save(DOCX_PATH)
    write_source_markdown()
    return DOCX_PATH


if __name__ == "__main__":
    path = build_docx()
    print(path)
