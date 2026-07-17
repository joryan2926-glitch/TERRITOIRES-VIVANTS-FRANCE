from pathlib import Path
from runpy import run_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

OUT_DIR = ROOT / "documents" / "courriers-lancement-saint-etienne"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "31-mission-locale-saint-etienne-lancement-jeunesse-tvf.docx"
MD_PATH = SRC_DIR / "courrier-mission-locale-saint-etienne-lancement-tvf.md"

txt = BASE["txt"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_number = BASE["add_number"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
BORDER = BASE["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        ("Destinataire", "Mission Locale de Saint-&Eacute;tienne - &agrave; l'attention de la direction, des conseillers insertion jeunesse et de l'&eacute;quipe relation entreprises"),
        ("Adresse", "Coordonn&eacute;es op&eacute;rationnelles &agrave; confirmer avant envoi - Saint-&Eacute;tienne"),
        ("Objet", "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les passerelles possibles entre revitalisation territoriale, engagement citoyen, d&eacute;couverte des m&eacute;tiers et parcours jeunesse"),
        ("Pi&egrave;ces jointes", "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche Solidarit&eacute; & Insertion, fiche chantiers encadr&eacute;s, fiche mat&eacute;riaux / r&eacute;emploi, fiche besoins Saint-&Eacute;tienne"),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def add_bridge_table(doc):
    doc.add_heading(txt("Passerelles possibles avec la Mission Locale"), level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.6, 6.3, 6.2])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Sujet", "Apport possible de TVF", "Point &agrave; cadrer avec la Mission Locale"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(head))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.1)
    rows = [
        ("D&eacute;couverte des m&eacute;tiers", "Faire d&eacute;couvrir les m&eacute;tiers li&eacute;s &agrave; la r&eacute;novation, au r&eacute;emploi, &agrave; la logistique, aux espaces verts, &agrave; l'animation locale et aux services de proximit&eacute;.", "D&eacute;finir les formats compatibles : information collective, visite, atelier, immersion courte, rencontre avec des professionnels."),
        ("Chantiers encadr&eacute;s", "Proposer des supports de terrain utiles : tri de mat&eacute;riaux, pr&eacute;paration logistique, entretien d'espaces, actions participatives et petites remises en &eacute;tat.", "S&eacute;curiser l'encadrement, les assurances, les limites de responsabilit&eacute; et l'intervention de professionnels habilit&eacute;s."),
        ("Engagement citoyen", "Orienter certains jeunes vers des actions b&eacute;n&eacute;voles, &eacute;ducatives ou citoyennes autour de la revitalisation de leur quartier.", "Clarifier le cadre : volontariat, dur&eacute;e, accompagnement, objectifs p&eacute;dagogiques et absence de substitution &agrave; un emploi."),
        ("Relation entreprises", "Mettre en lien les jeunes avec des artisans, entreprises, commerces, structures de l'ESS ou acteurs du r&eacute;emploi mobilis&eacute;s autour de TVF.", "Identifier les besoins de recrutement, stages, visites, t&eacute;moignages ou actions de sensibilisation possibles."),
        ("Suivi des parcours", "Conserver une trace des actions propos&eacute;es par TVF : dossier, objectif, partenaire, date, suite donn&eacute;e, retour d'exp&eacute;rience.", "D&eacute;finir ce qui rel&egrave;ve de TVF et ce qui reste exclusivement suivi par la Mission Locale."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.25)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - Mission Locale de Saint-&Eacute;tienne"))
    r.bold = True
    r.font.size = Pt(16.4)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Demande d'un premier &eacute;change sur les jeunes, la d&eacute;couverte des m&eacute;tiers, l'engagement citoyen et les projets de terrain TVF"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(doc, "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF porte une d&eacute;marche appel&eacute;e &agrave; se d&eacute;velopper &agrave; l'&eacute;chelle nationale, avec un premier ancrage op&eacute;rationnel progressif sur son territoire d'origine.")
    add_paragraph(doc, "L'association agit sur les ressources inutilis&eacute;es des territoires : logements vacants ou d&eacute;grad&eacute;s, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. Sa m&eacute;thode vise &agrave; transformer ces ressources en projets utiles pour les habitants, les associations, les collectivit&eacute;s, les entreprises et les acteurs de l'insertion.")
    add_paragraph(doc, "Les missions locales exercent une mission reconnue d'accueil, d'information, d'orientation et d'accompagnement des jeunes rencontrant des difficult&eacute;s d'insertion professionnelle et sociale. C'est pourquoi TVF souhaite vous rencontrer afin d'identifier les passerelles possibles entre les futurs projets de revitalisation et les parcours jeunesse du territoire st&eacute;phanois.")

    add_callout(doc, "Demande principale", "TVF sollicite un premier rendez-vous afin de pr&eacute;senter sa m&eacute;thode, d'identifier les formats utiles pour les jeunes accompagn&eacute;s et de d&eacute;terminer comment les futurs dossiers TVF pourraient servir de supports de d&eacute;couverte, d'engagement, de remobilisation ou de mise en relation avec les m&eacute;tiers locaux.")

    doc.add_heading(txt("Pourquoi associer la Mission Locale"), level=1)
    add_paragraph(doc, "TVF ne con&ccedil;oit pas la revitalisation territoriale comme une action uniquement immobili&egrave;re ou technique. Remettre en usage un lieu, trier des mat&eacute;riaux, r&eacute;activer un commerce ou animer un chantier participatif peut aussi devenir un support concret pour comprendre des m&eacute;tiers, reprendre confiance, rencontrer des professionnels et s'engager dans un projet utile.")
    add_paragraph(doc, "Le lien avec la Mission Locale permettrait de construire un cadre lisible pour les jeunes : quelles actions sont pertinentes, quelles limites respecter, quels partenaires mobiliser, comment &eacute;viter toute confusion entre b&eacute;n&eacute;volat, stage, immersion, emploi, formation ou insertion par l'activit&eacute; &eacute;conomique.")
    add_paragraph(doc, "TVF souhaite avancer avec prudence. Les actions propos&eacute;es ne devront jamais remplacer un accompagnement existant, un emploi, une formation ou une mission encadr&eacute;e par un organisme comp&eacute;tent. L'objectif est de cr&eacute;er des passerelles utiles, simples et s&eacute;curis&eacute;es.")

    add_bridge_table(doc)

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "Identifier le bon interlocuteur Mission Locale pour les sujets engagement, d&eacute;couverte des m&eacute;tiers, relation entreprises et actions de terrain.")
    add_bullet(doc, "D&eacute;terminer les formats compatibles : information collective, visite de chantier, atelier r&eacute;emploi, rencontre avec artisans, action citoyenne ou projet encadr&eacute;.")
    add_bullet(doc, "Cartographier les secteurs m&eacute;tiers pouvant int&eacute;resser les jeunes : b&acirc;timent, r&eacute;novation, logistique, collecte, tri, mobilier, espaces verts, commerce de proximit&eacute;, animation locale.")
    add_bullet(doc, "Pr&eacute;ciser les limites de responsabilit&eacute;, les conditions d'assurance, la protection des personnes, les autorisations et les donn&eacute;es pouvant &ecirc;tre partag&eacute;es.")
    add_bullet(doc, "Identifier un premier cas simple &agrave; Saint-&Eacute;tienne permettant de tester une rencontre ou une action de sensibilisation sans engagement institutionnel pr&eacute;matur&eacute;.")

    doc.add_heading(txt("Ce que TVF peut apporter"), level=1)
    add_number(doc, "Des supports concrets de terrain : lieux vacants, mat&eacute;riaux, commerces, friches, actions de quartier et besoins associatifs.")
    add_number(doc, "Un r&eacute;seau &agrave; construire avec des artisans, entreprises, associations, structures d'insertion, collectivit&eacute;s et habitants.")
    add_number(doc, "Des situations p&eacute;dagogiques pour comprendre les m&eacute;tiers de la r&eacute;novation, du r&eacute;emploi, de la logistique et de l'animation territoriale.")
    add_number(doc, "Une tra&ccedil;abilit&eacute; simple des actions propos&eacute;es : dossier, objectif, date, acteur mobilis&eacute;, suite donn&eacute;e.")
    add_number(doc, "Une contribution &agrave; l'engagement citoyen des jeunes autour de projets visibles et utiles &agrave; leur territoire.")

    add_callout(doc, "Principe de prudence", "TVF ne se substitue pas &agrave; la Mission Locale, aux conseillers, aux organismes de formation, aux employeurs, aux prescripteurs ou aux structures d'insertion. Aucune promesse d'emploi, de formation, de financement, de validation de parcours ou de partenariat public ne sera formul&eacute;e sans cadre &eacute;crit et accord pr&eacute;alable des acteurs comp&eacute;tents.", fill=LIGHT_BLUE)

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(doc, "Nous proposons un rendez-vous d'environ une heure afin de vous pr&eacute;senter TVF, de comprendre les besoins des jeunes accompagn&eacute;s, d'identifier les formats d'action pertinents et de s&eacute;curiser les limites de coop&eacute;ration avant toute initiative de terrain.")
    add_paragraph(doc, "&Agrave; l'issue de cet &eacute;change, TVF pourra formaliser une note courte pr&eacute;cisant les actions envisageables, les pi&egrave;ces utiles, les interlocuteurs &agrave; mobiliser, les responsabilit&eacute;s et les conditions minimales avant tout test local.")

    doc.add_heading("Conclusion", level=1)
    add_paragraph(doc, "Territoires Vivants France souhaite faire de Saint-&Eacute;tienne un territoire pilote dans lequel la revitalisation des lieux peut aussi devenir une opportunit&eacute; humaine, p&eacute;dagogique et citoyenne. Une rencontre avec la Mission Locale permettrait de construire cette dimension avec m&eacute;thode, prudence et utilit&eacute; concr&egrave;te pour les jeunes.")
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(head))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration officielle de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche Solidarit&eacute; & Insertion", "Expliquer les passerelles entre revitalisation, jeunes, chantiers et engagement citoyen.", "&Agrave; joindre"),
        ("Fiche chantiers encadr&eacute;s", "D&eacute;crire les actions de terrain possibles et les limites &agrave; respecter.", "&Agrave; joindre"),
        ("Fiche mat&eacute;riaux / r&eacute;emploi", "D&eacute;crire les besoins autour du tri, de la collecte, du stockage et de la valorisation.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(doc, "Objet d'e-mail recommand&eacute;", "Demande de rendez-vous - TERRITOIRES VIVANTS FRANCE / Mission Locale de Saint-&Eacute;tienne - jeunesse, engagement et d&eacute;couverte des m&eacute;tiers", fill=LIGHT_BLUE)
    doc.add_heading("Point &agrave; confirmer", level=1)
    add_paragraph(doc, "Les coordonn&eacute;es op&eacute;rationnelles exactes du destinataire doivent &ecirc;tre confirm&eacute;es avant envoi afin d'&eacute;viter toute erreur d'adressage. Aucune adresse non v&eacute;rifi&eacute;e n'a &eacute;t&eacute; invent&eacute;e dans ce document.")

    doc.save(DOCX_PATH)


def build_md():
    md = txt("""# Courrier - Mission Locale de Saint-&Eacute;tienne - lancement TVF

Objet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les passerelles possibles entre revitalisation territoriale, engagement citoyen, d&eacute;couverte des m&eacute;tiers et parcours jeunesse.

Intention : solliciter un premier rendez-vous afin d'identifier les formats utiles pour les jeunes accompagn&eacute;s par la Mission Locale.

Principe : TVF ne se substitue pas aux missions de la Mission Locale, aux employeurs, aux organismes de formation ou aux structures d'insertion.
""")
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
