from pathlib import Path
from runpy import run_path
from html import unescape

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

Document = BASE["Document"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
BLUE = BASE["BLUE"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
GOLD = BASE["GOLD"]
BORDER = BASE["BORDER"]

OUT_DIR = ROOT / "documents" / "brochures"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "brochure-entreprises-ressources-inutilisees-tvf.docx"
MD_PATH = SRC_DIR / "brochure-entreprises-ressources-inutilisees-tvf.md"
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"
PHOTO = ROOT / "assets" / "photos" / "materiaux-reemploi-echantillons.jpg"


def u(text: str) -> str:
    return unescape(text)


def set_run(run, size=None, color=None, bold=None, italic=None):
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(u(text))
    set_run(r, 8.5, MUTED, italic=True)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(u(text))
    set_run(r, 8.4, MUTED, italic=True)


def table(doc, headers, rows, widths, fill=LIGHT_GREEN):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, widths)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=BORDER, size="8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(u(header))
        set_run(r, 8.8, GREEN, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i], color="DDE6DE", size="6")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(u(value))
            set_run(r, 8.45, BLUE if i == 0 else RGBColor(31, 45, 38), bold=(i == 0))
    doc.add_paragraph()
    return t


def cover(doc):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Cm(5.8))
        p_logo.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(u("Brochure entreprises et acteurs &eacute;conomiques"))
    set_run(r, 10.2, GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(u("Stocks dormants, &eacute;quipements inutilis&eacute;s et ressources territoriales"))
    set_run(r, 24.0, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run(
        u(
            "Transformer ce qui p&egrave;se aujourd'hui sur l'entreprise en ressources utiles "
            "pour les territoires, les associations, les chantiers et les projets locaux."
        )
    )
    set_run(r, 12.7, GREEN, bold=True)

    meta = doc.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, [5.5, 5.5, 5.5])
    values = [
        ("Document", "information et partenariat"),
        ("Public", "entreprises, artisans, commerces"),
        ("Version", "juillet 2026"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = meta.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GREEN if idx != 1 else LIGHT_BLUE)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(u(label + "\n"))
        set_run(r1, 8.2, MUTED, bold=True)
        r2 = p.add_run(u(value))
        set_run(r2, 9.2, BLUE, bold=True)

    if PHOTO.exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(PHOTO), width=Cm(15.6))
        caption(
            doc,
            "Mat&eacute;riaux, &eacute;quipements et ressources inutilis&eacute;s peuvent &ecirc;tre "
            "qualifi&eacute;s, trac&eacute;s et r&eacute;affect&eacute;s &agrave; des projets utiles lorsque le cadre "
            "est clair.",
        )

    add_callout(
        doc,
        "Message central",
        "Vos stocks immobilisent de la tr&eacute;sorerie. Nos territoires manquent de ressources. "
        "TVF propose de faire le lien entre les entreprises qui disposent de mat&eacute;riaux, "
        "d'&eacute;quipements ou de moyens inutilis&eacute;s, et les projets locaux qui peuvent "
        "leur donner une seconde vie.",
        fill=LIGHT_GREEN,
    )


def toc(doc):
    doc.add_heading("Sommaire", level=1)
    items = [
        "1. Pourquoi les entreprises sont concern&eacute;es",
        "2. La solution TVF : valoriser les ressources inutilis&eacute;es",
        "3. Banque de Mat&eacute;riaux TVF : fonctionnement",
        "4. Ce que votre entreprise peut proposer",
        "5. B&eacute;n&eacute;fices pour l'entreprise",
        "6. Formes de partenariat et conventions possibles",
        "7. Parcours type d'une contribution",
        "8. Crit&egrave;res d'acceptation et points de vigilance",
        "9. Cadre juridique, fiscalit&eacute; et tra&ccedil;abilit&eacute;",
        "10. Proc&eacute;dure pas &agrave; pas",
        "11. Pi&egrave;ces &agrave; fournir et fiche entreprise",
    ]
    for item in items:
        add_bullet(doc, item)
    note(
        doc,
        "Cette brochure informe et oriente. Les aspects comptables, fiscaux, assurantiels "
        "ou juridiques doivent &ecirc;tre confirm&eacute;s avec les conseils habituels de l'entreprise.",
    )


def why(doc):
    doc.add_heading(u("Pourquoi les entreprises sont concern&eacute;es ?"), level=1)
    add_paragraph(
        doc,
        "Beaucoup d'entreprises disposent de ressources encore utiles qui restent immobilis&eacute;es : "
        "surplus de chantier, fins de s&eacute;rie, palettes de mat&eacute;riaux, mobilier professionnel, "
        "&eacute;quipements techniques, stocks d&eacute;class&eacute;s, invendus, outillage ou locaux temporairement "
        "sous-utilis&eacute;s. Ces ressources occupent de l'espace, immobilisent parfois de la "
        "tr&eacute;sorerie et peuvent finir par g&eacute;n&eacute;rer des co&ucirc;ts d'&eacute;vacuation ou de destruction.",
    )
    add_paragraph(
        doc,
        "Dans le m&ecirc;me temps, les territoires manquent de moyens pour remettre en usage des "
        "locaux, soutenir des associations, am&eacute;nager des lieux utiles, lancer des chantiers "
        "solidaires ou r&eacute;activer des espaces vacants. TVF propose donc une approche de "
        "coop&eacute;ration : ce qui repr&eacute;sente aujourd'hui un co&ucirc;t ou un stock dormant pour "
        "l'entreprise peut devenir une ressource pour un projet local.",
    )
    rows = [
        ("Pressions &eacute;conomiques", "Hausse du co&ucirc;t des mati&egrave;res premi&egrave;res, augmentation des charges, baisse de la demande sur certains segments, immobilisation de tr&eacute;sorerie."),
        ("Co&ucirc;ts logistiques", "Stocks qui occupent de l'espace, co&ucirc;t de stockage, inventaires complexes, besoin de lib&eacute;rer des surfaces."),
        ("Ressources non utilis&eacute;es", "Mat&eacute;riaux, mobilier, &eacute;quipements, invendus, fins de s&eacute;rie ou surplus qui restent disponibles sans usage imm&eacute;diat."),
        ("Co&ucirc;ts de sortie", "Evacuation, destruction, transport, tri ou traitement parfois co&ucirc;teux, lorsque le r&eacute;emploi n'a pas &eacute;t&eacute; anticip&eacute;."),
        ("Enjeux RSE", "Attentes croissantes des clients, salari&eacute;s, collectivit&eacute;s et partenaires sur l'impact local, la sobri&eacute;t&eacute; et l'&eacute;conomie circulaire."),
    ]
    table(doc, ["Enjeu", "Ce que cela signifie pour l'entreprise"], rows, [4.6, 11.9], fill=LIGHT_BLUE)


def solution(doc):
    doc.add_heading(u("La solution TVF : un outil de gestion des ressources inutilis&eacute;es"), level=1)
    add_paragraph(
        doc,
        "TVF ne se limite pas &agrave; solliciter des dons. L'association propose un dispositif "
        "d'organisation, de qualification et de r&eacute;affectation des ressources inutilis&eacute;es. "
        "L'objectif est de construire un cadre utile &agrave; l'entreprise et utile au territoire : "
        "identifier ce qui peut encore servir, v&eacute;rifier les conditions de r&eacute;emploi, formaliser "
        "le partenariat puis orienter les ressources vers des projets coh&eacute;rents.",
    )
    rows = [
        ("Don", "Mise &agrave; disposition gratuite de mat&eacute;riaux, &eacute;quipements ou mobilier pouvant &ecirc;tre r&eacute;utilis&eacute;s dans un projet TVF ou territorial."),
        ("Partenariat", "Coop&eacute;ration plus large autour d'un chantier, d'un local, d'un besoin logistique, d'un soutien technique ou d'une action territoriale."),
        ("Convention de mise &agrave; disposition", "Cadre &eacute;crit permettant de pr&eacute;ciser les ressources confi&eacute;es, leur usage, la dur&eacute;e, la tra&ccedil;abilit&eacute; et les responsabilit&eacute;s."),
        ("Liquidation solidaire", "Orientation encadr&eacute;e d'un stock dormant, d'un invendu ou d'un lot inutilis&eacute; vers des projets utiles, lorsque la situation le permet."),
        ("R&eacute;emploi", "Utilisation directe ou adapt&eacute;e de mat&eacute;riaux encore exploitables, apr&egrave;s qualification de leur &eacute;tat et de leur destination."),
        ("R&eacute;affectation locale", "Orientation de ressources vers des logements, locaux associatifs, chantiers, ateliers, espaces partag&eacute;s ou projets port&eacute;s avec des acteurs du territoire."),
    ]
    table(doc, ["Solution", "Utilisation possible"], rows, [4.2, 12.3], fill=LIGHT_GREEN)
    add_callout(
        doc,
        "Positionnement TVF",
        "TVF n'est pas une d&eacute;chetterie, un service d'enl&egrave;vement automatique ou une plateforme "
        "de revente classique. Chaque proposition est instruite, et chaque ressource accept&eacute;e "
        "doit pouvoir &ecirc;tre reli&eacute;e &agrave; un usage coh&eacute;rent.",
        fill=LIGHT_BLUE,
    )


def bank_process(doc):
    doc.add_heading(u("Banque de Mat&eacute;riaux TVF : fonctionnement"), level=1)
    rows = [
        ("1. Signalement", "L'entreprise transmet la nature des ressources, les volumes, les photos, le lieu, l'&eacute;tat et la disponibilit&eacute;."),
        ("2. Contr&ocirc;le administratif", "TVF v&eacute;rifie l'identit&eacute; de l'entreprise, l'origine des ressources, la propri&eacute;t&eacute; et les contraintes connues."),
        ("3. Qualification", "Les ressources sont class&eacute;es par cat&eacute;gorie, &eacute;tat, quantit&eacute;, potentiel de r&eacute;emploi, risques et conditions logistiques."),
        ("4. D&eacute;cision d'orientation", "TVF peut accepter, demander des pr&eacute;cisions, orienter vers une fili&egrave;re adapt&eacute;e, refuser ou r&eacute;server pour un projet identifi&eacute;."),
        ("5. Convention", "Lorsque le volume ou la valeur le justifie, une convention pr&eacute;cise les engagements, la tra&ccedil;abilit&eacute;, la communication et les responsabilit&eacute;s."),
        ("6. R&eacute;ception ou collecte", "Les modalit&eacute;s d'apport, d'enl&egrave;vement, de stockage, de tri et de manutention sont d&eacute;finies selon les moyens disponibles."),
        ("7. Affectation", "Les ressources sont orient&eacute;es vers des projets locaux valid&eacute;s : r&eacute;habilitation, association, chantier, local, formation ou espace utile."),
        ("8. Suivi", "TVF conserve une trace du dossier et peut produire un bilan qualitatif : origine, destination, projet, type de ressource et valorisation RSE."),
    ]
    table(doc, ["&Eacute;tape", "Fonction"], rows, [3.8, 12.7], fill=LIGHT_BLUE)


def resources(doc):
    doc.add_heading(u("Ce que votre entreprise peut proposer"), level=1)
    rows = [
        ("Mat&eacute;riaux de construction", "Bois, panneaux, carrelage, fa&iuml;ence, parquet, briques, pierres, isolants propres, &eacute;l&eacute;ments de second oeuvre."),
        ("Menuiseries", "Portes, fen&ecirc;tres, volets, escaliers, garde-corps, placards, quincaillerie, poign&eacute;es, charni&egrave;res."),
        ("Sanitaires et plomberie", "Lavabos, WC, douches, baignoires, robinetterie, radiateurs, &eacute;quipements sous r&eacute;serve de contr&ocirc;le."),
        ("Electricit&eacute; et &eacute;clairage", "Luminaires, prises, interrupteurs, chemins de c&acirc;bles, coffrets, armoires, c&acirc;bles identifi&eacute;s."),
        ("Mobilier professionnel", "Bureaux, tables, chaises, armoires, rayonnages, comptoirs, plans de travail, mobilier de commerce ou bureau."),
        ("Equipements techniques", "Outillage, &eacute;quipements d'atelier, mat&eacute;riel de chantier, EPI neufs ou non p&eacute;rim&eacute;s, supports de stockage."),
        ("Locaux ou surfaces", "Espace de stockage temporaire, local disponible, zone de tri, mise &agrave; disposition ponctuelle selon convention."),
        ("Comp&eacute;tences et services", "M&eacute;c&eacute;nat de comp&eacute;tences, diagnostic, transport, manutention, conseil technique, formation, appui logistique."),
        ("Soutien financier", "Parrainage, m&eacute;c&eacute;nat, contribution &agrave; un projet, aide &agrave; l'achat de fournitures manquantes ou &agrave; la logistique."),
    ]
    table(doc, ["Cat&eacute;gorie", "Exemples"], rows, [4.4, 12.1], fill=LIGHT_GREEN)


def benefits(doc):
    doc.add_heading(u("Les b&eacute;n&eacute;fices pour votre entreprise"), level=1)
    add_paragraph(
        doc,
        "Au lieu de supporter uniquement le co&ucirc;t d'un stock dormant, l'entreprise peut "
        "l'inscrire dans une d&eacute;marche de valorisation concr&egrave;te. La contribution peut "
        "r&eacute;pondre &agrave; des objectifs internes de sobri&eacute;t&eacute;, de RSE, de dialogue territorial "
        "et de mobilisation des salari&eacute;s.",
    )
    rows = [
        ("Optimisation", "R&eacute;duire les surfaces occup&eacute;es, clarifier les stocks, &eacute;viter certaines destructions inutiles et fluidifier la sortie des ressources sans usage."),
        ("RSE et impact", "Contribuer &agrave; l'&eacute;conomie circulaire, au r&eacute;emploi, &agrave; la r&eacute;duction du gaspillage et &agrave; des projets locaux visibles."),
        ("Relation territoriale", "Renforcer le lien avec les collectivit&eacute;s, associations, acteurs de l'ESS, artisans, habitants et projets de proximit&eacute;."),
        ("Image et marque employeur", "Donner du sens &agrave; des ressources inutilis&eacute;es, impliquer les &eacute;quipes et valoriser une contribution concr&egrave;te."),
        ("Tra&ccedil;abilit&eacute;", "B&eacute;n&eacute;ficier d'un cadre &eacute;crit, d'une destination identifi&eacute;e et d'un suivi qualitatif utilisable dans une d&eacute;marche RSE."),
        ("Cadre fiscal &agrave; &eacute;tudier", "Selon la r&eacute;glementation applicable et la situation de l'entreprise, certaines op&eacute;rations peuvent relever de dispositifs fiscaux &agrave; confirmer avec les conseils habituels."),
    ]
    table(doc, ["B&eacute;n&eacute;fice", "Apport potentiel"], rows, [4.4, 12.1], fill=LIGHT_BLUE)


def partnership_forms(doc):
    doc.add_heading(u("Formes de partenariat et conventions possibles"), level=1)
    rows = [
        ("Don de mat&eacute;riaux", "Transfert encadr&eacute; de ressources utilisables vers TVF ou un projet valid&eacute;, avec tra&ccedil;abilit&eacute; et accord &eacute;crit si n&eacute;cessaire."),
        ("Don d'&eacute;quipements", "Mobilier, outillage, machines simples, rayonnages, luminaires ou &eacute;quipements pouvant servir &agrave; un lieu ou une action."),
        ("Convention de liquidation solidaire", "Organisation d'une sortie de stock dormant, invendu ou fin de s&eacute;rie vers des usages territoriaux, sans assimilation &agrave; une liquidation commerciale classique."),
        ("Convention de stockage", "Mise &agrave; disposition d'un local, d'une zone, d'un entrep&ocirc;t ou d'un espace de tri pour les ressources TVF, avec r&egrave;gles d'acc&egrave;s et de s&eacute;curit&eacute;."),
        ("Parrainage de projet", "Soutien &agrave; un projet identifi&eacute; : local associatif, chantier, logement, espace de formation, jardin, commerce de proximit&eacute;."),
        ("M&eacute;c&eacute;nat financier", "Contribution financi&egrave;re &agrave; un projet ou au fonctionnement, dans un cadre juridique et fiscal &agrave; confirmer avec les conseils de l'entreprise."),
        ("M&eacute;c&eacute;nat de comp&eacute;tences", "Mise &agrave; disposition de temps, expertise, transport, manutention, diagnostic, communication, formation ou appui technique."),
        ("Partenariat op&eacute;rationnel", "Coop&eacute;ration plus large : mat&eacute;riaux, comp&eacute;tences, site pilote, logistique, financement, suivi d'impact et communication encadr&eacute;e."),
    ]
    table(doc, ["Cadre possible", "Objet"], rows, [4.6, 11.9], fill=LIGHT_GREEN)
    add_callout(
        doc,
        "Prudence contractuelle",
        "Aucune qualification fiscale, comptable ou juridique ne doit &ecirc;tre d&eacute;duite de cette "
        "brochure seule. Les modalit&eacute;s exactes sont &agrave; valider au cas par cas, avec les "
        "conseils habituels de l'entreprise.",
        fill=LIGHT_BLUE,
    )


def example(doc):
    doc.add_heading(u("Exemple de parcours"), level=1)
    rows = [
        ("Situation", "Une entreprise de BTP dispose de palettes de carrelage invendues, stock&eacute;es depuis plusieurs mois et sans destination imm&eacute;diate."),
        ("Contact TVF", "L'entreprise transmet la liste, les photos, la quantit&eacute;, les dimensions, le lieu, l'&eacute;tat et la date de disponibilit&eacute;."),
        ("Qualification", "TVF v&eacute;rifie l'origine, l'&eacute;tat, l'homog&eacute;n&eacute;it&eacute; des lots, la logistique, le stockage et les besoins de projets locaux."),
        ("Convention", "Une convention pr&eacute;cise la contribution, les modalit&eacute;s d'enl&egrave;vement, la destination possible, la communication et la tra&ccedil;abilit&eacute;."),
        ("R&eacute;affectation", "Les ressources peuvent &ecirc;tre orient&eacute;es vers plusieurs projets : local associatif, chantier participatif, commerce en r&eacute;activation, atelier ou am&eacute;nagement."),
        ("Valorisation", "L'entreprise peut valoriser une contribution RSE, un impact environnemental qualitatif et un ancrage local, sans inventer de chiffre non v&eacute;rifi&eacute;."),
    ]
    table(doc, ["&Eacute;tape", "Illustration"], rows, [4.2, 12.3], fill=LIGHT_BLUE)


def acceptance(doc):
    doc.add_heading(u("Crit&egrave;res d'acceptation et points de vigilance"), level=1)
    rows = [
        ("Origine", "L'entreprise doit pouvoir expliquer l'origine des ressources et confirmer qu'elle est autoris&eacute;e &agrave; les proposer."),
        ("Etat", "Les ressources doivent &ecirc;tre exploitables, propres, identifiables et compatibles avec un r&eacute;emploi raisonnable."),
        ("S&eacute;curit&eacute;", "Les risques sanitaires, chimiques, &eacute;lectriques, structurels ou de manutention doivent &ecirc;tre signal&eacute;s et trait&eacute;s."),
        ("Volume", "Le volume doit &ecirc;tre coh&eacute;rent avec les besoins, la capacit&eacute; de stockage, la logistique et la destination possible."),
        ("Tra&ccedil;abilit&eacute;", "Chaque contribution significative doit pouvoir &ecirc;tre rattach&eacute;e &agrave; un dossier, un proposant, une date et une destination."),
        ("Non-acceptation", "Chaque proposition ne vaut pas acceptation. TVF peut refuser, orienter vers une fili&egrave;re adapt&eacute;e ou demander des compl&eacute;ments."),
    ]
    table(doc, ["Crit&egrave;re", "Exigence TVF"], rows, [4.1, 12.4], fill=LIGHT_GREEN)
    rows = [
        ("Refus de principe", "Amiante, produits dangereux non identifi&eacute;s, d&eacute;chets m&eacute;lang&eacute;s, mat&eacute;riaux contamin&eacute;s, &eacute;l&eacute;ments insalubres ou dangereux."),
        ("Vigilance technique", "Electricit&eacute;, gaz, pression, chauffage, structure, machines, produits chimiques, isolants, vitrages anciens, peintures anciennes."),
        ("Orientation externe", "Lorsque TVF ne peut pas accepter, l'entreprise peut &ecirc;tre orient&eacute;e vers une fili&egrave;re r&eacute;glementaire ou professionnelle plus adapt&eacute;e."),
    ]
    table(doc, ["Situation", "Position"], rows, [4.1, 12.4], fill=LIGHT_BLUE)


def legal_fiscal(doc):
    doc.add_heading(u("Cadre juridique, fiscalit&eacute; et tra&ccedil;abilit&eacute;"), level=1)
    add_paragraph(
        doc,
        "Selon le cadre juridique et fiscal applicable, certaines op&eacute;rations pourront "
        "&eacute;ventuellement ouvrir droit aux dispositifs pr&eacute;vus par la r&eacute;glementation. "
        "TVF peut aider &agrave; identifier la nature de l'op&eacute;ration et les documents utiles, "
        "mais l'entreprise doit valider le traitement comptable, fiscal, assurantiel et "
        "juridique avec ses conseils habituels.",
    )
    rows = [
        ("Convention", "Identifier les parties, les ressources, la destination, la dur&eacute;e, les responsabilit&eacute;s, les conditions de transport et de communication."),
        ("Bordereau", "D&eacute;crire les ressources remises : cat&eacute;gorie, quantit&eacute;, &eacute;tat, date, lieu, photos, observations."),
        ("Attestation", "Selon le montage, formaliser une contribution, un don, une mise &agrave; disposition ou une coop&eacute;ration sans confusion avec un partenariat public non valid&eacute;."),
        ("Assurance", "V&eacute;rifier les responsabilit&eacute;s lors du stockage, de la manutention, du transport, de l'utilisation et de l'acc&egrave;s aux lieux."),
        ("Communication", "Toute mention publique de l'entreprise, du soutien ou du partenariat doit &ecirc;tre valid&eacute;e par les parties avant diffusion."),
    ]
    table(doc, ["Document / point", "Utilit&eacute;"], rows, [4.4, 12.1], fill=LIGHT_BLUE)


def procedure(doc):
    doc.add_heading(u("Proc&eacute;dure pas &agrave; pas"), level=1)
    rows = [
        ("1. Premier &eacute;change", "Identifier l'entreprise, le contact, la nature de la ressource et l'objectif : lib&eacute;rer un stock, soutenir un projet, valoriser une d&eacute;marche RSE."),
        ("2. Fiche entreprise", "Renseigner les informations de base : SIRET, adresse, contact, secteur, ressources, volumes, photos et contraintes."),
        ("3. Inventaire", "Classer les ressources par cat&eacute;gorie, quantit&eacute;, &eacute;tat, dimensions, valeur indicative si connue, disponibilit&eacute; et conditions d'acc&egrave;s."),
        ("4. Qualification TVF", "V&eacute;rifier la coh&eacute;rence avec l'objet TVF, les risques, la logistique, la faisabilit&eacute; de stockage et la destination possible."),
        ("5. Choix du cadre", "Don, convention, liquidation solidaire, stockage, partenariat, m&eacute;c&eacute;nat ou mise &agrave; disposition selon la situation."),
        ("6. Validation", "Confirmer par &eacute;crit les engagements, responsabilit&eacute;s, documents &agrave; produire et communication autoris&eacute;e."),
        ("7. Enl&egrave;vement ou apport", "Organiser transport, manutention, rendez-vous, s&eacute;curit&eacute;, bordereau et photos de remise."),
        ("8. Affectation et suivi", "Rattacher la contribution &agrave; un dossier, un projet, une destination ou une r&eacute;serve qualifi&eacute;e."),
    ]
    table(doc, ["&Eacute;tape", "Action"], rows, [3.8, 12.7], fill=LIGHT_GREEN)


def documents_needed(doc):
    doc.add_heading(u("Pi&egrave;ces et informations &agrave; fournir"), level=1)
    rows = [
        ("Identification", "Raison sociale, SIRET, adresse, contact r&eacute;f&eacute;rent, fonction, t&eacute;l&eacute;phone, e-mail."),
        ("Description", "Cat&eacute;gorie, nature des ressources, quantit&eacute;, dimensions, poids, &eacute;tat, photos, conditionnement."),
        ("Origine", "Chantier, stock dormant, invendu, d&eacute;pose, renouvellement de mobilier, fin de s&eacute;rie, liquidation, surplus."),
        ("Disponibilit&eacute;", "Date limite, horaires, urgence, contraintes de chantier, possibilit&eacute; d'apport ou besoin d'enl&egrave;vement."),
        ("Acc&egrave;s et logistique", "Adresse, quai, stationnement, palette, engin disponible, manutention, &eacute;tage, contraintes de s&eacute;curit&eacute;."),
        ("Risques", "Produits dangereux, &eacute;lectricit&eacute;, gaz, poussi&egrave;res, poids, fragilit&eacute;, humidit&eacute;, amiante ou plomb si concern&eacute;."),
        ("Documents", "Photos, inventaire, fiches techniques, notices, factures ou r&eacute;f&eacute;rences si disponibles, attestation de propri&eacute;t&eacute; si utile."),
        ("Objectif", "Don, partenariat, stockage, liquidation solidaire, m&eacute;c&eacute;nat, RSE, soutien &agrave; un projet local."),
    ]
    table(doc, ["Information", "D&eacute;tail attendu"], rows, [4.3, 12.2], fill=LIGHT_BLUE)


def company_form(doc):
    doc.add_page_break()
    doc.add_heading(u("Fiche entreprise - proposition de ressources"), level=1)
    add_paragraph(doc, "Cette fiche sert &agrave; ouvrir une instruction. Elle ne vaut pas acceptation, enl&egrave;vement, convention, attestation fiscale ou partenariat public.")
    rows = [
        ("Date", "................................................................................................"),
        ("Raison sociale", "................................................................................................"),
        ("SIRET", "................................................................................................"),
        ("Adresse", "................................................................................................"),
        ("Contact r&eacute;f&eacute;rent", "Nom : ........................................ Fonction : ........................................"),
        ("T&eacute;l&eacute;phone / e-mail", "................................................................................................"),
        ("Secteur d'activit&eacute;", "[ ] BTP   [ ] Industrie   [ ] Commerce   [ ] H&ocirc;tellerie-restauration   [ ] Bureaux   [ ] Autre"),
        ("Type de contribution", "[ ] Mat&eacute;riaux   [ ] Equipements   [ ] Mobilier   [ ] Local   [ ] Transport   [ ] Comp&eacute;tences   [ ] Financier"),
        ("Cat&eacute;gorie principale", "[ ] Bois   [ ] Menuiseries   [ ] Sanitaires   [ ] Eclairage   [ ] M&eacute;taux   [ ] Mobilier   [ ] Outillage   [ ] Autre"),
        ("Description", "................................................................................................\n................................................................................................"),
        ("Quantit&eacute; / dimensions", "................................................................................................"),
        ("Etat apparent", "[ ] Tr&egrave;s bon   [ ] Bon   [ ] Moyen   [ ] A v&eacute;rifier   Commentaire : ........................................"),
        ("Photos jointes", "[ ] Oui   [ ] Non   Nombre : .................."),
        ("Disponibilit&eacute;", "Date limite : .................................... Horaires / contraintes : ........................................"),
        ("Lieu de stockage actuel", "................................................................................................"),
        ("Logistique", "[ ] Apport possible   [ ] Enl&egrave;vement demand&eacute;   [ ] Palette   [ ] Manutention n&eacute;cessaire   [ ] Quai"),
        ("Cadre envisag&eacute;", "[ ] Don   [ ] Partenariat   [ ] Liquidation solidaire   [ ] Stockage   [ ] M&eacute;c&eacute;nat   [ ] A d&eacute;finir"),
        ("Risques connus", "[ ] Aucun signal&eacute;   [ ] A pr&eacute;ciser : ................................................................"),
        ("Objectif RSE / territorial", "................................................................................................"),
        ("D&eacute;claration", "Je comprends que cette proposition ne vaut pas acceptation, ne vaut pas enl&egrave;vement automatique et ne vaut pas attestation fiscale automatique."),
        ("Signature / cachet", "................................................................................................"),
    ]
    table(doc, ["Champ", "R&eacute;ponse / information &agrave; renseigner"], rows, [4.7, 11.8], fill=LIGHT_GREEN)
    add_callout(
        doc,
        "Contact TVF",
        "TERRITOIRES VIVANTS FRANCE - 25 rue &Eacute;lise Gervais, 42000 Saint-&Eacute;tienne - contact@territoiresvivantsfrance.fr - 04 65 81 54 69 - www.territoiresvivantsfrance.fr",
        fill=LIGHT_BLUE,
    )


def final_notes(doc):
    doc.add_heading(u("Points &agrave; verrouiller avant toute contribution"), level=1)
    for item in [
        "Origine et propri&eacute;t&eacute; des ressources propos&eacute;es.",
        "Etat, quantit&eacute;, dimensions, photos et risques connus.",
        "Cadre choisi : don, convention, stockage, liquidation solidaire, m&eacute;c&eacute;nat ou partenariat.",
        "Responsabilit&eacute;s li&eacute;es au transport, &agrave; la manutention, au stockage et &agrave; l'usage.",
        "Traitement comptable, fiscal, assurantiel et juridique valid&eacute; par l'entreprise.",
        "Validation &eacute;crite avant toute communication publique.",
        "Tra&ccedil;abilit&eacute; de la destination et suivi qualitatif lorsque la ressource est affect&eacute;e.",
    ]:
        add_bullet(doc, "[ ] " + item)
    add_callout(
        doc,
        "Conclusion",
        "TVF permet &agrave; une entreprise de transformer des ressources immobilis&eacute;es en contribution "
        "territoriale utile. Le dispositif associe qualification, convention, r&eacute;emploi, tra&ccedil;abilit&eacute; "
        "et valorisation RSE, sans confusion entre don libre, enl&egrave;vement automatique et partenariat "
        "institutionnel non valid&eacute;.",
        fill=LIGHT_GREEN,
    )


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    cover(doc)
    toc(doc)
    doc.add_page_break()
    why(doc)
    solution(doc)
    bank_process(doc)
    resources(doc)
    benefits(doc)
    partnership_forms(doc)
    example(doc)
    acceptance(doc)
    legal_fiscal(doc)
    procedure(doc)
    documents_needed(doc)
    company_form(doc)
    final_notes(doc)
    doc.save(DOCX_PATH)


def build_md():
    md = u(
        """# Brochure TVF - Entreprises

Document d'information destin&eacute; aux entreprises, artisans, commerces et acteurs &eacute;conomiques disposant de stocks dormants, mat&eacute;riaux, &eacute;quipements, mobiliers, locaux ou comp&eacute;tences pouvant contribuer &agrave; des projets locaux.

## Message central

Vos stocks immobilisent de la tr&eacute;sorerie. Nos territoires manquent de ressources. Faisons le lien.

Ce qui repr&eacute;sente aujourd'hui un co&ucirc;t pour votre entreprise peut devenir une ressource pour un territoire.

## Principes

- Chaque proposition ne vaut pas acceptation.
- TVF qualifie les ressources avant toute affectation.
- Les ressources ne sont pas distribu&eacute;es librement.
- Les cadres possibles sont : don, convention, liquidation solidaire, stockage, m&eacute;c&eacute;nat, partenariat op&eacute;rationnel.
- Les aspects fiscaux et juridiques doivent &ecirc;tre valid&eacute;s avec les conseils habituels de l'entreprise.

## Contact

TERRITOIRES VIVANTS FRANCE - contact@territoiresvivantsfrance.fr - 04 65 81 54 69 - www.territoiresvivantsfrance.fr
"""
    )
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
