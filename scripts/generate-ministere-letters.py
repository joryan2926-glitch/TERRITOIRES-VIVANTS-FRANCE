from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import re

ROOT = Path(r'C:\Users\jowst\Documents\TERRITOIRES VIVANTS FRANCE')
OUT = ROOT / 'documents' / 'courriers-ministeres-tvf'
OUT.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / 'assets' / 'logo-territoires-vivants-france-web.png'

TVF = {
    'name': 'Territoires Vivants France',
    'subtitle': 'Agence Territoriale de Revitalisation Immobilière',
    'address': '25 rue Élise Gervais, 42000 Saint-Étienne',
    'phone': '04 65 81 54 69',
    'email': 'contact@territoiresvivantsfrance.fr',
    'site': 'www.territoiresvivantsfrance.fr',
    'rna': 'RNA W922015538',
    'siret': 'SIRET 897 226 138 00018',
}

letters = [
    {
        'slug':'01-ville-logement-vincent-jeanbrun',
        'minister':'Monsieur Vincent Jeanbrun',
        'title':'Ministre de la Ville et du Logement',
        'salutation':'Monsieur le Ministre,',
        'object':'Proposition d’échange autour de la remise en usage des logements vacants et de l’accompagnement des propriétaires',
        'angle':'le logement vacant, la rénovation, la remise en usage du parc existant et l’accompagnement des propriétaires constituent le cœur opérationnel de notre démarche.',
        'programs':['lutte contre la vacance des logements','rénovation de l’habitat privé','remise en usage du parc existant','accompagnement des propriétaires et sécurisation des parcours','prévention de la dégradation du bâti'],
        'request':'Nous souhaiterions pouvoir présenter cette méthode à vos services afin d’identifier les cadres de travail, programmes ou interlocuteurs nationaux avec lesquels TVF pourrait construire une expérimentation utile, prudente et reproductible à partir de son territoire pilote.',
    },
    {
        'slug':'02-amenagement-territoire-francoise-gatel',
        'minister':'Madame Françoise Gatel',
        'title':'Ministre de l’Aménagement du territoire et de la Décentralisation',
        'salutation':'Madame la Ministre,',
        'object':'Proposition d’échange sur une méthode territoriale de repérage, qualification et remise en usage des biens vacants',
        'angle':'la vacance immobilière ne peut être traitée uniquement bien par bien : elle touche les centralités, les communes, les quartiers, l’attractivité locale et l’ingénierie dont disposent les collectivités.',
        'programs':['revitalisation des centralités','accompagnement des communes et intercommunalités','ingénierie territoriale','observation locale de la vacance','coordination entre propriétaires, collectivités et partenaires'],
        'request':'Nous sollicitons un échange avec vos services afin d’étudier comment TVF pourrait contribuer, comme outil associatif et opérationnel, aux démarches territoriales de repérage, de qualification et de suivi des situations immobilières dormantes.',
    },
    {
        'slug':'03-transition-ecologique-monique-barbut',
        'minister':'Madame Monique Barbut',
        'title':'Ministre de la Transition écologique, de la Biodiversité et du Climat',
        'salutation':'Madame la Ministre,',
        'object':'Proposition d’échange sur le réemploi, la sobriété foncière et la remise en usage du patrimoine existant',
        'angle':'remettre en usage un logement, un commerce, un bâtiment ou une friche déjà existants participe à une logique de sobriété foncière, de réemploi et de limitation du gaspillage de ressources.',
        'programs':['sobriété foncière','réemploi des matériaux','économie circulaire dans le bâtiment','transformation des friches','réduction du gaspillage et valorisation des ressources inutilisées'],
        'request':'Nous souhaitons présenter à vos services la Matériauthèque Solidaire de TVF et la méthode de suivi des biens vacants afin d’identifier les convergences possibles avec les politiques nationales de transition écologique et de réemploi.',
    },
    {
        'slug':'04-pme-commerce-serge-papin',
        'minister':'Monsieur Serge Papin',
        'title':'Ministre des PME, du Commerce, de l’Artisanat, du Tourisme et du Pouvoir d’achat',
        'salutation':'Monsieur le Ministre,',
        'object':'Proposition d’échange sur la revitalisation des locaux commerciaux vacants et la mobilisation des entreprises locales',
        'angle':'la vacance commerciale fragilise les rues, les centres-villes et l’installation de nouvelles activités ; elle nécessite un travail patient avec les propriétaires, les collectivités, les artisans et les porteurs de projet.',
        'programs':['revitalisation commerciale','soutien aux centralités','mobilisation des artisans et entreprises locales','réactivation de locaux vacants','réemploi de ressources et stocks dormants'],
        'request':'Nous souhaitons échanger avec vos services afin de présenter le pôle Commerce Vivant de TVF et d’identifier les coopérations possibles autour des locaux vacants, des artisans, des commerçants et des entreprises partenaires.',
    },
    {
        'slug':'05-economie-finances-roland-lescure',
        'minister':'Monsieur Roland Lescure',
        'title':'Ministre de l’Économie, des Finances, de la Souveraineté industrielle, énergétique et numérique',
        'salutation':'Monsieur le Ministre,',
        'object':'Proposition d’échange sur un outil territorial de valorisation immobilière, économique et partenariale',
        'angle':'les biens vacants, stocks dormants et ressources inutilisées représentent à la fois une perte de valeur économique et une opportunité pour les territoires lorsqu’ils sont qualifiés, orientés et mobilisés dans un cadre sérieux.',
        'programs':['valorisation du patrimoine existant','mobilisation des entreprises et fondations','financement de projets territoriaux','économie circulaire et ressources dormantes','outils numériques de suivi et d’observation'],
        'request':'Nous sollicitons un échange afin d’identifier les dispositifs économiques, numériques ou partenariaux susceptibles d’accompagner la structuration de TVF OS, de TVF Mobile et des premiers parcours territoriaux pilotés par l’association.',
    },
    {
        'slug':'06-travail-solidarites-jean-pierre-farandou',
        'minister':'Monsieur Jean-Pierre Farandou',
        'title':'Ministre du Travail et des Solidarités',
        'salutation':'Monsieur le Ministre,',
        'object':'Proposition d’échange sur l’utilité sociale des projets de revitalisation immobilière et territoriale',
        'angle':'la remise en usage de biens vacants peut aussi devenir un levier d’insertion, de formation, de transmission de compétences et d’amélioration du cadre de vie lorsque les projets sont correctement encadrés.',
        'programs':['insertion par l’activité','chantiers encadrés','transmission de compétences','mobilisation associative et partenariale','projets utiles aux publics et territoires fragilisés'],
        'request':'Nous souhaiterions présenter le pôle Solidarité et Insertion de TVF afin d’étudier les conditions dans lesquelles des projets immobiliers ou de réemploi pourraient être articulés avec les acteurs compétents de l’emploi, de l’insertion et des solidarités.',
    },
    {
        'slug':'07-vie-associative-marina-ferrari',
        'minister':'Madame Marina Ferrari',
        'title':'Ministre des Sports, de la Jeunesse et de la Vie associative',
        'salutation':'Madame la Ministre,',
        'object':'Proposition d’échange autour d’une association structurée comme outil territorial d’intérêt général',
        'angle':'TVF est une association loi 1901 qui cherche à structurer une action territoriale concrète, lisible et durable autour de la vacance immobilière, du bénévolat utile et de la mobilisation citoyenne.',
        'programs':['vie associative','engagement citoyen','participation des habitants','bénévolat de compétences','structuration d’outils associatifs au service des territoires'],
        'request':'Nous sollicitons un échange avec vos services pour présenter la trajectoire de TVF, son territoire pilote et ses outils en préparation, afin d’identifier les appuis possibles à la structuration associative et à la mobilisation citoyenne.',
    },
    {
        'slug':'08-culture-catherine-pegard',
        'minister':'Madame Catherine Pégard',
        'title':'Ministre de la Culture',
        'salutation':'Madame la Ministre,',
        'object':'Proposition d’échange sur la valorisation du patrimoine bâti vacant et des centres anciens',
        'angle':'de nombreux biens vacants s’inscrivent dans des tissus anciens, des rues historiques, des rez-de-chaussée patrimoniaux ou des bâtiments qui participent à l’identité d’un territoire.',
        'programs':['patrimoine bâti','centres anciens','valorisation des bâtiments délaissés','usages culturels ou associatifs temporaires','préservation de l’identité urbaine et locale'],
        'request':'Nous souhaiterions présenter la méthode TVF afin d’échanger sur la manière dont certains biens vacants ou bâtiments délaissés pourraient être mieux repérés, documentés et orientés dans le respect du patrimoine et des compétences professionnelles habilitées.',
    },
    {
        'slug':'09-justice-gerald-darmanin',
        'minister':'Monsieur Gérald Darmanin',
        'title':'Garde des Sceaux, ministre de la Justice',
        'salutation':'Monsieur le Ministre,',
        'object':'Proposition d’échange sur les situations juridiques complexes liées aux biens vacants',
        'angle':'les biens vacants peuvent être liés à des successions, indivisions, propriétaires non identifiés, situations contentieuses ou blocages juridiques qui rendent indispensable une approche prudente et respectueuse du droit de propriété.',
        'programs':['droit de propriété','successions et indivisions','biens présumés sans maître','sécurisation juridique des démarches','orientation vers les professionnels compétents'],
        'request':'Nous souhaitons échanger avec vos services afin de mieux encadrer la doctrine interne de TVF sur les situations juridiques sensibles et d’identifier les bonnes pratiques permettant d’éviter toute confusion entre signalement, présomption et situation juridiquement établie.',
    },
    {
        'slug':'10-agriculture-ruralite-annie-genevard',
        'minister':'Madame Annie Genevard',
        'title':'Ministre de l’Agriculture et de la Souveraineté alimentaire',
        'salutation':'Madame la Ministre,',
        'object':'Proposition d’échange sur la vacance, les bâtiments délaissés et la revitalisation des territoires ruraux',
        'angle':'la vacance immobilière et foncière concerne aussi les bourgs, villages, bâtiments ruraux, terrains délaissés et espaces dont l’usage doit être étudié avec prudence au regard des besoins locaux.',
        'programs':['revitalisation des bourgs et villages','bâtiments ruraux délaissés','foncier et usages locaux','projets de proximité','coopération avec les collectivités rurales'],
        'request':'Nous souhaiterions présenter l’approche TVF afin d’étudier les conditions dans lesquelles cette méthode pourrait, à terme, être adaptée à des territoires ruraux ou périurbains, en lien avec les collectivités et acteurs compétents.',
    },
]

intro = (
    "Territoires Vivants France est une association loi 1901 qui structure une Agence Territoriale de Revitalisation Immobilière. "
    "Son objet est de repérer, comprendre, qualifier et accompagner la remise en usage de logements, commerces, bâtiments, friches, terrains et ressources aujourd’hui dormants, dans le respect du droit de propriété, des compétences publiques et du rôle des professionnels habilités."
)
method = [
    "repérer les biens ou situations signalés par les propriétaires, habitants, collectivités, entreprises ou partenaires ;",
    "qualifier les informations disponibles sans conclure hâtivement à la vacance, à l’abandon ou à la disponibilité du bien ;",
    "ouvrir, lorsque cela est pertinent, un dossier opérationnel avec historique, pièces, photos, interlocuteurs et étapes de suivi ;",
    "mettre en relation les acteurs utiles : propriétaires, collectivités, entreprises, associations, professionnels, établissements fonciers ou financeurs ;",
    "orienter les projets vers une remise en usage réaliste : rénovation, réemploi, occupation encadrée, projet associatif, activité commerciale ou autre solution adaptée."
]
tools = (
    "Deux outils sont en préparation pour rendre cette méthode plus lisible et plus opérationnelle : TVF OS, système interne de suivi des demandes, dossiers, biens, documents, partenaires et indicateurs ; et TVF Mobile, application de terrain destinée au signalement, à la géolocalisation, aux photographies, aux visites et à la remontée d’informations vers TVF OS."
)
impact = (
    "À ce stade, TVF ne souhaite pas afficher de résultats chiffrés non consolidés. L’impact recherché est toutefois clair : réduire les situations dormantes, éviter la dispersion des informations, faciliter le dialogue avec les propriétaires, soutenir les collectivités, mobiliser les ressources locales et contribuer à la revitalisation des territoires à partir d’un premier territoire pilote à Saint-Étienne."
)

GREEN = RGBColor(24, 57, 47)
SAGE = RGBColor(83, 121, 91)
GOLD = RGBColor(180, 148, 90)
DARK = RGBColor(31, 41, 55)
GRAY = RGBColor(102, 112, 133)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = 'Inter'
    r.font.size = Pt(size)
    r.bold = bold
    if color: r.font.color.rgb = color


def add_para(doc, text='', style=None, bold_start=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.12
    if bold_start and text.startswith(bold_start):
        r1 = p.add_run(bold_start)
        r1.bold = True
        r1.font.color.rgb = GREEN
        r2 = p.add_run(text[len(bold_start):])
    else:
        r2 = p.add_run(text)
    for run in p.runs:
        run.font.name = 'Inter'
        run.font.size = Pt(10.2)
        run.font.color.rgb = DARK
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = 'Manrope'
    r.font.bold = True
    r.font.size = Pt(15 if level == 1 else 12.5)
    r.font.color.rgb = GREEN
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    r = p.add_run(text)
    r.font.name = 'Inter'
    r.font.size = Pt(9.7)
    r.font.color.rgb = DARK
    return p


def add_header(doc, item):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.58)
    sec.bottom_margin = Inches(0.62)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    header = sec.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(7.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.55)
    table.columns[1].width = Inches(4.45)
    left, right = table.rows[0].cells
    if LOGO.exists():
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(LOGO), width=Inches(1.72))
    else:
        set_cell_text(left, TVF['name'], True, GREEN, 11)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lines = [TVF['subtitle'], TVF['address'], f"{TVF['phone']} · {TVF['email']}", f"{TVF['rna']} · {TVF['siret']}", TVF['site']]
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = 'Inter'
        r.font.size = Pt(8.2 if idx else 8.7)
        r.font.bold = idx == 0
        r.font.color.rgb = GREEN if idx == 0 else GRAY
        if idx < len(lines)-1:
            p.add_run('\n')
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run('Territoires Vivants France — Courrier institutionnel — Document préparé pour transmission ministérielle')
    rf.font.name = 'Inter'
    rf.font.size = Pt(7.8)
    rf.font.color.rgb = GRAY


def build_letter(item):
    doc = Document()
    styles = doc.styles
    styles['Normal'].font.name = 'Inter'
    styles['Normal'].font.size = Pt(10.2)
    styles['Normal'].font.color.rgb = DARK
    if 'List Bullet' in styles:
        styles['List Bullet'].font.name = 'Inter'
        styles['List Bullet'].font.size = Pt(9.7)
    add_header(doc, item)

    # Dateline and recipient
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('Saint-Étienne, le 1 août 2026')
    r.font.name = 'Inter'; r.font.size = Pt(9.2); r.font.color.rgb = GRAY

    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(3.2)
    t.columns[1].width = Inches(3.8)
    c0, c1 = t.rows[0].cells
    shade_cell(c0, 'F3F7F2')
    shade_cell(c1, 'FFFFFF')
    set_cell_text(c0, 'Courrier institutionnel\nDemande d’échange et de mise en relation', True, GREEN, 10)
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line, bold in [(item['minister'], True), (item['title'], False), ('Hôtel de Matignon / ministère compétent', False), ('Paris', False)]:
        r = p.add_run(line)
        r.font.name = 'Inter'; r.font.size = Pt(9.2); r.bold = bold; r.font.color.rgb = DARK if bold else GRAY
        p.add_run('\n')

    add_heading(doc, 'Objet', 2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(11)
    r = p.add_run(item['object'])
    r.font.name = 'Manrope'; r.font.bold = True; r.font.size = Pt(12.2); r.font.color.rgb = GREEN

    add_para(doc, item['salutation'])
    add_para(doc, intro)
    add_para(doc, f"Si nous nous permettons de nous adresser à vous, c’est parce que {item['angle']}")

    add_heading(doc, 'Une réponse opérationnelle aux situations immobilières dormantes', 1)
    add_para(doc, "TVF part d’un constat simple : les informations relatives aux biens vacants ou délaissés sont souvent dispersées entre propriétaires, habitants, collectivités, entreprises et partenaires techniques. Cette dispersion ralentit la compréhension des situations et retarde la recherche de solutions adaptées.")
    add_para(doc, "Notre méthode vise donc à organiser une chaîne de travail lisible, depuis le signalement jusqu’à l’orientation ou l’ouverture d’un dossier opérationnel. Elle ne se substitue ni aux collectivités, ni aux services de l’État, ni aux professionnels habilités ; elle agit comme un outil de repérage, de qualification, de coordination et de suivi.")

    add_heading(doc, 'Méthode TVF', 2)
    for b in method:
        add_bullet(doc, b)

    add_heading(doc, 'Lien avec les politiques publiques nationales', 1)
    add_para(doc, "La démarche TVF s’inscrit dans un cadre national déjà fortement structuré autour de la lutte contre la vacance, de la revitalisation territoriale, de la rénovation de l’habitat, de la sobriété foncière, de l’économie circulaire et du soutien aux centralités. Nous souhaitons nous inscrire dans ces orientations sans revendiquer de compétence qui ne serait pas la nôtre.")
    add_para(doc, "Les axes de convergence avec votre ministère concernent notamment :")
    for b in item['programs']:
        add_bullet(doc, b + '.')

    add_heading(doc, 'Outils en préparation : TVF OS et TVF Mobile', 1)
    add_para(doc, tools)
    add_para(doc, "Ces outils ont vocation à sécuriser le traitement des demandes : référence unique, historique, pièces jointes, photos, carte, statut, responsable, prochaine étape, documents manquants et indicateurs. L’objectif n’est pas de créer un simple CRM, mais un système opérationnel de revitalisation immobilière et territoriale.")

    add_heading(doc, 'Territoire pilote et impact recherché', 1)
    add_para(doc, impact)
    add_para(doc, "TVF souhaite avancer avec prudence : les résultats, indicateurs et partenariats ne seront présentés comme acquis qu’après vérification et formalisation. Cette exigence est essentielle pour construire une relation de confiance avec les institutions publiques, les propriétaires, les collectivités et les partenaires privés.")

    add_heading(doc, 'Demande formulée', 1)
    add_para(doc, item['request'])
    add_para(doc, "Concrètement, nous souhaiterions pouvoir obtenir un temps d’échange, une orientation vers le service compétent, ou une mise en relation permettant d’identifier le bon cadre de dialogue institutionnel. TVF est disponible pour transmettre un dossier de présentation, présenter son territoire pilote et détailler les garanties méthodologiques prévues.")

    add_para(doc, "Nous vous prions d’agréer, " + item['salutation'].replace(',', '').lower() + ", l’expression de notre haute considération.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line, bold in [('Pour Territoires Vivants France', False), ('Le Président / représentant légal', True), ('Signature', False)]:
        r = p.add_run(line)
        r.font.name = 'Inter'; r.font.size = Pt(9.5); r.bold = bold; r.font.color.rgb = GREEN if bold else GRAY
        p.add_run('\n')

    # Contact callout
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0,0)
    shade_cell(cell, 'F3F7F2')
    set_cell_text(cell, f"Contact TVF : {TVF['phone']} · {TVF['email']} · {TVF['site']}", True, GREEN, 9.3)

    path = OUT / f"courrier-{item['slug']}.docx"
    doc.save(path)
    return path

paths = [build_letter(item) for item in letters]

# recap
recap = Document()
sec = recap.sections[0]
sec.top_margin = Inches(0.65); sec.bottom_margin = Inches(0.65); sec.left_margin = Inches(0.75); sec.right_margin = Inches(0.75)
if LOGO.exists():
    p = recap.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.add_run().add_picture(str(LOGO), width=Inches(1.65))
p = recap.add_paragraph(); r = p.add_run('Courriers ministériels personnalisés — Territoires Vivants France')
r.font.name='Manrope'; r.font.bold=True; r.font.size=Pt(18); r.font.color.rgb=GREEN
add_para(recap, 'Ce récapitulatif indique les courriers générés, le ministère ciblé et l’angle de personnalisation retenu. Les courriers ne présentent aucun partenariat comme acquis et ne contiennent aucun chiffre d’impact non consolidé.')
t = recap.add_table(rows=1, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
hdr = t.rows[0].cells
for cell, txt in zip(hdr, ['N°','Destinataire','Angle TVF','Fichier']):
    shade_cell(cell, '18392F'); set_cell_text(cell, txt, True, RGBColor(255,255,255), 8.7)
for i, item in enumerate(letters, 1):
    row = t.add_row().cells
    vals = [str(i), f"{item['minister']}\n{item['title']}", item['object'], f"courrier-{item['slug']}.docx"]
    for cell, val in zip(row, vals):
        set_cell_text(cell, val, False, DARK, 8.2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
recap_path = OUT / '00-recapitulatif-courriers-ministeres-tvf.docx'
recap.save(recap_path)
print('CREATED')
for p in [recap_path] + paths:
    print(p)
