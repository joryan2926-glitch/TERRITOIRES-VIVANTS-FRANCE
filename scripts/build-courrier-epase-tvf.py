from pathlib import Path
from runpy import run_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

OUT_DIR = ROOT / "documents" / "courriers-lancement-saint-etienne"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "29-epase-lancement-cooperation-amenagement-tvf.docx"
MD_PATH = SRC_DIR / "courrier-epase-lancement-tvf.md"

txt = BASE["txt"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_number = BASE["add_number"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
BORDER = BASE["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        ("Destinataire", "EPASE - &Eacute;tablissement Public d'Am&eacute;nagement de Saint-&Eacute;tienne - &agrave; l'attention de la direction g&eacute;n&eacute;rale et des services concern&eacute;s"),
        ("Adresse", "49 rue de la Montat, 42100 Saint-&Eacute;tienne - 04 77 34 43 60"),
        ("Objet", "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les compl&eacute;mentarit&eacute;s possibles autour du b&acirc;ti vacant, des friches, des commerces, du r&eacute;emploi de mat&eacute;riaux et des usages territoriaux"),
        ("Pi&egrave;ces jointes", "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche Habitat Vivant, fiche commerces inoccup&eacute;s, fiche friches / biens vacants, fiche mat&eacute;riaux et r&eacute;emploi"),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def add_complementarity_table(doc):
    doc.add_heading(txt("Compl&eacute;mentarit&eacute;s possibles &agrave; examiner"), level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.5, 6.4, 6.2])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Sujet", "Apport possible de TVF", "Point &agrave; cadrer avec l'EPASE"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(head))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.1)

    rows = [
        ("B&acirc;ti vacant et biens &agrave; qualifier", "Rep&eacute;rer, documenter et pr&eacute;qualifier des biens signal&eacute;s par des habitants, propri&eacute;taires, associations ou entreprises.", "V&eacute;rifier les cas utiles aux quartiers d'intervention et les limites li&eacute;es aux op&eacute;rations d'am&eacute;nagement."),
        ("Commerces et locaux inoccup&eacute;s", "Identifier les usages possibles : commerce de proximit&eacute;, atelier, local associatif, espace de formation ou activit&eacute; utile.", "Articuler la d&eacute;marche avec les offres immobili&egrave;res, implantations et besoins d'activit&eacute;s d&eacute;j&agrave; suivis."),
        ("Friches et espaces d&eacute;laiss&eacute;s", "Faire remonter des besoins d'usage, d'animation, de r&eacute;emploi ou de transformation l&eacute;g&egrave;re lorsque le cadre le permet.", "Distinguer ce qui rel&egrave;ve de l'am&eacute;nagement public, de la ma&icirc;trise d'ouvrage et de l'observation citoyenne."),
        ("Mat&eacute;riaux et r&eacute;emploi", "Valoriser des ressources inutilis&eacute;es dans des projets valid&eacute;s : bois, menuiseries, mobiliers, &eacute;quipements, mat&eacute;riaux de chantier.", "Pr&eacute;ciser les conditions techniques, juridiques, assurantielles et logistiques du r&eacute;emploi."),
        ("Insertion et usages de quartier", "Relier certains projets &agrave; des structures d'insertion, b&eacute;n&eacute;voles, associations et acteurs locaux.", "S&eacute;curiser les responsabilit&eacute;s, les autorisations, les conventions et la communication publique."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.35)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - EPASE"))
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Demande d'un premier &eacute;change sur les passerelles entre am&eacute;nagement, revitalisation, r&eacute;emploi et usages territoriaux"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(doc, "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un ancrage op&eacute;rationnel progressif sur son territoire d'origine.")
    add_paragraph(doc, "TVF intervient sur des ressources aujourd'hui insuffisamment mobilis&eacute;es : logements vacants ou d&eacute;grad&eacute;s, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains inutilis&eacute;s, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. La d&eacute;marche vise &agrave; observer, qualifier, orienter, mettre en relation et suivre les usages possibles, sans se substituer aux ma&icirc;tres d'ouvrage, aux collectivit&eacute;s, aux am&eacute;nageurs ou aux professionnels comp&eacute;tents.")
    add_paragraph(doc, "L'EPASE constitue un interlocuteur central pour ce premier ancrage st&eacute;phanois. Son site officiel pr&eacute;sente des missions de conception et d'am&eacute;nagement, de r&eacute;habilitation et gestion, d'innovation et exp&eacute;rimentation, ainsi qu'un r&ocirc;le d'information et d'accompagnement. Ses interventions concernent notamment Manufacture-Plaine Achille, Jacquard, Ch&acirc;teaucreux, Pont de l'Ane-Monthieu, Centre-ville et Chappe Ferdinand.")

    add_callout(doc, "Demande principale", "TVF sollicite un premier rendez-vous avec l'EPASE afin de pr&eacute;senter sa m&eacute;thode, de comprendre les cadres d'intervention de l'&eacute;tablissement public d'am&eacute;nagement et d'identifier les sujets sur lesquels une contribution associative de type observation, pr&eacute;qualification, r&eacute;emploi ou mobilisation locale pourrait compl&eacute;ter les dispositifs existants.")

    doc.add_heading(txt("Pourquoi solliciter l'EPASE"), level=1)
    add_paragraph(doc, "Le projet TVF na&icirc;t &agrave; Saint-&Eacute;tienne dans un territoire d&eacute;j&agrave; engag&eacute; dans une transformation urbaine importante. L'EPASE travaille sur des secteurs o&ugrave; les enjeux de logements, commerces, espaces publics, foncier, activit&eacute;s, image urbaine et attractivit&eacute; se croisent. Ces sujets rejoignent directement la logique de TVF : redonner un usage utile &agrave; ce qui existe d&eacute;j&agrave; avant de consommer de nouvelles ressources.")
    add_paragraph(doc, "TVF peut apporter une capacit&eacute; de terrain compl&eacute;mentaire : recevoir des signalements, aider &agrave; qualifier des besoins, identifier des ressources mat&eacute;rielles, documenter des propositions d'usages, mobiliser des associations ou structures d'insertion et pr&eacute;parer des dossiers lisibles avant toute orientation vers les acteurs publics ou techniques comp&eacute;tents.")
    add_paragraph(doc, "Cette contribution ne doit pas &ecirc;tre confondue avec une mission d'am&eacute;nagement. TVF ne revendique aucune comp&eacute;tence d'acquisition, de portage foncier, de ma&icirc;trise d'ouvrage, de commercialisation, de prescription urbaine ou d'autorisation administrative. L'objectif est de construire une passerelle utile entre besoins de terrain, ressources dormantes et projets de revitalisation.")

    add_complementarity_table(doc)

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "Organiser un rendez-vous de cadrage avec la direction ou les services concern&eacute;s afin de pr&eacute;senter TVF et de comprendre les conditions d'une contribution compl&eacute;mentaire.")
    add_bullet(doc, "Identifier les p&eacute;rim&egrave;tres ou sujets pour lesquels TVF pourrait faire remonter des observations, besoins, ressources ou propositions d'usages sans interférer avec les op&eacute;rations conduites par l'EPASE.")
    add_bullet(doc, "Pr&eacute;ciser les limites relatives aux donn&eacute;es, aux biens, aux autorisations, &agrave; la confidentialit&eacute;, &agrave; la responsabilit&eacute;, &agrave; la communication publique et aux conventions.")
    add_bullet(doc, "Examiner la possibilit&eacute; d'un premier cas d'&eacute;tude ou d'une premi&egrave;re m&eacute;thode de signalement / pr&eacute;qualification &agrave; Saint-&Eacute;tienne.")
    add_bullet(doc, "Orienter TVF vers les interlocuteurs adapt&eacute;s pour les sujets habitat, commerces, mat&eacute;riauth&egrave;que, insertion, usages temporaires ou r&eacute;emploi.")

    doc.add_heading(txt("Ce que TVF peut apporter"), level=1)
    add_number(doc, "Une lecture citoyenne et associative des ressources inutilis&eacute;es : biens, locaux, mat&eacute;riaux, &eacute;quipements, usages manquants.")
    add_number(doc, "Des dossiers pr&eacute;qualifi&eacute;s : localisation, statut apparent, photos, besoin exprim&eacute;, acteur demandeur, usage envisag&eacute;, limites connues.")
    add_number(doc, "Une mobilisation d'acteurs : propri&eacute;taires, associations, entreprises, artisans, structures d'insertion, b&eacute;n&eacute;voles et financeurs potentiels.")
    add_number(doc, "Une logique de r&eacute;emploi compatible avec les objectifs de sobri&eacute;t&eacute; fonci&egrave;re, de r&eacute;duction du gaspillage et de valorisation des ressources.")
    add_number(doc, "Un suivi d'utilit&eacute; territoriale : demandes re&ccedil;ues, ressources orient&eacute;es, dossiers instruits, acteurs mobilis&eacute;s, suites donn&eacute;es.")

    add_callout(doc, "Principe de prudence", "Aucune mention de partenariat, de soutien, de validation ou de financement par l'EPASE ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable. Ce courrier vise uniquement &agrave; solliciter un premier &eacute;change de cadrage.", fill=LIGHT_BLUE)

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(doc, "Nous proposons un rendez-vous d'environ une heure afin de vous pr&eacute;senter TVF, d'identifier les sujets compatibles, de fixer les limites de l'intervention associative et de d&eacute;terminer si un premier circuit de remont&eacute;e d'information ou de pr&eacute;qualification peut &ecirc;tre test&eacute; &agrave; Saint-&Eacute;tienne.")
    add_paragraph(doc, "&Agrave; l'issue de cet &eacute;change, TVF pourra transmettre une note de cadrage pr&eacute;cisant les sujets compatibles, les pi&egrave;ces utiles, les acteurs &agrave; mobiliser, les risques &agrave; pr&eacute;venir et les conditions de communication &agrave; respecter.")

    doc.add_heading("Conclusion", level=1)
    add_paragraph(doc, "Territoires Vivants France souhaite construire son lancement st&eacute;phanois avec s&eacute;rieux et lisibilit&eacute;. Un premier &eacute;change avec l'EPASE permettrait de mieux articuler la d&eacute;marche TVF avec les politiques et op&eacute;rations urbaines d&eacute;j&agrave; engag&eacute;es, sans confusion de r&ocirc;le ni effet d'annonce pr&eacute;matur&eacute;.")
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous &agrave; la date qui vous conviendra et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(head))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration officielle de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche friches / biens vacants", "D&eacute;crire les situations que TVF souhaite observer et qualifier.", "&Agrave; joindre"),
        ("Fiche mat&eacute;riaux et r&eacute;emploi", "Pr&eacute;senter le fonctionnement de la valorisation des ressources inutilis&eacute;es.", "&Agrave; joindre"),
        ("Fiche Habitat / Commerce", "Pr&eacute;senter les passerelles logements, locaux, commerces et usages de proximit&eacute;.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(doc, "Objet d'e-mail recommand&eacute;", "Demande de rendez-vous - TERRITOIRES VIVANTS FRANCE / EPASE - revitalisation, r&eacute;emploi et usages territoriaux", fill=LIGHT_BLUE)
    doc.add_heading("Sources de cadrage", level=1)
    add_paragraph(doc, "Les informations utilis&eacute;es pour orienter ce courrier proviennent du site officiel de l'EPASE, notamment les pages relatives aux missions, p&eacute;rim&egrave;tres, innovations, offres immobili&egrave;res, quartiers d'intervention et coordonn&eacute;es de l'&eacute;tablissement.")

    doc.save(DOCX_PATH)


def build_md():
    md = txt("""# Courrier - EPASE - lancement TVF

Destinataire : EPASE - &Eacute;tablissement Public d'Am&eacute;nagement de Saint-&Eacute;tienne.

Objet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les compl&eacute;mentarit&eacute;s possibles autour du b&acirc;ti vacant, des friches, des commerces, du r&eacute;emploi de mat&eacute;riaux et des usages territoriaux.

Date : Saint-&Eacute;tienne, le 14 juillet 2026.

## Intention

TVF sollicite un premier rendez-vous avec l'EPASE afin de pr&eacute;senter sa m&eacute;thode et d'identifier les sujets sur lesquels une contribution associative de type observation, pr&eacute;qualification, r&eacute;emploi ou mobilisation locale pourrait compl&eacute;ter les dispositifs existants.

## Principe de prudence

Aucune mention de partenariat, de soutien, de validation ou de financement par l'EPASE ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable.
""")
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
