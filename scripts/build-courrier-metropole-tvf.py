from pathlib import Path
from html import unescape
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents" / "courriers-lancement-saint-etienne"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "22-saint-etienne-metropole-lancement-cooperation-tvf.docx"
MD_PATH = SRC_DIR / "courrier-saint-etienne-metropole-lancement-tvf.md"
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"

GREEN = RGBColor(24, 63, 34)
BLUE = RGBColor(7, 30, 48)
MUTED = RGBColor(89, 100, 94)
LIGHT_GREEN = "EAF3EA"
LIGHT_BLUE = "EEF4F7"
GOLD = RGBColor(178, 132, 24)
BORDER = "B8C8BC"


def txt(value: str) -> str:
    return unescape(value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8E0D9", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = BLUE
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in [
        ("Heading 1", 16.5, GREEN, 14, 7),
        ("Heading 2", 12.8, GREEN, 11, 5),
        ("Heading 3", 11.5, BLUE, 7, 3),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Cm(17.3))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.3, 12.0])
    for cell in table.row_cells(0):
        set_cell_border(cell, color="FFFFFF", size="0")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left = table.cell(0, 0)
    p = left.paragraphs[0]
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(4.7))
    right = table.cell(0, 1)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("TERRITOIRES VIVANTS FRANCE\n"))
    r.bold = True
    r.font.size = Pt(9.8)
    r.font.color.rgb = GREEN
    r2 = p.add_run(txt("Association loi 1901 d&eacute;clar&eacute;e - RNA W423016361\nSIREN 107 226 128 - SIRET 107 226 128 00018\n25 rue &Eacute;lise Gervais, 42000 Saint-&Eacute;tienne\ncontact@territoiresvivantsfrance.fr - 04 65 81 54 69"))
    r2.font.size = Pt(8.2)
    r2.font.color.rgb = MUTED

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt("TERRITOIRES VIVANTS FRANCE - www.territoiresvivantsfrance.fr - Document de travail institutionnel"))
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED


def add_paragraph(doc, text, bold_first=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5.5)
    p.paragraph_format.line_spacing = 1.08
    if bold_first and text.startswith(bold_first):
        r = p.add_run(txt(bold_first))
        r.bold = True
        r.font.color.rgb = GREEN
        p.add_run(txt(text[len(bold_first):]))
    else:
        p.add_run(txt(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4.8)
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.add_run(txt(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4.8)
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.first_line_indent = Cm(-0.28)
    p.add_run(txt(text))
    return p


def add_callout(doc, title, body, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17.1)
    cell = table.cell(0, 0)
    cell.width = Cm(17.1)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=BORDER, size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(txt(title))
    r.bold = True
    r.font.color.rgb = GREEN
    r.font.size = Pt(11.2)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    p2.add_run(txt(body))
    return table


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        ("Destinataire", "Saint-&Eacute;tienne M&eacute;tropole - &agrave; l'attention de la Pr&eacute;sidence et des services comp&eacute;tents"),
        ("Adresse", "Cit&eacute; Gr&uuml;ner, 42000 Saint-&Eacute;tienne"),
        ("Objet", "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change pour une coop&eacute;ration pilote &agrave; Saint-&Eacute;tienne"),
        ("Pi&egrave;ces jointes", "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche besoins de lancement, fiche local de stockage / mat&eacute;riaux"),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - Saint-&Eacute;tienne M&eacute;tropole"))
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'un premier rendez-vous de cadrage"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(doc, "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e, dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un ancrage op&eacute;rationnel et progressif sur son territoire d'origine.")
    add_paragraph(doc, "Notre objet est de contribuer &agrave; remettre en usage des ressources aujourd'hui insuffisamment mobilis&eacute;es : logements vacants ou d&eacute;grad&eacute;s, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains inutilis&eacute;s, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. La d&eacute;marche repose sur une logique simple : identifier, qualifier, mettre en relation, conventionner, suivre et mesurer l'utilit&eacute; territoriale.")

    add_callout(doc, "Demande principale", "TVF sollicite un premier rendez-vous de cadrage avec Saint-&Eacute;tienne M&eacute;tropole afin de pr&eacute;senter sa m&eacute;thode, d'identifier les services concern&eacute;s et d'&eacute;tudier les conditions d'un lancement pilote utile au territoire.")

    doc.add_heading(txt("Pourquoi s'adresser &agrave; Saint-&Eacute;tienne M&eacute;tropole"), level=1)
    add_paragraph(doc, "La M&eacute;tropole occupe une place centrale dans les sujets que TVF souhaite traiter : habitat et logement, politique de la ville, d&eacute;veloppement &eacute;conomique, &eacute;conomie sociale et solidaire, d&eacute;chets, transition &eacute;cologique, am&eacute;nagement, PLUi, foncier, insertion et coop&eacute;ration avec les communes. TVF ne se positionne pas comme un dispositif concurrent, mais comme un outil d'appui capable de faciliter le passage du constat &agrave; l'action.")
    add_paragraph(doc, "L'association peut notamment contribuer &agrave; faire remonter des situations de terrain, pr&eacute;qualifier des besoins, orienter des ressources mat&eacute;rielles vers des projets utiles, structurer des premiers dossiers et mobiliser des partenaires publics, priv&eacute;s, associatifs ou citoyens autour d'une m&ecirc;me logique de valorisation territoriale.")

    doc.add_heading(txt("Axes de coop&eacute;ration propos&eacute;s"), level=1)
    add_number(doc, "Habitat et patrimoine vacant : rep&eacute;rer des situations, recueillir les informations utiles, orienter les propri&eacute;taires ou acteurs concern&eacute;s vers un cadre d'&eacute;tude et pr&eacute;parer des dossiers qualifi&eacute;s.")
    add_number(doc, "Mat&eacute;riaux et &eacute;conomie circulaire : identifier des ressources encore utilisables, organiser leur qualification, &eacute;viter le gaspillage et les affecter &agrave; des projets valid&eacute;s par TVF.")
    add_number(doc, "Commerces, locaux et activit&eacute;s de proximit&eacute; : contribuer &agrave; la lecture des besoins locaux et &agrave; l'&eacute;mergence d'usages utiles pour les habitants, associations, artisans ou porteurs de projets.")
    add_number(doc, "Friches, terrains et espaces sous-utilis&eacute;s : participer &agrave; la remont&eacute;e d'id&eacute;es, &agrave; la pr&eacute;qualification et &agrave; la recherche de partenaires capables de transformer ces lieux en ressources territoriales.")
    add_number(doc, "Solidarit&eacute;, insertion et engagement citoyen : orienter les projets vers des chantiers encadr&eacute;s, des structures d'insertion, des b&eacute;n&eacute;voles et des comp&eacute;tences locales.")

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "La d&eacute;signation d'un interlocuteur ou d'un service r&eacute;f&eacute;rent pour orienter TVF vers les bons services m&eacute;tropolitains.")
    add_bullet(doc, "Une mise en relation avec les services habitat, d&eacute;chets, &eacute;conomie circulaire, d&eacute;veloppement &eacute;conomique, ESS, politique de la ville, insertion, am&eacute;nagement et foncier.")
    add_bullet(doc, "L'identification d'un local ou espace de stockage temporaire, s&eacute;curis&eacute; et accessible, permettant de trier et qualifier des mat&eacute;riaux r&eacute;employables avant affectation &agrave; des projets utiles.")
    add_bullet(doc, "L'orientation vers des entreprises, bailleurs, acteurs de la construction, ressourceries, structures d'insertion, associations et communes susceptibles de contribuer au lancement pilote.")
    add_bullet(doc, "L'&eacute;tude d'un cadre de coop&eacute;ration territoriale progressif, sans communication de partenariat ni engagement public avant validation formelle par la M&eacute;tropole et par TVF.")

    add_callout(doc, "Principe de prudence", "Aucune mention de partenariat, de soutien, de financement ou de validation par Saint-&Eacute;tienne M&eacute;tropole ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable. Le premier &eacute;change demand&eacute; vise uniquement &agrave; cadrer, orienter et identifier les suites possibles.", fill=LIGHT_BLUE)

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(doc, "Nous proposons un rendez-vous d'une dur&eacute;e d'environ une heure afin de pr&eacute;senter TVF, de comprendre les priorit&eacute;s m&eacute;tropolitaines concern&eacute;es, d'identifier les services &agrave; associer et de d&eacute;terminer si un premier dossier pilote peut &ecirc;tre instruit &agrave; Saint-&Eacute;tienne.")
    add_paragraph(doc, "&Agrave; l'issue de ce rendez-vous, TVF pourra transmettre une note de cadrage plus pr&eacute;cise comprenant les besoins identifi&eacute;s, les acteurs &agrave; mobiliser, les pi&egrave;ces utiles, le calendrier envisag&eacute; et les limites d'intervention de l'association.")

    doc.add_heading("Conclusion", level=1)
    add_paragraph(doc, "Territoires Vivants France souhaite avancer de mani&egrave;re s&eacute;rieuse, progressive et tra&ccedil;able. Le lancement &agrave; Saint-&Eacute;tienne n'a pas pour objectif de multiplier les annonces, mais de construire une premi&egrave;re m&eacute;thode de terrain, utile aux habitants, compatible avec les politiques publiques locales et reproductible ensuite sur d'autres territoires.")
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous &agrave; la date qui vous conviendra et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    headers = ["Document", "Utilit&eacute;", "Statut"]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(text))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche local de stockage", "D&eacute;crire le besoin minimal : accessibilit&eacute;, s&eacute;curit&eacute;, stockage, tri, dur&eacute;e.", "&Agrave; joindre"),
        ("Fiche mat&eacute;riaux", "Pr&eacute;senter le fonctionnement de la valorisation des ressources inutilis&eacute;es.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(doc, "Objet d'e-mail recommand&eacute;", "Demande de rendez-vous - lancement pilote TERRITOIRES VIVANTS FRANCE avec Saint-&Eacute;tienne M&eacute;tropole", fill=LIGHT_BLUE)
    doc.save(DOCX_PATH)


def build_md():
    md = txt("""# Courrier - Saint-&Eacute;tienne M&eacute;tropole - lancement TVF\n\nDestinataire : Saint-&Eacute;tienne M&eacute;tropole, &agrave; l'attention de la Pr&eacute;sidence et des services comp&eacute;tents.\n\nObjet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change pour une coop&eacute;ration pilote &agrave; Saint-&Eacute;tienne.\n\nDate : Saint-&Eacute;tienne, le 14 juillet 2026.\n\n## Corps du courrier\n\nMadame, Monsieur,\n\nJ'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e, dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un ancrage op&eacute;rationnel et progressif sur son territoire d'origine.\n\nNotre objet est de contribuer &agrave; remettre en usage des ressources aujourd'hui insuffisamment mobilis&eacute;es : logements vacants ou d&eacute;grad&eacute;s, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains inutilis&eacute;s, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. La d&eacute;marche repose sur une logique simple : identifier, qualifier, mettre en relation, conventionner, suivre et mesurer l'utilit&eacute; territoriale.\n\nTVF sollicite un premier rendez-vous de cadrage avec Saint-&Eacute;tienne M&eacute;tropole afin de pr&eacute;senter sa m&eacute;thode, d'identifier les services concern&eacute;s et d'&eacute;tudier les conditions d'un lancement pilote utile au territoire.\n\n## Demandes op&eacute;rationnelles\n\n- D&eacute;signation d'un interlocuteur ou d'un service r&eacute;f&eacute;rent.\n- Mise en relation avec les services habitat, d&eacute;chets, &eacute;conomie circulaire, d&eacute;veloppement &eacute;conomique, ESS, politique de la ville, insertion, am&eacute;nagement et foncier.\n- Identification d'un local ou espace de stockage temporaire pour trier et qualifier des mat&eacute;riaux r&eacute;employables.\n- Orientation vers entreprises, bailleurs, acteurs de la construction, ressourceries, structures d'insertion, associations et communes.\n- &Eacute;tude d'un cadre de coop&eacute;ration territoriale progressif, sans communication de partenariat ni engagement public avant validation formelle.\n\n## Signature\n\nEdryan Rangoly  \nPr&eacute;sident fondateur  \nTERRITOIRES VIVANTS FRANCE\n""")
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)