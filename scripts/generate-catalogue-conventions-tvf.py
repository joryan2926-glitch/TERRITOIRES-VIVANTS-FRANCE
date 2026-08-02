from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image

ROOT = Path(r'C:\Users\jowst\Documents\TERRITOIRES VIVANTS FRANCE')
OUT = ROOT / 'documents' / 'brochures-tvf'
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / 'catalogue-conventions-cadres-tvf.docx'
PDF = OUT / 'catalogue-conventions-cadres-tvf.pdf'
LOGO = ROOT / 'assets' / 'logo-territoires-vivants-france-web.png'
GREEN = RGBColor(24,57,47)
DARK = RGBColor(31,41,55)
GRAY = RGBColor(102,112,133)
GOLD = RGBColor(180,148,90)
PALE = 'F3F7F2'
NOTE = 'FBF8EF'
BORDER = 'D9E0DC'
TVF = {
 'name':'Territoires Vivants France',
 'subtitle':'Agence Territoriale de Revitalisation Immobilière',
 'address':'25 rue Élise Gervais, 42000 Saint-Étienne',
 'phone':'04 65 81 54 69',
 'email':'contact@territoiresvivantsfrance.fr',
 'site':'www.territoiresvivantsfrance.fr',
 'rna':'RNA W922015538',
 'siret':'SIRET 897 226 138 00018',
}

CONVENTIONS = [
 ('Convention cadre propriétaire','Propriétaires, indivisions, représentants habilités','Étudier un bien, organiser la visite, encadrer les documents, photos, confidentialité et accompagnement TVF.','Le propriétaire reste propriétaire. TVF n’obtient aucun droit de gestion, vente, location ou perception financière sans convention spécifique.'),
 ('Convention de mise à disposition d’un bien','Propriétaires et porteurs de projets autorisés','Autoriser un usage défini du bien : stockage, action associative, expérimentation, occupation temporaire ou projet territorial.','Elle traite la durée, les charges, assurances, accès, entretien, restitution, indemnité ou participation éventuelle.'),
 ('Convention cadre collectivité','Communes, intercommunalités, départements, établissements publics','Encadrer le repérage, l’observatoire, les signalements, la qualification et le suivi territorial.','TVF agit en appui : elle ne remplace pas les pouvoirs publics, les services instructeurs ou les décisions de la collectivité.'),
 ('Convention cadre entreprise / partenaire','Entreprises, fondations, artisans, experts, chambres consulaires','Organiser expertise, mécénat, prestation, partenariat territorial, communication et action commune.','Le cadre distingue clairement don, mécénat, prestation, parrainage et partenariat opérationnel.'),
 ('Convention ressources et matériaux','Entreprises, distributeurs, artisans, collectivités, propriétaires','Encadrer don, collecte, stockage, contrôle, traçabilité et affectation de ressources réemployables.','Aucune réduction d’impôt n’est automatique : elle dépend du cadre légal, de l’éligibilité et des justificatifs.'),
 ('Convention projet solidaire / association','Associations, structures d’insertion, organismes de formation, porteurs de projets','Encadrer un projet social, éducatif, solidaire, d’insertion, d’apprentissage ou d’utilité territoriale.','Elle précise sécurité, encadrement, bénévoles, ressources, communication et limites de responsabilité.'),
 ('Convention confidentialité et données','Tous partenaires ayant accès à des informations sensibles','Protéger données personnelles, informations patrimoniales, documents, photos, accès TVF OS et confidentialité.','Elle peut être autonome ou intégrée comme annexe obligatoire à une autre convention.'),
]

UNIVERSAL = [
 ('Objet et périmètre','Chaque convention doit définir précisément pourquoi elle est signée, pour quel bien, quel projet, quel territoire, quelles ressources et quelles limites.'),
 ('Identification des parties','TVF, propriétaire, collectivité, entreprise, association, mandataire ou partenaire doivent être identifiés avec leur qualité et leur capacité à signer.'),
 ('Absence de mandat immobilier','Tant que le cadre légal ne l’autorise pas, une convention TVF ne doit pas être confondue avec un mandat de vente, location, gestion locative, administration de biens ou perception de loyers.'),
 ('Autorisation de visite','Ce n’est pas forcément une convention séparée : c’est une clause ou une annexe précisant l’accès, les personnes présentes, les zones visitées, la sécurité et les photos.'),
 ('Photos et communication','Les photos peuvent servir au dossier interne. Toute publication, communication ou diffusion externe doit être autorisée séparément.'),
 ('Confidentialité','Les données de propriétaire, adresse précise, documents, photos, devis, difficultés personnelles ou patrimoniales ne doivent pas circuler librement.'),
 ('RGPD et données','La finalité, les données collectées, les destinataires, les durées de conservation, la sécurité et les droits des personnes doivent être indiqués.'),
 ('Responsabilités et assurances','Chaque partie reste responsable de son rôle : propriétaire pour son bien, entreprise pour son intervention, association pour ses participants, collectivité pour ses compétences.'),
 ('Fiscalité prudente','Un don ou mécénat peut ouvrir des droits seulement si les conditions légales sont remplies. TVF ne doit jamais promettre une réduction d’impôt automatique.'),
 ('Fin, restitution et classement','Toute convention doit prévoir comment s’arrête l’intervention, comment le bien ou les ressources sont restitués, et comment le dossier est classé.'),
]

PUBLICS = [
 ('Propriétaire','Faire étudier son bien, conserver sa propriété, autoriser une visite, transmettre des documents, examiner une remise en usage ou une mise à disposition encadrée.','La convention doit rassurer : TVF accompagne et oriente, mais ne prend pas le contrôle du bien.'),
 ('Entreprise','Proposer matériaux, équipements, expertise, mécénat, prestation ou partenariat.','La convention doit clarifier l’intérêt territorial, la traçabilité, les responsabilités, la fiscalité et la communication.'),
 ('Collectivité','Organiser un cadre de coopération, repérage, observatoire, qualification, suivi et restitution.','La convention doit être institutionnelle : TVF agit comme appui opérationnel, pas comme autorité publique.'),
 ('Association / projet solidaire','Développer une action utile : insertion, atelier, formation, réemploi, activation temporaire ou projet local.','La convention doit encadrer sécurité, encadrement, assurances, usage des lieux et ressources mobilisées.'),
]

CHOICE = [
 ('Étudier un logement vacant avec un propriétaire','Convention cadre propriétaire','Annexe visite/photos + liste documents'),
 ('Autoriser TVF ou un projet à utiliser un local','Convention de mise à disposition d’un bien','État des lieux + assurance + durée + restitution'),
 ('Créer une coopération avec une mairie','Convention cadre collectivité','Périmètre territorial + référents + règles données'),
 ('Récupérer des matériaux auprès d’une entreprise','Convention ressources et matériaux','Inventaire + bon de remise + transport + stockage'),
 ('Recevoir un soutien financier ou expertise entreprise','Convention cadre entreprise / partenaire','Clause mécénat/prestation + communication'),
 ('Mettre en place un projet d’insertion ou association','Convention projet solidaire / association','Fiche projet + sécurité + participants'),
 ('Partager des documents sensibles avec un partenaire','Convention confidentialité et données','Liste personnes habilitées + durée conservation'),
]

REFERENCES = [
 'Service-Public Entreprendre, mécénat d’entreprise : dons en faveur d’organismes sans but lucratif, vérifié le 20 juillet 2026.',
 'Impots.gouv.fr, dons et réductions d’impôt, modifié le 24 novembre 2025.',
 'CNIL, les six grands principes du RGPD : finalité, minimisation, transparence, droits, durée de conservation et sécurité.',
 'CNIL, information des personnes et transparence : identité du responsable, finalités, base légale, destinataires et durée de conservation.',
]


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = tcPr.first_child_found_in('w:tcBorders')
    if tb is None:
        tb = OxmlElement('w:tcBorders')
        tcPr.append(tb)
    for e in ('top','left','bottom','right'):
        el = tb.find(qn('w:'+e))
        if el is None:
            el = OxmlElement('w:'+e)
            tb.append(el)
        el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),'6')
        el.set(qn('w:color'),BORDER)

def run(p, text, size=10, bold=False, color=DARK, font='Inter'):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    return r

def p(doc, text='', size=10, bold=False, color=DARK, before=0, after=5, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.14
    if align:
        para.alignment = align
    run(para, text, size, bold, color)
    return para

def h1(doc, text):
    para = p(doc, text, 17, True, GREEN, before=14, after=7)
    return para

def h2(doc, text):
    return p(doc, text, 13, True, GREEN, before=10, after=4)

def callout(doc, title, text, fill=NOTE):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    borders(cell)
    cell.text = ''
    pp = cell.paragraphs[0]
    pp.paragraph_format.space_after = Pt(2)
    run(pp, title + ' ', 9.2, True, GREEN)
    run(pp, text, 9.2, False, DARK)

def table_docx(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if widths:
        for i,w in enumerate(widths):
            t.columns[i].width = Inches(w)
    for i,head in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c, PALE); borders(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; c.text=''; run(c.paragraphs[0], head, 8.4, True, GREEN)
    for row in rows:
        cells = t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; borders(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; c.text=''; run(c.paragraphs[0], val, 8.1, False, DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def cover(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(.65); sec.bottom_margin = Inches(.65); sec.left_margin = Inches(.7); sec.right_margin = Inches(.7)
    if LOGO.exists():
        para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(LOGO), width=Inches(1.55))
    p(doc, 'CATALOGUE DES CONVENTIONS CADRES TVF', 23, True, GREEN, before=18, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'Guide interne d’utilisation, clauses universelles et règles de prudence', 12, False, GRAY, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    callout(doc, 'Objectif.', 'Donner à TVF une bibliothèque contractuelle claire : peu de grandes conventions cadres, enrichies par des clauses et annexes modulables selon les propriétaires, collectivités, entreprises, associations et partenaires.')
    p(doc, TVF['name'], 12, True, GREEN, before=18, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, TVF['subtitle'], 10, False, DARK, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, f"{TVF['address']} | {TVF['phone']} | {TVF['email']}\n{TVF['rna']} | {TVF['siret']}", 8.3, False, GRAY, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    callout(doc, 'Avertissement.', 'Ce catalogue est un outil interne d’organisation. Il ne remplace pas une consultation juridique. Chaque convention réelle doit être adaptée au bien, aux parties, aux responsabilités, aux assurances, aux règles fiscales et au cadre réglementaire applicable.')
    doc.add_page_break()

def build_docx():
    doc = Document()
    cover(doc)
    h1(doc, '1. La logique générale des conventions TVF')
    p(doc, 'TVF doit éviter de multiplier les petits documents isolés. La bonne méthode consiste à utiliser quelques conventions cadres solides, capables d’intégrer les autorisations, règles, annexes et clauses spécifiques nécessaires à chaque situation.')
    p(doc, 'Cette logique permet de sécuriser le travail sans rendre la relation trop lourde : une convention principale fixe le cadre, des annexes décrivent le bien, les matériaux, la visite, les documents, le projet ou les personnes habilitées.')
    callout(doc, 'Principe directeur.', 'Une convention cadre doit expliquer qui agit, sur quoi, pourquoi, pendant combien de temps, avec quelles responsabilités, quelles limites, quelles données et quelles suites possibles.')

    h1(doc, '2. Les 7 grandes conventions cadres')
    table_docx(doc, ['Convention', 'Public principal', 'Usage', 'Point de vigilance'], CONVENTIONS, [1.7,1.45,2.05,1.65])

    h1(doc, '3. Les publics concernés')
    for name, use, care in PUBLICS:
        h2(doc, name)
        p(doc, use)
        callout(doc, 'À sécuriser.', care, PALE)

    h1(doc, '4. Convention cadre propriétaire')
    p(doc, 'Cette convention est la porte d’entrée principale lorsqu’un propriétaire souhaite faire étudier un logement, commerce, bâtiment, terrain ou local vacant. Elle doit rassurer : le propriétaire conserve son bien, ses droits et sa décision finale.')
    p(doc, 'TVF peut étudier, visiter, recueillir des documents, orienter, mettre en relation et proposer une suite. En revanche, TVF ne doit pas donner l’impression de gérer le bien, de percevoir des loyers, de choisir un locataire, de vendre, de louer ou d’administrer le bien sans cadre professionnel adapté.')
    table_docx(doc, ['Bloc de clauses', 'Contenu attendu'], [
        ('Étude préalable','Description du bien, situation déclarée, objectifs du propriétaire, documents transmis, pistes de remise en usage.'),
        ('Visite et photos','Autorisation d’accès, zones visitées, sécurité, photos internes, diffusion publique interdite sans accord.'),
        ('Documents','Titre, taxe foncière, diagnostics existants, plans, devis, photos, courriers, informations notariées.'),
        ('Limites TVF','Aucune promesse de financement, travaux, location, vente, rendement, aide publique ou remise en usage automatique.'),
        ('Absence de mandat','Clause indiquant que la convention n’est pas un mandat immobilier ou de gestion locative.'),
        ('Fin d’accompagnement','Droit d’arrêt par le propriétaire ou TVF, classement sans suite, absence de réponse, impossibilité ou risque.'),
    ], [2.0,4.8])

    h1(doc, '5. Mise à disposition d’un bien : propriétaire toujours propriétaire')
    p(doc, 'La mise à disposition est différente de l’accompagnement. Ici, le propriétaire autorise un usage concret du bien : stockage, action associative, occupation temporaire, atelier, expérimentation ou projet territorial. Le propriétaire reste propriétaire. La convention ne transfère pas la propriété du bien.')
    p(doc, 'La convention doit préciser l’usage autorisé, la durée, les accès, les charges, les assurances, l’entretien, les interdictions, les travaux autorisés ou interdits et la restitution du bien. Elle doit éviter toute ambiguïté avec un bail, une gestion locative ou un mandat immobilier lorsque ce n’est pas le cadre choisi.')
    callout(doc, 'Rendement ou participation.', 'Si une participation financière, indemnité d’occupation, redevance, remboursement de charges ou rendement est envisagé, il faut le qualifier juridiquement avant signature. TVF ne doit pas promettre un rendement automatique ni organiser une gestion locative sans cadre professionnel adapté.')
    table_docx(doc, ['Sujet', 'Clause recommandée'], [
        ('Usage autorisé','Définir précisément ce qui peut être fait dans le bien, par qui, à quelles dates et avec quelles limites.'),
        ('Durée','Prévoir une durée fixe, renouvelable ou expérimentale, avec modalités de fin anticipée.'),
        ('Charges','Répartir fluides, entretien courant, nettoyage, déchets, assurance, taxes ou frais éventuels.'),
        ('Travaux','Interdire toute transformation sans accord écrit ; distinguer entretien, aménagement léger et travaux.'),
        ('Restitution','État des lieux d’entrée/sortie, clés, remise en état, retrait du matériel ajouté.'),
        ('Responsabilité','Clarifier dommage au bien, dommage aux tiers, responsabilité des utilisateurs, assurance obligatoire.'),
    ], [1.7,5.1])

    h1(doc, '6. Ressources et matériaux : entreprises, dons, traçabilité')
    p(doc, 'La convention ressources et matériaux permet de sécuriser les dons, mises à disposition, collectes, stockages et affectations de matériaux ou équipements. Elle protège l’entreprise, TVF et le projet bénéficiaire.')
    p(doc, 'L’entreprise peut contribuer par des matériaux neufs, surplus, fins de série, équipements, mobilier, outillage ou ressources techniques. TVF doit garder un droit de refus si les matériaux sont dangereux, trop dégradés, non conformes, inutiles, impossibles à transporter ou impossibles à stocker.')
    callout(doc, 'Fiscalité.', 'Le mécénat d’entreprise peut ouvrir droit à réduction d’impôt uniquement si les conditions légales sont remplies. Un don, même utile, ne donne pas automatiquement droit à un reçu fiscal. L’entreprise reste responsable de son appréciation fiscale.')
    table_docx(doc, ['Étape', 'Règle professionnelle'], [
        ('Proposition','Identifier donateur, nature des ressources, quantité, état, photos, lieu et contraintes.'),
        ('Contrôle','TVF vérifie utilité, sécurité, conformité, stockage et besoins réels.'),
        ('Collecte','Définir responsable, transport, chargement, assurances et coût éventuel.'),
        ('Stockage','Préciser lieu, durée, accès, inventaire, responsabilité et conditions de conservation.'),
        ('Affectation','Réserver aux projets TVF, partenaires conventionnés ou actions territoriales validées.'),
        ('Justificatifs','Bon de remise, inventaire, fiche de sortie, photos, affectation projet, réserve fiscale.'),
    ], [1.45,5.35])

    h1(doc, '7. Collectivités : le cadre d’intervention de TVF')
    p(doc, 'La convention collectivité doit expliquer que TVF intervient comme outil complémentaire de repérage, observation, qualification, coordination et suivi. TVF ne remplace pas les pouvoirs de police, les services instructeurs, les décisions publiques, les procédures d’urbanisme ou les professionnels réglementés.')
    p(doc, 'Elle fixe le périmètre géographique, les sujets concernés, les référents, les données partageables, la confidentialité, les restitutions, les indicateurs et les limites de publication. Elle doit distinguer les signalements non vérifiés, les situations qualifiées et les données consolidées.')
    table_docx(doc, ['Cadre', 'Contenu à préciser'], [
        ('Périmètre','Commune, quartier, secteur, territoire pilote, types de biens observés.'),
        ('Données','Ce qui peut être partagé, avec qui, pour quelle finalité, pendant combien de temps.'),
        ('Observatoire','Cartographie, fiches, indicateurs, rapports, données anonymisées ou agrégées.'),
        ('Propriétaires','Qui contacte, dans quel cadre, avec quelles limites et quelle traçabilité.'),
        ('TVF OS','Droits d’accès, dossiers visibles, besoin d’en connaître, journalisation.'),
        ('Restitution','Bilans, réunions, cartes, rapports, limites de communication publique.'),
    ], [1.5,5.3])

    h1(doc, '8. Entreprises, partenaires et projets solidaires')
    p(doc, 'La convention entreprise / partenaire sert à encadrer l’expertise, le mécénat, la prestation, le soutien financier, le mécénat de compétences ou le partenariat opérationnel. Elle doit distinguer clairement ce qui est gratuit, ce qui est rémunéré, ce qui relève du don, ce qui relève d’une prestation et ce qui relève d’un simple partenariat d’image.')
    p(doc, 'La convention projet solidaire / association sert à encadrer les actions utiles : insertion, formation, ateliers, activités associatives, chantiers encadrés, réemploi, animation locale ou usage temporaire. Elle doit traiter la sécurité, les participants, les assurances, l’encadrement et l’usage éventuel d’un bien.')

    h1(doc, '9. Confidentialité et données : clause transversale obligatoire')
    p(doc, 'La confidentialité n’est pas un détail. TVF manipule des adresses, photos, coordonnées, documents de propriété, situations personnelles, données patrimoniales, devis, informations sur des biens et échanges avec des partenaires. Ces informations doivent être protégées.')
    p(doc, 'Les principes à reprendre dans toutes les conventions sont : finalité déterminée, minimisation des données, information des personnes, droits d’accès et rectification, durée de conservation limitée, sécurité, accès restreint et interdiction de réutilisation non autorisée.')
    table_docx(doc, ['Principe', 'Application TVF'], UNIVERSAL[5:8], [1.55,5.25])

    h1(doc, '10. Clauses universelles TVF')
    table_docx(doc, ['Clause universelle', 'Rôle dans les conventions'], UNIVERSAL, [1.85,4.95])

    h1(doc, '11. Quelle convention utiliser ?')
    table_docx(doc, ['Situation', 'Convention principale', 'Annexes / articles à ajouter'], CHOICE, [2.35,2.05,2.4])

    h1(doc, '12. Ce que TVF doit éviter de promettre')
    for item in [
        'Un rendement automatique pour le propriétaire.',
        'Une réduction d’impôt automatique pour une entreprise ou un donateur.',
        'Une aide publique certaine ou un financement garanti.',
        'Une remise en usage certaine du bien.',
        'La réalisation de travaux sans convention, devis, assurance et accord écrit.',
        'Une gestion locative, recherche de locataire, perception de loyers ou mandat immobilier sans cadre professionnel adapté.',
        'La possibilité d’entrer dans un bien sans autorisation claire.',
        'La diffusion publique de photos ou données sans accord.'
    ]:
        p(doc, '• ' + item, 9.8, False, DARK, after=2)

    h1(doc, '13. Références de prudence')
    for ref in REFERENCES:
        p(doc, '• ' + ref, 8.8, False, GRAY, after=2)
    callout(doc, 'Conclusion.', 'Ce catalogue doit servir de guide interne pour choisir le bon modèle, assembler les bonnes clauses et éviter les promesses juridiquement fragiles. La signature réelle d’une convention doit toujours être adaptée au cas concret.')

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(footer, f"{TVF['name']} — Catalogue des conventions cadres TVF — document interne", 7, False, GRAY)
    doc.save(DOCX)

# PDF generation mirrors content in a controlled institutional layout.
def pdf_styles():
    st = getSampleStyleSheet()
    st.add(ParagraphStyle(name='T',fontName='Helvetica-Bold',fontSize=21,leading=25,textColor=colors.HexColor('#18392F'),alignment=1,spaceAfter=4))
    st.add(ParagraphStyle(name='Sub',fontName='Helvetica',fontSize=10.5,leading=14,textColor=colors.HexColor('#667085'),alignment=1,spaceAfter=8))
    st.add(ParagraphStyle(name='H1',fontName='Helvetica-Bold',fontSize=14.5,leading=18,textColor=colors.HexColor('#18392F'),spaceBefore=10,spaceAfter=5))
    st.add(ParagraphStyle(name='H2',fontName='Helvetica-Bold',fontSize=11.5,leading=14,textColor=colors.HexColor('#18392F'),spaceBefore=7,spaceAfter=3))
    st.add(ParagraphStyle(name='B',fontName='Helvetica',fontSize=9,leading=12,textColor=colors.HexColor('#1F2937'),spaceAfter=5))
    st.add(ParagraphStyle(name='Small',fontName='Helvetica',fontSize=7.5,leading=9.5,textColor=colors.HexColor('#667085'),spaceAfter=2))
    st.add(ParagraphStyle(name='Note',fontName='Helvetica',fontSize=8.2,leading=10.8,textColor=colors.HexColor('#475467'),backColor=colors.HexColor('#FBF8EF'),borderColor=colors.HexColor('#E4D7B8'),borderWidth=.5,borderPadding=6,spaceAfter=7))
    return st

def pdf_table(story, headers, rows, widths, st):
    data = [[Paragraph(h, st['Small']) for h in headers]]
    for row in rows:
        data.append([Paragraph(str(c), st['Small']) for c in row])
    t = Table(data, colWidths=[w*mm for w in widths], repeatRows=1)
    t.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),0.35,colors.HexColor('#D9E0DC')),
        ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#F3F7F2')),
        ('VALIGN',(0,0),(-1,-1),'TOP'),
        ('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4),
        ('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),4),
    ]))
    story.append(t); story.append(Spacer(1, 4*mm))

def build_pdf():
    st = pdf_styles(); story=[]
    if LOGO.exists():
        logo = Image(str(LOGO), width=32*mm, height=24*mm)
        story.append(logo)
    story.append(Spacer(1, 8*mm))
    story.append(Paragraph('CATALOGUE DES CONVENTIONS CADRES TVF', st['T']))
    story.append(Paragraph('Guide interne d’utilisation, clauses universelles et règles de prudence', st['Sub']))
    story.append(Paragraph('Objectif. Donner à TVF une bibliothèque contractuelle claire : peu de grandes conventions cadres, enrichies par des clauses et annexes modulables selon les propriétaires, collectivités, entreprises, associations et partenaires.', st['Note']))
    story.append(Paragraph(f"<b>{TVF['name']}</b><br/>{TVF['subtitle']}<br/>{TVF['address']} | {TVF['phone']} | {TVF['email']}<br/>{TVF['rna']} | {TVF['siret']}", st['Small']))
    story.append(PageBreak())
    for title, blocks in [
        ('1. La logique générale des conventions TVF', ['TVF doit éviter de multiplier les petits documents isolés. La bonne méthode consiste à utiliser quelques conventions cadres solides, capables d’intégrer les autorisations, règles, annexes et clauses spécifiques nécessaires à chaque situation.', 'Une convention principale fixe le cadre ; des annexes décrivent le bien, les matériaux, la visite, les documents, le projet ou les personnes habilitées.']),
    ]:
        story.append(Paragraph(title, st['H1']))
        for b in blocks: story.append(Paragraph(b, st['B']))
    story.append(Paragraph('2. Les 7 grandes conventions cadres', st['H1']))
    pdf_table(story, ['Convention','Public','Usage','Vigilance'], CONVENTIONS, [38,31,52,48], st)
    story.append(Paragraph('3. Les publics concernés', st['H1']))
    for name,use,care in PUBLICS:
        story.append(Paragraph(name, st['H2'])); story.append(Paragraph(use, st['B'])); story.append(Paragraph('<b>À sécuriser.</b> '+care, st['Note']))
    sections = [
        ('4. Convention cadre propriétaire', ['Cette convention est la porte d’entrée principale lorsqu’un propriétaire souhaite faire étudier un logement, commerce, bâtiment, terrain ou local vacant. Le propriétaire conserve son bien, ses droits et sa décision finale.', 'TVF peut étudier, visiter, recueillir des documents, orienter, mettre en relation et proposer une suite. TVF ne doit pas donner l’impression de gérer, vendre, louer ou administrer le bien sans cadre professionnel adapté.']),
        ('5. Mise à disposition d’un bien : propriétaire toujours propriétaire', ['La mise à disposition autorise un usage concret du bien : stockage, action associative, occupation temporaire, atelier, expérimentation ou projet territorial. Elle ne transfère pas la propriété du bien.', 'Si participation financière, indemnité d’occupation, redevance, remboursement de charges ou rendement sont envisagés, il faut qualifier juridiquement le dispositif avant signature.']),
        ('6. Ressources et matériaux : entreprises, dons, traçabilité', ['La convention ressources sécurise dons, mises à disposition, collectes, stockages et affectations de matériaux ou équipements.', 'Le mécénat d’entreprise peut ouvrir droit à réduction d’impôt uniquement si les conditions légales sont remplies. Aucun avantage fiscal ne doit être promis automatiquement.']),
        ('7. Collectivités : le cadre d’intervention de TVF', ['TVF intervient comme outil complémentaire de repérage, observation, qualification, coordination et suivi.', 'TVF ne remplace pas les pouvoirs de police, services instructeurs, décisions publiques, procédures d’urbanisme ou professionnels réglementés.']),
        ('8. Entreprises, partenaires et projets solidaires', ['La convention entreprise distingue don, mécénat, prestation, parrainage et partenariat opérationnel.', 'La convention projet solidaire traite insertion, formation, ateliers, sécurité, participants, assurances, encadrement et usage éventuel d’un bien.']),
        ('9. Confidentialité et données', ['TVF manipule adresses, photos, coordonnées, documents de propriété, situations personnelles, données patrimoniales, devis et échanges avec partenaires. Ces informations doivent être protégées.', 'Les principes à reprendre sont : finalité, minimisation, information, droits des personnes, durée de conservation limitée, sécurité, accès restreint et interdiction de réutilisation non autorisée.']),
    ]
    for title, paras in sections:
        story.append(Paragraph(title, st['H1']))
        for para in paras: story.append(Paragraph(para, st['B']))
    story.append(Paragraph('10. Clauses universelles TVF', st['H1']))
    pdf_table(story, ['Clause','Rôle'], UNIVERSAL, [45,124], st)
    story.append(Paragraph('11. Quelle convention utiliser ?', st['H1']))
    pdf_table(story, ['Situation','Convention','Annexes / articles'], CHOICE, [58,49,62], st)
    story.append(Paragraph('12. Ce que TVF doit éviter de promettre', st['H1']))
    for item in ['Un rendement automatique pour le propriétaire.','Une réduction d’impôt automatique pour une entreprise ou un donateur.','Une aide publique certaine ou un financement garanti.','Une remise en usage certaine du bien.','La réalisation de travaux sans convention, devis, assurance et accord écrit.','Une gestion locative, recherche de locataire, perception de loyers ou mandat immobilier sans cadre professionnel adapté.','La possibilité d’entrer dans un bien sans autorisation claire.','La diffusion publique de photos ou données sans accord.']:
        story.append(Paragraph('• '+item, st['B']))
    story.append(Paragraph('13. Références de prudence', st['H1']))
    for ref in REFERENCES: story.append(Paragraph('• '+ref, st['Small']))
    story.append(Paragraph('Conclusion. Ce catalogue sert à choisir le bon modèle, assembler les bonnes clauses et éviter les promesses juridiquement fragiles. Chaque signature réelle doit être adaptée au cas concret.', st['Note']))
    SimpleDocTemplate(str(PDF), pagesize=A4, rightMargin=14*mm, leftMargin=14*mm, topMargin=13*mm, bottomMargin=13*mm).build(story)

build_docx(); build_pdf()
print(DOCX); print(PDF)
