from pathlib import Path
from runpy import run_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

OUT_DIR = ROOT / "documents" / "courriers-lancement-saint-etienne"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "26-renovactions42-lancement-cooperation-renovation-tvf.docx"
MD_PATH = SRC_DIR / "courrier-renovactions42-lancement-tvf.md"

txt = BASE["txt"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_number = BASE["add_number"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
BORDER = BASE["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        (
            "Destinataire",
            "R&eacute;nov'actions42 - Service public de la r&eacute;novation de l'habitat - &agrave; l'attention de la coordination du dispositif et de l'ALEC42",
        ),
        (
            "Adresse",
            "Maison de l'Habitat et du Logement - 20 A rue Bala&yuml;, 42000 Saint-&Eacute;tienne - 04 77 41 41 25",
        ),
        (
            "Objet",
            "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les compl&eacute;mentarit&eacute;s possibles autour des logements vacants, de l'habitat d&eacute;grad&eacute;, des propri&eacute;taires et du r&eacute;emploi de mat&eacute;riaux",
        ),
        (
            "Pi&egrave;ces jointes",
            "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche Habitat Vivant, fiche propri&eacute;taires, fiche logement vacant / bien &agrave; qualifier, fiche mat&eacute;riaux et r&eacute;emploi",
        ),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def add_complementarity_table(doc):
    doc.add_heading(txt("Compl&eacute;mentarit&eacute;s possibles &agrave; cadrer"), level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.5, 6.4, 6.2])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Besoin constat&eacute;", "Apport possible de TVF", "Lien avec R&eacute;nov'actions42"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(head))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.1)

    rows = [
        (
            "Logements vacants ou d&eacute;grad&eacute;s",
            "Rep&eacute;rer les situations, recueillir les premiers &eacute;l&eacute;ments et orienter les propri&eacute;taires vers un cadre d'&eacute;tude.",
            "V&eacute;rifier les modalit&eacute;s d'orientation vers le service public comp&eacute;tent lorsque le sujet rel&egrave;ve de la r&eacute;novation &eacute;nerg&eacute;tique.",
        ),
        (
            "Propri&eacute;taires peu inform&eacute;s",
            "Expliquer les premi&egrave;res &eacute;tapes, aider &agrave; constituer une fiche de demande et clarifier les pi&egrave;ces utiles.",
            "S&eacute;curiser le passage vers les conseillers R&eacute;nov'actions42 pour les aides, le diagnostic, le plan de travaux et les professionnels.",
        ),
        (
            "Mat&eacute;riaux r&eacute;employables",
            "Identifier des ressources issues de travaux, surplus, mobiliers ou &eacute;quipements pouvant &ecirc;tre valoris&eacute;s dans des projets valid&eacute;s.",
            "D&eacute;terminer les limites techniques, assurantielles et r&eacute;glementaires du r&eacute;emploi dans les projets de r&eacute;novation.",
        ),
        (
            "Habitat ancien",
            "Mettre en relation des besoins sociaux, associatifs ou territoriaux avec des biens &agrave; remettre en usage.",
            "Compl&eacute;ter l'accompagnement technique par une lecture territoriale : usage futur, acteurs &agrave; mobiliser, b&eacute;n&eacute;fices locaux.",
        ),
        (
            "Suivi de dossier",
            "Documenter les demandes TVF : num&eacute;ro de dossier, statut, pi&egrave;ces, acteurs, suites donn&eacute;es, limites d'intervention.",
            "Cr&eacute;er un circuit clair afin d'&eacute;viter les doublons, les promesses non valid&eacute;es et les orientations incompl&egrave;tes.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.5)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - R&eacute;nov'actions42"))
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Demande d'un premier &eacute;change sur les passerelles entre revitalisation territoriale et r&eacute;novation de l'habitat"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame, Monsieur,")
    add_paragraph(
        doc,
        "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un ancrage op&eacute;rationnel progressif sur son territoire d'origine.",
    )
    add_paragraph(
        doc,
        "Notre d&eacute;marche vise &agrave; remettre en usage des ressources aujourd'hui insuffisamment mobilis&eacute;es : logements vacants ou d&eacute;grad&eacute;s, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains inutilis&eacute;s, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. Dans le champ de l'habitat, TVF ne se substitue pas aux services publics de r&eacute;novation : l'association peut rep&eacute;rer, qualifier, orienter et accompagner la constitution de premiers dossiers lorsque les situations rencontr&eacute;es touchent &agrave; la vacance, &agrave; la d&eacute;gradation du b&acirc;ti ou &agrave; la recherche d'un usage utile.",
    )
    add_paragraph(
        doc,
        "R&eacute;nov'actions42 repr&eacute;sente un interlocuteur essentiel pour cette approche. Le dispositif accompagne gratuitement les acheteurs, propri&eacute;taires et propri&eacute;taires bailleurs dans leurs d&eacute;marches de r&eacute;novation &eacute;nerg&eacute;tique, depuis le conseil et le diagnostic jusqu'au programme de travaux, au financement, &agrave; la mise en relation avec des entreprises locales et au suivi de la consommation apr&egrave;s travaux.",
    )

    add_callout(
        doc,
        "Demande principale",
        "TVF sollicite un premier rendez-vous avec R&eacute;nov'actions42 afin de pr&eacute;senter sa m&eacute;thode, de comprendre les circuits d'orientation existants et de d&eacute;finir comment les demandes rep&eacute;r&eacute;es par TVF peuvent &ecirc;tre dirig&eacute;es vers le bon service public sans cr&eacute;er de doublon ni de confusion pour les propri&eacute;taires.",
    )

    doc.add_heading(txt("Pourquoi solliciter R&eacute;nov'actions42"), level=1)
    add_paragraph(
        doc,
        "Le territoire de la Loire est concern&eacute; par un parc de logements anciens et &eacute;nergivores. R&eacute;nov'actions42 rappelle notamment que 60 % des logements du parc priv&eacute; ont &eacute;t&eacute; construits avant 1975, ce qui confirme l'importance des enjeux d'information, de conseil, d'orientation et d'accompagnement des propri&eacute;taires.",
    )
    add_paragraph(
        doc,
        "TVF intervient en amont ou en compl&eacute;ment de cette logique : rep&eacute;rer des logements vacants ou d&eacute;grad&eacute;s, entrer en contact avec des propri&eacute;taires, comprendre les freins &agrave; la remise en usage, identifier les ressources mat&eacute;rielles mobilisables et orienter les situations vers les interlocuteurs comp&eacute;tents. L'objectif est d'&eacute;viter les parcours dispers&eacute;s et de faciliter le passage du constat &agrave; une instruction structur&eacute;e.",
    )
    add_paragraph(
        doc,
        "Cette compl&eacute;mentarit&eacute; est particuli&egrave;rement utile pour les biens qui ne sont pas seulement des projets de travaux, mais aussi des situations territoriales : logement vacant depuis plusieurs ann&eacute;es, propri&eacute;taire isol&eacute;, besoin d'usage solidaire, mobilisation de mat&eacute;riaux, participation d'artisans ou d'acteurs d'insertion, articulation avec des associations ou collectivit&eacute;s.",
    )

    add_complementarity_table(doc)

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "Organiser un rendez-vous avec la coordination R&eacute;nov'actions42 ou l'ALEC42 afin de pr&eacute;senter TVF et de comprendre les conditions d'orientation vers le service public de la r&eacute;novation de l'habitat.")
    add_bullet(doc, "D&eacute;terminer les informations minimales qu'un dossier TVF devrait collecter avant toute orientation : adresse, statut du bien, propri&eacute;taire, usage envisag&eacute;, niveau de d&eacute;gradation, urgence, photos, besoins, pi&egrave;ces disponibles.")
    add_bullet(doc, "Identifier les cas dans lesquels TVF doit simplement orienter vers R&eacute;nov'actions42, et ceux dans lesquels TVF peut continuer &agrave; instruire un volet territorial compl&eacute;mentaire : usage solidaire, r&eacute;emploi, b&eacute;n&eacute;volat, partenaires, recherche de local ou projet associatif.")
    add_bullet(doc, "Construire une fiche d'orientation commune ou compatible, afin d'&eacute;viter les doublons, les demandes incompl&egrave;tes et les mauvaises attentes des propri&eacute;taires.")
    add_bullet(doc, "Examiner si un premier circuit pilote peut &ecirc;tre test&eacute; sur Saint-&Eacute;tienne, sans communication de partenariat ni engagement public avant validation &eacute;crite.")

    doc.add_heading(txt("Ce que TVF peut apporter aux propri&eacute;taires et au territoire"), level=1)
    add_number(doc, "Un premier rep&eacute;rage des biens vacants ou d&eacute;grad&eacute;s, avec une fiche claire et tra&ccedil;able.")
    add_number(doc, "Une aide &agrave; la formulation du besoin : remettre en location, vendre, r&eacute;nover, proposer un usage temporaire, envisager un projet solidaire ou rechercher une solution progressive.")
    add_number(doc, "Une orientation vers les bons interlocuteurs pour les sujets techniques, &eacute;nerg&eacute;tiques, financiers, juridiques ou administratifs.")
    add_number(doc, "Une capacit&eacute; &agrave; relier le projet de r&eacute;novation &agrave; des b&eacute;n&eacute;fices territoriaux : cadre de vie, logement, insertion, r&eacute;emploi de mat&eacute;riaux, activit&eacute; locale, mobilisation associative.")
    add_number(doc, "Une m&eacute;thode de suivi permettant de conserver les pi&egrave;ces, les contacts, les &eacute;tapes, les d&eacute;cisions et les limites de chaque dossier.")

    add_callout(
        doc,
        "Principe de prudence",
        "Aucune mention de partenariat, de soutien, de validation ou de financement par R&eacute;nov'actions42, l'ALEC42 ou les collectivit&eacute;s porteuses ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable. Ce courrier vise uniquement &agrave; solliciter un premier &eacute;change de cadrage.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(
        doc,
        "Nous proposons un rendez-vous d'environ une heure afin de vous pr&eacute;senter TVF, d'identifier les sujets compatibles avec R&eacute;nov'actions42, de pr&eacute;ciser les informations &agrave; collecter et de d&eacute;finir un circuit simple d'orientation pour les propri&eacute;taires ou porteurs de projets rencontr&eacute;s par l'association.",
    )
    add_paragraph(
        doc,
        "&Agrave; l'issue de cet &eacute;change, TVF pourra produire une note de cadrage indiquant les cas d'orientation, les limites d'intervention, les pi&egrave;ces &agrave; demander, les interlocuteurs &agrave; mobiliser et les pr&eacute;cautions de communication &agrave; respecter.",
    )

    doc.add_heading("Conclusion", level=1)
    add_paragraph(
        doc,
        "Territoires Vivants France souhaite construire son lancement &agrave; Saint-&Eacute;tienne de mani&egrave;re s&eacute;rieuse, utile et compl&eacute;mentaire des dispositifs existants. Le rapprochement avec R&eacute;nov'actions42 permettrait de mieux orienter les propri&eacute;taires concern&eacute;s, de renforcer la qualit&eacute; des dossiers transmis et de relier la r&eacute;novation de l'habitat &agrave; une logique plus large de revitalisation territoriale.",
    )
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous &agrave; la date qui vous conviendra et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame, Monsieur, l'expression de ma consid&eacute;ration distingu&eacute;e.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(head))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration officielle de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche Habitat Vivant", "Pr&eacute;senter le p&ocirc;le TVF li&eacute; aux logements vacants, d&eacute;grad&eacute;s et &agrave; remettre en usage.", "&Agrave; joindre"),
        ("Fiche propri&eacute;taire", "D&eacute;crire les informations &agrave; collecter aupr&egrave;s d'un propri&eacute;taire avant orientation ou instruction.", "&Agrave; joindre"),
        ("Fiche logement vacant / bien &agrave; qualifier", "Structurer les premiers &eacute;l&eacute;ments : adresse, &eacute;tat, usage, photos, besoin, contraintes.", "&Agrave; joindre"),
        ("Fiche mat&eacute;riaux et r&eacute;emploi", "Pr&eacute;senter le fonctionnement de la valorisation des ressources inutilis&eacute;es.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)

    doc.add_paragraph()
    add_callout(
        doc,
        "Objet d'e-mail recommand&eacute;",
        "Demande de rendez-vous - TERRITOIRES VIVANTS FRANCE / R&eacute;nov'actions42 - habitat vacant, propri&eacute;taires et r&eacute;novation",
        fill=LIGHT_BLUE,
    )
    doc.add_heading("Sources de cadrage", level=1)
    add_paragraph(
        doc,
        "Les informations utilis&eacute;es pour orienter ce courrier proviennent du site officiel R&eacute;nov'actions42, notamment les mentions relatives au service public de la r&eacute;novation de l'habitat, &agrave; l'accompagnement gratuit, aux publics concern&eacute;s, aux &eacute;tapes du projet, aux coordonn&eacute;es et &agrave; la mise en oeuvre op&eacute;rationnelle par l'ALEC42.",
    )

    doc.save(DOCX_PATH)


def build_md():
    md = txt(
        """# Courrier - R&eacute;nov'actions42 - lancement TVF

Destinataire : R&eacute;nov'actions42, service public de la r&eacute;novation de l'habitat, &agrave; l'attention de la coordination du dispositif et de l'ALEC42.

Objet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'&eacute;change sur les compl&eacute;mentarit&eacute;s possibles autour des logements vacants, de l'habitat d&eacute;grad&eacute;, des propri&eacute;taires et du r&eacute;emploi de mat&eacute;riaux.

Date : Saint-&Eacute;tienne, le 14 juillet 2026.

## Intention du courrier

Ce courrier sollicite un premier rendez-vous avec R&eacute;nov'actions42 afin de pr&eacute;senter TVF, de comprendre les circuits d'orientation existants et de d&eacute;finir comment les demandes rep&eacute;r&eacute;es par TVF peuvent &ecirc;tre dirig&eacute;es vers le bon service public sans doublon ni confusion.

## Demandes op&eacute;rationnelles

- Organiser un rendez-vous avec la coordination R&eacute;nov'actions42 ou l'ALEC42.
- D&eacute;terminer les informations minimales qu'un dossier TVF devrait collecter avant orientation.
- Identifier les cas d'orientation simple vers R&eacute;nov'actions42 et les cas o&ugrave; TVF peut instruire un volet territorial compl&eacute;mentaire.
- Construire une fiche d'orientation compatible.
- Examiner un premier circuit pilote &agrave; Saint-&Eacute;tienne.

## Principe de prudence

Aucune mention de partenariat, de soutien, de validation ou de financement par R&eacute;nov'actions42, l'ALEC42 ou les collectivit&eacute;s porteuses ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable.

## Signature

Edryan Rangoly  
Pr&eacute;sident fondateur  
TERRITOIRES VIVANTS FRANCE
"""
    )
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
