from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image

ROOT = Path(r'C:\Users\jowst\Documents\TERRITOIRES VIVANTS FRANCE')
OUT = ROOT / 'documents' / 'conventions-tvf'
OUT.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / 'assets' / 'logo-territoires-vivants-france-web.png'
GREEN = RGBColor(24,57,47); DARK = RGBColor(31,41,55); GRAY = RGBColor(102,112,133)
PALE='F3F7F2'; NOTE='FBF8EF'; BORDER='D9E0DC'
TVF = {'name':'Territoires Vivants France','subtitle':'Agence Territoriale de Revitalisation Immobilière','address':'25 rue Élise Gervais, 42000 Saint-Étienne','phone':'04 65 81 54 69','email':'contact@territoiresvivantsfrance.fr','site':'www.territoiresvivantsfrance.fr','rna':'RNA W922015538','siret':'SIRET 897 226 138 00018'}

CONVENTIONS = [
('convention-cadre-proprietaire-tvf','CONVENTION CADRE PROPRIÉTAIRE','Étude préalable, visite, documents, confidentialité et accompagnement',[
('1. Parties et identification du bien',["La convention est conclue entre Territoires Vivants France, association loi 1901 agissant comme Agence Territoriale de Revitalisation Immobilière, et le propriétaire, représentant ou mandataire habilité du bien concerné.","Le bien est désigné par son adresse, sa commune, sa nature, ses références cadastrales si connues et tout élément permettant son identification. Une annexe peut détailler plusieurs biens." ]),
('2. Objet de la convention',["La convention encadre l’étude préalable du bien, l’analyse des informations transmises, la visite éventuelle, la prise de photographies, la réception de documents et l’accompagnement possible par TVF.","Elle ne constitue ni mandat immobilier, ni mandat de gestion locative, ni mandat de vente ou location, ni promesse de travaux, ni garantie de financement, ni engagement automatique de remise en usage." ]),
('3. Déclarations du propriétaire',["Le signataire déclare être propriétaire, représentant habilité ou disposer d’un droit suffisant pour échanger sur le bien. Il s’engage à transmettre des informations sincères et à signaler toute difficulté juridique, successorale, locative, technique ou administrative connue.","Lorsque la qualité du signataire reste à confirmer, TVF peut demander tout justificatif utile avant de poursuivre l’instruction." ]),
('4. Étude préalable et rôle de TVF',["TVF peut ouvrir une fiche interne, classer les informations, analyser la situation apparente, étudier les usages envisageables, rechercher les interlocuteurs utiles et formuler des pistes d’orientation.","TVF agit comme structure d’observation, d’accompagnement, d’orientation et de coordination. Elle ne se substitue pas aux professionnels réglementés, aux autorités compétentes, aux diagnostiqueurs, architectes, notaires, entreprises ou organismes financeurs." ]),
('5. Visite, accès et sécurité',["Toute visite nécessite l’accord préalable du propriétaire ou de son représentant. Les zones accessibles, la date, les personnes présentes et les conditions de sécurité sont définies avant l’intervention.","TVF peut refuser ou interrompre une visite en cas d’accès incertain, de danger apparent, d’opposition d’un tiers, d’absence d’autorisation claire ou de conditions matérielles insuffisantes." ]),
('6. Photographies et documents',["Le propriétaire autorise TVF à prendre des photographies strictement nécessaires au suivi interne du dossier. Toute diffusion publique, usage de communication ou transmission extérieure non nécessaire fera l’objet d’un accord spécifique.","TVF peut demander des documents utiles : titre de propriété, taxe foncière, diagnostics, plans, photos, devis, courriers, documents notariaux, assurance, coordonnées de mandataire ou toute pièce permettant de comprendre la situation." ]),
('7. Confidentialité et données personnelles',["Les parties s’engagent à préserver la confidentialité des documents, coordonnées, photos, informations patrimoniales, échanges et difficultés personnelles ou juridiques portées à leur connaissance.","Les données personnelles sont traitées uniquement pour l’étude, l’orientation et le suivi du dossier. Les personnes concernées peuvent demander l’accès, la rectification ou l’exercice de leurs droits auprès de TVF." ]),
('8. Responsabilités et assurances',["Le propriétaire conserve la responsabilité de son bien, de son état, de ses assurances et des informations transmises. TVF intervient dans les limites de sa mission d’accompagnement.","Chaque partenaire ou professionnel intervenant conserve la responsabilité de sa mission et doit disposer des assurances, qualifications et autorisations nécessaires." ]),
('9. Durée, fin et suites possibles',["La convention est conclue pour une durée à compléter. Elle peut prendre fin à la demande d’une partie, en cas d’absence de réponse, d’impossibilité juridique ou technique, de risque non maîtrisé ou de situation ne relevant pas de TVF.","Les suites possibles sont : poursuite de l’instruction, visite complémentaire, orientation partenaire, proposition TVF, convention spécifique, mise à disposition, classement sans suite ou autre solution adaptée." ])], ['Désignation du bien','Liste des documents transmis','Autorisation de visite et photos','Fin d’accompagnement ou classement sans suite']),
('convention-ressources-materiaux-tvf','CONVENTION RESSOURCES ET MATÉRIAUX','Don, collecte, stockage, réemploi et affectation des ressources',[
('1. Parties et objet',["La convention est conclue entre TVF et le donateur, partenaire, entreprise, collectivité, association ou propriétaire proposant des matériaux, équipements, mobiliers, stocks ou ressources.","Elle organise la proposition, l’étude, l’acceptation éventuelle, la collecte, le stockage, le référencement, la mise à disposition ou l’affectation de ressources dans le cadre des missions de TVF." ]),
('2. Ressources concernées',["Les ressources peuvent comprendre notamment bois, carrelage, portes, fenêtres, sanitaires, mobilier, luminaires, outillage, équipements de bureau, matériaux de finition, fins de série, stocks inutilisés ou éléments de réemploi.","Une annexe inventaire précise la nature, la quantité, l’état apparent, la localisation, les photos, la disponibilité, la valeur estimative si connue et les contraintes de manutention ou transport." ]),
('3. Mode de contribution',["La contribution peut prendre la forme d’un don, d’une mise à disposition, d’un prêt, d’un dépôt temporaire, d’une liquidation solidaire ou d’un mécénat en nature lorsque les conditions sont réunies.","Le transfert ou non de propriété doit être clairement indiqué. À défaut de précision contraire, une ressource acceptée et remise à TVF est destinée à être affectée par TVF selon les besoins de ses projets." ]),
('4. Contrôle préalable et refus',["TVF se réserve le droit d’accepter ou refuser toute proposition selon l’état, la sécurité, l’utilité, la conformité, les possibilités de transport, les capacités de stockage et les besoins réellement identifiés.","Sont notamment refusables les ressources dangereuses, polluées, non conformes, trop dégradées, inutilisables, sans utilité identifiée ou impossibles à collecter ou stocker dans de bonnes conditions." ]),
('5. Collecte, transport et stockage',["Les modalités de collecte sont définies avant intervention : lieu, date, responsable, accès, chargement, moyens matériels, sécurité, coût éventuel et responsabilité pendant le transport.","Le stockage peut être assuré par TVF, le proposant ou un partenaire conventionné. Le lieu, la durée, les accès, les assurances et les conditions de conservation sont précisés lorsque nécessaire." ]),
('6. Affectation et traçabilité',["Les ressources acceptées sont affectées prioritairement à des projets suivis par TVF, actions territoriales identifiées, partenaires conventionnés ou opérations correspondant à l’objet de l’association.","TVF peut établir bon de remise, fiche d’entrée, inventaire, fiche de sortie ou justificatif d’affectation. La traçabilité reste proportionnée à la nature et à la valeur des ressources." ]),
('7. Fiscalité, mécénat et communication',["La contribution ne donne pas automatiquement droit à un avantage fiscal ou à un reçu fiscal. Toute qualification de mécénat doit être vérifiée selon la réglementation applicable et la situation des parties.","Le nom, le logo, l’image ou les informations commerciales du partenaire ne peuvent être utilisés en communication publique qu’avec accord préalable." ]),
('8. Responsabilités, assurances et confidentialité',["Chaque partie demeure responsable des opérations qu’elle réalise directement. Les intervenants professionnels doivent disposer des assurances et qualifications nécessaires.","Les informations non publiques relatives aux stocks, prix, volumes, contraintes internes, coordonnées ou projets suivis sont confidentielles, sauf nécessité d’exécution ou obligation légale." ]),
('9. Durée et résiliation',["La convention peut être ponctuelle ou conclue pour une durée déterminée. Elle peut être modifiée par écrit si les conditions de contribution, stockage, transport ou affectation évoluent.","TVF peut cesser la prise en charge lorsqu’une ressource devient inutilisable, non conforme, impossible à stocker, non affectable ou contraire à ses missions." ])], ['Inventaire des ressources','Bon de remise ou collecte','Fiche de stockage','Fiche de sortie ou affectation projet']),
('convention-cadre-collectivite-tvf','CONVENTION CADRE COLLECTIVITÉ','Coopération territoriale, observatoire, repérage et suivi',[
('1. Parties, contexte et finalité',["La convention est conclue entre TVF et la collectivité, l’établissement public ou la structure territoriale partenaire représenté par son représentant habilité.","Elle répond aux enjeux de vacance immobilière, commerces fermés, friches, terrains sans usage, bâtiments délaissés, ressources non mobilisées et revitalisation durable du territoire." ]),
('2. Objet et périmètre',["La convention organise une coopération autour du repérage, de l’observation, de la qualification, du suivi et de l’orientation des situations immobilières ou territoriales.","Le périmètre géographique et thématique est défini en annexe : commune, quartier, secteur, territoire pilote, habitat, commerce, friches, foncier, matériaux, solidarité ou observatoire." ]),
('3. Missions de TVF',["TVF peut contribuer à l’organisation des signalements, fiches de qualification, repérages, suivis, orientations, mises en relation, réunions et restitutions adaptées au périmètre convenu.","TVF agit comme outil complémentaire d’observation, coordination, accompagnement et suivi. Elle ne remplace pas les pouvoirs de police, services instructeurs, décisions publiques ou professionnels réglementés." ]),
('4. Missions de la collectivité',["La collectivité peut désigner un référent, préciser ses priorités, transmettre les informations légalement communicables, faciliter l’identification des services compétents et participer au suivi.","Elle conserve la responsabilité de ses compétences, décisions administratives, procédures internes, obligations réglementaires et vérification du cadre de partage des données." ]),
('5. Données et confidentialité',["Les parties définissent les informations échangées, leurs finalités, les accès autorisés, les mesures de sécurité, les durées de conservation et les règles de confidentialité.","La convention ne peut autoriser l’accès à des données cadastrales, fiscales, foncières, administratives ou personnelles dont la communication serait contraire au cadre légal applicable." ]),
('6. Signalements et qualification',["Un signalement ne permet pas à lui seul d’affirmer juridiquement qu’un bien est vacant, abandonné, dangereux ou sans propriétaire. Les situations doivent être qualifiées avec prudence.","Chaque situation peut être classée selon son type, son niveau d’information, son état d’avancement, son degré de priorité, les vérifications nécessaires et les suites envisageables." ]),
('7. Observatoire territorial et TVF OS',["TVF peut contribuer à un observatoire territorial : fiches, cartes, tableaux de bord, indicateurs, rapports et restitutions. Les données publiques doivent être anonymisées, agrégées ou limitées si nécessaire.","Un accès limité à TVF OS ou aux outils de suivi peut être ouvert selon les droits accordés, le périmètre de mission et le principe du besoin d’en connaître." ]),
('8. Propriétaires, partenaires et communication',["Les relations avec les propriétaires se font dans le respect du droit de propriété et du cadre convenu. Aucune action sur un bien privé ne peut être présentée comme obligatoire sans fondement légal ou accord de la personne concernée.","Toute communication publique, usage de logo, chiffre, résultat ou annonce de partenariat doit être validé au préalable par les parties." ]),
('9. Restitution, durée et conditions financières',["Les parties peuvent prévoir réunions, comptes rendus, bilans intermédiaires, rapport final, cartes anonymisées, indicateurs et calendrier de restitution.","La convention précise la durée, les modalités de renouvellement, les conditions financières éventuelles, l’évaluation et les cas de résiliation ou suspension." ])], ['Périmètre géographique','Référents et contacts','Règles de partage des données','Modèle de fiche signalement','Indicateurs et calendrier de restitution'])
]

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def borders(cell):
    tcPr = cell._tc.get_or_add_tcPr(); tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None: tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right'):
        el = tcBorders.find(qn('w:'+edge))
        if el is None: el = OxmlElement('w:'+edge); tcBorders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'8'); el.set(qn('w:color'),BORDER)
def run(p, text, size=10, bold=False, color=DARK, font='Inter'):
    r=p.add_run(text); r.font.name=font; r.font.size=Pt(size); r.bold=bold; r.font.color.rgb=color; return r
def cell_text(cell, text, bold=False, fill=None):
    cell.text='';
    if fill: shade(cell, fill)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); run(p,text,8.5,bold,GREEN if bold else DARK); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; borders(cell)
def note(doc, text):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; cell_text(t.rows[0].cells[0], text, False, NOTE)
    for r in t.rows[0].cells[0].paragraphs[0].runs: r.font.size=Pt(8); r.font.color.rgb=GRAY
def fields(doc, rows):
    t=doc.add_table(rows=len(rows), cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(2.25); t.columns[1].width=Inches(4.5)
    for i,(a,b) in enumerate(rows): cell_text(t.rows[i].cells[0],a,True,PALE); cell_text(t.rows[i].cells[1],b)
def make_docx(slug,title,subtitle,articles,annexes):
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.65); sec.right_margin=Inches(.65)
    h=doc.add_table(rows=1, cols=2); h.alignment=WD_TABLE_ALIGNMENT.CENTER; h.columns[0].width=Inches(1.25); h.columns[1].width=Inches(5.5)
    if LOGO.exists(): h.rows[0].cells[0].paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.0))
    p=h.rows[0].cells[1].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run(p,TVF['name']+'\n',12,True,GREEN,'Manrope'); run(p,TVF['subtitle']+'\n',8.5,False,DARK); run(p,f"{TVF['address']} | {TVF['phone']} | {TVF['email']}\n{TVF['rna']} | {TVF['siret']}",7.2,False,GRAY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(10); run(p,title,18,True,GREEN,'Manrope')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(p,subtitle,10,False,GRAY)
    note(doc,"MODÈLE INTERNE À ADAPTER — Document de travail institutionnel. Toute signature réelle doit être précédée d’une vérification juridique adaptée à la situation, aux parties, au bien, aux responsabilités et au cadre réglementaire applicable.")
    fields(doc,[('Référence interne TVF','[à compléter]'),('Parties signataires','[à compléter]'),('Date d’effet','[à compléter]'),('Durée','[à compléter]')])
    for head, paras in articles:
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(9); run(p,head,12,True,GREEN,'Manrope')
        for para in paras:
            q=doc.add_paragraph(); q.paragraph_format.line_spacing=1.12; q.paragraph_format.space_after=Pt(5); run(q,para,9.6,False,DARK)
    p=doc.add_paragraph(); run(p,'Annexes recommandées',12,True,GREEN,'Manrope')
    for a in annexes:
        q=doc.add_paragraph(); q.paragraph_format.left_indent=Inches(.18); run(q,'□ '+a,9.5,False,DARK)
    p=doc.add_paragraph(); run(p,'Signatures',12,True,GREEN,'Manrope')
    sig=doc.add_table(rows=3, cols=2); sig.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,row in enumerate([('Pour Territoires Vivants France','Pour la partie cocontractante'),('Nom, qualité, date','Nom, qualité, date'),('Signature','Signature')]):
        for c,t in enumerate(row): cell_text(sig.rows[i].cells[c],t,i==0,PALE if i==0 else None)
    f=doc.sections[0].footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(f,f"{TVF['name']} — {title} — modèle interne confidentiel",7,False,GRAY)
    path=OUT/(slug+'.docx'); doc.save(path); return path

def make_pdf(slug,title,subtitle,articles,annexes):
    st=getSampleStyleSheet(); st.add(ParagraphStyle(name='T',fontName='Helvetica-Bold',fontSize=17,leading=21,textColor=colors.HexColor('#18392F'),alignment=1)); st.add(ParagraphStyle(name='S',fontSize=9.5,leading=12,textColor=colors.HexColor('#667085'),alignment=1,spaceAfter=8)); st.add(ParagraphStyle(name='H',fontName='Helvetica-Bold',fontSize=11.5,leading=14,textColor=colors.HexColor('#18392F'),spaceBefore=9,spaceAfter=3)); st.add(ParagraphStyle(name='B',fontSize=9.1,leading=12,textColor=colors.HexColor('#1F2937'),spaceAfter=5)); st.add(ParagraphStyle(name='N',fontSize=8,leading=10.5,textColor=colors.HexColor('#475467'),backColor=colors.HexColor('#FBF8EF'),borderColor=colors.HexColor('#E4D7B8'),borderWidth=.5,borderPadding=6,spaceAfter=8)); st.add(ParagraphStyle(name='Small',fontSize=7.2,leading=9,textColor=colors.HexColor('#667085')))
    story=[]; logo=Image(str(LOGO), width=24*mm, height=18*mm) if LOGO.exists() else ''
    ht=Table([[logo,Paragraph(f"<b>{TVF['name']}</b><br/>{TVF['subtitle']}<br/>{TVF['address']} | {TVF['phone']} | {TVF['email']}<br/>{TVF['rna']} | {TVF['siret']}",st['Small'])]],colWidths=[30*mm,140*mm]); ht.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'MIDDLE'),('ALIGN',(1,0),(1,0),'RIGHT')])) ; story.append(ht); story.append(Spacer(1,7*mm)); story.append(Paragraph(title,st['T'])); story.append(Paragraph(subtitle,st['S'])); story.append(Paragraph('MODÈLE INTERNE À ADAPTER — Document de travail institutionnel. Toute signature réelle doit être précédée d’une vérification juridique adaptée.',st['N']))
    tbl=Table([('Référence interne TVF','[à compléter]'),('Parties signataires','[à compléter]'),('Date d’effet','[à compléter]'),('Durée','[à compléter]')],colWidths=[50*mm,120*mm]); tbl.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.45,colors.HexColor('#D9E0DC')),('BACKGROUND',(0,0),(0,-1),colors.HexColor('#F3F7F2')),('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,-1),7.8),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)])); story.append(tbl)
    for h,paras in articles:
        story.append(Paragraph(h,st['H']))
        for para in paras: story.append(Paragraph(para,st['B']))
    story.append(Paragraph('Annexes recommandées',st['H']))
    for a in annexes: story.append(Paragraph('□ '+a,st['B']))
    story.append(Paragraph('Signatures',st['H'])); sig=Table([['Pour Territoires Vivants France','Pour la partie cocontractante'],['Nom, qualité, date','Nom, qualité, date'],['Signature','Signature']],colWidths=[85*mm,85*mm],rowHeights=[9*mm,9*mm,22*mm]); sig.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.45,colors.HexColor('#D9E0DC')),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#F3F7F2')),('ALIGN',(0,0),(-1,-1),'CENTER'),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('FONTSIZE',(0,0),(-1,-1),8)])); story.append(sig)
    path=OUT/(slug+'.pdf'); SimpleDocTemplate(str(path),pagesize=A4,rightMargin=14*mm,leftMargin=14*mm,topMargin=11*mm,bottomMargin=12*mm).build(story); return path

for slug,title,sub,arts,annexes in CONVENTIONS:
    print(make_docx(slug,title,sub,arts,annexes)); print(make_pdf(slug,title,sub,arts,annexes))
