from pathlib import Path
import runpy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
TPL = runpy.run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

DOCX_PATH = ROOT / "documents" / "courriers-lancement-saint-etienne" / "24-prefecture-loire-lancement-cadrage-tvf.docx"
MD_PATH = ROOT / "documents" / "sources" / "courrier-prefecture-loire-lancement-tvf.md"
DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
MD_PATH.parent.mkdir(parents=True, exist_ok=True)

txt = TPL["txt"]
style_document = TPL["style_document"]
add_header_footer = TPL["add_header_footer"]
add_paragraph = TPL["add_paragraph"]
add_bullet = TPL["add_bullet"]
add_number = TPL["add_number"]
add_callout = TPL["add_callout"]
set_table_width = TPL["set_table_width"]
set_cell_border = TPL["set_cell_border"]
set_cell_shading = TPL["set_cell_shading"]
set_repeat_table_header = TPL["set_repeat_table_header"]
GREEN = TPL["GREEN"]
MUTED = TPL["MUTED"]
LIGHT_GREEN = TPL["LIGHT_GREEN"]
LIGHT_BLUE = TPL["LIGHT_BLUE"]
BORDER = TPL["BORDER"]


def add_metadata_block(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.8, 13.3])
    rows = [
        ("Destinataire", "Pr&eacute;fecture de la Loire - &agrave; l'attention de Madame la Pr&eacute;f&egrave;te de la Loire et des services comp&eacute;tents"),
        ("Adresse", "2 rue Charles de Gaulle, 42022 Saint-&Eacute;tienne Cedex 1"),
        ("Objet", "Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'orientation pour un lancement pilote encadr&eacute; dans la Loire"),
        ("Pi&egrave;ces jointes", "Dossier de pr&eacute;sentation TVF, r&eacute;c&eacute;piss&eacute; RNA, avis SIRENE, statuts, fiche besoins de lancement, fiche m&eacute;thode / instruction, fiche local de stockage / mat&eacute;riaux"),
    ]
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt(label))
        r.bold = True
        r.font.color.rgb = GREEN
        r.font.size = Pt(9.4)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(txt(value)).font.size = Pt(9.4)
    doc.add_paragraph()


def add_annex_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.0, 8.2, 3.9])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(["Document", "Utilit&eacute;", "Statut"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(txt(text))
        rr.bold = True
        rr.font.color.rgb = GREEN
        rr.font.size = Pt(9.2)
    rows = [
        ("Dossier de pr&eacute;sentation TVF", "Pr&eacute;senter l'objet, la m&eacute;thode, les p&ocirc;les et le positionnement de l'association.", "&Agrave; joindre"),
        ("R&eacute;c&eacute;piss&eacute; RNA", "Justifier la d&eacute;claration de l'association.", "&Agrave; joindre"),
        ("Avis SIRENE", "Justifier le SIREN, le SIRET et l'identification administrative.", "&Agrave; joindre"),
        ("Statuts", "Permettre la lecture officielle de l'objet associatif et du cadre de fonctionnement.", "&Agrave; joindre si utile"),
        ("Fiche m&eacute;thode / instruction", "Expliquer comment TVF re&ccedil;oit, qualifie, trace et oriente les demandes.", "&Agrave; joindre"),
        ("Fiche local de stockage / mat&eacute;riaux", "D&eacute;crire le besoin minimal : accessibilit&eacute;, s&eacute;curit&eacute;, stockage, tri, dur&eacute;e.", "&Agrave; joindre"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(txt(value)).font.size = Pt(8.8)


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(txt("Saint-&Eacute;tienne, le 14 juillet 2026"))
    r.font.size = Pt(10.2)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(txt("Courrier de lancement - Pr&eacute;fecture de la Loire"))
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run(txt("Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'orientation des services de l'&Eacute;tat"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    add_metadata_block(doc)

    add_paragraph(doc, "Madame la Pr&eacute;f&egrave;te,")
    add_paragraph(doc, "J'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e, dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un lancement pilote progressif dans la Loire.")
    add_paragraph(doc, "Notre objet est de contribuer &agrave; remettre en usage des ressources aujourd'hui insuffisamment mobilis&eacute;es : logements vacants ou d&eacute;grad&eacute;s, commerces inoccup&eacute;s, b&acirc;timents d&eacute;laiss&eacute;s, friches, terrains inutilis&eacute;s, mobiliers, &eacute;quipements et mat&eacute;riaux r&eacute;employables. La d&eacute;marche repose sur une logique de qualification, de mise en relation, de conventionnement, de suivi et de mesure de l'utilit&eacute; territoriale.")

    add_callout(doc, "Demande principale", "TVF sollicite un premier &eacute;change avec la Pr&eacute;fecture de la Loire afin de pr&eacute;senter sa m&eacute;thode, d'identifier les services de l'&Eacute;tat concern&eacute;s et de s'assurer que le lancement pilote s'inscrive dans un cadre institutionnel clair, prudent et compl&eacute;mentaire des politiques publiques existantes.")

    doc.add_heading(txt("Pourquoi s'adresser &agrave; la Pr&eacute;fecture de la Loire"), level=1)
    add_paragraph(doc, "Les services de l'&Eacute;tat occupent une place essentielle dans les sujets que TVF souhaite traiter : am&eacute;nagement du territoire, urbanisme, logement, environnement, politique de la ville, coh&eacute;sion sociale, emploi, insertion, vie associative, accompagnement des collectivit&eacute;s et coordination des dispositifs publics. TVF souhaite donc se pr&eacute;senter en amont, avec transparence, afin d'&ecirc;tre correctement orient&eacute;e.")
    add_paragraph(doc, "L'association ne cherche pas &agrave; se substituer aux dispositifs publics, aux collectivit&eacute;s ou aux services techniques comp&eacute;tents. Elle souhaite agir comme une structure de terrain capable de recevoir des situations, les qualifier, les documenter, mobiliser les bons interlocuteurs et contribuer &agrave; transformer des ressources inutilis&eacute;es en projets utiles, dans le respect des cadres juridiques, administratifs et op&eacute;rationnels applicables.")

    doc.add_heading(txt("Axes de dialogue propos&eacute;s"), level=1)
    add_number(doc, "Orientation institutionnelle : identifier les services de l'&Eacute;tat &agrave; rencontrer selon les sujets trait&eacute;s, notamment logement, urbanisme, coh&eacute;sion sociale, insertion, environnement et associations.")
    add_number(doc, "Habitat et patrimoine vacant : comprendre les limites d'intervention d'une association, les points de vigilance juridiques et les interlocuteurs comp&eacute;tents avant toute action de terrain.")
    add_number(doc, "Mat&eacute;riaux et r&eacute;emploi : cadrer les conditions de collecte, stockage, tri, tra&ccedil;abilit&eacute;, s&eacute;curit&eacute; et affectation des ressources r&eacute;employables.")
    add_number(doc, "Insertion et chantiers solidaires : identifier les structures, dispositifs et pr&eacute;cautions n&eacute;cessaires pour envisager des actions encadr&eacute;es et utiles.")
    add_number(doc, "Territoire pilote : organiser une premi&egrave;re phase sobre, document&eacute;e et tra&ccedil;able, sans annonce de partenariat public avant validation formelle.")

    doc.add_heading(txt("Demandes op&eacute;rationnelles &agrave; examiner"), level=1)
    add_bullet(doc, "Un premier rendez-vous d'orientation avec les services comp&eacute;tents de la Pr&eacute;fecture ou des directions d&eacute;partementales concern&eacute;es.")
    add_bullet(doc, "Une orientation vers la DDETS pour les sujets emploi, insertion, solidarit&eacute;, vie associative et structures d'insertion.")
    add_bullet(doc, "Une orientation vers les services comp&eacute;tents en mati&egrave;re d'am&eacute;nagement, urbanisme, logement, environnement, risques, foncier ou politiques territoriales.")
    add_bullet(doc, "Un avis de cadrage sur les points de vigilance &agrave; respecter avant de solliciter des collectivit&eacute;s, entreprises, propri&eacute;taires ou structures d'insertion.")
    add_bullet(doc, "Une indication sur les dispositifs, appels &agrave; projets, cadres de coop&eacute;ration ou interlocuteurs publics auxquels TVF pourrait utilement se rattacher sans confusion institutionnelle.")

    add_callout(doc, "Principe de prudence", "Aucune mention de partenariat, de soutien, de financement ou de validation par la Pr&eacute;fecture de la Loire ne sera utilis&eacute;e publiquement par TVF sans accord &eacute;crit pr&eacute;alable. Le premier &eacute;change demand&eacute; vise uniquement &agrave; pr&eacute;senter l'association, cadrer la d&eacute;marche et identifier les suites possibles.", fill=LIGHT_BLUE)

    doc.add_heading(txt("Premi&egrave;re &eacute;tape propos&eacute;e"), level=1)
    add_paragraph(doc, "Nous proposons un rendez-vous d'environ une heure afin de pr&eacute;senter TVF, de recueillir les orientations utiles, d'identifier les services &agrave; solliciter et de s&eacute;curiser la phase de lancement. Cette &eacute;tape permettrait d'&eacute;viter toute confusion avec les dispositifs existants et de construire une m&eacute;thode de travail sobre, progressive et compatible avec les priorit&eacute;s publiques.")
    add_paragraph(doc, "&Agrave; l'issue de cet &eacute;change, TVF pourra transmettre une note de cadrage plus pr&eacute;cise comprenant les interlocuteurs &agrave; rencontrer, les sujets &agrave; instruire, les limites d'intervention de l'association et les conditions d'une premi&egrave;re action pilote &agrave; Saint-&Eacute;tienne ou dans la Loire.")

    doc.add_heading("Conclusion", level=1)
    add_paragraph(doc, "Territoires Vivants France souhaite avancer de mani&egrave;re s&eacute;rieuse, progressive et tra&ccedil;able. Le lancement dans la Loire n'a pas vocation &agrave; produire une communication pr&eacute;matur&eacute;e, mais &agrave; construire une premi&egrave;re m&eacute;thode de terrain conforme aux exigences institutionnelles, utile aux habitants et compl&eacute;mentaire des acteurs publics et locaux.")
    add_paragraph(doc, "Je me tiens &agrave; votre disposition pour convenir d'un rendez-vous &agrave; la date qui vous conviendra et vous remercie par avance de l'attention port&eacute;e &agrave; cette demande.")
    add_paragraph(doc, "Je vous prie d'agr&eacute;er, Madame la Pr&eacute;f&egrave;te, l'expression de ma haute consid&eacute;ration.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(txt("Edryan Rangoly\nPr&eacute;sident fondateur\nTERRITOIRES VIVANTS FRANCE"))
    r.bold = True
    r.font.color.rgb = GREEN

    doc.add_page_break()
    doc.add_heading(txt("Annexe - Pi&egrave;ces et informations &agrave; joindre"), level=1)
    add_paragraph(doc, "Cette annexe peut &ecirc;tre conserv&eacute;e avec le courrier ou utilis&eacute;e comme check-list avant l'envoi.")
    add_annex_table(doc)
    doc.add_paragraph()
    add_callout(doc, "Objet d'e-mail recommand&eacute;", "Demande d'orientation - lancement pilote TERRITOIRES VIVANTS FRANCE dans la Loire", fill=LIGHT_BLUE)
    doc.save(DOCX_PATH)


def build_md():
    md = txt("""# Courrier - Pr&eacute;fecture de la Loire - lancement TVF\n\nDestinataire : Pr&eacute;fecture de la Loire, &agrave; l'attention de Madame la Pr&eacute;f&egrave;te de la Loire et des services comp&eacute;tents.\n\nAdresse : 2 rue Charles de Gaulle, 42022 Saint-&Eacute;tienne Cedex 1.\n\nObjet : Pr&eacute;sentation de TERRITOIRES VIVANTS FRANCE et demande d'orientation pour un lancement pilote encadr&eacute; dans la Loire.\n\nDate : Saint-&Eacute;tienne, le 14 juillet 2026.\n\n## Corps du courrier\n\nMadame la Pr&eacute;f&egrave;te,\n\nJ'ai l'honneur de vous pr&eacute;senter TERRITOIRES VIVANTS FRANCE, association loi 1901 d&eacute;clar&eacute;e, dont le si&egrave;ge est situ&eacute; &agrave; Saint-&Eacute;tienne. TVF a vocation &agrave; devenir une plateforme nationale de coop&eacute;ration au service de la revitalisation territoriale, en commen&ccedil;ant par un lancement pilote progressif dans la Loire.\n\nTVF sollicite un premier &eacute;change avec la Pr&eacute;fecture de la Loire afin de pr&eacute;senter sa m&eacute;thode, d'identifier les services de l'&Eacute;tat concern&eacute;s et de s'assurer que le lancement pilote s'inscrive dans un cadre institutionnel clair, prudent et compl&eacute;mentaire des politiques publiques existantes.\n\n## Demandes op&eacute;rationnelles\n\n- Premier rendez-vous d'orientation avec les services comp&eacute;tents.\n- Orientation vers la DDETS pour les sujets emploi, insertion, solidarit&eacute;, vie associative et structures d'insertion.\n- Orientation vers les services comp&eacute;tents en mati&egrave;re d'am&eacute;nagement, urbanisme, logement, environnement, risques, foncier ou politiques territoriales.\n- Avis de cadrage sur les points de vigilance &agrave; respecter avant de solliciter collectivit&eacute;s, entreprises, propri&eacute;taires ou structures d'insertion.\n- Indication sur les dispositifs, appels &agrave; projets, cadres de coop&eacute;ration ou interlocuteurs publics auxquels TVF pourrait utilement se rattacher.\n\n## Signature\n\nEdryan Rangoly  \nPr&eacute;sident fondateur  \nTERRITOIRES VIVANTS FRANCE\n""")
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)