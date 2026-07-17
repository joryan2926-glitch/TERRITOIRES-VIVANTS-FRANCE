from pathlib import Path
from runpy import run_path
from html import unescape

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

Document = BASE["Document"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
BLUE = BASE["BLUE"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
GOLD = BASE["GOLD"]
BORDER = BASE["BORDER"]

OUT_DIR = ROOT / "documents" / "courriers-fondations-mecenat"
SRC_DIR = ROOT / "documents" / "sources" / "courriers-fondations-mecenat"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"

TVF = {
    "name": "TERRITOIRES VIVANTS FRANCE",
    "status": "Association loi 1901 déclarée",
    "rna": "W423016361",
    "siren": "107 226 128",
    "siret": "107 226 128 00018",
    "address": "25 rue Élise Gervais, 42000 Saint-Étienne",
    "email": "contact@territoiresvivantsfrance.fr",
    "phone": "04 65 81 54 69",
    "site": "www.territoiresvivantsfrance.fr",
}


def u(text: str) -> str:
    return unescape(text)


def set_run(run, size=None, color=None, bold=None, italic=None):
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph(doc, text, size=10.2, color=None, bold=False, italic=False, after=7, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    r = p.add_run(u(text))
    set_run(r, size=size, color=color or RGBColor(31, 45, 38), bold=bold, italic=italic)
    return p


def bullet(doc, text):
    add_bullet(doc, text)


def kv_table(doc, rows, widths=(4.2, 12.0), fill=LIGHT_GREEN):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, list(widths))
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, title in enumerate(["Rubrique", "Contenu"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=BORDER, size="8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(title)
        set_run(r, 8.8, GREEN, bold=True)
    for label, value in rows:
        cells = t.add_row().cells
        for i, text in enumerate((label, value)):
            set_cell_border(cells[i], color="DDE6DE", size="6")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(u(text))
            set_run(r, 8.15, BLUE if i == 0 else RGBColor(31, 45, 38), bold=(i == 0))
    doc.add_paragraph()
    return t


def top_block(doc, letter):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Cm(5.1))
        p_logo.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(TVF["name"])
    set_run(r, 11.2, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{TVF['status']} - RNA {TVF['rna']} - SIREN {TVF['siren']} - SIRET {TVF['siret']}\n{TVF['address']} - {TVF['email']} - {TVF['phone']} - {TVF['site']}")
    set_run(r, 8.2, MUTED)

    paragraph(doc, "Saint-Étienne, le [date à compléter]", size=9.5, color=MUTED, after=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    kv_table(
        doc,
        [
            ("Destinataire", f"{letter['recipient']}\nÀ l'attention de [service mécénat / direction des programmes]\n[Adresse ou dépôt en ligne à compléter]"),
            ("Objet", letter["subject"]),
            ("Pièces jointes", "Dossier de présentation TVF, brochure Territoire pilote Saint-Étienne, brochure Banque de Matériaux TVF, statuts, récépissé RNA, avis SIRENE, budget prévisionnel à compléter."),
        ],
        fill=LIGHT_BLUE,
    )


def common_intro(doc):
    paragraph(doc, "Madame, Monsieur,", after=8)
    paragraph(
        doc,
        "Je me permets de vous contacter afin de vous présenter TERRITOIRES VIVANTS FRANCE, association loi 1901 déclarée, implantée à Saint-Étienne. TVF a pour objet de contribuer à la revitalisation des territoires en travaillant sur les logements vacants, les commerces inoccupés, les bâtiments inutilisés, les friches, les terrains délaissés, les matériaux réemployables et les actions d'insertion.",
    )
    paragraph(
        doc,
        "L'association est en phase de lancement opérationnel sur le territoire stéphanois. Ce premier ancrage doit permettre de structurer une méthode reproductible : repérer les besoins, qualifier les dossiers, mobiliser les acteurs locaux, formaliser les conventions, sécuriser les ressources et suivre les impacts.",
    )
    add_callout(
        doc,
        "Positionnement",
        "TVF ne se présente pas comme un dispositif isolé. Sa vocation est de devenir un outil de coopération territoriale capable de relier collectivités, propriétaires, entreprises, associations, artisans, bénévoles, mécènes et financeurs autour de projets concrets.",
        fill=LIGHT_GREEN,
    )


def common_project(doc):
    doc.add_heading("Projet présenté au soutien", level=2)
    paragraph(
        doc,
        "La demande porte sur l'amorçage du pilote TVF à Saint-Étienne : création d'un cadre d'instruction des demandes, premières actions de repérage, structuration d'une Banque de Matériaux, mobilisation de propriétaires et d'entreprises, préparation de conventions de coopération et lancement de premiers dossiers tests.",
    )
    for item in [
        "repérage de logements, locaux, commerces, friches et ressources inutilisées ;",
        "qualification des situations avec une grille de faisabilité et de sécurité ;",
        "mise en relation des propriétaires, collectivités, entreprises, associations et acteurs techniques ;",
        "organisation d'une Banque de Matériaux : inventaire, traçabilité, stockage, affectation à des projets validés ;",
        "préparation de projets de rénovation légère, d'aménagement utile, d'usage temporaire ou de chantier encadré ;",
        "suivi d'indicateurs : dossiers ouverts, ressources réemployées, lieux qualifiés, acteurs mobilisés, freins identifiés.",
    ]:
        bullet(doc, item)


def common_ask(doc, letter):
    doc.add_heading("Demande formulée", level=2)
    paragraph(doc, letter["ask"])
    kv_table(
        doc,
        [
            ("Soutien recherché", letter["support"]),
            ("Usage envisagé", letter["usage"]),
            ("Montant", "À préciser selon l'appel à projets, les critères de la fondation et le budget validé du dossier."),
            ("Durée", "Amorçage sur 12 mois, avec possibilité de bilan intermédiaire et de prolongement si le dispositif est pertinent."),
            ("Livrables", "Note de cadrage, fiches dossiers, registre matériaux, conventions types, suivi d'indicateurs, bilan de phase pilote."),
        ],
        fill=LIGHT_GREEN,
    )


def guarantees(doc):
    doc.add_heading("Garanties et méthode", level=2)
    for item in [
        "aucun projet n'est présenté comme réalisé tant qu'il n'a pas été effectivement mis en œuvre ;",
        "aucun partenaire, financeur ou soutien n'est cité publiquement sans accord écrit ;",
        "chaque dossier est instruit avant acceptation : besoin, propriété, sécurité, assurance, faisabilité, utilité territoriale ;",
        "les matériaux ne sont pas distribués librement : ils sont affectés à des projets validés et traçables ;",
        "les travaux lourds ou dangereux relèvent de professionnels qualifiés et d'un cadre adapté ;",
        "un bilan peut être transmis au financeur selon les indicateurs retenus ensemble.",
    ]:
        bullet(doc, item)


def signature(doc):
    paragraph(doc, "Nous souhaiterions pouvoir vous présenter plus précisément cette démarche et vérifier si elle peut entrer dans vos priorités d'action ou dans un dispositif de mécénat adapté.", after=8)
    paragraph(doc, "Dans l'attente de votre retour, je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées.", after=14)
    paragraph(doc, "Pour TERRITOIRES VIVANTS FRANCE", size=10.3, color=BLUE, bold=True, after=3)
    paragraph(doc, "Edryan Rangoly\nPrésident fondateur", size=9.8, after=14)
    paragraph(doc, "Signature :", size=9.2, color=MUTED, after=2)
    paragraph(doc, "\n\n", after=2)
    add_callout(
        doc,
        "Contact TVF",
        f"{TVF['name']} - {TVF['address']} - {TVF['email']} - {TVF['phone']} - {TVF['site']}",
        fill=LIGHT_GREEN,
    )


LETTERS = [
    {
        "slug": "01-courrier-fondation-de-france",
        "recipient": "Fondation de France",
        "subject": "Demande d'étude d'un soutien - pilote territorial TVF à Saint-Étienne",
        "angle_title": "Pourquoi solliciter la Fondation de France ?",
        "angle": "La Fondation de France met en avant l'intérêt général, l'ancrage territorial et la possibilité pour les porteurs de projet de soumettre une initiative utile. Le pilote TVF s'inscrit dans cette logique : partir d'un besoin local documenté, mobiliser des acteurs de proximité et construire une réponse durable autour de biens et ressources inutilisés.",
        "ask": "Nous sollicitons un échange afin d'étudier l'éligibilité du pilote TVF à un soutien de la Fondation de France, notamment pour la phase de structuration, de qualification des dossiers et de mobilisation territoriale.",
        "support": "Soutien financier d'amorçage, orientation vers un programme adapté, conseil de dépôt de projet ou mise en relation avec une fondation abritée pertinente.",
        "usage": "Structuration du pilote, outils de suivi, premières qualifications terrain, documentation des dossiers, communication institutionnelle et coordination locale.",
    },
    {
        "slug": "02-courrier-fondation-vinci-pour-la-cite",
        "recipient": "Fondation d'entreprise VINCI pour la Cité",
        "subject": "Demande d'étude d'un soutien - insertion, logement et chantiers encadrés TVF",
        "angle_title": "Pourquoi solliciter la Fondation VINCI pour la Cité ?",
        "angle": "La Fondation VINCI pour la Cité soutient des initiatives d'insertion sociale et professionnelle, d'égalité des chances, d'insertion par le logement et de lien social en territoires fragilisés. TVF peut rejoindre ces enjeux par des chantiers encadrés, des actions de remobilisation, la remise en usage de lieux utiles et la valorisation de compétences liées aux métiers du bâtiment.",
        "ask": "Nous sollicitons un rendez-vous afin d'étudier un soutien au pilote TVF sur les dimensions insertion, logement, remobilisation et mécénat de compétences, notamment autour de premiers chantiers encadrés à Saint-Étienne.",
        "support": "Soutien financier ciblé, mécénat de compétences, parrainage associatif, expertise chantier, mise en relation avec collaborateurs ou entités locales du groupe.",
        "usage": "Préparation de chantiers légers, sécurisation des procédures, mobilisation d'acteurs de l'insertion, qualification de biens et coordination des premiers dossiers.",
    },
    {
        "slug": "03-courrier-fondation-bouygues-terre-plurielle",
        "recipient": "Fondation Terre Plurielle / Groupe Bouygues",
        "subject": "Demande d'étude d'un soutien - revitalisation territoriale, emploi et ressources locales",
        "angle_title": "Pourquoi solliciter Bouygues / Terre Plurielle ?",
        "angle": "La Fondation Terre Plurielle de Bouygues Construction agit en faveur de l'éducation, de l'emploi, de la lutte contre la précarité, de la santé et du handicap, avec une implication possible de collaborateurs. TVF rejoint particulièrement l'axe emploi, précarité et territoire par la création de projets locaux autour de la rénovation légère, du réemploi et de l'insertion.",
        "ask": "Nous sollicitons un échange afin d'étudier une coopération ou un soutien d'amorçage autour du pilote TVF à Saint-Étienne, en particulier sur la mobilisation de compétences bâtiment, la structuration de chantiers encadrés et la valorisation de ressources inutilisées.",
        "support": "Soutien financier, mécénat de compétences, conseil technique, parrainage de projet, mise en relation avec acteurs construction ou collaborateurs engagés.",
        "usage": "Diagnostic de faisabilité, préparation de fiches projets, outillage méthodologique, qualification technique de ressources, premiers ateliers de remise en usage.",
    },
    {
        "slug": "04-courrier-fondation-groupe-edf",
        "recipient": "Fondation groupe EDF",
        "subject": "Demande d'étude d'un soutien - transition sociale, écologique et insertion territoriale",
        "angle_title": "Pourquoi solliciter la Fondation groupe EDF ?",
        "angle": "La Fondation groupe EDF porte une mission d'intérêt général associant transition écologique et sociale, insertion par l'éducation et la formation, et accompagnement des associations pour intégrer les enjeux écologiques à leurs missions sociales. TVF propose précisément de relier transition écologique, réemploi, utilité sociale, formation de terrain et revitalisation locale.",
        "ask": "Nous sollicitons un échange afin d'étudier l'éligibilité du pilote TVF à un soutien de la Fondation groupe EDF, notamment sur la dimension transition écologique, inclusion, formation de terrain et réemploi territorial.",
        "support": "Soutien financier, accompagnement méthodologique, mécénat de compétences, orientation vers un appel à projets ou un dispositif territorial adapté.",
        "usage": "Structuration d'indicateurs environnementaux, Banque de Matériaux, qualification de ressources, ateliers de sensibilisation, projets de remise en usage à utilité sociale.",
    },
    {
        "slug": "05-courrier-fondation-credit-agricole-solidarite-developpement",
        "recipient": "Fondation Crédit Agricole Solidarité et Développement",
        "subject": "Demande d'étude d'un soutien - logement, insertion et autonomie territoriale",
        "angle_title": "Pourquoi solliciter la Fondation Crédit Agricole Solidarité et Développement ?",
        "angle": "La Fondation Crédit Agricole Solidarité et Développement agit autour de l'insertion sociale, de l'aide au logement, de l'insertion économique et professionnelle, et de l'autonomie socioéconomique des personnes en difficulté. TVF rejoint ces axes par la remise en usage de biens dormants, la mobilisation de ressources locales et la construction de parcours de coopération utiles aux habitants.",
        "ask": "Nous sollicitons un rendez-vous afin d'étudier un soutien au pilote TVF à Saint-Étienne, notamment sur les volets logement, insertion économique et professionnelle, et structuration d'une méthode locale reproductible.",
        "support": "Soutien financier d'amorçage, orientation vers appel à projets, relais territorial, mise en relation avec acteurs locaux ou dispositifs de mécénat du groupe.",
        "usage": "Qualification de biens, accompagnement de propriétaires, préparation de conventions, mobilisation de ressources, premiers dossiers logement/local/insertion.",
    },
]


def build_letter(letter):
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    top_block(doc, letter)
    common_intro(doc)
    doc.add_heading(letter["angle_title"], level=1)
    paragraph(doc, letter["angle"])
    common_project(doc)
    common_ask(doc, letter)
    guarantees(doc)
    doc.add_heading("Prochaine étape proposée", level=2)
    paragraph(
        doc,
        "Nous proposons un premier rendez-vous de trente à quarante-cinq minutes afin de présenter le pilote, préciser les besoins de financement, vérifier l'adéquation avec vos critères et identifier le bon canal de dépôt ou d'instruction.",
    )
    signature(doc)
    return doc


def source_markdown(letter):
    lines = [
        f"# {letter['recipient']}",
        "",
        f"Objet : {letter['subject']}",
        "",
        "Madame, Monsieur,",
        "",
        "Je me permets de vous contacter afin de vous présenter TERRITOIRES VIVANTS FRANCE, association loi 1901 déclarée, implantée à Saint-Étienne.",
        "",
        f"## {letter['angle_title']}",
        "",
        letter["angle"],
        "",
        "## Demande",
        "",
        letter["ask"],
        "",
        f"Soutien recherché : {letter['support']}",
        "",
        f"Usage envisagé : {letter['usage']}",
        "",
        "## Contact",
        "",
        f"{TVF['name']} - {TVF['email']} - {TVF['phone']} - {TVF['site']}",
    ]
    return "\n".join(lines)


def build_all():
    paths = []
    for letter in LETTERS:
        doc = build_letter(letter)
        docx_path = OUT_DIR / f"{letter['slug']}.docx"
        md_path = SRC_DIR / f"{letter['slug']}.md"
        doc.save(docx_path)
        md_path.write_text(source_markdown(letter), encoding="utf-8")
        paths.append(docx_path)
    return paths


if __name__ == "__main__":
    for path in build_all():
        print(path)
