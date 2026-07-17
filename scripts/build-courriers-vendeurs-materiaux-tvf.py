from pathlib import Path
from runpy import run_path
from html import unescape
from datetime import date

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

Document = BASE["Document"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
BLUE = BASE["BLUE"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
GOLD = BASE["GOLD"]
BORDER = BASE["BORDER"]

OUT_DIR = ROOT / "documents" / "courriers-materiaux-vendeurs"
SRC_DIR = ROOT / "documents" / "sources" / "courriers-materiaux-vendeurs"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"

TVF = {
    "name": "TERRITOIRES VIVANTS FRANCE",
    "status": "Association loi 1901 déclarée",
    "rna": "W423016361",
    "siren": "107 226 128",
    "siret": "107 226 128 00018",
    "address": "25 rue Élise Gervais, 42000 Saint-Étienne",
    "email": "contact@territoiresvivantsfrance.fr",
    "phone": "04 65 81 54 69",
    "site": "www.territoiresvivantsfrance.fr",
}


def u(text: str) -> str:
    return unescape(text)


def set_run(run, size=None, color=None, bold=None, italic=None):
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph(doc, text, size=10.2, color=None, bold=False, italic=False, after=7, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    r = p.add_run(u(text))
    set_run(r, size=size, color=color or RGBColor(31, 45, 38), bold=bold, italic=italic)
    return p


def bullet(doc, text):
    add_bullet(doc, text)


def kv_table(doc, rows, widths=(4.2, 12.0), fill=LIGHT_GREEN):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, list(widths))
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, title in enumerate(["Rubrique", "Contenu"]):
        cell = hdr.cells[i]
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=BORDER, size="8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(title)
        set_run(r, 8.8, GREEN, bold=True)
    for label, value in rows:
        cells = t.add_row().cells
        for i, text in enumerate((label, value)):
            set_cell_border(cells[i], color="DDE6DE", size="6")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(u(text))
            set_run(r, 8.3, BLUE if i == 0 else RGBColor(31, 45, 38), bold=(i == 0))
    doc.add_paragraph()
    return t


def top_block(doc, title):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Cm(5.1))
        p_logo.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(TVF["name"])
    set_run(r, 11.2, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{TVF['status']} - RNA {TVF['rna']} - SIREN {TVF['siren']} - SIRET {TVF['siret']}\n{TVF['address']} - {TVF['email']} - {TVF['phone']} - {TVF['site']}")
    set_run(r, 8.2, MUTED)

    paragraph(doc, "Saint-Étienne, le [date à compléter]", size=9.5, color=MUTED, after=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    kv_table(
        doc,
        [
            ("Destinataire", "[Nom de l'entreprise]\nÀ l'attention de [responsable d'agence / direction / responsable magasin]\n[Adresse]\n[Code postal] [Ville]"),
            ("Objet", title),
            ("Pièces jointes", "Dossier de présentation TVF, brochure Banque de Matériaux TVF, fiche contribution entreprise, bordereau de proposition de matériaux, fiche contact."),
        ],
        fill=LIGHT_BLUE,
    )


def signature(doc):
    paragraph(doc, "Dans l'attente de votre retour, je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées.", after=14)
    paragraph(doc, "Pour TERRITOIRES VIVANTS FRANCE", size=10.3, color=BLUE, bold=True, after=3)
    paragraph(doc, "Edryan Rangoly\nPrésident fondateur", size=9.8, after=14)
    paragraph(doc, "Signature :", size=9.2, color=MUTED, after=2)
    paragraph(doc, "\n\n", after=2)
    add_callout(
        doc,
        "Contact TVF",
        f"{TVF['name']} - {TVF['address']} - {TVF['email']} - {TVF['phone']} - {TVF['site']}",
        fill=LIGHT_GREEN,
    )


def common_intro(doc):
    paragraph(doc, "Madame, Monsieur,", after=8)
    paragraph(
        doc,
        "Je me permets de vous contacter afin de vous présenter TERRITOIRES VIVANTS FRANCE, association loi 1901 déclarée, implantée à Saint-Étienne. TVF a pour objet de contribuer à la revitalisation des territoires en travaillant sur les logements vacants, les commerces inoccupés, les bâtiments inutilisés, les friches, les terrains délaissés, les matériaux réemployables et les actions d'insertion.",
    )
    paragraph(
        doc,
        "Notre démarche vise à relier les propriétaires, collectivités, entreprises, artisans, associations, bénévoles et financeurs autour de projets concrets. L'objectif est simple : transformer des ressources aujourd'hui inutilisées en solutions utiles aux habitants, aux quartiers et aux territoires.",
    )


def common_materials_frame(doc):
    paragraph(
        doc,
        "Dans ce cadre, TVF structure une Banque de Matériaux. Il ne s'agit ni d'une déchetterie, ni d'une plateforme de dons libres, ni d'un système de distribution automatique. Les ressources proposées sont identifiées, qualifiées, tracées, puis affectées à des projets validés par TVF selon leur état, leur utilité, leur disponibilité et les besoins du territoire.",
    )
    add_callout(
        doc,
        "Principe de coopération",
        "Ce qui représente parfois un coût, un stock dormant ou une ressource immobilisée pour une entreprise peut devenir une ressource utile pour un projet local : remise en usage d'un local, aménagement associatif, chantier encadré, action de réemploi, rénovation légère ou équipement d'un lieu utile.",
        fill=LIGHT_GREEN,
    )


def common_guarantees(doc):
    doc.add_heading("Cadre proposé", level=2)
    for item in [
        "un premier échange pour comprendre vos contraintes, vos volumes éventuels et les types de ressources concernées ;",
        "un inventaire simple des matériaux ou équipements proposés, avec photos, quantités, état et conditions de retrait ;",
        "une validation TVF avant toute affectation à un projet ;",
        "une convention ou un accord écrit lorsque la situation le nécessite ;",
        "une traçabilité de l'origine, de la destination et de l'usage des ressources ;",
        "une communication sobre, uniquement après accord écrit de votre structure.",
    ]:
        bullet(doc, item)

    doc.add_heading("Demande formulée", level=2)
    paragraph(
        doc,
        "Nous souhaiterions pouvoir échanger avec un responsable afin d'étudier si une coopération ponctuelle ou régulière peut être envisagée. L'objectif du premier rendez-vous n'est pas de vous demander un engagement immédiat, mais de comprendre vos pratiques, vos contraintes et les conditions dans lesquelles certains matériaux ou équipements pourraient être valorisés plutôt que détruits, stockés durablement ou écartés.",
    )


LETTERS = [
    {
        "slug": "01-courrier-negoces-materiaux-generalistes",
        "title": "Demande de rendez-vous - coopération TVF pour la valorisation de matériaux réutilisables",
        "sector_title": "Négoces et vendeurs de matériaux généralistes",
        "target": "POINT.P, Chausson Matériaux, Gedimat, BigMat, SAMSE, négoces professionnels.",
        "resources": [
            "bois, panneaux, plaques, isolants, menuiseries, quincaillerie ;",
            "carrelage, faïence, revêtements, sanitaires, équipements de second œuvre ;",
            "fins de série, retours non remis en rayon, matériaux déclassés mais utilisables ;",
            "échantillons, produits d'exposition, surplus ou lots incomplets exploitables.",
        ],
        "angle": "Votre activité vous place au cœur de la chaîne des matériaux. TVF souhaite construire avec les acteurs locaux une solution de valorisation territoriale des ressources encore utiles, en particulier lorsque certains produits ne trouvent plus leur place dans les circuits commerciaux classiques.",
    },
    {
        "slug": "02-courrier-magasins-bricolage-grandes-surfaces",
        "title": "Demande de rendez-vous - valorisation d'invendus, retours et matériaux de bricolage réutilisables",
        "sector_title": "Magasins de bricolage et grandes surfaces spécialisées",
        "target": "Leroy Merlin, Castorama, Brico Dépôt, Weldom, magasins de bricolage et aménagement.",
        "resources": [
            "produits d'exposition, fins de gamme, retours non remis en rayon ;",
            "outillage, peintures utilisables, luminaires, quincaillerie, petits équipements ;",
            "meubles, étagères, plans de travail, portes, revêtements et articles de rénovation ;",
            "lots incomplets ou déclassés pouvant être affectés à des projets validés.",
        ],
        "angle": "Les magasins de bricolage disposent parfois de ressources encore exploitables mais difficiles à réintégrer dans le circuit de vente. TVF propose d'étudier une solution encadrée, locale et traçable pour orienter ces ressources vers des projets d'intérêt territorial.",
    },
    {
        "slug": "03-courrier-sanitaire-plomberie-chauffage",
        "title": "Demande de rendez-vous - réemploi encadré de sanitaires, plomberie et équipements techniques",
        "sector_title": "Sanitaire, plomberie et chauffage",
        "target": "CEDEO, Richardson, PUM, grossistes plomberie, sanitaire et chauffage.",
        "resources": [
            "lavabos, WC, receveurs, douches, robinetterie, meubles de salle d'eau ;",
            "radiateurs, raccords, tubes, gaines, accessoires, équipements non dangereux ;",
            "fins de série, produits d'exposition, lots incomplets, retours qualifiables ;",
            "équipements uniquement réutilisables dans un cadre compatible avec la sécurité et la réglementation.",
        ],
        "angle": "Les équipements sanitaires et techniques peuvent être très utiles pour remettre en état des locaux associatifs, logements ou espaces partagés. TVF souhaite étudier avec vous les conditions d'un réemploi prudent, contrôlé et réservé aux produits adaptés.",
    },
    {
        "slug": "04-courrier-electricite-luminaires",
        "title": "Demande de rendez-vous - valorisation encadrée de matériel électrique et luminaires réutilisables",
        "sector_title": "Électricité, luminaires et équipements techniques",
        "target": "Rexel, Sonepar, YESSS Électrique, distributeurs de matériel électrique.",
        "resources": [
            "luminaires, appareillage, câbles, gaines, armoires, coffrets et accessoires ;",
            "produits d'exposition, fins de gamme, lots incomplets ou emballages abîmés ;",
            "matériel réutilisable uniquement après qualification de l'état et des conditions de sécurité ;",
            "éléments pouvant servir à des aménagements associatifs ou locaux réhabilités.",
        ],
        "angle": "Le matériel électrique exige un cadre sérieux. TVF ne recherche pas une récupération non contrôlée, mais une coopération permettant d'identifier uniquement les équipements réutilisables dans des conditions sûres, documentées et compatibles avec les règles applicables.",
    },
    {
        "slug": "05-courrier-bois-menuiserie-panneaux",
        "title": "Demande de rendez-vous - valorisation de bois, menuiseries, panneaux et éléments d'agencement",
        "sector_title": "Bois, menuiserie, panneaux et agencement",
        "target": "Dispano, Panofrance, menuisiers, agenceurs, fabricants, négoces bois et panneaux.",
        "resources": [
            "bois, panneaux, tasseaux, plans de travail, portes, fenêtres, quincaillerie ;",
            "chutes exploitables, produits d'exposition, erreurs de commande, lots incomplets ;",
            "éléments d'agencement pouvant servir à des locaux associatifs ou ateliers partagés ;",
            "menuiseries réutilisables après vérification de l'état et de l'usage envisagé.",
        ],
        "angle": "Le bois et les éléments de menuiserie peuvent avoir une seconde vie très concrète dans des projets locaux : aménagements, ateliers, mobilier, locaux associatifs ou petites rénovations. TVF souhaite construire une méthode de récupération propre, tracée et utile.",
    },
    {
        "slug": "06-courrier-carrelage-revetements-peinture",
        "title": "Demande de rendez-vous - valorisation de carrelage, revêtements et matériaux de finition",
        "sector_title": "Carrelage, revêtements, peinture et finitions",
        "target": "Fournisseurs carrelage, sols, peinture, revêtements muraux, carreleurs, peintres.",
        "resources": [
            "carrelage, faïence, parquet, sols souples, plinthes, revêtements muraux ;",
            "peintures et produits utilisables non périmés, accessoires, consommables ;",
            "fins de lots, erreurs de commande, surplus de chantier, produits d'exposition ;",
            "matériaux pouvant servir à des remises en état légères ou à des locaux utiles.",
        ],
        "angle": "Les matériaux de finition sont souvent décisifs pour rendre un lieu propre, digne et utilisable. TVF souhaite étudier comment certains lots encore exploitables pourraient être orientés vers des projets de proximité plutôt que rester immobilisés ou être écartés.",
    },
    {
        "slug": "07-courrier-entreprises-btp-artisans",
        "title": "Demande de rendez-vous - surplus de chantier, matériaux déposés et réemploi territorial",
        "sector_title": "Entreprises BTP, artisans et chantiers",
        "target": "Entreprises générales, démolisseurs, artisans du second œuvre, menuisiers, plombiers, électriciens, carreleurs, peintres.",
        "resources": [
            "surplus de chantier, matériaux déposés proprement, fins de travaux ;",
            "portes, fenêtres, bois, sanitaires, carrelage, luminaires, mobilier, outillage ;",
            "matériaux réutilisables avant évacuation, destruction ou stockage durable ;",
            "compétences, conseils techniques ou participation encadrée à des projets locaux.",
        ],
        "angle": "Les chantiers génèrent parfois des ressources encore utilisables. TVF souhaite proposer aux entreprises et artisans une solution simple pour signaler, qualifier et orienter ces ressources vers des projets validés, sans perturber l'organisation des chantiers ni créer de charge excessive.",
    },
]


def build_letter(letter):
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    top_block(doc, letter["title"])

    paragraph(doc, "Madame, Monsieur,", after=8)
    common_intro(doc)

    doc.add_heading(letter["sector_title"], level=1)
    paragraph(doc, letter["angle"])
    paragraph(doc, f"Ce courrier s'adresse notamment aux acteurs suivants : {letter['target']}", size=9.5, color=MUTED, italic=True)

    doc.add_heading("Ressources pouvant être étudiées", level=2)
    for item in letter["resources"]:
        bullet(doc, item)

    common_materials_frame(doc)
    common_guarantees(doc)

    doc.add_heading("Intérêt pour votre structure", level=2)
    for item in [
        "réduire certains coûts liés au stockage ou à la destruction de ressources encore utilisables ;",
        "valoriser une démarche RSE concrète et locale ;",
        "contribuer à des projets de revitalisation territoriale à Saint-Étienne et dans la Loire ;",
        "renforcer le lien avec les collectivités, associations, artisans et acteurs locaux ;",
        "disposer d'un cadre plus lisible qu'un don ponctuel non suivi.",
    ]:
        bullet(doc, item)

    doc.add_heading("Point important", level=2)
    paragraph(
        doc,
        "Toute proposition de matériaux ou d'équipements fera l'objet d'une vérification préalable. TVF se réserve la possibilité de refuser une ressource si son état, sa dangerosité, son volume, son coût logistique, son absence de traçabilité ou son inadéquation avec les projets identifiés ne permettent pas une valorisation responsable.",
    )

    signature(doc)
    return doc


def source_markdown(letter):
    lines = [
        f"# {letter['sector_title']}",
        "",
        f"Objet : {letter['title']}",
        "",
        "Madame, Monsieur,",
        "",
        "Je me permets de vous contacter afin de vous présenter TERRITOIRES VIVANTS FRANCE, association loi 1901 déclarée, implantée à Saint-Étienne.",
        "",
        letter["angle"],
        "",
        "## Ressources pouvant être étudiées",
        "",
    ]
    lines.extend([f"- {x}" for x in letter["resources"]])
    lines.extend(
        [
            "",
            "## Cadre",
            "",
            "TVF n'est ni une déchetterie, ni une plateforme de dons libres, ni un système de distribution automatique. Les ressources proposées sont identifiées, qualifiées, tracées et affectées à des projets validés.",
            "",
            "## Contact",
            "",
            f"{TVF['name']} - {TVF['email']} - {TVF['phone']} - {TVF['site']}",
        ]
    )
    return "\n".join(lines)


def build_all():
    paths = []
    for letter in LETTERS:
        doc = build_letter(letter)
        docx_path = OUT_DIR / f"{letter['slug']}.docx"
        md_path = SRC_DIR / f"{letter['slug']}.md"
        doc.save(docx_path)
        md_path.write_text(source_markdown(letter), encoding="utf-8")
        paths.append(docx_path)
    return paths


if __name__ == "__main__":
    for path in build_all():
        print(path)
