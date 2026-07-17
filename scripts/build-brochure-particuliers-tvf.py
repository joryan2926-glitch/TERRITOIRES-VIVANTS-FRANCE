from pathlib import Path
from runpy import run_path
from html import unescape

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

Document = BASE["Document"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
BLUE = BASE["BLUE"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
GOLD = BASE["GOLD"]
BORDER = BASE["BORDER"]

OUT_DIR = ROOT / "documents" / "brochures"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "brochure-particuliers-biens-dormants-tvf.docx"
MD_PATH = SRC_DIR / "brochure-particuliers-biens-dormants-tvf.md"
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"
PHOTO = ROOT / "assets" / "photos" / "immeuble-renovation-meudon.jpg"


def u(text: str) -> str:
    return unescape(text)


def set_run(run, size=None, color=None, bold=None, italic=None):
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(u(text))
    set_run(r, 8.5, MUTED, italic=True)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(u(text))
    set_run(r, 8.4, MUTED, italic=True)


def table(doc, headers, rows, widths, fill=LIGHT_GREEN):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, widths)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=BORDER, size="8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(u(header))
        set_run(r, 8.8, GREEN, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i], color="DDE6DE", size="6")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(u(value))
            set_run(r, 8.45, BLUE if i == 0 else RGBColor(31, 45, 38), bold=(i == 0))
    doc.add_paragraph()
    return t


def cover(doc):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Cm(5.8))
        p_logo.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(u("Brochure propri&eacute;taires et particuliers"))
    set_run(r, 10.2, GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(u("Biens dormants, logements vacants et usages solidaires"))
    set_run(r, 24.5, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run(
        u(
            "Comprendre comment un bien inutilis&eacute; peut devenir un projet utile, "
            "encadr&eacute; et valorisant pour le propri&eacute;taire comme pour le territoire."
        )
    )
    set_run(r, 12.8, GREEN, bold=True)

    meta = doc.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, [5.5, 5.5, 5.5])
    values = [
        ("Document", "information et pr&eacute;-orientation"),
        ("Public", "particuliers et propri&eacute;taires"),
        ("Version", "juillet 2026"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = meta.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GREEN if idx != 1 else LIGHT_BLUE)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(u(label + "\n"))
        set_run(r1, 8.2, MUTED, bold=True)
        r2 = p.add_run(u(value))
        set_run(r2, 9.2, BLUE, bold=True)

    if PHOTO.exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(PHOTO), width=Cm(15.6))
        caption(
            doc,
            "Exemple de b&acirc;timent ancien r&eacute;habilit&eacute; : un bien dormant peut redevenir "
            "une ressource territoriale lorsque le projet est techniquement, juridiquement "
            "et financi&egrave;rement cadr&eacute;.",
        )

    add_callout(
        doc,
        "Id&eacute;e centrale",
        "Le propri&eacute;taire conserve la propri&eacute;t&eacute; de son bien. TVF peut &eacute;tudier "
        "avec lui un cadre de coop&eacute;ration permettant de remettre le bien en usage, "
        "lorsque le dossier est compatible avec l'objet de l'association, les moyens "
        "mobilisables et l'utilit&eacute; territoriale recherch&eacute;e.",
        fill=LIGHT_GREEN,
    )


def toc(doc):
    doc.add_heading("Sommaire", level=1)
    items = [
        "1. Pourquoi un bien dormant devient un enjeu",
        "2. Ce que TVF peut faire pour un particulier",
        "3. Biens concern&eacute;s et usages possibles",
        "4. Comment TVF choisit les biens et les dossiers",
        "5. Diagnostic, &eacute;tude et orientation technique",
        "6. Conventions possibles et dur&eacute;es indicatives",
        "7. Sc&eacute;narios &eacute;conomiques et revenus possibles",
        "8. Travaux possibles, limites et mat&eacute;riaux de r&eacute;emploi",
        "9. Proc&eacute;dure pas &agrave; pas",
        "10. Pi&egrave;ces &agrave; fournir et fiche propri&eacute;taire",
        "11. FAQ et points de vigilance",
    ]
    for item in items:
        add_bullet(doc, item)
    note(
        doc,
        "Cette brochure informe et oriente. Elle ne remplace pas un conseil juridique, "
        "fiscal, notarial, architectural ou technique adapt&eacute; au dossier.",
    )


def why(doc):
    doc.add_heading(u("Pourquoi agir sur un bien dormant ?"), level=1)
    add_paragraph(
        doc,
        "Un logement, un local ou un terrain inutilis&eacute; n'est jamais totalement neutre. "
        "Il peut perdre de la valeur, se d&eacute;grader, g&eacute;n&eacute;rer des charges, poser des "
        "difficult&eacute;s d'assurance ou d'entretien, et donner une image d'abandon dans "
        "un quartier. Pour un propri&eacute;taire, la vacance peut aussi devenir un sujet "
        "lourd &agrave; g&eacute;rer : travaux trop co&ucirc;teux, manque de temps, h&eacute;ritage complexe, "
        "absence de projet, peur des impay&eacute;s ou incertitude sur le bon montage.",
    )
    add_paragraph(
        doc,
        "L'approche TVF consiste &agrave; regarder le bien non comme un probl&egrave;me isol&eacute;, "
        "mais comme une ressource potentielle. Un bien vacant peut devenir un logement "
        "&eacute;tudiant, un logement associatif, un atelier partag&eacute;, un lieu de formation, "
        "un commerce de proximit&eacute;, un espace temporaire ou un support de projet "
        "solidaire, &agrave; condition que le cadre soit s&eacute;rieux, &eacute;crit et compatible "
        "avec les moyens disponibles.",
    )
    add_callout(
        doc,
        "Ce que le propri&eacute;taire doit retenir",
        "Une demande adress&eacute;e &agrave; TVF ne vaut pas acceptation automatique, ne vaut pas "
        "engagement de travaux imm&eacute;diats et ne garantit pas un revenu. Elle ouvre "
        "une phase d'instruction permettant de v&eacute;rifier si le bien peut entrer dans "
        "un projet r&eacute;aliste.",
        fill=LIGHT_BLUE,
    )


def tvf_role(doc):
    doc.add_heading(u("Ce que TVF peut faire pour un particulier"), level=1)
    rows = [
        ("&Eacute;couter et qualifier", "Comprendre la situation du propri&eacute;taire, l'histoire du bien, les contraintes, les objectifs et le niveau d'ouverture &agrave; une coop&eacute;ration."),
        ("Pr&eacute;diagnostiquer", "R&eacute;aliser une premi&egrave;re lecture du bien : usage possible, &eacute;tat apparent, points bloquants, documents manquants, besoins de visite ou d'expertise."),
        ("Orienter techniquement", "Lorsque le sujet d&eacute;passe la pr&eacute;qualification, TVF peut orienter vers une personne comp&eacute;tente : artisan, diagnostiqueur, architecte, bureau d'&eacute;tudes, notaire ou autre professionnel selon le besoin."),
        ("Mobiliser des ressources", "Identifier des mat&eacute;riaux de r&eacute;emploi disponibles, des comp&eacute;tences, des b&eacute;n&eacute;voles, des entreprises ou des partenaires pouvant contribuer au projet, sans garantie de disponibilit&eacute;."),
        ("Structurer le montage", "Proposer un cadre de conventionnement adapt&eacute; : mise &agrave; disposition, usage temporaire, loyer solidaire, partage de revenus ou coop&eacute;ration territoriale."),
        ("Suivre le projet", "Garder une trace du dossier, des d&eacute;cisions, des pi&egrave;ces, des engagements et des &eacute;tapes jusqu'&agrave; la mise en usage ou &agrave; l'abandon motiv&eacute; du dossier."),
    ]
    table(doc, ["R&ocirc;le TVF", "Ce que cela signifie"], rows, [4.2, 12.3], fill=LIGHT_GREEN)

    doc.add_heading(u("Ce que TVF ne promet pas"), level=2)
    for item in [
        "TVF ne promet pas une r&eacute;novation gratuite ou imm&eacute;diate.",
        "TVF ne s'engage pas &agrave; accepter tous les biens propos&eacute;s.",
        "TVF ne garantit pas un rendement financier, un loyer ou une rentabilit&eacute;.",
        "TVF ne se substitue pas &agrave; un notaire, un avocat, un architecte, un diagnostiqueur ou un ma&icirc;tre d'oeuvre lorsque leur intervention est n&eacute;cessaire.",
        "TVF n'intervient pas sur les grosses ma&ccedil;onneries, les reprises structurelles lourdes, les situations de p&eacute;ril ou les travaux dangereux non encadr&eacute;s par des professionnels qualifi&eacute;s.",
    ]:
        add_bullet(doc, item)


def goods_and_uses(doc):
    doc.add_heading(u("Biens concern&eacute;s et usages possibles"), level=1)
    rows = [
        ("Logement vacant", "Studio, appartement, maison, logement au-dessus d'un commerce, logement familial inoccup&eacute;.", "Logement &eacute;tudiant, logement temporaire, logement associatif, habitat interg&eacute;n&eacute;rationnel, logement pour travailleur ou personne accompagn&eacute;e."),
        ("Immeuble d&eacute;grad&eacute;", "Petit immeuble ancien, bien h&eacute;rit&eacute;, immeuble partiellement vide, parties communes &agrave; remettre en &eacute;tat.", "R&eacute;habilitation progressive, occupation temporaire encadr&eacute;e, projet mixte logement/association, chantier solidaire si faisable."),
        ("Commerce ou local ferm&eacute;", "Boutique, cellule commerciale, ancien atelier, rez-de-chauss&eacute;e inutilis&eacute;.", "Commerce de proximit&eacute;, atelier partag&eacute;, lieu associatif, espace culturel, local de formation, boutique &eacute;ph&eacute;m&egrave;re."),
        ("B&acirc;timent inutilis&eacute;", "Garage, d&eacute;p&ocirc;t, annexe, petite halle, b&acirc;timent agricole ou local professionnel vacant.", "Stockage de mat&eacute;riaux, atelier, activit&eacute; ESS, espace technique, local mutualis&eacute; selon s&eacute;curit&eacute;."),
        ("Terrain ou espace d&eacute;laiss&eacute;", "Cour, friche, jardin abandonn&eacute;, parcelle non utilis&eacute;e, terrain en attente.", "Jardin partag&eacute;, micro-verger, espace p&eacute;dagogique, base chantier, espace de rencontre ou am&eacute;nagement temporaire."),
    ]
    table(doc, ["Type de bien", "Exemples", "Usages envisageables"], rows, [3.8, 5.7, 7.0])


def selection(doc):
    doc.add_heading(u("Comment TVF choisit les biens et les dossiers"), level=1)
    add_paragraph(
        doc,
        "TVF ne s&eacute;lectionne pas un bien uniquement parce qu'il est vacant. Le dossier "
        "doit entrer dans une logique de d&eacute;veloppement territorial : utilit&eacute; pour les "
        "habitants, faisabilit&eacute; technique, cadre juridique clair, possibilit&eacute; de "
        "financement, compatibilit&eacute; avec les p&ocirc;les TVF et capacit&eacute; &agrave; produire "
        "un usage concret.",
    )
    rows = [
        ("Propri&eacute;t&eacute; et autorisation", "Le propri&eacute;taire ou son repr&eacute;sentant doit pouvoir justifier qu'il a le droit de proposer le bien et de signer une convention."),
        ("Utilit&eacute; territoriale", "Le bien doit pouvoir servir un besoin : logement, association, commerce de proximit&eacute;, formation, stockage, insertion, espace vert ou activit&eacute; locale."),
        ("Faisabilit&eacute; technique", "L'&eacute;tat apparent doit permettre une remise en usage raisonnable. Les travaux lourds ou dangereux ne sont pas accept&eacute;s sans professionnels qualifi&eacute;s et financement adapt&eacute;."),
        ("Niveau de travaux", "TVF privil&eacute;gie la r&eacute;novation, la remise en &eacute;tat, le second oeuvre, l'am&eacute;nagement et le r&eacute;emploi. Les grosses ma&ccedil;onneries, reprises structurelles lourdes, p&eacute;rils ou d&eacute;samiantages complexes ne rel&egrave;vent pas d'une intervention TVF simple."),
        ("Ressources disponibles", "Le dossier est plus solide si des mat&eacute;riaux, comp&eacute;tences, entreprises, b&eacute;n&eacute;voles ou financements peuvent &ecirc;tre mobilis&eacute;s. Leur disponibilit&eacute; n'est jamais automatique."),
        ("Dur&eacute;e acceptable", "La dur&eacute;e de convention doit &ecirc;tre coh&eacute;rente avec l'investissement engag&eacute;, l'usage pr&eacute;vu et l'int&eacute;r&ecirc;t du propri&eacute;taire."),
        ("Risques et assurances", "TVF doit pouvoir identifier les risques, les conditions d'assurance, la s&eacute;curit&eacute; des personnes et les responsabilit&eacute;s de chacun."),
        ("Impact et suivi", "Un dossier prioritaire doit permettre de suivre ce qu'il produit : bien remis en usage, service cr&eacute;&eacute;, local ouvert, mat&eacute;riaux r&eacute;employ&eacute;s ou activit&eacute; utile."),
    ]
    table(doc, ["Crit&egrave;re", "Ce que TVF v&eacute;rifie"], rows, [4.1, 12.4], fill=LIGHT_BLUE)

    doc.add_heading(u("Niveaux d'orientation possibles"), level=2)
    rows = [
        ("Niveau 1 - &Eacute;tude possible", "Le bien est clair juridiquement, accessible, document&eacute;, et semble compatible avec une remise en usage raisonnable."),
        ("Niveau 2 - &Agrave; compl&eacute;ter", "Le dossier peut &ecirc;tre int&eacute;ressant, mais il manque des pi&egrave;ces, photos, diagnostics, devis, autorisations ou informations sur les risques."),
        ("Niveau 3 - Orientation externe", "Le bien demande d'abord une expertise professionnelle, un diagnostic technique, une intervention juridique, une d&eacute;cision de copropri&eacute;t&eacute; ou un montage financier."),
        ("Niveau 4 - Non retenu &agrave; ce stade", "Le bien pr&eacute;sente des risques trop importants, une impossibilit&eacute; d'acc&egrave;s, une situation juridique incertaine, des travaux trop lourds ou une absence d'usage territorial identifiable."),
    ]
    table(doc, ["Niveau", "Signification"], rows, [4.4, 12.1])


def diagnosis(doc):
    doc.add_heading(u("Diagnostic, &eacute;tude et orientation technique"), level=1)
    add_paragraph(
        doc,
        "Le diagnostic peut se faire en plusieurs niveaux. TVF peut assurer une premi&egrave;re "
        "lecture de terrain et de dossier : photos, usage possible, coh&eacute;rence avec l'objet "
        "associatif, documents manquants, risques apparents. Cette lecture ne remplace pas "
        "un diagnostic technique r&eacute;glementaire ou une expertise professionnelle.",
    )
    add_paragraph(
        doc,
        "Lorsque le dossier l'exige, TVF peut orienter le propri&eacute;taire vers une personne "
        "comp&eacute;tente : diagnostiqueur, artisan, architecte, bureau d'&eacute;tudes, expert "
        "b&acirc;timent, notaire, avocat ou autre professionnel. Selon les cas, le professionnel "
        "peut &ecirc;tre propos&eacute; par TVF, choisi par le propri&eacute;taire ou retenu d'un commun "
        "accord. Les modalit&eacute;s, les co&ucirc;ts et les responsabilit&eacute;s doivent &ecirc;tre "
        "clarifi&eacute;s avant toute intervention.",
    )
    rows = [
        ("Pr&eacute;diagnostic TVF", "Lecture initiale, photos, &eacute;change propri&eacute;taire, coh&eacute;rence du projet, identification des points bloquants."),
        ("Visite de rep&eacute;rage", "Visite non intrusive lorsque le bien est accessible et que la s&eacute;curit&eacute; minimale est assur&eacute;e."),
        ("Diagnostic professionnel", "Intervention d'un professionnel qualifi&eacute; lorsque la structure, les fluides, la s&eacute;curit&eacute;, l'amiante, le plomb, l'&eacute;lectricit&eacute; ou le gaz doivent &ecirc;tre v&eacute;rifi&eacute;s."),
        ("Sc&eacute;nario d'usage", "D&eacute;finition d'un usage possible : logement, association, commerce, atelier, stockage, formation, occupation temporaire."),
        ("Sc&eacute;nario &eacute;conomique", "Lecture des co&ucirc;ts, apports possibles, mat&eacute;riaux, dur&eacute;e de convention, loyer solidaire ou partage de revenus si pertinent."),
    ]
    table(doc, ["Niveau", "Contenu"], rows, [4.4, 12.1], fill=LIGHT_BLUE)


def conventions(doc):
    doc.add_heading(u("Conventions possibles"), level=1)
    add_paragraph(
        doc,
        "Le conventionnement est le coeur du dispositif. Il sert &agrave; &eacute;crire clairement "
        "ce que chaque partie accepte : dur&eacute;e, usage, travaux, acc&egrave;s, assurance, "
        "responsabilit&eacute;s, participation financi&egrave;re, revenus &eacute;ventuels, communication "
        "et conditions de restitution du bien.",
    )
    rows = [
        ("Autorisation de visite et d'&eacute;tude", "Permet &agrave; TVF ou &agrave; des personnes mandat&eacute;es d'&eacute;tudier le bien, sans engagement de travaux.", "Ponctuelle ou quelques semaines"),
        ("Lettre d'intention", "Cadre pr&eacute;alable indiquant que les parties acceptent d'&eacute;tudier un projet sans engagement d&eacute;finitif.", "1 &agrave; 6 mois selon dossier"),
        ("Convention de valorisation", "Cadre souple pour qualifier le bien, mobiliser des ressources, rechercher des partenaires et construire un sc&eacute;nario.", "6 &agrave; 18 mois"),
        ("Mise &agrave; disposition temporaire", "Le propri&eacute;taire autorise un usage encadr&eacute; du bien pour une dur&eacute;e limit&eacute;e, avec obligations d'entretien, assurance et restitution.", "1 &agrave; 3 ans selon usage"),
        ("Bien Solidaire &agrave; Usage Partag&eacute;", "TVF coordonne une remise en usage en &eacute;change d'un droit d'usage temporaire proportionn&eacute; &agrave; l'investissement et au projet.", "3, 5, 10 ou 15 ans selon investissement"),
        ("Loyer solidaire", "Le propri&eacute;taire per&ccedil;oit un loyer mod&eacute;r&eacute; ou progressif compatible avec l'usage social, associatif ou territorial.", "Dur&eacute;e &agrave; fixer selon cadre juridique"),
        ("Partage de revenus", "Lorsque le bien g&eacute;n&egrave;re une activit&eacute;, une part des recettes peut &ecirc;tre affect&eacute;e au propri&eacute;taire, au projet ou &agrave; l'entretien.", "Dur&eacute;e &agrave; fixer par convention"),
        ("Convention longue sp&eacute;cifique", "Pour les dossiers importants, un montage plus long peut &ecirc;tre &eacute;tudi&eacute; avec un professionnel du droit.", "Au cas par cas"),
    ]
    table(doc, ["Convention", "Utilit&eacute;", "Dur&eacute;e indicative"], rows, [4.5, 8.1, 3.9])
    note(doc, "Les dur&eacute;es sont indicatives. Elles doivent &ecirc;tre adapt&eacute;es au droit applicable, au type de bien, &agrave; l'usage, &agrave; l'investissement et aux conseils professionnels mobilis&eacute;s.")


def economics(doc):
    doc.add_heading(u("Sc&eacute;narios &eacute;conomiques et revenus possibles"), level=1)
    add_paragraph(
        doc,
        "Un projet TVF peut &ecirc;tre utile au territoire tout en pr&eacute;servant l'int&eacute;r&ecirc;t du "
        "propri&eacute;taire. Selon le type de bien, les travaux, les financements et l'usage "
        "retenu, plusieurs sc&eacute;narios peuvent &ecirc;tre &eacute;tudi&eacute;s. Aucun ne constitue une "
        "promesse de rendement : tout doit &ecirc;tre calcul&eacute;, &eacute;crit et valid&eacute; avant engagement.",
    )
    rows = [
        ("Formule A - R&eacute;novation contre usage", "TVF ou les partenaires mobilis&eacute;s contribuent &agrave; remettre le bien en &eacute;tat. En retour, une dur&eacute;e d'usage temporaire est convenue. Le propri&eacute;taire r&eacute;cup&egrave;re ensuite un bien am&eacute;lior&eacute; selon les termes pr&eacute;vus.", "Propri&eacute;taire sans budget imm&eacute;diat mais ouvert &agrave; un usage temporaire utile."),
        ("Formule B - Loyer solidaire", "Le bien est utilis&eacute; contre un loyer mod&eacute;r&eacute;, progressif ou adapt&eacute; &agrave; l'activit&eacute;. Le niveau de loyer doit rester compatible avec le projet social ou territorial.", "Logement, local associatif, atelier, commerce de proximit&eacute;."),
        ("Formule C - Partage des revenus", "Lorsque le bien permet une activit&eacute; &eacute;conomique, une cl&eacute; de partage peut &ecirc;tre pr&eacute;vue : propri&eacute;taire, projet, entretien, fonds de travaux ou fonctionnement.", "Commerce, atelier, tiers-lieu, espace de formation, activit&eacute; mutualis&eacute;e."),
        ("Formule D - Contribution aux charges", "Le projet peut prendre en charge une partie des charges, fluides, entretien courant, assurance d'usage ou petits travaux selon convention.", "Occupation temporaire ou mise &agrave; disposition encadr&eacute;e."),
        ("Formule E - Montage mixte", "Combinaison possible : mat&eacute;riaux de r&eacute;emploi, b&eacute;n&eacute;volat encadr&eacute;, entreprise partenaire, subvention, loyer solidaire et usage temporaire.", "Dossiers plus complexes n&eacute;cessitant une instruction compl&egrave;te."),
    ]
    table(doc, ["Sc&eacute;nario", "Principe", "Cas adapt&eacute;"], rows, [4.3, 7.5, 4.7], fill=LIGHT_BLUE)
    add_callout(
        doc,
        "Rendement et prudence",
        "Un revenu peut &ecirc;tre envisag&eacute; lorsque le projet le permet : loyer solidaire, "
        "indemnit&eacute;, part de recettes, prise en charge de charges ou valorisation du "
        "patrimoine. Mais TVF ne garantit jamais un rendement. Le calcul d&eacute;pend du "
        "co&ucirc;t des travaux, de la dur&eacute;e, de l'usage, des risques, des financements "
        "et du cadre juridique retenu.",
        fill=LIGHT_GREEN,
    )


def works_materials(doc):
    doc.add_heading(u("Travaux, limites et mat&eacute;riaux de r&eacute;emploi"), level=1)
    add_paragraph(
        doc,
        "TVF intervient dans une logique de r&eacute;novation utile et raisonnable. L'association "
        "peut contribuer &agrave; organiser des travaux de remise en &eacute;tat, d'am&eacute;nagement, "
        "de second oeuvre ou de valorisation, mais les travaux lourds doivent &ecirc;tre "
        "trait&eacute;s par des professionnels qualifi&eacute;s et dans un cadre distinct.",
    )
    rows = [
        ("Travaux envisageables", "Nettoyage, d&eacute;barras, peinture, rev&ecirc;tements, petits am&eacute;nagements, menuiseries simples, mobilier, agencement, r&eacute;emploi de mat&eacute;riaux, remise en usage progressive."),
        ("Travaux &agrave; encadrer fortement", "&Eacute;lectricit&eacute;, plomberie, chauffage, isolation, s&eacute;curit&eacute;, accessibilit&eacute;, diagnostics, travaux soumis &agrave; autorisation ou assurance professionnelle."),
        ("Travaux non pris en charge simplement", "Grosse ma&ccedil;onnerie, reprise structurelle, toiture lourde, p&eacute;ril, d&eacute;samiantage complexe, d&eacute;molition, travaux dangereux ou non assurables."),
        ("Mat&eacute;riaux de r&eacute;emploi", "TVF peut &eacute;tudier l'utilisation de mat&eacute;riaux disponibles : bois, menuiseries, sanitaires, &eacute;clairage, mobilier, rev&ecirc;tements, &eacute;quipements. Leur disponibilit&eacute; et leur compatibilit&eacute; ne sont pas garanties."),
        ("Niveau d'acceptation", "Plus le bien n&eacute;cessite des travaux lourds, urgents ou dangereux, plus le dossier devra &ecirc;tre orient&eacute; vers une expertise externe ou report&eacute; jusqu'&agrave; financement adapt&eacute;."),
    ]
    table(doc, ["Sujet", "Position TVF"], rows, [4.2, 12.3])


def procedure(doc):
    doc.add_heading(u("Proc&eacute;dure pas &agrave; pas"), level=1)
    rows = [
        ("1. Premier contact", "Le propri&eacute;taire contacte TVF et d&eacute;crit bri&egrave;vement son bien, sa situation et son objectif."),
        ("2. Fiche propri&eacute;taire", "TVF demande les informations de base : adresse, type de bien, &eacute;tat, photos, disponibilit&eacute;, contraintes, documents existants."),
        ("3. Pr&eacute;qualification", "TVF v&eacute;rifie si le dossier peut entrer dans son objet, s'il est suffisamment document&eacute; et s'il existe un usage territorial potentiel."),
        ("4. Visite ou diagnostic", "Une visite peut &ecirc;tre organis&eacute;e si le bien est accessible et s&eacute;curis&eacute;. Si besoin, TVF oriente vers un professionnel comp&eacute;tent."),
        ("5. Sc&eacute;nario de projet", "TVF propose ou &eacute;tudie un ou plusieurs usages : logement, association, commerce, atelier, stockage, espace de formation, jardin, etc."),
        ("6. Lecture &eacute;conomique", "Les co&ucirc;ts, ressources, mat&eacute;riaux, contributions, financements possibles, dur&eacute;e de convention et revenus &eacute;ventuels sont mis en regard."),
        ("7. D&eacute;cision d'instruction", "TVF peut accepter l'instruction, demander des pi&egrave;ces, orienter le dossier, le mettre en attente ou refuser avec motif."),
        ("8. Convention", "Si le dossier avance, une convention &eacute;crite fixe les engagements, limites, dur&eacute;es, usages, assurances et modalit&eacute;s de restitution."),
        ("9. Mise en oeuvre", "Travaux raisonnables, collecte de mat&eacute;riaux, mobilisation de partenaires, coordination, communication encadr&eacute;e et suivi."),
        ("10. Restitution ou poursuite", "En fin de convention : restitution du bien, renouvellement, changement d'usage ou nouveau montage selon accord."),
    ]
    table(doc, ["&Eacute;tape", "Objectif"], rows, [3.8, 12.7], fill=LIGHT_BLUE)


def documents_needed(doc):
    doc.add_heading(u("Pi&egrave;ces et informations &agrave; fournir"), level=1)
    rows = [
        ("Identit&eacute; du propri&eacute;taire", "Nom, pr&eacute;nom, adresse, t&eacute;l&eacute;phone, e-mail, situation : propri&eacute;taire unique, indivision, SCI, bailleur, mandataire."),
        ("Justificatif de propri&eacute;t&eacute;", "Titre de propri&eacute;t&eacute;, attestation notariale, taxe fonci&egrave;re ou document permettant de confirmer le lien avec le bien."),
        ("Adresse et description", "Adresse exacte, type de bien, surface approximative, nombre de pi&egrave;ces, niveaux, acc&egrave;s, stationnement, &eacute;tat g&eacute;n&eacute;ral."),
        ("Photos", "Fa&ccedil;ade, entr&eacute;e, pi&egrave;ces principales, toiture apparente si visible, sanitaires, &eacute;lectricit&eacute;, sols, murs, acc&egrave;s, points de vigilance."),
        ("Diagnostics disponibles", "DPE, &eacute;lectricit&eacute;, gaz, amiante, plomb, termites, ERP, diagnostics copropri&eacute;t&eacute; ou tout document d&eacute;j&agrave; disponible."),
        ("Situation juridique", "Occupation, bail en cours, indivision, servitude, copropri&eacute;t&eacute;, autorisations, litige, assurance, sinistre, arr&ecirc;t&eacute; ou contrainte connue."),
        ("Objectif du propri&eacute;taire", "Vendre, conserver, valoriser, louer, mettre &agrave; disposition, obtenir un revenu, soutenir un projet local, &eacute;viter la d&eacute;gradation."),
        ("Contraintes", "Budget, calendrier, travaux refus&eacute;s ou accept&eacute;s, dur&eacute;e maximale d'usage, zones interdites, acc&egrave;s, voisinage, disponibilit&eacute;."),
    ]
    table(doc, ["Pi&egrave;ce / information", "Utilit&eacute; pour l'instruction"], rows, [4.5, 12.0])


def faq(doc):
    doc.add_heading(u("FAQ propri&eacute;taires"), level=1)
    rows = [
        ("Est-ce que TVF accepte tous les biens ?", "Non. Chaque demande est instruite selon l'&eacute;tat du bien, l'usage possible, la s&eacute;curit&eacute;, les moyens mobilisables, les documents disponibles et l'int&eacute;r&ecirc;t territorial."),
        ("Est-ce que TVF r&eacute;nove gratuitement mon bien ?", "Non. TVF peut construire un montage, rechercher des ressources, mobiliser des partenaires ou des mat&eacute;riaux, mais aucune r&eacute;novation gratuite ou imm&eacute;diate n'est garantie."),
        ("Puis-je gagner de l'argent ?", "Un revenu peut &ecirc;tre &eacute;tudi&eacute; selon le montage : loyer solidaire, partage de recettes, indemnit&eacute;, prise en charge de charges. Rien n'est garanti avant convention."),
        ("Qui choisit l'usage du bien ?", "L'usage est construit avec le propri&eacute;taire et TVF. Il doit &ecirc;tre compatible avec le bien, les autorisations, le territoire, les moyens et la convention."),
        ("Puis-je r&eacute;cup&eacute;rer mon bien ?", "Oui, selon les conditions de la convention. La dur&eacute;e, le pr&eacute;avis, la restitution et l'&eacute;tat attendu doivent &ecirc;tre &eacute;crits."),
        ("Que se passe-t-il si le dossier est trop lourd ?", "TVF peut demander des pi&egrave;ces, orienter vers des professionnels, mettre le dossier en attente ou refuser l'instruction si les risques sont trop importants."),
        ("Les mat&eacute;riaux de r&eacute;emploi sont-ils garantis ?", "Non. Ils d&eacute;pendent des ressources disponibles, de leur qualit&eacute;, de leur compatibilit&eacute; et de la capacit&eacute; &agrave; les stocker ou les transporter."),
    ]
    table(doc, ["Question", "R&eacute;ponse"], rows, [5.2, 11.3], fill=LIGHT_BLUE)


def owner_form(doc):
    doc.add_page_break()
    doc.add_heading(u("Fiche propri&eacute;taire - proposition de bien dormant"), level=1)
    add_paragraph(doc, "Cette fiche permet de transmettre une premi&egrave;re demande &agrave; TVF. Elle ne vaut pas acceptation du bien ni engagement de travaux. Elle sert &agrave; ouvrir une pr&eacute;qualification.")
    rows = [
        ("Date", "................................................................................................"),
        ("Nom / pr&eacute;nom", "................................................................................................"),
        ("T&eacute;l&eacute;phone / e-mail", "................................................................................................"),
        ("Statut", "[ ] Propri&eacute;taire unique   [ ] Indivision   [ ] SCI   [ ] Mandataire   [ ] Autre : ........................"),
        ("Adresse du bien", "................................................................................................"),
        ("Type de bien", "[ ] Logement   [ ] Immeuble   [ ] Commerce   [ ] Local   [ ] B&acirc;timent   [ ] Terrain   [ ] Autre"),
        ("Surface / niveaux", "................................................................................................"),
        ("Situation actuelle", "[ ] Vacant   [ ] D&eacute;grad&eacute;   [ ] Partiellement occup&eacute;   [ ] Ferm&eacute;   [ ] En attente de travaux"),
        ("Depuis quand ?", "................................................................................................"),
        ("&Eacute;tat apparent", "[ ] Bon   [ ] Moyen   [ ] D&eacute;grad&eacute;   [ ] Tr&egrave;s d&eacute;grad&eacute;   [ ] &Agrave; diagnostiquer"),
        ("Photos jointes", "[ ] Oui   [ ] Non   Nombre : .................."),
        ("Diagnostics disponibles", "[ ] DPE   [ ] &Eacute;lectricit&eacute;   [ ] Gaz   [ ] Amiante   [ ] Plomb   [ ] Aucun   [ ] Autre"),
        ("Objectif du propri&eacute;taire", "[ ] Valoriser   [ ] Louer   [ ] Soutenir un projet   [ ] &Eacute;viter la d&eacute;gradation   [ ] Obtenir un revenu"),
        ("Usage envisag&eacute;", "[ ] Logement &eacute;tudiant   [ ] Association   [ ] Commerce   [ ] Atelier   [ ] Tiers-lieu   [ ] Autre"),
        ("Dur&eacute;e envisageable", "[ ] 6-12 mois   [ ] 1-3 ans   [ ] 3-5 ans   [ ] 5-10 ans   [ ] 10 ans et plus   [ ] &Agrave; discuter"),
        ("Travaux acceptables", "[ ] Nettoyage   [ ] Peinture   [ ] Second oeuvre   [ ] Am&eacute;nagement   [ ] &Agrave; d&eacute;finir"),
        ("Contraintes connues", "................................................................................................\n................................................................................................"),
        ("Acc&egrave;s / logistique", "................................................................................................"),
        ("Documents &agrave; joindre", "[ ] Justificatif propri&eacute;t&eacute;   [ ] Taxe fonci&egrave;re   [ ] Photos   [ ] Diagnostics   [ ] Plans   [ ] Assurance"),
        ("D&eacute;claration", "Je comprends que cette demande ne vaut pas acceptation, ne vaut pas convention et ne vaut pas r&eacute;novation imm&eacute;diate."),
        ("Signature", "................................................................................................"),
    ]
    table(doc, ["Champ", "R&eacute;ponse / information &agrave; renseigner"], rows, [4.7, 11.8])
    add_callout(
        doc,
        "Contact TVF",
        "TERRITOIRES VIVANTS FRANCE - 25 rue &Eacute;lise Gervais, 42000 Saint-&Eacute;tienne - contact@territoiresvivantsfrance.fr - 04 65 81 54 69 - www.territoiresvivantsfrance.fr",
        fill=LIGHT_BLUE,
    )


def final_notes(doc):
    doc.add_heading(u("Points &agrave; verrouiller avant toute action"), level=1)
    for item in [
        "Identification du propri&eacute;taire et droit de proposer le bien.",
        "Accord &eacute;crit avant visite, intervention ou communication.",
        "V&eacute;rification des risques : structure, amiante, plomb, &eacute;lectricit&eacute;, gaz, humidit&eacute;, assurance.",
        "D&eacute;finition &eacute;crite de l'usage, de la dur&eacute;e, des responsabilit&eacute;s et de la restitution.",
        "Clarification du financement, des mat&eacute;riaux disponibles et des travaux r&eacute;ellement faisables.",
        "Validation de toute mention publique du propri&eacute;taire ou du bien.",
        "Recours &agrave; un professionnel comp&eacute;tent lorsque le sujet technique, juridique ou fiscal l'exige.",
    ]:
        add_bullet(doc, "[ ] " + item)
    add_callout(
        doc,
        "Conclusion",
        "TVF peut aider un propri&eacute;taire &agrave; transformer un bien dormant en ressource utile. "
        "La m&eacute;thode repose sur une instruction prudente : comprendre, qualifier, diagnostiquer, "
        "choisir un usage, conventionner, mobiliser les moyens, suivre l'impact et restituer le "
        "bien selon les conditions pr&eacute;vues.",
        fill=LIGHT_GREEN,
    )


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    cover(doc)
    toc(doc)
    doc.add_page_break()
    why(doc)
    tvf_role(doc)
    goods_and_uses(doc)
    selection(doc)
    diagnosis(doc)
    conventions(doc)
    economics(doc)
    works_materials(doc)
    procedure(doc)
    documents_needed(doc)
    faq(doc)
    owner_form(doc)
    final_notes(doc)
    doc.save(DOCX_PATH)


def build_md():
    md = u(
        """# Brochure TVF - Particuliers et propri&eacute;taires

Document d'information destin&eacute; aux particuliers poss&eacute;dant un bien dormant : logement vacant, commerce ferm&eacute;, local inutilis&eacute;, immeuble d&eacute;grad&eacute;, terrain ou b&acirc;timent sous-utilis&eacute;.

## Principes

- Le propri&eacute;taire conserve la propri&eacute;t&eacute; du bien.
- Toute demande ouvre une phase d'instruction ; elle ne vaut pas acceptation ni r&eacute;novation imm&eacute;diate.
- TVF privil&eacute;gie la remise en usage, la r&eacute;novation raisonnable, le r&eacute;emploi de mat&eacute;riaux et les projets utiles au territoire.
- Les grosses ma&ccedil;onneries, reprises structurelles lourdes, p&eacute;rils ou travaux dangereux doivent &ecirc;tre trait&eacute;s par des professionnels qualifi&eacute;s et ne rel&egrave;vent pas d'une intervention simple TVF.
- Des revenus peuvent &ecirc;tre &eacute;tudi&eacute;s selon convention : loyer solidaire, partage de recettes, indemnisation, prise en charge de charges ou valorisation patrimoniale. Aucun rendement n'est garanti.

## Contact

TERRITOIRES VIVANTS FRANCE - contact@territoiresvivantsfrance.fr - 04 65 81 54 69 - www.territoiresvivantsfrance.fr
"""
    )
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)

