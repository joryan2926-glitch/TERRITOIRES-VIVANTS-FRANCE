from pathlib import Path
from runpy import run_path
from html import unescape

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = run_path(str(ROOT / "scripts" / "build-courrier-metropole-tvf.py"))

Document = BASE["Document"]
style_document = BASE["style_document"]
add_header_footer = BASE["add_header_footer"]
add_paragraph = BASE["add_paragraph"]
add_bullet = BASE["add_bullet"]
add_callout = BASE["add_callout"]
set_table_width = BASE["set_table_width"]
set_cell_border = BASE["set_cell_border"]
set_cell_shading = BASE["set_cell_shading"]
set_repeat_table_header = BASE["set_repeat_table_header"]
GREEN = BASE["GREEN"]
BLUE = BASE["BLUE"]
MUTED = BASE["MUTED"]
LIGHT_GREEN = BASE["LIGHT_GREEN"]
LIGHT_BLUE = BASE["LIGHT_BLUE"]
GOLD = BASE["GOLD"]
BORDER = BASE["BORDER"]

OUT_DIR = ROOT / "documents" / "brochures"
SRC_DIR = ROOT / "documents" / "sources"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SRC_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "brochure-collectivites-cooperation-territoriale-tvf.docx"
MD_PATH = SRC_DIR / "brochure-collectivites-cooperation-territoriale-tvf.md"
LOGO = ROOT / "assets" / "logo-territoires-vivants-france-web.png"
PHOTO = ROOT / "assets" / "photos" / "france-saint-etienne-chateaucreux.jpg"


def u(text: str) -> str:
    return unescape(text)


def set_run(run, size=None, color=None, bold=None, italic=None):
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(u(text))
    set_run(r, 8.5, MUTED, italic=True)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(u(text))
    set_run(r, 8.4, MUTED, italic=True)


def table(doc, headers, rows, widths, fill=LIGHT_GREEN):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, widths)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=BORDER, size="8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(u(header))
        set_run(r, 8.8, GREEN, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i], color="DDE6DE", size="6")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(u(value))
            set_run(r, 8.45, BLUE if i == 0 else RGBColor(31, 45, 38), bold=(i == 0))
    doc.add_paragraph()
    return t


def cover(doc):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Cm(5.8))
        p_logo.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(u("Brochure collectivit&eacute;s territoriales"))
    set_run(r, 10.2, GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(u("Coop&eacute;ration territoriale, patrimoine vacant et ressources locales"))
    set_run(r, 23.5, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run(
        u(
            "Faire de TVF un outil op&eacute;rationnel au service des communes, intercommunalit&eacute;s, "
            "D&eacute;partements et R&eacute;gions qui souhaitent r&eacute;activer des lieux, des mat&eacute;riaux "
            "et des projets utiles."
        )
    )
    set_run(r, 12.5, GREEN, bold=True)

    meta = doc.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, [5.5, 5.5, 5.5])
    values = [
        ("Document", "information et cadrage"),
        ("Public", "collectivit&eacute;s et services publics"),
        ("Version", "juillet 2026"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = meta.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GREEN if idx != 1 else LIGHT_BLUE)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(u(label + "\n"))
        set_run(r1, 8.2, MUTED, bold=True)
        r2 = p.add_run(u(value))
        set_run(r2, 9.2, BLUE, bold=True)

    if PHOTO.exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(PHOTO), width=Cm(15.6))
        caption(
            doc,
            "Territoire urbain en transformation : TVF intervient comme outil de mise en relation, "
            "d'instruction et de coop&eacute;ration, sans se substituer aux comp&eacute;tences publiques.",
        )

    add_callout(
        doc,
        "Message central",
        "TVF ne remplace pas les politiques publiques locales. L'association peut aider &agrave; "
        "rep&eacute;rer, qualifier, relier et structurer des projets autour du patrimoine vacant, "
        "des commerces inoccup&eacute;s, des friches, des ressources inutilis&eacute;es et du r&eacute;emploi.",
        fill=LIGHT_GREEN,
    )


def toc(doc):
    doc.add_heading("Sommaire", level=1)
    items = [
        "1. Pourquoi les collectivit&eacute;s sont concern&eacute;es",
        "2. Ce que TVF peut apporter &agrave; une collectivit&eacute;",
        "3. Politiques publiques et besoins compatibles",
        "4. Ce que la collectivit&eacute; peut proposer",
        "5. B&eacute;n&eacute;fices pour le territoire",
        "6. Formes de coop&eacute;ration et conventions possibles",
        "7. Parcours type d'un dossier territorial",
        "8. Crit&egrave;res de s&eacute;lection et limites d'intervention",
        "9. Stockage, mat&eacute;riaux et locaux techniques",
        "10. Pi&egrave;ces &agrave; fournir et fiche collectivit&eacute;",
        "11. Points de vigilance avant signature",
    ]
    for item in items:
        add_bullet(doc, item)
    note(
        doc,
        "Cette brochure informe et oriente. Toute coop&eacute;ration avec une collectivit&eacute; doit "
        "respecter ses comp&eacute;tences, ses proc&eacute;dures internes, ses r&egrave;gles budg&eacute;taires, "
        "ses obligations de commande publique et les d&eacute;cisions de ses organes comp&eacute;tents.",
    )


def why(doc):
    doc.add_heading(u("Pourquoi les collectivit&eacute;s sont concern&eacute;es ?"), level=1)
    add_paragraph(
        doc,
        "Les collectivit&eacute;s sont directement confront&eacute;es aux effets visibles de la vacance : "
        "logements ferm&eacute;s, rez-de-chauss&eacute;e commerciaux inoccup&eacute;s, friches, espaces d&eacute;laiss&eacute;s, "
        "b&acirc;timents publics sous-utilis&eacute;s, besoins associatifs non satisfaits et difficult&eacute;s "
        "de financement pour certains projets de proximit&eacute;. Ces situations p&egrave;sent sur "
        "l'image du territoire, la qualit&eacute; de vie, l'attractivit&eacute; et la capacit&eacute; &agrave; r&eacute;pondre "
        "aux besoins des habitants.",
    )
    add_paragraph(
        doc,
        "TVF propose une m&eacute;thode de terrain pour transformer ces constats en dossiers "
        "exploitables : identifier les biens et ressources, qualifier les usages possibles, "
        "mobiliser les bons acteurs, rechercher des mat&eacute;riaux ou comp&eacute;tences disponibles, "
        "puis cadrer les suites par convention lorsque le projet est r&eacute;aliste.",
    )
    rows = [
        ("Habitat vacant", "Rep&eacute;rer les logements ou immeubles dormants, identifier les propri&eacute;taires, comprendre les blocages et orienter vers des solutions progressives."),
        ("Commerces inoccup&eacute;s", "R&eacute;fl&eacute;chir &agrave; des usages temporaires, commerces de proximit&eacute;, ateliers, locaux associatifs ou occupations utiles."),
        ("Friches et terrains", "Pr&eacute;qualifier des espaces d&eacute;laiss&eacute;s et imaginer des usages sobres : jardins, lieux partag&eacute;s, activit&eacute;s ESS, espaces p&eacute;dagogiques."),
        ("Mat&eacute;riaux et mobilier", "Valoriser des ressources issues de travaux, locaux publics, entreprises ou dons pour soutenir des projets territoriaux."),
        ("Insertion et engagement", "Relier les projets &agrave; des chantiers participatifs, structures d'insertion, b&eacute;n&eacute;voles ou formations lorsque le cadre le permet."),
    ]
    table(doc, ["Enjeu territorial", "Contribution possible de TVF"], rows, [4.5, 12.0], fill=LIGHT_BLUE)


def tvf_role(doc):
    doc.add_heading(u("Ce que TVF peut apporter &agrave; une collectivit&eacute;"), level=1)
    rows = [
        ("Observation de terrain", "Aider &agrave; structurer une lecture locale des biens vacants, locaux ferm&eacute;s, friches, terrains, mat&eacute;riaux et besoins associatifs."),
        ("Qualification des dossiers", "Transformer un signalement ou une id&eacute;e en dossier lisible : adresse, propri&eacute;t&eacute;, usage possible, pi&egrave;ces, risques, besoins et prochaines &eacute;tapes."),
        ("Mise en relation", "Relier propri&eacute;taires, entreprises, associations, artisans, b&eacute;n&eacute;voles, financeurs potentiels et services concern&eacute;s."),
        ("Cadre de coop&eacute;ration", "Proposer des conventions adapt&eacute;es : diagnostic, stockage, mise &agrave; disposition, projet pilote, r&eacute;emploi ou action territoriale."),
        ("Appui op&eacute;rationnel", "Aider &agrave; mobiliser des mat&eacute;riaux, comp&eacute;tences, locaux, b&eacute;n&eacute;voles ou ressources utiles selon les moyens disponibles."),
        ("Suivi et restitution", "Produire une tra&ccedil;abilit&eacute; minimale : dossiers ouverts, ressources mobilis&eacute;es, projets suivis, points bloquants et r&eacute;sultats qualitatifs."),
    ]
    table(doc, ["Apport TVF", "Utilit&eacute; pour la collectivit&eacute;"], rows, [4.4, 12.1], fill=LIGHT_GREEN)
    add_callout(
        doc,
        "Principe de compl&eacute;mentarit&eacute;",
        "TVF agit en appui et en coordination. L'association ne remplace pas les services "
        "publics, les bureaux d'&eacute;tudes, les bailleurs, les op&eacute;rateurs, les entreprises ou "
        "les dispositifs existants. Elle aide &agrave; relier les acteurs et &agrave; rendre les dossiers "
        "plus faciles &agrave; instruire.",
        fill=LIGHT_BLUE,
    )


def public_policy(doc):
    doc.add_heading(u("Politiques publiques et besoins compatibles"), level=1)
    add_paragraph(
        doc,
        "Selon le territoire, l'action de TVF peut s'inscrire en coh&eacute;rence avec des priorit&eacute;s "
        "publiques d&eacute;j&agrave; engag&eacute;es : revitalisation de centres-villes, habitat ancien, "
        "vacance commerciale, recyclage foncier, sobri&eacute;t&eacute; fonci&egrave;re, &eacute;conomie circulaire, "
        "transition &eacute;cologique, insertion professionnelle, participation citoyenne ou "
        "soutien aux associations. La compatibilit&eacute; doit toujours &ecirc;tre v&eacute;rifi&eacute;e au cas par cas.",
    )
    rows = [
        ("Habitat et logement", "Logements vacants, habitat d&eacute;grad&eacute;, besoins temporaires, interg&eacute;n&eacute;rationnel, h&eacute;bergement associatif ou &eacute;tudiant."),
        ("Revitalisation commerciale", "Locaux ferm&eacute;s, rez-de-chauss&eacute;e vacants, commerces de proximit&eacute;, artisans, boutiques temporaires, tiers-lieux."),
        ("Foncier et friches", "Recyclage d'espaces d&eacute;laiss&eacute;s, usages transitoires, jardins partag&eacute;s, espaces p&eacute;dagogiques, pr&eacute;figuration de projets."),
        ("Economie circulaire", "R&eacute;emploi de mat&eacute;riaux, r&eacute;duction du gaspillage, collecte, tri, stockage, r&eacute;affectation &agrave; des projets locaux."),
        ("Insertion et ESS", "Chantiers participatifs, d&eacute;couverte des m&eacute;tiers, b&eacute;n&eacute;volat, mobilisation associative, parcours vers l'emploi."),
        ("Participation citoyenne", "Signalements, id&eacute;es d'usage, implication d'habitants, rencontres de quartier, actions collectives encadr&eacute;es."),
    ]
    table(doc, ["Axe public", "Contribution possible"], rows, [4.3, 12.2], fill=LIGHT_GREEN)


def what_collectivity_can_offer(doc):
    doc.add_heading(u("Ce que la collectivit&eacute; peut proposer"), level=1)
    rows = [
        ("Interlocuteur r&eacute;f&eacute;rent", "D&eacute;signer un service ou une personne capable d'orienter TVF vers les bons contacts : habitat, foncier, commerce, ESS, d&eacute;chets, insertion, technique."),
        ("Donn&eacute;es et constats", "Partager des constats non sensibles, diagnostics existants, p&eacute;rim&egrave;tres, besoins locaux, signalements ou priorit&eacute;s territoriales."),
        ("Biens ou espaces", "Identifier des locaux vacants, b&acirc;timents sous-utilis&eacute;s, terrains, friches, espaces de stockage, ateliers ou sites pilotes possibles."),
        ("Mat&eacute;riaux et mobilier", "Proposer du mobilier public inutilis&eacute;, &eacute;quipements, mat&eacute;riaux issus de travaux ou ressources encore exploitables."),
        ("Mise en relation", "Orienter vers communes membres, services, bailleurs, associations, entreprises, structures d'insertion, chambres consulaires ou acteurs locaux."),
        ("Cadre administratif", "Etudier les conventions, autorisations, mises &agrave; disposition, proc&eacute;dures, d&eacute;cisions et limites juridiques applicables."),
        ("Soutien op&eacute;rationnel", "Appui logistique, local temporaire, relais de communication, coordination de r&eacute;union, aide &agrave; l'identification de financements."),
    ]
    table(doc, ["Contribution possible", "Utilit&eacute;"], rows, [4.5, 12.0], fill=LIGHT_BLUE)


def benefits(doc):
    doc.add_heading(u("B&eacute;n&eacute;fices pour le territoire"), level=1)
    rows = [
        ("Acc&eacute;l&eacute;ration", "Passer plus vite du constat au dossier : lieu identifi&eacute;, besoin qualifi&eacute;, acteurs list&eacute;s, premi&egrave;res options cadr&eacute;es."),
        ("Coordination", "Eviter que chaque sujet reste isol&eacute; entre habitat, commerce, d&eacute;chets, foncier, insertion, associations et propri&eacute;taires."),
        ("Sobri&eacute;t&eacute;", "R&eacute;utiliser des lieux, mat&eacute;riaux et &eacute;quipements existants avant de consommer de nouvelles ressources."),
        ("Attractivit&eacute;", "R&eacute;duire l'image d'abandon, soutenir les centralit&eacute;s, r&eacute;activer des locaux et donner des usages visibles aux espaces d&eacute;laiss&eacute;s."),
        ("Utilit&eacute; sociale", "Orienter les projets vers des besoins concrets : logement, association, formation, insertion, services de proximit&eacute;, espaces partag&eacute;s."),
        ("Tra&ccedil;abilit&eacute;", "Suivre les dossiers, ressources mobilis&eacute;es, blocages, usages cr&eacute;&eacute;s et points &agrave; arbitrer."),
    ]
    table(doc, ["B&eacute;n&eacute;fice", "Effet attendu"], rows, [4.4, 12.1], fill=LIGHT_GREEN)


def conventions(doc):
    doc.add_heading(u("Formes de coop&eacute;ration et conventions possibles"), level=1)
    rows = [
        ("Convention de coop&eacute;ration territoriale", "Cadre g&eacute;n&eacute;ral pour organiser les relations, les objectifs, le p&eacute;rim&egrave;tre, les r&ocirc;les, la dur&eacute;e et les limites de communication.", "6 &agrave; 36 mois selon projet"),
        ("Autorisation de diagnostic ou rep&eacute;rage", "Permettre &agrave; TVF d'&eacute;tudier un site, un besoin ou une ressource, sans engagement de r&eacute;alisation.", "Ponctuelle ou quelques mois"),
        ("Convention de mise &agrave; disposition", "Mettre temporairement &agrave; disposition un local, terrain, espace de stockage ou &eacute;quipement, avec responsabilit&eacute;s et usages pr&eacute;cis.", "Au cas par cas"),
        ("Convention de stockage", "Cadrer un espace de tri, stockage ou qualification des mat&eacute;riaux : acc&egrave;s, s&eacute;curit&eacute;, assurance, dur&eacute;e, volume, gestion.", "3 &agrave; 24 mois selon besoin"),
        ("Convention de projet pilote", "Tester un mode d'action sur un p&eacute;rim&egrave;tre limit&eacute; : quartier, rue, centre-bourg, local, friche ou ressource identifi&eacute;e.", "6 &agrave; 18 mois"),
        ("Convention de soutien financier", "Soutenir une action ou une mission lorsque le cadre budg&eacute;taire, juridique et d&eacute;cisionnel de la collectivit&eacute; le permet.", "Selon d&eacute;cision comp&eacute;tente"),
        ("Convention de communication encadr&eacute;e", "D&eacute;finir ce qui peut &ecirc;tre rendu public : logo, mention, photo, bilan, &eacute;v&eacute;nement, partenaires, sans annonce non valid&eacute;e.", "Li&eacute;e au projet"),
        ("Protocole de suivi", "Organiser les points d'&eacute;tape, indicateurs, comptes rendus, documents, arbitrages et sortie de convention.", "Pendant toute l'action"),
    ]
    table(doc, ["Cadre", "Objet", "Dur&eacute;e indicative"], rows, [4.6, 8.0, 3.9], fill=LIGHT_BLUE)
    note(doc, "Les dur&eacute;es sont indicatives. Elles doivent &ecirc;tre adapt&eacute;es au droit applicable, aux d&eacute;cisions de la collectivit&eacute;, aux biens concern&eacute;s, aux assurances et aux moyens mobilisables.")


def pathway(doc):
    doc.add_heading(u("Parcours type d'un dossier territorial"), level=1)
    rows = [
        ("1. Premier &eacute;change", "Identifier le besoin public, le service concern&eacute;, le p&eacute;rim&egrave;tre et les attentes de la collectivit&eacute;."),
        ("2. Fiche collectivit&eacute;", "Renseigner les informations de base : collectivit&eacute;, service, interlocuteur, sujet, lieu, documents disponibles."),
        ("3. Qualification", "Analyser le type de sujet : logement, commerce, friche, mat&eacute;riaux, local, association, insertion, stockage, projet pilote."),
        ("4. Pi&egrave;ces et contraintes", "Rassembler photos, plans, diagnostics, propri&eacute;t&eacute;, contraintes juridiques, accessibilit&eacute;, assurance, risques et calendrier."),
        ("5. Sc&eacute;nario d'action", "Proposer une ou plusieurs pistes : diagnostic, usage temporaire, collecte de mat&eacute;riaux, mise en relation, convention, chantier."),
        ("6. Validation interne", "La collectivit&eacute; v&eacute;rifie ses proc&eacute;dures, comp&eacute;tences, autorisations, d&eacute;cisions, budget et contraintes de commande publique."),
        ("7. Convention", "Formaliser les engagements, limites, responsabilit&eacute;s, dur&eacute;e, communication, suivi et conditions de sortie."),
        ("8. Mise en oeuvre", "Coordonner les acteurs, suivre les ressources, organiser les rendez-vous, documenter les d&eacute;cisions et les actions."),
        ("9. Bilan", "Restituer les r&eacute;sultats qualitatifs, les blocages, les suites possibles et les documents du dossier."),
    ]
    table(doc, ["Etape", "Action"], rows, [3.8, 12.7], fill=LIGHT_GREEN)


def criteria_limits(doc):
    doc.add_heading(u("Crit&egrave;res de s&eacute;lection et limites d'intervention"), level=1)
    rows = [
        ("Int&eacute;r&ecirc;t territorial", "Le dossier doit r&eacute;pondre &agrave; un besoin r&eacute;el : logement, local, commerce, association, friche, mat&eacute;riaux, insertion ou cadre de vie."),
        ("Clart&eacute; juridique", "Propri&eacute;t&eacute;, comp&eacute;tence, autorisation, assurance, responsabilit&eacute;s et usage doivent &ecirc;tre suffisamment clarifi&eacute;s."),
        ("Faisabilit&eacute;", "Le projet doit &ecirc;tre proportionn&eacute; aux moyens disponibles : ressources, partenaires, budget, temps, s&eacute;curit&eacute;, logistique."),
        ("Travaux raisonnables", "TVF privil&eacute;gie la remise en usage, l'am&eacute;nagement, le second oeuvre et le r&eacute;emploi. Les travaux lourds doivent &ecirc;tre port&eacute;s par des professionnels et un financement adapt&eacute;."),
        ("Tra&ccedil;abilit&eacute;", "Chaque action doit pouvoir &ecirc;tre document&eacute;e : dossier, photos, pi&egrave;ces, convention, d&eacute;cisions, destination, suivi."),
        ("Non-acceptation", "Une demande de collectivit&eacute; ne vaut pas acceptation automatique, partenariat public, financement ou engagement de r&eacute;alisation."),
    ]
    table(doc, ["Crit&egrave;re", "Lecture TVF"], rows, [4.3, 12.2], fill=LIGHT_BLUE)

    doc.add_heading(u("Ce que TVF ne fait pas"), level=2)
    for item in [
        "TVF ne remplace pas une ma&icirc;trise d'ouvrage publique.",
        "TVF ne se substitue pas aux services techniques, juridiques, urbanisme, habitat ou commande publique.",
        "TVF ne garantit pas l'obtention de subventions, partenaires, locaux, mat&eacute;riaux ou financements.",
        "TVF ne communique pas sur un soutien public sans accord &eacute;crit pr&eacute;alable.",
        "TVF ne prend pas en charge les travaux dangereux, structurels ou non assurables sans professionnels qualifi&eacute;s.",
    ]:
        add_bullet(doc, item)


def storage_materials(doc):
    doc.add_heading(u("Stockage, mat&eacute;riaux et locaux techniques"), level=1)
    add_paragraph(
        doc,
        "Le lancement d'un dispositif de r&eacute;emploi n&eacute;cessite souvent un lieu de stockage "
        "temporaire, s&eacute;curis&eacute; et accessible. Une collectivit&eacute; peut aider TVF &agrave; identifier "
        "un espace adapt&eacute; ou &agrave; mettre en relation l'association avec un acteur public, "
        "associatif ou priv&eacute; disposant d'une solution logistique.",
    )
    rows = [
        ("Local de stockage", "Espace couvert ou s&eacute;curis&eacute;, acc&egrave;s facile, possibilit&eacute; de tri, dur&eacute;e d'occupation, assurance et responsabilit&eacute;s &agrave; cadrer."),
        ("Espace de tri temporaire", "Zone ponctuelle pour qualifier, photographier, mesurer, classer et pr&eacute;parer les ressources avant orientation."),
        ("Mat&eacute;riaux publics inutilis&eacute;s", "Mobilier, &eacute;quipements, surplus de travaux, &eacute;l&eacute;ments d'am&eacute;nagement encore utilisables, sous r&eacute;serve de propri&eacute;t&eacute; et d'&eacute;tat."),
        ("Conditions d'usage", "Acc&egrave;s, horaires, cl&eacute;s, s&eacute;curit&eacute;, stationnement, manutention, volumes, interdictions, inventaire, sortie de site."),
        ("Tra&ccedil;abilit&eacute;", "Registre minimal : date, origine, cat&eacute;gorie, quantit&eacute;, &eacute;tat, photos, destination, observations."),
    ]
    table(doc, ["Sujet", "Conditions &agrave; cadrer"], rows, [4.2, 12.3], fill=LIGHT_GREEN)


def documents_needed(doc):
    doc.add_heading(u("Pi&egrave;ces et informations &agrave; fournir"), level=1)
    rows = [
        ("Identification", "Nom de la collectivit&eacute;, service concern&eacute;, adresse, interlocuteur, fonction, t&eacute;l&eacute;phone, e-mail."),
        ("Objet de la demande", "Diagnostic, local de stockage, bien vacant, commerce, friche, mat&eacute;riaux, mise en relation, projet pilote, autre besoin."),
        ("P&eacute;rim&egrave;tre", "Adresse, quartier, commune, parcelle, b&acirc;timent, rue, centre-bourg ou zone concern&eacute;e."),
        ("Documents disponibles", "Photos, plans, diagnostics, fiches techniques, notes internes, constats, documents publics, contraintes connues."),
        ("Propri&eacute;t&eacute; et comp&eacute;tence", "Pr&eacute;ciser si le bien ou l'espace appartient &agrave; la collectivit&eacute;, &agrave; un tiers, &agrave; un bailleur, &agrave; une copropri&eacute;t&eacute; ou &agrave; un propri&eacute;taire priv&eacute;."),
        ("Contraintes", "S&eacute;curit&eacute;, assurance, occupation, d&eacute;lais, travaux, risques, accessibilit&eacute;, servitudes, proc&eacute;dures, budget, calendrier."),
        ("Acteurs &agrave; associer", "Services internes, communes, bailleurs, associations, entreprises, structures d'insertion, acteurs ESS, chambres consulaires, propri&eacute;taires."),
        ("Validation attendue", "Rendez-vous, diagnostic, note de cadrage, convention, mise en relation, rep&eacute;rage, fiche projet, plan d'action."),
    ]
    table(doc, ["Information", "D&eacute;tail attendu"], rows, [4.3, 12.2], fill=LIGHT_BLUE)


def collectivity_form(doc):
    doc.add_page_break()
    doc.add_heading(u("Fiche collectivit&eacute; - demande de coop&eacute;ration territoriale"), level=1)
    add_paragraph(doc, "Cette fiche permet d'ouvrir une instruction avec TVF. Elle ne vaut pas partenariat public, convention, financement, validation politique ou engagement de r&eacute;alisation.")
    rows = [
        ("Date", "................................................................................................"),
        ("Collectivit&eacute;", "................................................................................................"),
        ("Type", "[ ] Commune   [ ] Intercommunalit&eacute;   [ ] D&eacute;partement   [ ] R&eacute;gion   [ ] Syndicat   [ ] Autre"),
        ("Service r&eacute;f&eacute;rent", "................................................................................................"),
        ("Interlocuteur", "Nom : ........................................ Fonction : ........................................"),
        ("T&eacute;l&eacute;phone / e-mail", "................................................................................................"),
        ("Objet principal", "[ ] Habitat vacant   [ ] Commerce   [ ] Friche   [ ] Mat&eacute;riaux   [ ] Stockage   [ ] Insertion   [ ] Projet pilote"),
        ("Adresse / p&eacute;rim&egrave;tre", "................................................................................................"),
        ("Description du besoin", "................................................................................................\n................................................................................................"),
        ("Documents disponibles", "[ ] Photos   [ ] Plans   [ ] Diagnostics   [ ] Note interne   [ ] Donn&eacute;es publiques   [ ] Autre"),
        ("Propri&eacute;t&eacute; du bien", "[ ] Collectivit&eacute;   [ ] Propri&eacute;taire priv&eacute;   [ ] Bailleur   [ ] Copropri&eacute;t&eacute;   [ ] A v&eacute;rifier"),
        ("Niveau de maturit&eacute;", "[ ] Id&eacute;e   [ ] Besoin identifi&eacute;   [ ] Site rep&eacute;r&eacute;   [ ] Projet en cadrage   [ ] Convention &agrave; &eacute;tudier"),
        ("Besoin TVF attendu", "[ ] Diagnostic   [ ] Mise en relation   [ ] Mat&eacute;riaux   [ ] Local stockage   [ ] Fiche projet   [ ] Convention"),
        ("Contraintes connues", "................................................................................................"),
        ("D&eacute;lai souhait&eacute;", "................................................................................................"),
        ("Acteurs &agrave; associer", "................................................................................................"),
        ("Validation interne", "[ ] Echange technique   [ ] Arbitrage service   [ ] D&eacute;cision &eacute;lu/service   [ ] D&eacute;lib&eacute;ration si n&eacute;cessaire"),
        ("D&eacute;claration", "Je comprends que cette demande ne vaut pas acceptation, ne vaut pas convention et ne vaut pas communication publique autoris&eacute;e."),
        ("Signature / cachet", "................................................................................................"),
    ]
    table(doc, ["Champ", "R&eacute;ponse / information &agrave; renseigner"], rows, [4.7, 11.8], fill=LIGHT_GREEN)
    add_callout(
        doc,
        "Contact TVF",
        "TERRITOIRES VIVANTS FRANCE - 25 rue &Eacute;lise Gervais, 42000 Saint-&Eacute;tienne - contact@territoiresvivantsfrance.fr - 04 65 81 54 69 - www.territoiresvivantsfrance.fr",
        fill=LIGHT_BLUE,
    )


def final_notes(doc):
    doc.add_heading(u("Points &agrave; verrouiller avant signature"), level=1)
    for item in [
        "Comp&eacute;tence de la collectivit&eacute; et service responsable.",
        "Propri&eacute;t&eacute; du bien, autorisation d'acc&egrave;s ou droit de mise &agrave; disposition.",
        "Assurances, responsabilit&eacute;s, s&eacute;curit&eacute; des lieux et conditions d'acc&egrave;s.",
        "Cadre de convention : objet, dur&eacute;e, r&ocirc;les, limites, suivi, sortie.",
        "Budget, soutien financier ou mise &agrave; disposition valid&eacute;s selon les proc&eacute;dures internes.",
        "Commande publique, subventions, communication et logo valid&eacute;s par les services comp&eacute;tents.",
        "Absence de communication publique avant accord &eacute;crit.",
    ]:
        add_bullet(doc, "[ ] " + item)
    add_callout(
        doc,
        "Conclusion",
        "TVF peut devenir un outil simple de passage &agrave; l'action : partir d'un besoin local, "
        "qualifier un dossier, relier les acteurs, mobiliser des ressources, cadrer une convention "
        "et suivre les r&eacute;sultats. La force du dispositif repose sur une coop&eacute;ration claire, "
        "prudente et utile au territoire.",
        fill=LIGHT_GREEN,
    )


def build_docx():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    cover(doc)
    toc(doc)
    doc.add_page_break()
    why(doc)
    tvf_role(doc)
    public_policy(doc)
    what_collectivity_can_offer(doc)
    benefits(doc)
    conventions(doc)
    pathway(doc)
    criteria_limits(doc)
    storage_materials(doc)
    documents_needed(doc)
    collectivity_form(doc)
    final_notes(doc)
    doc.save(DOCX_PATH)


def build_md():
    md = u(
        """# Brochure TVF - Collectivit&eacute;s

Document d'information destin&eacute; aux communes, intercommunalit&eacute;s, D&eacute;partements, R&eacute;gions et services publics souhaitant &eacute;tudier une coop&eacute;ration avec TERRITOIRES VIVANTS FRANCE.

## Message central

TVF ne remplace pas les politiques publiques locales. L'association peut aider &agrave; rep&eacute;rer, qualifier, relier et structurer des projets autour du patrimoine vacant, des commerces inoccup&eacute;s, des friches, des ressources inutilis&eacute;es et du r&eacute;emploi.

## Principes

- Chaque demande ne vaut pas acceptation.
- Aucune communication publique ne doit &ecirc;tre faite sans accord &eacute;crit.
- Les conventions doivent respecter les comp&eacute;tences, proc&eacute;dures internes, r&egrave;gles budg&eacute;taires et obligations de la collectivit&eacute;.
- TVF agit en coordination, sans se substituer aux services publics, ma&icirc;tres d'ouvrage, bureaux d'&eacute;tudes ou op&eacute;rateurs.

## Contact

TERRITOIRES VIVANTS FRANCE - contact@territoiresvivantsfrance.fr - 04 65 81 54 69 - www.territoiresvivantsfrance.fr
"""
    )
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(DOCX_PATH)
    print(MD_PATH)
