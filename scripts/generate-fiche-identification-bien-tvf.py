from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

ROOT = Path(r'C:\Users\jowst\Documents\TERRITOIRES VIVANTS FRANCE')
OUT = ROOT / 'documents' / 'formulaires-internes-tvf'
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / 'fiche-interne-identification-bien-tvf.docx'
PDF = OUT / 'fiche-interne-identification-bien-tvf.pdf'
LOGO = ROOT / 'assets' / 'logo-territoires-vivants-france-web.png'

GREEN = RGBColor(24,57,47)
GRAY = RGBColor(102,112,133)
DARK = RGBColor(31,41,55)
PALE = 'F3F7F2'
BORDER = 'D9E0DC'

TVF = {
    'name':'Territoires Vivants France',
    'subtitle':'Agence Territoriale de Revitalisation Immobilière',
    'address':'25 rue Élise Gervais, 42000 Saint-Étienne',
    'phone':'04 65 81 54 69',
    'email':'contact@territoiresvivantsfrance.fr',
    'site':'www.territoiresvivantsfrance.fr',
    'rna':'RNA W922015538',
    'siret':'SIRET 897 226 138 00018',
}

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right'):
        tag = 'w:' + edge
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), BORDER)

def set_text(cell, text, bold=False, color=DARK, size=9.2):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Inter'
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(cell)

def label_line(label, height=0.32):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(2.15)
    t.columns[1].width = Inches(4.6)
    c1, c2 = t.rows[0].cells
    set_text(c1, label, True, GREEN, 8.8)
    set_text(c2, ' ', False, DARK, 8.8)
    t.rows[0].height = Inches(height)
    return t

def checkbox_grid(title, options, cols=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.name = 'Manrope'; r.font.size = Pt(10.5); r.bold = True; r.font.color.rgb = GREEN
    rows = (len(options) + cols - 1)//cols
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i in range(cols): t.columns[i].width = Inches(6.75/cols)
    idx = 0
    for row in t.rows:
        for cell in row.cells:
            text = '☐ ' + options[idx] if idx < len(options) else ''
            set_text(cell, text, False, DARK, 8.7)
            idx += 1
    return t

def section(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = 'Manrope'; r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = GREEN
    return p

# DOCX
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.42); sec.bottom_margin = Inches(0.45); sec.left_margin = Inches(0.55); sec.right_margin = Inches(0.55)

h = doc.add_table(rows=1, cols=2)
h.alignment = WD_TABLE_ALIGNMENT.CENTER
h.autofit = False
h.columns[0].width = Inches(2.25); h.columns[1].width = Inches(4.5)
left, right = h.rows[0].cells
if LOGO.exists():
    p = left.paragraphs[0]
    p.add_run().add_picture(str(LOGO), width=Inches(1.55))
else:
    set_text(left, TVF['name'], True, GREEN, 11)
right.text = ''
p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for line, bold, size, color in [
    (TVF['name'], True, 10, GREEN), (TVF['subtitle'], True, 8.6, GREEN),
    (TVF['address'], False, 8, GRAY), (f"{TVF['phone']} · {TVF['email']}", False, 8, GRAY),
    (f"{TVF['rna']} · {TVF['siret']}", False, 8, GRAY), (TVF['site'], False, 8, GRAY)
]:
    r = p.add_run(line); r.font.name='Inter'; r.bold=bold; r.font.size=Pt(size); r.font.color.rgb=color; p.add_run('\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FICHE INTERNE TVF — IDENTIFICATION D’UN BIEN')
r.font.name='Manrope'; r.font.size=Pt(17); r.bold=True; r.font.color.rgb=GREEN
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Stade : identification et recherche initiale du bien — usage interne TVF')
r.font.name='Inter'; r.font.size=Pt(9.5); r.font.color.rgb=GRAY

note = doc.add_table(rows=1, cols=1)
note.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = note.cell(0,0); shade(cell, PALE); borders(cell)
set_text(cell, "Information interne : cette fiche sert uniquement au repérage et à l’identification initiale. Elle ne constitue ni une preuve juridique de vacance, ni une décision d’accompagnement, ni une autorisation d’accès au bien.", False, DARK, 8.5)

section('1. Référence interne')
label_line('Numéro interne TVF')
label_line('Date de création')
label_line('Agent / bénévole TVF')

section('2. Localisation du bien')
label_line('Adresse complète')
label_line('Commune')
label_line('Code postal')
label_line('Quartier / secteur')
label_line('Coordonnées GPS si connues')

section('Informations cadastrales si trouvées')
label_line('Référence cadastrale')
label_line('Section / parcelle')
label_line('Source de l’information')

section('3. Nature du bien')
checkbox_grid('', ['Maison', 'Appartement', 'Immeuble', 'Local commercial', 'Bâtiment d’activité', 'Friche', 'Terrain', 'Autre : ____________________'], 2)

section('4. État apparent')
checkbox_grid('', ['Bon état apparent', 'État à vérifier', 'Dégradé', 'Très dégradé', 'Risque visible', 'Photos disponibles'], 2)

section('5. Situation observée')
checkbox_grid('', ['Bien apparemment vacant', 'Volets fermés', 'Local sans activité', 'Dégradation visible', 'Terrain délaissé', 'Occupation inconnue', 'À confirmer', 'Autre : ____________________'], 2)

section('6. Origine de l’information')
checkbox_grid('', ['Propriétaire', 'Habitant', 'Collectivité', 'Association', 'Entreprise', 'Observation terrain', 'TVF Mobile', 'Formulaire site', 'E-mail', 'Téléphone', 'Courrier papier', 'Autre : ____________________'], 3)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(7)
r = p.add_run('Rappel interne : ne pas pénétrer dans une propriété privée sans autorisation. Les photos doivent être prises légalement depuis l’espace public ou avec accord du propriétaire ou de son représentant.')
r.font.name='Inter'; r.font.size=Pt(8); r.font.color.rgb=GRAY

doc.save(DOCX)

# PDF
styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name='TitleTVF', fontName='Helvetica-Bold', fontSize=17, leading=20, textColor=colors.HexColor('#18392F'), alignment=1, spaceAfter=4))
styles.add(ParagraphStyle(name='SubTVF', fontName='Helvetica', fontSize=8.8, leading=11, textColor=colors.HexColor('#667085'), alignment=1, spaceAfter=6))
styles.add(ParagraphStyle(name='H', fontName='Helvetica-Bold', fontSize=10.8, leading=13, textColor=colors.HexColor('#18392F'), spaceBefore=7, spaceAfter=4))
styles.add(ParagraphStyle(name='Cell', fontName='Helvetica', fontSize=8, leading=9.5, textColor=colors.HexColor('#1F2937')))
styles.add(ParagraphStyle(name='Head', fontName='Helvetica-Bold', fontSize=8, leading=9.5, textColor=colors.HexColor('#18392F')))
styles.add(ParagraphStyle(name='Small', fontName='Helvetica', fontSize=7.4, leading=9, textColor=colors.HexColor('#667085')))

story=[]
story.append(Paragraph('FICHE INTERNE TVF — IDENTIFICATION D’UN BIEN', styles['TitleTVF']))
story.append(Paragraph('Territoires Vivants France · Agence Territoriale de Revitalisation Immobilière<br/>25 rue Élise Gervais, 42000 Saint-Étienne · 04 65 81 54 69 · contact@territoiresvivantsfrance.fr<br/>RNA W922015538 · SIRET 897 226 138 00018 · www.territoiresvivantsfrance.fr', styles['SubTVF']))

def pdf_table(rows, widths=None, fill_header=False):
    data=[]
    for row in rows:
        data.append([Paragraph(str(x), styles['Head' if fill_header else 'Cell']) for x in row])
    table=Table(data, colWidths=widths or [45*mm, 125*mm])
    cmds=[('GRID',(0,0),(-1,-1),0.4,colors.HexColor('#D9E0DC')),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LEFTPADDING',(0,0),(-1,-1),5),('RIGHTPADDING',(0,0),(-1,-1),5),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)]
    if fill_header: cmds.append(('BACKGROUND',(0,0),(-1,-1),colors.HexColor('#F3F7F2')))
    table.setStyle(TableStyle(cmds))
    story.append(table)

story.append(Paragraph('Information interne : cette fiche sert uniquement au repérage et à l’identification initiale. Elle ne constitue ni une preuve juridique de vacance, ni une décision d’accompagnement, ni une autorisation d’accès au bien.', styles['Small']))
story.append(Spacer(1,4))
for title, labels in [
    ('1. Référence interne', ['Numéro interne TVF','Date de création','Agent / bénévole TVF']),
    ('2. Localisation du bien', ['Adresse complète','Commune','Code postal','Quartier / secteur','Coordonnées GPS si connues']),
    ('Informations cadastrales si trouvées', ['Référence cadastrale','Section / parcelle','Source de l’information'])
]:
    story.append(Paragraph(title, styles['H']))
    pdf_table([[l,''] for l in labels])

def pdf_checks(title, options, cols):
    story.append(Paragraph(title, styles['H']))
    rows=[]
    for i in range(0,len(options),cols):
        rows.append(['☐ '+x for x in options[i:i+cols]] + ['']*(cols-len(options[i:i+cols])))
    pdf_table(rows, [170*mm/cols]*cols)

pdf_checks('3. Nature du bien', ['Maison','Appartement','Immeuble','Local commercial','Bâtiment d’activité','Friche','Terrain','Autre : ____________________'], 2)
pdf_checks('4. État apparent', ['Bon état apparent','État à vérifier','Dégradé','Très dégradé','Risque visible','Photos disponibles'], 2)
pdf_checks('5. Situation observée', ['Bien apparemment vacant','Volets fermés','Local sans activité','Dégradation visible','Terrain délaissé','Occupation inconnue','À confirmer','Autre : ____________________'], 2)
pdf_checks('6. Origine de l’information', ['Propriétaire','Habitant','Collectivité','Association','Entreprise','Observation terrain','TVF Mobile','Formulaire site','E-mail','Téléphone','Courrier papier','Autre : ____________________'], 3)
story.append(Paragraph('Rappel interne : ne pas pénétrer dans une propriété privée sans autorisation. Les photos doivent être prises légalement depuis l’espace public ou avec accord du propriétaire ou de son représentant.', styles['Small']))

docpdf = SimpleDocTemplate(str(PDF), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=12*mm, bottomMargin=12*mm)
docpdf.build(story)
print(DOCX)
print(PDF)


